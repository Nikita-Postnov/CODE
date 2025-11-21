import os
import sys
import json
import re
import string
import html
import random
import difflib
import telethon
import requests
import tkinter as tk
from tkinter import ttk
import wave
import traceback
import base64
import html as html_lib
from uuid import uuid4
from cryptography.fernet import Fernet, InvalidToken
from cryptography.hazmat.primitives import hashes
from cryptography.hazmat.primitives.kdf.pbkdf2 import PBKDF2HMAC
from cryptography.hazmat.backends import default_backend
import time
import math
import tempfile
import uuid
import shutil
import datetime
import importlib.util
from urllib.parse import quote, unquote
from typing import Callable
import tkinter.messagebox as tk_messagebox
from tkinter import filedialog as tk_filedialog
from functools import partial
import pyperclip
import subprocess
import importlib
from PySide6.QtCore import (
    Qt,
    QMimeData,
    QTimer,
    QUrl,
    QDate,
    QSettings,
    QSize,
    QEvent,
    QDateTime,
    QPoint,
    QRect,
    QThread,
    Signal,
    QBuffer,
    QPointF,
    QIODevice,
    QFileSystemWatcher,
    QProcess,
)
from PySide6.QtGui import (
    QIcon,
    QBrush,
    QTextTableCellFormat,
    QTextFrameFormat,
    QFontMetrics,
    QDragEnterEvent,
    QDropEvent,
    QTextTableFormat,
    QMouseEvent,
    QContextMenuEvent,
    QCursor,
    QShortcut,
    QPainterPath,
    QImageReader,
    QFont,
    QImage,
    QPixmap,
    QKeySequence,
    QCloseEvent,
    QDrag,
    QAction,
    QPen,
    QColor,
    QTextCharFormat,
    QDesktopServices,
    QTextCursor,
    QPolygonF,
    QTransform,
    QTextListFormat,
    QTextBlockFormat,
    QPainter,
    QPalette,
    QTextDocument,
    QSyntaxHighlighter,
)
from PySide6.QtWidgets import (
    QApplication,
    QSlider,
    QPlainTextEdit,
    QLayoutItem,
    QStyle,
    QGraphicsView,
    QGraphicsScene,
    QGraphicsPixmapItem,
    QMainWindow,
    QTextEdit,
    QProgressBar,
    QVBoxLayout,
    QHBoxLayout,
    QFormLayout,
    QTextBrowser,
    QFileDialog,
    QPushButton,
    QListWidget,
    QListWidgetItem,
    QWidget,
    QLabel,
    QSizePolicy,
    QLineEdit,
    QMessageBox,
    QInputDialog,
    QColorDialog,
    QSpinBox,
    QScrollArea,
    QLayout,
    QStyleFactory,
    QToolTip,
    QFontComboBox,
    QDialog,
    QButtonGroup,
    QDialogButtonBox,
    QComboBox,
    QCheckBox,
    QSystemTrayIcon,
    QAbstractItemView,
    QMenu,
    QDateTimeEdit,
    QSplitter,
    QToolBar,
    QDockWidget,
    QToolButton,
    QFrame,
    QSizeGrip,
)

keyboard = None
_keyboard_spec = importlib.util.find_spec("keyboard")
if _keyboard_spec is not None:
    keyboard = importlib.util.module_from_spec(_keyboard_spec)
    _keyboard_spec.loader.exec_module(keyboard)


class MessageBox:
    @staticmethod
    def showerror(title, message):
        QMessageBox.critical(None, title, message)

    @staticmethod
    def showinfo(title, message):
        QMessageBox.information(None, title, message)

    @staticmethod
    def askyesno(title, message):
        return (
            QMessageBox.question(None, title, message, QMessageBox.Yes | QMessageBox.No)
            == QMessageBox.Yes
        )

    @staticmethod
    def showwarning(title, message):
        QMessageBox.warning(None, title, message)


messagebox = MessageBox
if getattr(sys, "frozen", False):
    APPDIR = os.path.dirname(sys.executable)
    try:
        os.chdir(APPDIR)
    except Exception:
        pass
else:
    APPDIR = os.path.abspath(os.path.dirname(__file__))

DATA_DIR = os.path.join(APPDIR, "Data")
NOTES_DIR = os.path.join(APPDIR, "Notes")
PASSWORDS_DIR = os.path.join(APPDIR, "Passwords")
SALT_PATH = os.path.join(DATA_DIR, "salt.bin")
SETTINGS_PATH = os.path.join(DATA_DIR, "settings.ini")
ICON_PATH = os.path.join(DATA_DIR, "icon.ico")
TRAY_ICON_PATH = os.path.join(DATA_DIR, "tray_icon.ico")
FILE_ICON_PATH = os.path.join(DATA_DIR, "file.ico")
USER_DICT_PATH = os.path.join(DATA_DIR, "user_dictionary.txt")
NOTES_WIDGET_PATH = os.path.join(DATA_DIR, "desktop_notes.json")
NOTES_WIDGET_BAK = os.path.join(DATA_DIR, "desktop_notes.json.bak")
RU_LAYOUT = "ёйцукенгшщзхъфывапролджэячсмитьбю"
EN_LAYOUT = "`qwertyuiop[]asdfghjkl;'zxcvbnm,."
RU_TO_EN = str.maketrans(RU_LAYOUT + RU_LAYOUT.upper(), EN_LAYOUT + EN_LAYOUT.upper())
EN_TO_RU = str.maketrans(EN_LAYOUT + EN_LAYOUT.upper(), RU_LAYOUT + RU_LAYOUT.upper())
RU_REFLEXIVE = ("ся", "сь")
RU_SUFFIXES_RU = tuple(
    sorted(
        {
            "иями",
            "ями",
            "ами",
            "ов",
            "ев",
            "ёв",
            "ам",
            "ям",
            "ах",
            "ях",
            "ою",
            "ею",
            "ой",
            "ей",
            "ый",
            "ий",
            "ой",
            "ая",
            "яя",
            "ое",
            "ее",
            "ые",
            "ие",
            "ого",
            "его",
            "ому",
            "ему",
            "ым",
            "им",
            "ых",
            "их",
            "ую",
            "юю",
            "а",
            "я",
            "ы",
            "и",
            "е",
            "у",
            "ю",
            "о",
            "ешь",
            "ешься",
            "ет",
            "ется",
            "ем",
            "емся",
            "ете",
            "етесь",
            "ут",
            "ют",
            "ат",
            "ят",
            "ишь",
            "ит",
            "им",
            "ите",
            "ите-ка",
            "ил",
            "ила",
            "ило",
            "или",
            "лся",
            "лась",
            "лось",
            "лись",
            "вший",
            "щий",
            "вшая",
            "вшее",
            "вшие",
        },
        key=len,
        reverse=True,
    )
)

RU_VARIANT_RULES: tuple[tuple[str, tuple[str]]] = tuple(
    sorted(
        [
            ("ого", ("ый", "ий", "ой")),
            ("ему", ("ый", "ий", "ой")),
            ("ому", ("ый", "ий", "ой")),
            ("ыми", ("ый", "ий", "ой")),
            ("ими", ("ий",)),
            ("ых", ("ые", "ие")),
            ("их", ("ые", "ие")),
            ("ая", ("ый", "ий", "ой")),
            ("яя", ("ий",)),
            ("ое", ("ый", "ий", "ой")),
            ("ее", ("ий",)),
            ("ые", ("ый", "ий", "ой")),
            ("ие", ("ий",)),
            ("ою", ("ая", "яя")),
            ("ею", ("ая", "яя")),
            ("ую", ("ая",)),
            ("юю", ("яя",)),
            ("ами", ("а", "")),
            ("ями", ("я", "")),
            ("ам", ("а", "")),
            ("ям", ("я", "")),
            ("ах", ("а", "")),
            ("ях", ("я", "")),
            ("ов", ("",)),
            ("ев", ("",)),
            ("ёв", ("",)),
            ("ей", ("я", "е", "")),
            ("ой", ("а", "о", "")),
            ("ешься", ("ться",)),
            ("ется", ("ться",)),
            ("етесь", ("ться",)),
            ("лся", ("ться",)),
            ("лась", ("ться",)),
            ("лось", ("ться",)),
            ("лись", ("ться",)),
            ("ешь", ("ть",)),
            ("ет", ("ть",)),
            ("ем", ("ть",)),
            ("ете", ("ть",)),
            ("ут", ("ть",)),
            ("ют", ("ть",)),
            ("ат", ("ть",)),
            ("ят", ("ть",)),
            ("ишь", ("ить", "ть")),
            ("ит", ("ить", "ть")),
            ("им", ("ить",)),
            ("ите", ("ить", "ть")),
            ("ил", ("ить", "ть")),
            ("ила", ("ить", "ть")),
            ("ило", ("ить", "ть")),
            ("или", ("ить", "ть")),
            ("вшийся", ("ться",)),
            ("вшая", ("ть",)),
            ("вшее", ("ть",)),
            ("вшие", ("ть",)),
            ("вшись", ("ться",)),
            ("вши", ("ть",)),
            ("ши", ("ть",)),
            ("в", ("ть",)),
        ],
        key=lambda p: len(p[0]),
        reverse=True,
    )
)
RU_SUFFIXES_RU = tuple(
    sorted(
        {
            "иями",
            "ями",
            "ами",
            "ями",
            "ями",
            "ов",
            "ев",
            "ёв",
            "ам",
            "ям",
            "ах",
            "ях",
            "ою",
            "ею",
            "ой",
            "ей",
            "ою",
            "ею",
            "ою",
            "ею",
            "ый",
            "ий",
            "ой",
            "ая",
            "яя",
            "ое",
            "ее",
            "ые",
            "ие",
            "ого",
            "его",
            "ому",
            "ему",
            "ым",
            "им",
            "ых",
            "их",
            "ую",
            "юю",
            "а",
            "я",
            "ы",
            "и",
            "е",
            "у",
            "ю",
            "о",
            "ешь",
            "ешься",
            "ет",
            "ется",
            "ем",
            "емся",
            "ете",
            "етесь",
            "ут",
            "ют",
            "ат",
            "ят",
            "ишь",
            "ит",
            "им",
            "ите",
            "ите-ка",
            "ил",
            "ила",
            "ило",
            "или",
            "лся",
            "лась",
            "лось",
            "лись",
            "вший",
            "щий",
            "вшая",
            "вшее",
            "вшие",
        },
        key=len,
        reverse=True,
    )
)
os.makedirs(DATA_DIR, exist_ok=True)
os.makedirs(NOTES_DIR, exist_ok=True)
os.makedirs(PASSWORDS_DIR, exist_ok=True)
if not os.path.exists(USER_DICT_PATH):
    with open(USER_DICT_PATH, "a", encoding="utf-8"):
        pass
MAX_HISTORY = 250
IMAGE_EXTENSIONS = [".png", ".jpg", ".jpeg", ".bmp", ".gif"]
AUDIO_EXTENSIONS = [".wav", ".mp3", ".ogg"]
ATTACH_FILE_FILTER = (
    f"Изображения ({' '.join('*' + ext for ext in IMAGE_EXTENSIONS)})"
    f";;Аудио ({' '.join('*' + ext for ext in AUDIO_EXTENSIONS)})"
    f";;Все файлы (*)"
)


def create_list(*items):
    return list(items)


EXAMPLE_NUMBERS = create_list(1, 2, 3, 4)
EXAMPLE_WORDS = create_list("alpha", "beta", "gamma")
EXAMPLE_MIXED = create_list("hello", 42, 3.14)


class DesktopNotesWidget(QWidget):
    def __init__(self):
        super().__init__()
        self.setWindowTitle("Заметки на рабочем столе")
        self.setWindowFlags(Qt.WindowStaysOnTopHint | Qt.FramelessWindowHint | Qt.Tool)
        self.setAttribute(Qt.WA_TranslucentBackground)
        self.resize(340, 420)
        self.setMinimumWidth(260)
        self.setMinimumHeight(120)
        self._resize_active = False
        self._resize_edge = 0
        self._RESIZE_MARGIN = 8
        self._drag_active = False
        self._drag_offset = QPoint()
        root = QVBoxLayout(self)
        root.setContentsMargins(10, 10, 10, 10)
        root.setSpacing(8)
        header_layout = QHBoxLayout()
        header_layout.setContentsMargins(0, 0, 0, 0)
        header_layout.setSpacing(6)

        self.title_lbl = QLabel("📋 Заметки", self)
        self.title_lbl.setStyleSheet("font-size: 18px; font-weight: 600; color: white;")

        self.btn_close = QToolButton(self)
        self.btn_close.setObjectName("dnw_close")
        self.btn_close.setText("✕")
        self.btn_close.setCursor(Qt.PointingHandCursor)
        self.btn_close.setToolTip("Закрыть")
        self.btn_close.clicked.connect(self.close)

        header_layout.addWidget(self.title_lbl)
        header_layout.addStretch(1)
        header_layout.addWidget(self.btn_close)

        self.header_widget = QWidget(self)
        self.header_widget.setLayout(header_layout)
        self.header_widget.setProperty("drag_handle", True)

        root.addWidget(self.header_widget)
        self.notes_list = QListWidget(self)
        self.notes_list.setSelectionMode(QAbstractItemView.SingleSelection)
        root.addWidget(self.notes_list, 1)
        btn_row = QHBoxLayout()
        add_btn = QPushButton("+ Добавить", self)
        add_btn.clicked.connect(self.add_note)
        btn_row.addWidget(add_btn)
        refresh_btn = QPushButton("Обновить", self)
        refresh_btn.clicked.connect(self.reload_notes)
        btn_row.addWidget(refresh_btn)
        btn_row.addStretch(1)
        self._buttons_widget = QWidget(self)
        self._buttons_widget.setLayout(btn_row)
        root.addWidget(self._buttons_widget)
        self.setStyleSheet("""
            QWidget {
                background-color: rgba(30, 30, 30, 252);
                border-radius: 12px;
                color: white;
                font-family: Segoe UI;
            }
            QListWidget {
                background-color: rgba(30, 30, 30, 255);
                border: none;
                color: white;
                font-size: 14px;
            }
            QPushButton {
                background-color: #3a3f47;
                border-radius: 8px;
                padding: 6px 12px;
                color: white;
            }
            QPushButton:hover {
                background-color: #505860;
            }
            QToolButton#dnw_close {
                background: transparent;
                border: none;
                padding: 2px 8px;
                font-weight: 700;
                font-size: 14px;
                color: #ffffff;
                border-radius: 8px;
            }
            QToolButton#dnw_close:hover { background-color: rgba(255, 255, 255, 0.18); }
            QToolButton#dnw_close:pressed { background-color: rgba(255, 255, 255, 0.28); }
        """)
        grip_row = QHBoxLayout()
        grip_row.setContentsMargins(0, 0, 0, 0)
        grip_row.addStretch(1)
        self._grip = QSizeGrip(self)
        self._grip.setToolTip("Потяните за угол, чтобы изменить размер")
        grip_row.addWidget(self._grip, 0, Qt.AlignRight | Qt.AlignBottom)
        self._grip_widget = QWidget(self)
        self._grip_widget.setLayout(grip_row)
        root.addWidget(self._grip_widget)
        self._grip_top_right = QSizeGrip(self)
        self._grip_top_right.setToolTip("Потяните за верхний правый угол, чтобы изменить размер")
        self._grip_top_right.setFixedSize(16, 16)
        self._grip_top_right.raise_()
        self.notes_list.itemDoubleClicked.connect(self._on_item_activated)
        self.notes_list.setContextMenuPolicy(Qt.CustomContextMenu)
        self.notes_list.customContextMenuRequested.connect(self.show_context_menu)
        self.title_lbl.setProperty("drag_handle", True)
        self.installEventFilter(self)
        self.title_lbl.installEventFilter(self)
        self.notes_list.installEventFilter(self)
        self.header_widget.installEventFilter(self)
        self.title_lbl.installEventFilter(self)
        os.makedirs(DATA_DIR, exist_ok=True)
        self._load_pos()
        self.reload_notes()

    def resizeEvent(self, e):
        super().resizeEvent(e)
        if not hasattr(self, "_resize_deb"):
            self._resize_deb = QTimer(self)
            self._resize_deb.setSingleShot(True)
            self._resize_deb.setInterval(0)
            self._resize_deb.timeout.connect(self._apply_dock_ratios)
        self._resize_deb.start()

    def changeEvent(self, e):
        if e.type() == QEvent.WindowStateChange:
            self._capture_dock_ratios()
            QTimer.singleShot(0, self._apply_dock_ratios)
        super().changeEvent(e)


    def _calc_needed_height(self) -> int:
        root: QVBoxLayout = self.layout()
        header_h   = self.header_widget.sizeHint().height() if hasattr(self, "header_widget") else 0
        buttons_h  = self._buttons_widget.sizeHint().height() if hasattr(self, "_buttons_widget") else 0
        grip_h     = self._grip_widget.sizeHint().height() if hasattr(self, "_grip_widget") else 0
        count = self.notes_list.count()
        if count > 0:
            row_h = max(1, self.notes_list.sizeHintForRow(0))
        else:
            row_h = self.notes_list.fontMetrics().height() + 8
        list_h = row_h * max(1, count) + 2 * self.notes_list.frameWidth()
        m = root.contentsMargins()
        margins = m.top() + m.bottom()
        spacings = root.spacing() * 3 
        total = margins + header_h + list_h + buttons_h + grip_h + spacings
        try:
            avail_h = (self.screen().availableGeometry().height() if self.screen() else 1000) - 20
            total = min(total, max(200, avail_h))
        except Exception:
            pass
        return max(total, self.minimumSizeHint().height(), 120)

    def _apply_auto_min_height(self) -> None:
        h_min = int(self._calc_needed_height())
        self.setMinimumHeight(h_min)
        if self.height() < h_min:
            self.resize(self.width(), h_min)

    def _edge_at(self, pos: QPoint) -> int:
        m = getattr(self, "_RESIZE_MARGIN", 8)
        r = self.rect()
        edge = 0
        if pos.x() <= r.left() + m:  edge |= 1
        if pos.x() >= r.right() - m: edge |= 2
        if pos.y() <= r.top() + m:   edge |= 4
        if pos.y() >= r.bottom() - m:edge |= 8
        return edge

    def _set_cursor_for_edge(self, edge: int) -> None:
        if edge in (1, 2):
            self.setCursor(Qt.SizeHorCursor)
        elif edge in (4, 8):
            self.setCursor(Qt.SizeVerCursor)
        elif edge in (1 | 4, 2 | 8):
            self.setCursor(Qt.SizeFDiagCursor)
        elif edge in (1 | 8, 2 | 4):
            self.setCursor(Qt.SizeBDiagCursor)
        else:
            self.unsetCursor()

    def _read_json_any(self, path: str) -> list[dict]:
        if not os.path.exists(path):
            return []

        try:
            raw = open(path, "r", encoding="utf-8").read().strip()
            if not raw:
                return []
            data = json.loads(raw)
        except Exception as e:
            print("DesktopNotesWidget: load error:", e)
            return []

        def _now_iso():
            return datetime.datetime.now().isoformat(timespec="seconds")
        notes: list[dict] = []
        if isinstance(data, list) and all(isinstance(x, dict) for x in data):
            for x in data:
                txt = (x.get("text") or "").strip()
                if not txt:
                    continue
                nid = x.get("id") or str(uuid.uuid4())
                created = x.get("created") or _now_iso()
                updated = x.get("updated") or created
                notes.append({"id": nid, "text": txt, "created": created, "updated": updated})
            return notes
        if isinstance(data, list) and all(isinstance(x, str) for x in data):
            for s in data:
                txt = (s or "").strip()
                if txt:
                    ts = _now_iso()
                    notes.append({"id": str(uuid.uuid4()), "text": txt, "created": ts, "updated": ts})
            return notes
        if isinstance(data, dict) and "notes" in data:
            return self._read_json_any_from_list_like(data.get("notes"))
        if isinstance(data, dict):
            for k, v in data.items():
                txt = (v or "").strip() if isinstance(v, str) else ""
                if txt:
                    ts = _now_iso()
                    nid = k if isinstance(k, str) else str(uuid.uuid4())
                    notes.append({"id": nid, "text": txt, "created": ts, "updated": ts})
            return notes
        return []
    
    def _is_drag_source(self, obj, ev) -> bool:
        try:
            if getattr(obj, "property") and obj.property("drag_handle") is True:
                return True
        except Exception:
            pass
        return bool(ev.modifiers() & Qt.AltModifier)

    def _create_app_note_from_text(self, text: str) -> None:
        lines = [ln.strip() for ln in (text or "").splitlines()]
        title = next((ln for ln in lines if ln), "Новая заметка")[:120]
        parts = []
        for ln in (text or "").splitlines():
            if ln.strip():
                parts.append(f"<p>{html.escape(ln)}</p>")
            else:
                parts.append("<p><br/></p>")
        content_html = (
            "<!DOCTYPE html><html><head><meta charset='utf-8'></head><body>"
            + "".join(parts)
            + "</body></html>"
        )
        note_uuid = str(uuid.uuid4())
        timestamp = QDateTime.currentDateTime().toString(Qt.ISODate)
        note = Note(
            title=title,
            content=content_html,
            tags=[],
            favorite=False,
            timestamp=timestamp,
            reminder=None,
            uuid=note_uuid,
        )
        note.history = [content_html]
        note.history_index = 0
        folder_name = NotesApp.safe_folder_name(title, note_uuid, timestamp)
        note_dir = os.path.join(NOTES_DIR, folder_name)
        os.makedirs(note_dir, exist_ok=True)
        note_path = os.path.join(note_dir, "note.json")
        with open(note_path, "w", encoding="utf-8") as f:
            json.dump(note.to_dict(), f, ensure_ascii=False, indent=2)

    def eventFilter(self, obj, ev):
        et = ev.type()
        if hasattr(ev, "globalPosition"):
            gp = ev.globalPosition().toPoint()
        else:
            gp = QCursor.pos()
        lp = self.mapFromGlobal(gp)
        if et == QEvent.MouseButtonPress and ev.button() == Qt.LeftButton:
            edge = self._edge_at(lp)
            if edge:
                self._resize_active = True
                self._resize_edge = edge
                self._resize_start_geom = self.frameGeometry()
                self._resize_start_pos = gp
                return True
            if self._is_drag_source(obj, ev):
                self._drag_active = True
                self._drag_offset = gp - self.frameGeometry().topLeft()
                return True
        elif et == QEvent.MouseMove:
            if self._resize_active:
                dx = gp.x() - self._resize_start_pos.x()
                dy = gp.y() - self._resize_start_pos.y()
                g = QRect(self._resize_start_geom)
                if self._resize_edge & 1:
                    g.setLeft(g.left() + dx)
                if self._resize_edge & 2:
                    g.setRight(g.right() + dx)
                if self._resize_edge & 4:
                    g.setTop(g.top() + dy)
                if self._resize_edge & 8:
                    g.setBottom(g.bottom() + dy)
                minw, minh = self.minimumWidth(), self.minimumHeight()
                if g.width() < minw:
                    if self._resize_edge & 1:
                        g.setLeft(g.right() - minw)
                    else:
                        g.setRight(g.left() + minw)
                if g.height() < minh:
                    if self._resize_edge & 4:
                        g.setTop(g.bottom() - minh)
                    else:
                        g.setBottom(g.top() + minh)
                self.setGeometry(g)
                return True
            edge = self._edge_at(lp)
            self._set_cursor_for_edge(edge)
            if self._drag_active:
                self.move(gp - self._drag_offset)
                return True
        elif et == QEvent.MouseButtonRelease and ev.button() == Qt.LeftButton:
            if self._resize_active:
                self._resize_active = False
                self._resize_edge = 0
                self._save_pos()
                self.unsetCursor()
                return True
            if self._drag_active:
                self._drag_active = False
                self._save_pos()
                return True
        return False

    def _load_pos(self):
        try:
            s = QSettings(SETTINGS_PATH, QSettings.IniFormat)
            s.beginGroup("DesktopNotesWidget")
            if s.contains("x") and s.contains("y"):
                self.move(int(s.value("x", type=int)), int(s.value("y", type=int)))
            if s.contains("w") and s.contains("h"):
                w = int(s.value("w", type=int)); h = int(s.value("h", type=int))
                if w > 100 and h > 100:
                    self.resize(w, h)
            s.endGroup()
        except Exception:
            pass

    def _save_pos(self):
        try:
            s = QSettings(SETTINGS_PATH, QSettings.IniFormat)
            s.beginGroup("DesktopNotesWidget")
            s.setValue("x", self.x()); s.setValue("y", self.y())
            s.setValue("w", self.width()); s.setValue("h", self.height())
            s.endGroup()
        except Exception:
            pass

    def _read_json_any_from_list_like(self, value) -> list[dict]:
        if isinstance(value, list):
            out = []
            for item in value:
                if isinstance(item, str):
                    ts = datetime.datetime.now().isoformat(timespec="seconds")
                    out.append({"id": str(uuid.uuid4()), "text": item.strip(), "created": ts, "updated": ts})
                elif isinstance(item, dict):
                    txt = (item.get("text") or "").strip()
                    if not txt:
                        continue
                    nid = item.get("id") or str(uuid.uuid4())
                    created = item.get("created") or datetime.datetime.now().isoformat(timespec="seconds")
                    updated = item.get("updated") or created
                    out.append({"id": nid, "text": txt, "created": created, "updated": updated})
            return out
        return []

    def _scan_app_notes(self) -> list[dict]:
        items: list[dict] = []
        try:
            if not os.path.isdir(NOTES_DIR):
                return items
            for folder in os.listdir(NOTES_DIR):
                if folder == "Trash":
                    continue
                folder_path = os.path.join(NOTES_DIR, folder)
                if not os.path.isdir(folder_path):
                    continue
                file_path = os.path.join(folder_path, "note.json")
                if not os.path.isfile(file_path):
                    continue
                try:
                    with open(file_path, "r", encoding="utf-8") as f:
                        data = json.load(f)
                    title = (data.get("title") or folder).strip()
                    created = (data.get("timestamp") or "").strip()
                    nid = (data.get("uuid") or str(uuid.uuid4())).strip()
                    items.append({
                        "id": nid,
                        "text": title,
                        "created": created,
                        "updated": created or "",
                        "source": "app",
                        "path": file_path,
                    })
                except Exception as e:
                    print("scan app notes error:", e)
        finally:
            return items

    def _atomic_write(self, path: str, payload: list[dict]) -> None:
        data = json.dumps(payload, ensure_ascii=False, indent=2)
        tmp = path + ".tmp"
        with open(tmp, "w", encoding="utf-8") as f:
            f.write(data)
            f.flush()
            os.fsync(f.fileno())
        try:
            if os.path.exists(path):
                shutil.copy2(path, NOTES_WIDGET_BAK)
        except Exception:
            pass
        os.replace(tmp, path)

    def _model_from_ui(self) -> list[dict]:
        out = []
        now_iso = datetime.datetime.now().isoformat(timespec="seconds")
        for i in range(self.notes_list.count()):
            it = self.notes_list.item(i)
            src = it.data(Qt.UserRole + 2) or "widget"
            if src != "widget":
                continue
            txt = (it.text() or "").lstrip("📌 ").strip()
            if not txt:
                continue
            nid = it.data(Qt.UserRole) or str(uuid.uuid4())
            created = it.data(Qt.UserRole + 1) or now_iso
            updated = now_iso
            out.append({"id": nid, "text": txt, "created": created, "updated": updated})
        return out

    def _ui_from_model(self, notes: list[dict]) -> None:
        self.notes_list.clear()
        for n in notes:
            txt = n.get("text", "")
            src = n.get("source", "widget")
            prefix = "📝 " if src == "app" else ""
            item = QListWidgetItem(prefix + txt)
            item.setData(Qt.UserRole, n.get("id"))
            item.setData(Qt.UserRole + 1, n.get("created"))
            item.setData(Qt.UserRole + 2, src)
            item.setData(Qt.UserRole + 3, n.get("path"))
            self.notes_list.addItem(item)

    def reload_notes(self):
        widget_notes = self._read_json_any(NOTES_WIDGET_PATH)
        for n in widget_notes:
            n["source"] = "widget"
            n["path"] = NOTES_WIDGET_PATH
        app_notes = self._scan_app_notes()
        combined = app_notes + widget_notes
        self._ui_from_model(combined)
        self._apply_auto_min_height()

    def save_notes(self):
        notes = self._model_from_ui()
        self._atomic_write(NOTES_WIDGET_PATH, notes)

    def add_note(self):
        text, ok = QInputDialog.getMultiLineText(
            self, "Новая заметка (виджет)", "Введите текст:"
        )
        if ok:
            text = (text or "").strip()
            if not text:
                return
            self._create_app_note_from_text(text)
            self.reload_notes()
            self._apply_auto_min_height()
            QMessageBox.information(self, "Создано", "Заметка добавлена.")

    def _on_item_activated(self, item: QListWidgetItem):
        src = item.data(Qt.UserRole + 2) or "widget"
        if src == "app":
            self.open_app_note_in_app(item)
        else:
            self.edit_note(item)

    def edit_note(self, item: QListWidgetItem):
        src = item.data(Qt.UserRole + 2) or "widget"
        if src == "app":
            self.open_app_note_in_app(item)
            return
        cur_text = item.text().lstrip("📌 ").strip()
        text, ok = QInputDialog.getMultiLineText(self, "Редактировать заметку", "Текст:", cur_text)
        if ok:
            item.setText("📌 " + (text or "").strip())
            self.save_notes()

    def delete_selected(self):
        row = self.notes_list.currentRow()
        if row < 0:
            return
        it = self.notes_list.item(row)
        src = it.data(Qt.UserRole + 2) or "widget"
        if src == "app":
            QMessageBox.information(self, "Удаление запрещено", "Удаление системных заметок из виджета отключено.")
            return
        self.notes_list.takeItem(row)
        self.save_notes()
        self._apply_auto_min_height()

    def open_app_note_folder(self, item: QListWidgetItem):
        path = item.data(Qt.UserRole + 3)
        if path and os.path.isfile(path):
            QDesktopServices.openUrl(QUrl.fromLocalFile(os.path.dirname(path)))

    def open_app_note_in_app(self, item: QListWidgetItem) -> None:
        try:
            target_uuid = item.data(Qt.UserRole) or None
            path = item.data(Qt.UserRole + 3) or None
            title_text = (item.text() or "").lstrip("📝 ").strip()
            app = QApplication.instance()
            launcher = getattr(app, "launcher_window", None) if app else None
            notes_win = None
            if launcher is not None and hasattr(launcher, "launch_notes"):
                launcher.launch_notes()
                notes_win = launcher.notes_window
            else:
                notes_win = NotesApp()
                notes_win.show()
                notes_win.raise_()
                notes_win.activateWindow()
            if hasattr(notes_win, "load_notes_from_disk"):
                notes_win.load_notes_from_disk()
            if hasattr(notes_win, "refresh_notes_list"):
                notes_win.refresh_notes_list()
            if not target_uuid and path and os.path.isfile(path):
                try:
                    with open(path, "r", encoding="utf-8") as f:
                        data = json.load(f)
                    target_uuid = data.get("uuid")
                except Exception:
                    target_uuid = None
            matched_item = None
            for i in range(notes_win.notes_list.count()):
                it = notes_win.notes_list.item(i)
                note_obj = it.data(Qt.UserRole)
                if not note_obj:
                    continue
                if (target_uuid and getattr(note_obj, "uuid", None) == target_uuid) or (
                    not target_uuid and title_text and getattr(note_obj, "title", "") == title_text
                ):
                    matched_item = it
                    break
            if matched_item is not None:
                notes_win.notes_list.setCurrentItem(matched_item)
                notes_win.select_note(matched_item.data(Qt.UserRole))
                notes_win.raise_()
                notes_win.activateWindow()
            else:
                QMessageBox.information(self, "Заметка не найдена",
                                        "Не удалось найти заметку в приложении. Список был обновлён.")
        except Exception as e:
            QMessageBox.warning(self, "Ошибка", f"Не удалось открыть заметку в приложении: {e}")

    def open_notes_app(self):
        app = QApplication.instance()
        launcher = getattr(app, "launcher_window", None) if app else None
        try:
            if launcher is not None and hasattr(launcher, "launch_notes"):
                launcher.launch_notes()
            else:
                win = NotesApp()
                win.show()
                win.raise_()
                win.activateWindow()
        except Exception as e:
            QMessageBox.warning(self, "Ошибка", f"Не удалось открыть приложение 'Заметки': {e}")

    def show_context_menu(self, pos):
        item = self.notes_list.itemAt(pos)
        menu = QMenu(self)
        if item is not None and (item.data(Qt.UserRole + 2) or "widget") == "app":
            menu.addAction("Открыть в «Заметках»", lambda: self.open_app_note_in_app(item))
            menu.addAction("Открыть папку заметки", lambda: self.open_app_note_folder(item))
            menu.addSeparator()
            menu.addAction("Обновить", self.reload_notes)
        else:
            menu.addAction("Удалить", self.delete_selected)
            menu.addSeparator()
            menu.addAction("Обновить", self.reload_notes)
            menu.addAction("Открыть файл хранилища", lambda: QDesktopServices.openUrl(QUrl.fromLocalFile(NOTES_WIDGET_PATH)))
            menu.addAction("Очистить всё", self.clear_all)
        menu.exec(self.notes_list.mapToGlobal(pos))

    def clear_all(self):
        if QMessageBox.question(self, "Подтверждение", "Удалить все виджет-заметки?") == QMessageBox.Yes:
            for i in reversed(range(self.notes_list.count())):
                it = self.notes_list.item(i)
                if (it.data(Qt.UserRole + 2) or "widget") == "widget":
                    self.notes_list.takeItem(i)
            self.save_notes()
            self._apply_auto_min_height()

    def closeEvent(self, e: QCloseEvent):
        self.save_notes()
        self._save_pos()
        super().closeEvent(e)

    def showEvent(self, e):
        super().showEvent(e)
        QTimer.singleShot(0, self._apply_auto_min_height)


class SpellCheckHighlighter(QSyntaxHighlighter):
    def __init__(self, parent):
        super().__init__(parent)
        self.err_fmt = QTextCharFormat()
        self.err_fmt.setUnderlineColor(Qt.red)
        self.err_fmt.setUnderlineStyle(QTextCharFormat.SpellCheckUnderline)
        try:
            from spellchecker import SpellChecker as _SpellChecker
            try:
                spell = _SpellChecker(language='en')
            except Exception:
                spell = _SpellChecker(language=None)
            try:
                self.spell_checker = _SpellChecker(language="ru")
            except Exception:
                self.spell_checker = _SpellChecker()
        except ImportError:
            self.spell_checker = None

        self._checked_cache = {}
        self.online_enabled = True
        self._load_user_dictionary()
        try:
            self._dict_watcher = QFileSystemWatcher([USER_DICT_PATH])
            self._dict_watcher.fileChanged.connect(self._reload_user_dictionary)
        except Exception:
            self._dict_watcher = None

    def _check_online(self, word: str, lang: str) -> bool:
        key = f"{lang}:{word.lower()}"
        if key in self._checked_cache:
            return self._checked_cache[key]
        try:
            if lang == "ru":
                resp = requests.get(
                    "https://speller.yandex.net/services/spellservice.json/checkText",
                    params={"text": word, "lang": "ru"},
                    timeout=2,
                )
                if resp.status_code == 200:
                    data = resp.json()
                    ok = len(data) == 0
                    self._checked_cache[key] = ok
                    return ok
            elif lang == "en":
                resp = requests.post(
                    "https://api.languagetool.org/v2/check",
                    data={"text": word, "language": "en-US"},
                    timeout=2,
                )
                if resp.status_code == 200:
                    data = resp.json()
                    ok = len(data.get("matches", [])) == 0
                    self._checked_cache[key] = ok
                    return ok
        except Exception:
            pass
        self._checked_cache[key] = True
        return True
    
    def set_online_enabled(self, enabled: bool) -> None:
        self.online_enabled = bool(enabled)
        self._checked_cache.clear()
        self.rehighlight()    

    def _is_known(self, w: str) -> bool:
        if not w:
            return False
        lw = self.normalize_e(w.strip().lower())
        if lw in getattr(self, "user_words", set()) or lw in getattr(self, "local_ignored", set()):
            return True
        if not self.spell_checker:
            return True
        if len(self.spell_checker.unknown([lw])) == 0:
            return True
        if not getattr(self, "online_enabled", True):
            return False
        if re.search(r"[А-Яа-яЁё]", lw):
            return self._check_online(lw, "ru")
        elif re.search(r"[A-Za-z]", lw):
            return self._check_online(lw, "en")

        return False

    @staticmethod
    def normalize_e(word: str) -> str:
        return word.replace("ё", "е").replace("Ё", "Е")

    def _ru_base_forms(self, w: str) -> list[str]:
        lw = self.normalize_e((w or "").lower())
        out: set[str] = {lw}
        stems = {lw}
        for r in RU_REFLEXIVE:
            if lw.endswith(r) and len(lw) - len(r) >= 3:
                stems.add(lw[: -len(r)])
                break
        generated = set()
        for s in list(stems):
            for suf in RU_SUFFIXES_RU:
                if s.endswith(suf) and len(s) - len(suf) >= 3:
                    generated.add(s[: -len(suf)])
            for suf, repls in RU_VARIANT_RULES:
                if s.endswith(suf) and len(s) - len(suf) >= 3:
                    stem = s[: -len(suf)]
                    for rep in repls:
                        generated.add(stem + rep)
        out |= stems
        out |= generated
        if len(out) > 32:
            out = set(list(out)[:32])
        return list(out)

    def _load_user_dictionary(self) -> None:
        prev_local = getattr(self, "local_ignored", set())
        self.user_words = set()
        self.local_ignored = set()
        try:
            with open(USER_DICT_PATH, "r", encoding="utf-8") as f:
                words = [line.strip().lower() for line in f if line.strip()]
            self.user_words.update(self.normalize_e(w) for w in words)
            if self.spell_checker and words:
                self.spell_checker.word_frequency.load_words(words)
        except Exception:
            pass
        self.local_ignored = prev_local

    def add_to_dictionary(self, word: str) -> None:
        w = (word or "").strip().lower()
        if not w or w in getattr(self, "user_words", set()):
            return
        try:
            with open(USER_DICT_PATH, "a", encoding="utf-8") as f:
                f.write(w + "\n")
                f.flush()
                os.fsync(f.fileno())
        except Exception:
            return
        self._load_user_dictionary()
        self.rehighlight()

    def set_local_ignored(self, words: list[str] | set[str]) -> None:
        self.local_ignored = {self.normalize_e(w.strip().lower()) for w in words if w.strip()}
        self.rehighlight()

    def _reload_user_dictionary(self, path: str) -> None:
        if self._dict_watcher and path not in self._dict_watcher.files():
            try:
                self._dict_watcher.addPath(path)
            except Exception:
                pass
        self._load_user_dictionary()
        self.rehighlight()

    def highlightBlock(self, text: str) -> None:
        if not text or len(text) < 3:
            return
        if len(text) > 1500:
            return
        if not self.spell_checker:
            return
        matches = list(re.finditer(r"[A-Za-zА-Яа-яЁё']+", text))
        if not matches:
            return
        words = [self.normalize_e(m.group().lower()) for m in matches if len(m.group()) >= 3]
        misspelled = self.spell_checker.unknown(words)
        misspelled -= getattr(self, "user_words", set())
        misspelled -= getattr(self, "local_ignored", set())
        to_keep = set()
        for w in list(misspelled):
            if re.search(r"[А-Яа-яЁё]", w):
                bases = self._ru_base_forms(w)
                if any(self._is_known(b) for b in bases):
                    continue
            elif self._is_known(w):
                continue
            to_keep.add(w)
        misspelled = to_keep
        for match in matches:
            word_norm = self.normalize_e(match.group().lower())
            if word_norm in misspelled:
                self.setFormat(match.start(), match.end() - match.start(), self.err_fmt)

def create_dropdown_combo(*items, parent=None):
    combo = QComboBox(parent)
    combo.setEditable(True)
    flat_items = []
    for item in items:
        if isinstance(item, (list, tuple, set)):
            flat_items.extend(item)
        else:
            flat_items.append(item)
    combo.addItems([str(i) for i in flat_items])

    def _commit_new_item() -> None:
        text = combo.currentText().strip()
        if text and combo.findText(text) == -1:
            combo.addItem(text)
            combo.setCurrentText(text)

    combo.lineEdit().editingFinished.connect(_commit_new_item)
    return combo


def handle_exception(exc_type, exc_value, exc_traceback, exc, tb):
    traceback.print_exception(exc_type, exc_value, exc_traceback)


sys.excepthook = lambda t, v, tb: print("Uncaught exception:", t, v)


def copy_default_icons():
    app_root = (
        os.path.dirname(sys.executable)
        if getattr(sys, "frozen", False)
        else os.path.abspath(os.path.dirname(__file__))
    )
    src_root = os.path.join(app_root, "Data")
    icon_names = ["icon.ico", "tray_icon.ico", "file.ico"]
    destination_paths = [ICON_PATH, TRAY_ICON_PATH, FILE_ICON_PATH]
    for name, dst in zip(icon_names, destination_paths):
        src = os.path.join(src_root, name)
        if os.path.isfile(src) and not os.path.isfile(dst):
            os.makedirs(os.path.dirname(dst), exist_ok=True)
            shutil.copy(src, dst)


copy_default_icons()


def paste_from_clipboard(widget):
    if pyperclip is None:
        messagebox.showerror("Ошибка", "pyperclip не установлен.")
        return
    if not widget or not hasattr(widget, "delete") or not hasattr(widget, "insert"):
        return
    try:
        text = pyperclip.paste()
        if isinstance(widget, (tk.Entry, ttk.Entry)):
            widget.delete(0, tk.END)
            widget.insert(0, text)
        else:
            if hasattr(widget, "selection_present") and widget.selection_present():
                widget.delete("sel.first", "sel.last")
            widget.insert(tk.INSERT, text)
    except (tk.TclError, pyperclip.PyperclipException) as exc:
        messagebox.showerror("Ошибка", f"Не удалось вставить из буфера: {exc}")


CUSTOM_MENU_STYLE = """
QMenu {
    background-color: #282828;
    color: #fff;
    border-radius: 10px;
    font-size: 15px;
    padding: 6px;
}
QMenu::item {
    background-color: transparent;
    color: #fff;
    padding: 8px 32px 8px 24px;
    border-radius: 6px;
}
QMenu::item:selected {
    background-color: #fff;
    color: #111;
    border-radius: 6px;
    border: 2px solid #7a7a7a;
}
QMenu::separator {
    height: 1px;
    background: #444;
    margin: 8px 0;
}
"""


class DrawingCanvas(QWidget):
    def __init__(self, parent=None):
        super().__init__(parent)
        self.viewport_pad = 16
        self._undo_stack: list[tuple[QImage, list[dict]]] = []
        self._redo_stack: list[tuple[QImage, list[dict]]] = []
        self._drag_changed = False
        self.canvas_bg = QColor("#f3f4f6")
        self.canvas_border_dark = QColor("#2f3338")
        self.canvas_border_light = QColor("#ffffff")
        self.viewport_pad = 16
        self._move_bbox = None
        self.setAttribute(Qt.WidgetAttribute.WA_StaticContents)
        self.image = QImage(1200, 800, QImage.Format.Format_ARGB32)
        self.image.fill(Qt.white)
        self._event_zoom = 1.0
        self._pending_zoom = None
        self.objects = []
        self.selected_index = None
        self._text_input_active = False
        self.drag_start = QPoint()
        self.drag_last = QPoint()
        self.hit_tolerance = 8
        self.show_layout_frame = True
        self.safe_margin = 16
        self.pen_color = QColor("black")
        self.pen_width = 3
        self.zoom = 1.0
        self.tool = "pen"
        self.is_drawing = False
        self.last_pos = QPoint()
        self.start_pos = QPoint()
        self.preview_pos = QPoint()
        self.fill_enabled = False
        self.fill_color = QColor(0, 0, 0, 60)
        self.text_font = QFont("Segoe UI", 18)
        self.move_active = False
        self.move_start = QPoint()
        self.move_delta = QPoint()
        self.image_before_move = None

    def set_color(self, color: QColor):
        self.pen_color = QColor(color)
        self.update()

    def set_show_frame(self, on: bool):
        self.show_layout_frame = bool(on)
        self.update()

    def set_safe_margin(self, px: int):
        self.safe_margin = max(0, int(px))
        self.update()

    def set_width(self, w: int):
        self.pen_width = max(1, int(w))
        self.update()

    def _bbox_union(self, a: QRect | None, b: QRect | None) -> QRect | None:
        if a is None or a.isNull():
            return b
        if b is None or b.isNull():
            return a
        return a.united(b)

    def _clone_objects(self) -> list[dict]:
        out = []
        for o in self.objects:
            c = {"type": o.get("type")}
            if "rect" in o:
                c["rect"] = QRect(o["rect"])
            if "p1" in o:
                c["p1"] = QPoint(o["p1"])
            if "p2" in o:
                c["p2"] = QPoint(o["p2"])
            if "pos" in o:
                c["pos"] = QPoint(o["pos"])
            if "text" in o:
                c["text"] = str(o["text"])
            if "font" in o:
                c["font"] = QFont(o["font"])
            if "pen" in o:
                c["pen"] = QPen(o["pen"])
            if "brush" in o:
                c["brush"] = QBrush(o["brush"])
            out.append(c)
        return out

    def _push_undo_snapshot(self) -> None:
        self._undo_stack.append((self.image.copy(), self._clone_objects()))
        if len(self._undo_stack) > 50:
            self._undo_stack.pop(0)
        self._redo_stack.clear()

    def undo(self) -> None:
        if not self._undo_stack:
            return
        self._redo_stack.append((self.image.copy(), self._clone_objects()))
        img, objs = self._undo_stack.pop()
        self.image = img
        self.objects = objs
        self.update()

    def redo(self) -> None:
        if not self._redo_stack:
            return
        self._undo_stack.append((self.image.copy(), self._clone_objects()))
        img, objs = self._redo_stack.pop()
        self.image = img
        self.objects = objs
        self.update()

    def _object_bbox(self, obj: dict) -> QRect:
        t = obj.get("type")
        if t in ("rect", "ellipse"):
            return QRect(obj["rect"])
        elif t in ("line", "arrow"):
            x1, y1 = obj["p1"].x(), obj["p1"].y()
            x2, y2 = obj["p2"].x(), obj["p2"].y()
            r = QRect(min(x1, x2), min(y1, y2), abs(x1 - x2), abs(y1 - y2))
            w = max(1, obj.get("pen", QPen()).width())
            return r.adjusted(-w, -w, w, w)
        elif t == "text":
            fm = QFontMetrics(obj.get("font", self.text_font))
            pos = obj.get("pos", QPoint())
            return QRect(
                pos, QSize(fm.horizontalAdvance(obj.get("text", "")), fm.height())
            )
        return QRect()

    def _objects_bbox(self) -> QRect | None:
        box = None
        for o in self.objects:
            box = self._bbox_union(box, self._object_bbox(o))
        return box

    def _raster_bbox(self, img: QImage) -> QRect | None:
        w, h = img.width(), img.height()
        left, right, top, bottom = w, -1, h, -1
        white = QColor(Qt.white).rgba()
        for y in range(h):
            for x in range(w):
                if img.pixel(x, y) != white:
                    if x < left:
                        left = x
                    if x > right:
                        right = x
                    if y < top:
                        top = y
                    if y > bottom:
                        bottom = y
        if right < left or bottom < top:
            return None
        return QRect(left, top, right - left + 1, bottom - top + 1)

    def _content_bbox(self) -> QRect | None:
        box = self._raster_bbox(self.image)
        box = self._bbox_union(box, self._objects_bbox())
        return box

    def _clamp_delta_to_box(self, dx: int, dy: int, box: QRect | None) -> QPoint:
        if box is None or box.isNull():
            return QPoint(0, 0)
        min_dx = -box.left()
        max_dx = (self.image.width() - 1) - box.right()
        min_dy = -box.top()
        max_dy = (self.image.height() - 1) - box.bottom()
        dx = max(min_dx, min(max_dx, dx))
        dy = max(min_dy, min(max_dy, dy))
        return QPoint(dx, dy)

    def set_tool(self, tool: str):
        self.tool = tool
        self.is_drawing = False
        self.move_active = False
        self.move_delta = QPoint()
        self.image_before_move = None
        self.selected_index = None
        self.update()

    def _draw_object(self, painter: QPainter, obj: dict) -> None:
        t = obj.get("type")
        pen: QPen = obj.get("pen", QPen(Qt.black, 3))
        painter.setPen(pen)
        painter.setBrush(obj.get("brush", Qt.NoBrush))
        if t == "rect":
            painter.drawRect(obj["rect"])
        elif t == "ellipse":
            painter.drawEllipse(obj["rect"])
        elif t == "line":
            painter.drawLine(obj["p1"], obj["p2"])
        elif t == "arrow":
            p1, p2 = QPointF(obj["p1"]), QPointF(obj["p2"])
            painter.drawLine(p1, p2)
            head = self._arrow_head(p1, p2, max(10, pen.width() * 3))
            painter.setBrush(QBrush(pen.color()))
            painter.drawPolygon(head)
        elif t == "text":
            painter.setFont(obj["font"])
            painter.drawText(obj["pos"], obj["text"])

    zoomChanged = Signal(float)

    def set_zoom(self, z: float):
        z = max(0.25, min(4.0, float(z)))
        if self.is_drawing or self.move_active:
            self._pending_zoom = z
            return
        if abs(z - self.zoom) > 1e-3:
            self.zoom = z
            self.update()
            try:
                self.zoomChanged.emit(self.zoom)
            except Exception:
                pass

    def set_zoom_percent(self, p: int):
        self.set_zoom(max(25, min(400, int(p))) / 100.0)

    def _to_image_pos(self, p: QPoint, zoom: float | None = None) -> QPoint:
        z = float(self.zoom if zoom is None else zoom)
        pad = getattr(self, "viewport_pad", 0)
        return QPoint(int((p.x() - pad) / z), int((p.y() - pad) / z))

    def wheelEvent(self, e):
        if e.modifiers() & Qt.ControlModifier:
            if self.is_drawing or self.move_active:
                e.ignore()
                return
            delta = e.angleDelta().y()
            step = 0.1 if delta > 0 else -0.1
            self.set_zoom(self.zoom + step)
            e.accept()
        else:
            super().wheelEvent(e)

    def _draw_objects(self, painter: QPainter) -> None:
        for obj in self.objects:
            self._draw_object(painter, obj)

    def _shape_from_current(self, end_pos: QPoint) -> dict:
        pen = QPen(
            self.pen_color, self.pen_width, Qt.SolidLine, Qt.RoundCap, Qt.RoundJoin
        )
        t = self.tool
        if t in ("rect", "ellipse"):
            return {
                "type": t,
                "rect": self._norm_rect(self.start_pos, end_pos),
                "pen": pen,
                "brush": QBrush(self.fill_color) if self.fill_enabled else Qt.NoBrush,
            }
        elif t in ("line", "arrow"):
            return {
                "type": t,
                "p1": QPoint(self.start_pos),
                "p2": QPoint(end_pos),
                "pen": pen,
            }
        return {}

    def _translate_object(self, obj: dict, delta: QPoint) -> None:
        if "rect" in obj:
            obj["rect"].translate(delta)
        if "p1" in obj:
            obj["p1"] += delta
        if "p2" in obj:
            obj["p2"] += delta
        if "pos" in obj:
            obj["pos"] += delta

    def _distance_to_segment(self, p: QPoint, a: QPoint, b: QPoint) -> float:
        ax, ay = a.x(), a.y()
        bx, by = b.x(), b.y()
        px, py = p.x(), p.y()
        vx, vy = bx - ax, by - ay
        wx, wy = px - ax, py - ay
        vv = vx * vx + vy * vy or 1.0
        t = max(0.0, min(1.0, (wx * vx + wy * vy) / vv))
        cx, cy = ax + t * vx, ay + t * vy
        dx, dy = px - cx, py - cy
        return (dx * dx + dy * dy) ** 0.5

    def _object_at(self, pos: QPoint) -> int | None:
        for i in range(len(self.objects) - 1, -1, -1):
            obj = self.objects[i]
            t = obj.get("type")
            if t in ("rect", "ellipse"):
                if (
                    obj["rect"]
                    .adjusted(
                        -self.hit_tolerance,
                        -self.hit_tolerance,
                        self.hit_tolerance,
                        self.hit_tolerance,
                    )
                    .contains(pos)
                ):
                    return i
            elif t in ("line", "arrow"):
                if (
                    self._distance_to_segment(pos, obj["p1"], obj["p2"])
                    <= self.hit_tolerance
                ):
                    return i
            elif t == "text":
                fm = QFontMetrics(obj["font"])
                br = QRect(
                    obj["pos"], QSize(fm.horizontalAdvance(obj["text"]), fm.height())
                )
                if br.adjusted(
                    -self.hit_tolerance,
                    -self.hit_tolerance,
                    self.hit_tolerance,
                    self.hit_tolerance,
                ).contains(pos):
                    return i
        return None

    def set_fill_enabled(self, enabled: bool):
        self.fill_enabled = bool(enabled)
        self.update()

    def set_fill_color(self, color: QColor):
        self.fill_color = QColor(color)
        self.update()

    def set_text_size(self, px: int):
        self.text_font.setPointSize(int(px))
        self.update()

    def clear(self):
        self.image.fill(Qt.white)
        self.update()

    @staticmethod
    def _norm_rect(a: QPoint, b: QPoint) -> QRect:
        return QRect(
            min(a.x(), b.x()), min(a.y(), b.y()), abs(a.x() - b.x()), abs(a.y() - b.y())
        )

    @staticmethod
    def _arrow_head(p1: QPointF, p2: QPointF, size: float) -> QPolygonF:
        vec = QPointF(p2.x() - p1.x(), p2.y() - p1.y())
        length = (vec.x() ** 2 + vec.y() ** 2) ** 0.5 or 1.0
        ux, uy = vec.x() / length, vec.y() / length
        nx, ny = -uy, ux
        back = QPointF(p2.x() - ux * size, p2.y() - uy * size)
        left = QPointF(back.x() + nx * size * 0.5, back.y() + ny * size * 0.5)
        right = QPointF(back.x() - nx * size * 0.5, back.y() - ny * size * 0.5)
        return QPolygonF([p2, left, right])

    def mousePressEvent(self, e):
        if e.button() != Qt.LeftButton:
            return
        self._event_zoom = self.zoom
        tools_need_grab = {
            "pen",
            "eraser",
            "rect",
            "ellipse",
            "line",
            "arrow",
            "move",
            "select",
        }
        if self.tool in tools_need_grab:
            w = self.window()
            can_grab = self.isVisible() and (w is not None and w.isVisible())
            wh = w.windowHandle() if (w is not None) else None
            if wh is not None and hasattr(wh, "isExposed"):
                can_grab = can_grab and wh.isExposed()
            if can_grab:
                try:
                    self.grabMouse()
                except Exception:
                    pass
        if self.tool == "text":
            try:
                self.releaseMouse()
            except Exception:
                pass
            w = self.window()
            if (
                self._text_input_active
                or not self.isVisible()
                or w is None
                or not w.isVisible()
            ):
                return
            self._text_input_active = True
            try:
                dlg = QInputDialog(w)
                dlg.setModal(True)
                dlg.setWindowTitle("Текст")
                dlg.setLabelText("Введите текст:")
                dlg.setInputMode(QInputDialog.TextInput)
                dlg.setWindowFlag(Qt.WindowStaysOnTopHint, True)
                fg = dlg.frameGeometry()
                fg.moveCenter(w.frameGeometry().center())
                dlg.move(fg.topLeft())
                if dlg.exec() == QDialog.Accepted:
                    if self.tool == "text" and self.isVisible() and w.isVisible():
                        text = dlg.textValue().strip()
                        if text:
                            img_pos = self._to_image_pos(
                                e.position().toPoint(), zoom=self._event_zoom
                            )
                            obj = {
                                "type": "text",
                                "text": text,
                                "pos": img_pos,
                                "font": QFont(self.text_font),
                                "pen": QPen(
                                    self.pen_color,
                                    self.pen_width,
                                    Qt.SolidLine,
                                    Qt.RoundCap,
                                    Qt.RoundJoin,
                                ),
                            }
                            self._push_undo_snapshot()
                            self.objects.append(obj)
                            self.update()
            finally:
                self._text_input_active = False
            return
        if self.tool == "select":
            self.is_drawing = True
            self.last_pos = self._to_image_pos(
                e.position().toPoint(), zoom=self._event_zoom
            )
            self.selected_index = self._object_at(self.last_pos)
            self.drag_start = self.last_pos
            self.drag_last = self.last_pos
            return
        if self.tool == "move":
            self.move_active = True
            self.is_drawing = True
            self.move_start = self._to_image_pos(
                e.position().toPoint(), zoom=self._event_zoom
            )
            self.image_before_move = self.image.copy()
            self._move_bbox = self._content_bbox()
            self.move_delta = QPoint()
            return
        self.is_drawing = True
        self.last_pos = self._to_image_pos(
            e.position().toPoint(), zoom=self._event_zoom
        )
        self.start_pos = self.last_pos
        self.preview_pos = self.last_pos
        if self.tool in ("pen", "eraser"):
            self._push_undo_snapshot()
            painter = QPainter(self.image)
            pen = QPen(
                Qt.white if self.tool == "eraser" else self.pen_color,
                self.pen_width,
                Qt.SolidLine,
                Qt.RoundCap,
                Qt.RoundJoin,
            )
            painter.setPen(pen)
            painter.drawPoint(self.last_pos)
            painter.end()
            self.update()

    def mouseMoveEvent(self, e):
        if not self.is_drawing:
            return
        if self.tool == "move" and self.move_active:
            pos_img = self._to_image_pos(e.position().toPoint(), zoom=self._event_zoom)
            raw = pos_img - self.move_start
            self.move_delta = self._clamp_delta_to_box(
                raw.x(), raw.y(), self._move_bbox
            )
            self.update()
            return
        if self.tool == "select" and self.selected_index is not None:
            pos_img = self._to_image_pos(e.position().toPoint(), zoom=self._event_zoom)
            raw = pos_img - self.drag_last
            if not raw.isNull():
                obj = self.objects[self.selected_index]
                box = self._object_bbox(obj)
                allowed = self._clamp_delta_to_box(raw.x(), raw.y(), box)
                if not self._drag_changed:
                    self._push_undo_snapshot()
                    self._drag_changed = True
                self._translate_object(obj, allowed)
                self.update()
            return
        pos_img = self._to_image_pos(e.position().toPoint(), zoom=self._event_zoom)
        if self.tool in ("pen", "eraser"):
            painter = QPainter(self.image)
            pen = QPen(
                Qt.white if self.tool == "eraser" else self.pen_color,
                self.pen_width,
                Qt.SolidLine,
                Qt.RoundCap,
                Qt.RoundJoin,
            )
            painter.setPen(pen)
            painter.drawLine(self.last_pos, pos_img)
            painter.end()
            self.last_pos = pos_img
            self.update()
        else:
            self.preview_pos = pos_img
            self.update()

    def mouseReleaseEvent(self, e):
        if e.button() != Qt.MouseButton.LeftButton or not self.is_drawing:
            return
        try:
            self.releaseMouse()
        except Exception:
            pass
        self.is_drawing = False
        end_img = self._to_image_pos(e.position().toPoint(), zoom=self._event_zoom)
        if self.tool == "move" and self.move_active:
            self.move_active = False
            dx, dy = self.move_delta.x(), self.move_delta.y()
            if dx or dy:
                self._push_undo_snapshot()
                newimg = QImage(self.image.size(), self.image.format())
                newimg.fill(Qt.white)
                p = QPainter(newimg)
                p.setRenderHint(QPainter.Antialiasing, True)
                p.drawImage(dx, dy, self.image_before_move)
                p.end()
                self.image = newimg
                shift = QPoint(dx, dy)
                for obj in self.objects:
                    self._translate_object(obj, shift)
            self.image_before_move = None
            self.move_delta = QPoint()
            self._move_bbox = None
            self.update()
            if self._pending_zoom is not None:
                z = self._pending_zoom
                self._pending_zoom = None
                self.set_zoom(z)
            return
        if self.tool == "select":
            self._drag_changed = False
            self.selected_index = None
            self.update()
            if self._pending_zoom is not None:
                z = self._pending_zoom
                self._pending_zoom = None
                self.set_zoom(z)
            return
        if self.tool in ("rect", "ellipse", "line", "arrow"):
            self._push_undo_snapshot()
            painter = QPainter(self.image)
            pen = QPen(
                self.pen_color,
                self.pen_width,
                Qt.PenStyle.SolidLine,
                Qt.PenCapStyle.RoundCap,
                Qt.PenJoinStyle.RoundJoin,
            )
            painter.setPen(pen)
            brush = Qt.BrushStyle.NoBrush
            if self.tool in ("rect", "ellipse") and self.fill_enabled:
                brush = QBrush(self.fill_color)
            painter.setBrush(brush)
            if self.tool == "rect":
                painter.drawRect(self._norm_rect(self.start_pos, end_img))
            elif self.tool == "ellipse":
                painter.drawEllipse(self._norm_rect(self.start_pos, end_img))
            elif self.tool == "line":
                painter.drawLine(self.start_pos, end_img)
            elif self.tool == "arrow":
                p1, p2 = QPointF(self.start_pos), QPointF(end_img)
                painter.drawLine(p1, p2)
                head = self._arrow_head(p1, p2, max(10, self.pen_width * 3))
                painter.setBrush(QBrush(self.pen_color))
                painter.drawPolygon(head)
            painter.end()
            self.update()
            if self._pending_zoom is not None:
                z = self._pending_zoom
                self._pending_zoom = None
                self.set_zoom(z)

    def paintEvent(self, e):
        painter = QPainter(self)
        painter.fillRect(self.rect(), Qt.white)
        pad = getattr(self, "viewport_pad", 0)
        canvas_w = int(self.image.width() * self.zoom)
        canvas_h = int(self.image.height() * self.zoom)
        painter.fillRect(
            QRect(pad, pad, canvas_w, canvas_h),
            getattr(self, "canvas_bg", QColor("#f3f4f6")),
        )
        painter.translate(pad / self.zoom, pad / self.zoom)
        painter.scale(self.zoom, self.zoom)
        if (
            self.tool == "move"
            and self.move_active
            and self.image_before_move is not None
        ):
            painter.drawImage(self.move_delta, self.image_before_move)
            painter.save()
            painter.translate(self.move_delta)
            self._draw_objects(painter)
            painter.restore()
        else:
            painter.drawImage(0, 0, self.image)
            self._draw_objects(painter)
        if self.is_drawing and self.tool in ("rect", "ellipse", "line", "arrow"):
            pen = QPen(
                self.pen_color, self.pen_width, Qt.SolidLine, Qt.RoundCap, Qt.RoundJoin
            )
            painter.setPen(pen)
            painter.setBrush(
                QBrush(self.fill_color)
                if (self.tool in ("rect", "ellipse") and self.fill_enabled)
                else Qt.NoBrush
            )
            if self.tool == "rect":
                painter.drawRect(self._norm_rect(self.start_pos, self.preview_pos))
            elif self.tool == "ellipse":
                painter.drawEllipse(self._norm_rect(self.start_pos, self.preview_pos))
            elif self.tool == "line":
                painter.drawLine(self.start_pos, self.preview_pos)
            elif self.tool == "arrow":
                p1, p2 = QPointF(self.start_pos), QPointF(self.preview_pos)
                painter.drawLine(p1, p2)
                head = self._arrow_head(p1, p2, max(10, self.pen_width * 3))
                painter.setBrush(QBrush(self.pen_color))
                painter.drawPolygon(head)
        painter.resetTransform()
        outer = QRect(pad - 1, pad - 1, canvas_w + 2, canvas_h + 2)
        inner = QRect(pad, pad, canvas_w, canvas_h)
        if self.show_layout_frame:
            dash = QPen(QColor("#9aa0a6"))
            dash.setStyle(Qt.DashLine)
            dash.setCosmetic(True)
            painter.setPen(dash)
            painter.setBrush(Qt.NoBrush)
            painter.drawRect(QRect(pad, pad, canvas_w, canvas_h))
            m = int(round(self.safe_margin * self.zoom))
            if m > 0 and 2 * m < min(canvas_w, canvas_h):
                painter.drawRect(
                    QRect(pad + m, pad + m, canvas_w - 2 * m, canvas_h - 2 * m)
                )
        pen = QPen(getattr(self, "canvas_border_dark", QColor("#2f3338")), 1)
        pen.setCosmetic(True)
        painter.setPen(pen)
        painter.setBrush(Qt.NoBrush)
        painter.drawRect(outer)
        pen2 = QPen(getattr(self, "canvas_border_light", QColor("#ffffff")), 1)
        pen2.setCosmetic(True)
        painter.setPen(pen2)
        painter.drawRect(inner)
        painter.end()

    def sizeHint(self):
        return QSize(self.image.width(), self.image.height())

    def get_image(self) -> QImage:
        if not self.objects:
            return self.image
        out = QImage(self.image.size(), self.image.format())
        out.fill(Qt.white)
        p = QPainter(out)
        p.drawImage(0, 0, self.image)
        self._draw_objects(p)
        p.end()
        return out


class DrawingDialog(QDialog):
    def __init__(self, parent=None):
        super().__init__(parent)
        self.setWindowTitle("Рисование")
        self.setModal(True)
        self.setWindowFlag(Qt.WindowStaysOnTopHint, True)
        self.resize(1100, 700)
        self.canvas = DrawingCanvas(self)
        self.canvas.setMinimumSize(800, 500)
        sc_undo = QShortcut(QKeySequence.Undo, self)
        sc_undo.setContext(Qt.WidgetWithChildrenShortcut)
        sc_undo.activated.connect(self.canvas.undo)
        sc_redo = QShortcut(QKeySequence.Redo, self)
        sc_redo.setContext(Qt.WidgetWithChildrenShortcut)
        sc_redo.activated.connect(self.canvas.redo)
        tools_layout = QHBoxLayout()
        tools_layout.setContentsMargins(6, 6, 6, 6)
        tools_layout.setSpacing(8)

        def _mk_tool(text: str, tool_name: str):
            btn = QToolButton(self)
            btn.setText(text)
            btn.setCheckable(True)
            btn.clicked.connect(lambda: self.canvas.set_tool(tool_name))
            return btn

        self.btn_select = _mk_tool("☰ Выбор", "select")
        self.btn_pencil = _mk_tool("✎ Карандаш", "pen")
        self.btn_move = _mk_tool("🖐 Перемещение", "move")
        self.btn_rect = _mk_tool("▭ Прямоугольник", "rect")
        self.btn_ellipse = _mk_tool("◯ Эллипс", "ellipse")
        self.btn_line = _mk_tool("— Линия", "line")
        self.btn_arrow = _mk_tool("➤ Стрелка", "arrow")
        self.btn_text = _mk_tool("Т Текст", "text")
        self.btn_eraser = _mk_tool("⌫ Ластик", "eraser")
        grp = QButtonGroup(self)
        grp.setExclusive(True)
        for b in (
            self.btn_select,
            self.btn_pencil,
            self.btn_move,
            self.btn_rect,
            self.btn_ellipse,
            self.btn_line,
            self.btn_arrow,
            self.btn_text,
            self.btn_eraser,
        ):
            grp.addButton(b)
            tools_layout.addWidget(b)
        self.btn_select.setChecked(True)
        self.canvas.set_tool("select")
        tools_layout.addStretch(1)
        opts = QFormLayout()
        opts.setContentsMargins(6, 0, 6, 6)
        opts.setSpacing(8)
        self.cb_frame = QCheckBox("Границы макета", self)
        self.cb_frame.setChecked(True)
        self.cb_frame.toggled.connect(self.canvas.set_show_frame)
        opts.addRow(self.cb_frame)
        self.spin_margin = QSpinBox(self)
        self.spin_margin.setRange(0, 400)
        self.spin_margin.setValue(getattr(self.canvas, "safe_margin", 16))
        self.spin_margin.valueChanged.connect(self.canvas.set_safe_margin)
        opts.addRow("Отступ, px:", self.spin_margin)
        zoom_row = QHBoxLayout()
        self.btn_zoom_minus = QToolButton(self)
        self.btn_zoom_minus.setText("−")
        self.btn_zoom_plus = QToolButton(self)
        self.btn_zoom_plus.setText("+")
        self.lbl_zoom = QLabel(self)
        self.lbl_zoom.setMinimumWidth(48)
        self.lbl_zoom.setAlignment(Qt.AlignCenter)

        def _update_zoom_label():
            self.lbl_zoom.setText(f"{int(round(self.canvas.zoom * 100))}%")

        def _step_zoom(sign: int):
            z = self.canvas.zoom
            new_z = max(0.1, min(5.0, round((z + sign * 0.05), 2)))
            self.canvas.set_zoom(new_z)
            _update_zoom_label()

        self.btn_zoom_minus.clicked.connect(lambda: _step_zoom(-1))
        self.btn_zoom_plus.clicked.connect(lambda: _step_zoom(+1))
        _update_zoom_label()
        zoom_row.addWidget(self.btn_zoom_minus)
        zoom_row.addWidget(self.lbl_zoom)
        zoom_row.addWidget(self.btn_zoom_plus)
        opts.addRow("Масштаб:", zoom_row)
        self.spin_width = QSpinBox(self)
        self.spin_width.setRange(1, 40)
        self.spin_width.setValue(getattr(self.canvas, "pen_width", 3))
        self.spin_width.valueChanged.connect(self.canvas.set_width)
        opts.addRow("Толщина:", self.spin_width)

        def _btn_color(title: str, getter, setter):
            btn = QPushButton(title, self)

            def pick():
                c = QColorDialog.getColor(getter(), self, title)
                if c.isValid():
                    setter(c)
                    btn.setStyleSheet(f"background: {c.name()};")

            btn.clicked.connect(pick)
            try:
                btn.setStyleSheet(f"background: {getter().name()};")
            except Exception:
                pass
            return btn

        line_color_btn = _btn_color(
            "Цвет линии…",
            lambda: getattr(self.canvas, "pen_color"),
            self.canvas.set_color,
        )
        self.btn_fill = _btn_color(
            "Цвет заливки…",
            lambda: getattr(self.canvas, "fill_color"),
            self.canvas.set_fill_color,
        )

        self.cb_fill = QCheckBox("Заливка", self)
        self.cb_fill.setChecked(getattr(self.canvas, "fill_enabled", False))
        self.cb_fill.toggled.connect(self.on_fill_toggled)
        self.btn_fill.setEnabled(self.cb_fill.isChecked())
        opts.addRow(line_color_btn)
        opts.addRow(self.btn_fill)
        opts.addRow(self.cb_fill)
        self.spin_text = QSpinBox(self)
        self.spin_text.setRange(6, 96)
        self.spin_text.setValue(getattr(self.canvas, "text_size", 18))
        self.spin_text.valueChanged.connect(self.canvas.set_text_size)
        opts.addRow("Размер текста:", self.spin_text)
        self.btn_clear = QPushButton("Очистить", self)
        self.btn_clear.clicked.connect(self.canvas.clear)
        opts.addRow(self.btn_clear)
        buttons = QDialogButtonBox(QDialogButtonBox.Ok | QDialogButtonBox.Cancel, self)
        buttons.accepted.connect(self.accept)
        buttons.rejected.connect(self.reject)
        root = QVBoxLayout(self)
        tools_wrap = QWidget(self)
        tools_wrap.setLayout(tools_layout)
        tools_wrap.setSizePolicy(QSizePolicy.Expanding, QSizePolicy.Fixed)
        root.addWidget(tools_wrap)
        root.addWidget(self.canvas, 1)
        opts_wrap = QWidget(self)
        opts_wrap.setLayout(opts)
        root.addWidget(opts_wrap)
        root.addWidget(buttons)
        _update_zoom_label()

    def _select_tool(self, btn: QPushButton, tool: str):
        for b in (
            self.btn_pencil,
            self.btn_rect,
            self.btn_ellipse,
            self.btn_line,
            self.btn_arrow,
            self.btn_text,
            self.btn_eraser,
            self.btn_move,
            self.btn_select,
        ):
            b.setChecked(b is btn)
        self.canvas.set_tool(tool)

    def _choose_pen_color(self):
        col = QColorDialog.getColor(self.canvas.pen_color, self, "Цвет линии")
        if col.isValid():
            self.canvas.set_color(col)

    def on_fill_toggled(self, on: bool) -> None:
        self.btn_fill.setEnabled(on)
        if hasattr(self.canvas, "set_fill_enabled"):
            self.canvas.set_fill_enabled(on)

    def _choose_fill_color(self):
        col = QColorDialog.getColor(
            self.canvas.fill_color,
            self,
            "Цвет заливки",
            QColorDialog.ColorDialogOption.ShowAlphaChannel,
        )
        if col.isValid():
            self.canvas.set_fill_color(col)

    def get_image(self) -> QImage:
        return self.canvas.get_image()

    def showEvent(self, e):
        super().showEvent(e)
        p = self.parent()
        if p and isinstance(p, QWidget):
            fg = self.frameGeometry()
            fg.moveCenter(p.frameGeometry().center())
            self.move(fg.topLeft())


class AudioRecorderThread(QThread):
    recording_finished = Signal(str)

    def __init__(self, file_path):
        super().__init__()
        self.file_path = file_path
        self._running = True
        self._wavefile = None

    def callback(self, indata, frames, time, status):
        if self._running and self._wavefile:
            self._wavefile.writeframes(indata.tobytes())

    def run(self):
        try:
            import sounddevice as sd

            with wave.open(self.file_path, "wb") as wf:
                self._wavefile = wf
                wf.setnchannels(1)
                wf.setsampwidth(2)
                wf.setframerate(44100)
                with sd.InputStream(
                    samplerate=44100,
                    channels=1,
                    dtype="int16",
                    callback=self.callback,
                ):
                    while self._running:
                        sd.sleep(100)
            self.recording_finished.emit(self.file_path)
        except Exception as e:
            print("Ошибка записи:", e)
        finally:
            self._wavefile = None

    def stop(self):
        self._running = False


class FlowLayout(QLayout):
    def __init__(self, parent=None, margin=0, spacing=6):
        super().__init__(parent)
        self.setContentsMargins(margin, margin, margin, margin)
        self.setSpacing(spacing)
        self.itemList = []

    def addItem(self, item: QLayoutItem) -> None:
        self.itemList.append(item)

    def count(self) -> int:
        return len(self.itemList)

    def itemAt(self, index: int) -> QLayoutItem | None:
        return self.itemList[index] if index < len(self.itemList) else None

    def takeAt(self, index: int) -> QLayoutItem | None:
        if 0 <= index < len(self.itemList):
            return self.itemList.pop(index)
        return None

    def expandingDirections(self):
        return Qt.Orientations()

    def hasHeightForWidth(self):
        return True

    def heightForWidth(self, width: int) -> int:
        height = self.doLayout(QRect(0, 0, width, 0), True)
        return height

    def setGeometry(self, rect: QRect) -> None:
        super().setGeometry(rect)
        self.doLayout(rect, False)

    def doLayout(self, rect: QRect, testOnly: bool) -> int:
        x = rect.x()
        y = rect.y()
        lineHeight = 0

        for item in self.itemList:
            wid = item.widget()
            spaceX = self.spacing() + wid.style().layoutSpacing(
                QSizePolicy.PushButton, QSizePolicy.PushButton, Qt.Horizontal
            )
            spaceY = self.spacing() + wid.style().layoutSpacing(
                QSizePolicy.PushButton, QSizePolicy.PushButton, Qt.Vertical
            )
            if wid is not None and wid.property("flow_break") is True:
                x = rect.x()
                y = y + lineHeight + spaceY
                lineHeight = 0
                continue
            nextX = x + item.sizeHint().width() + spaceX
            if nextX - spaceX > rect.right() and lineHeight > 0:
                x = rect.x()
                y = y + lineHeight + spaceY
                nextX = x + item.sizeHint().width() + spaceX
                lineHeight = 0
            if not testOnly:
                item.setGeometry(QRect(QPoint(x, y), item.sizeHint()))
            x = nextX
            lineHeight = max(lineHeight, item.sizeHint().height())
        return y + lineHeight - rect.y()


class CustomTextEdit(QTextEdit):
    def __init__(
        self,
        parent: QWidget | None = None,
        paste_image_callback: Callable | None = None,
    ):
        super().__init__(parent)
        self.paste_image_callback = paste_image_callback
        self.setFont(QFont("Times New Roman", 14))
        self.setCurrentFont(QFont("Times New Roman"))
        self.setFontPointSize(14)
        self.document().setDefaultFont(QFont("Times New Roman", 14))
        self.setAcceptDrops(True)

    def insertFromMimeData(self, source: QMimeData) -> None:
        if source.hasImage() and self.paste_image_callback:
            image = source.imageData()
            self.paste_image_callback(image)
        elif source.hasHtml():
            html = source.html()
            html = re.sub(r"font-family:[^;\"]*;?", "", html, flags=re.IGNORECASE)
            html = re.sub(r"font-size:[^;\"]*;?", "", html, flags=re.IGNORECASE)
            wrapped = f"<span style=\"font-family:'Times New Roman'; font-size:14pt;\">{html}</span>"
            self.textCursor().insertHtml(wrapped)
        elif source.hasText():
            text = source.text()
            try:
                html = self._linkify_plain_text(text)
            except Exception:
                html = None
            if html is not None:
                wrapped = f"<span style=\"font-family:'Times New Roman'; font-size:14pt;\">{html}</span>"
                self.textCursor().insertHtml(wrapped)
            else:
                fmt = QTextCharFormat()
                fmt.setFont(QFont("Times New Roman", 14))
                self.textCursor().insertText(text, fmt)
        else:
            super().insertFromMimeData(source)

    def _linkify_plain_text(self, text: str) -> str | None:
        if not text:
            return None
        token_re = re.compile(
            r"(?:https?://|ftp://)[^\s<>\"']+"
            r"|www\.[^\s<>\"']+"
            r"|[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Za-z]{2,}"
            r"|[A-Za-z]:[\\/][^\s<>\"']+"
            r"|/[^\s<>\"']+"
        )
        found_any = False
        out_parts: list[str] = []
        pos = 0
        for m in token_re.finditer(text):
            if m.start() > pos:
                out_parts.append(html_lib.escape(text[pos : m.start()]))
            raw = m.group(0)
            href = None
            label = html_lib.escape(raw)
            if re.match(r"^(?:https?://|ftp://)", raw, flags=re.I):
                href = raw
            elif raw.lower().startswith("www."):
                href = "https://" + raw
            elif re.match(r"^[A-Za-z0-9._%+-]+@", raw):
                href = f"mailto:{raw}"
            elif os.path.isabs(raw):
                try:
                    href = QUrl.fromLocalFile(os.path.abspath(raw)).toString()
                except Exception:
                    href = "file://" + quote(os.path.abspath(raw))
            if href:
                found_any = True
                out_parts.append(
                    f'<a href="{html_lib.escape(href)}" style="text-decoration: underline;">{label}</a>'
                )
            else:
                out_parts.append(label)
            pos = m.end()
        if pos < len(text):
            out_parts.append(html_lib.escape(text[pos:]))
        if not found_any:
            return None
        html = "".join(out_parts).replace("\r\n", "<br>").replace("\n", "<br>")
        html = f"<span style=\"font-family:'Times New Roman'; font-size:14pt;\">{html}</span>"
        return html

    def ignore_in_this_note(self, word: str) -> None:
        mw = self.window()
        if hasattr(mw, "add_word_to_note_ignore"):
            mw.add_word_to_note_ignore(word)

    def copy_without_formatting(self) -> None:
        cur = self.textCursor()
        plain = (
            cur.selection().toPlainText() if cur.hasSelection() else self.toPlainText()
        )
        md = QMimeData()
        md.setText(self._sanitize_plain_for_copy(plain))
        QApplication.clipboard().setMimeData(md)

    def copy_with_formatting(self) -> None:
        md = self._build_rich_mime_from_selection()
        QApplication.clipboard().setMimeData(md)

    def dragEnterEvent(self, event: QDragEnterEvent) -> None:
        if event.mimeData().hasUrls():
            event.acceptProposedAction()
        else:
            super().dragEnterEvent(event)

    def keyPressEvent(self, event):
        if event.key() in (Qt.Key_Return, Qt.Key_Enter):
            cursor = self.textCursor()
            block_text = cursor.block().text()
            normalized = block_text.lstrip("\u202a\u202c\ufeff\u200b\xa0 ").replace(
                "\xa0", " "
            )
            m = re.match(r"^[☐☑][ \u00A0]", normalized)
            if m:
                after_marker = normalized[m.end() :]
                if after_marker.strip() == "":
                    super().keyPressEvent(event)
                    return
                else:
                    super().keyPressEvent(event)
                    cursor = self.textCursor()
                    cursor.insertText("☐ ")
                    self.setTextCursor(cursor)
                    return
            if block_text.strip().startswith("• "):
                if block_text.strip() == "•":
                    cursor.movePosition(QTextCursor.StartOfBlock)
                    cursor.movePosition(QTextCursor.Right, QTextCursor.KeepAnchor, 2)
                    cursor.removeSelectedText()
                    super().keyPressEvent(event)
                    return
                else:
                    super().keyPressEvent(event)
                    cursor = self.textCursor()
                    cursor.insertText("• ")
                    self.setTextCursor(cursor)
                    return
            num_match = re.match(r"^(\d+)\. ", block_text.strip())
            if num_match:
                number = int(num_match.group(1))
                if block_text.strip() == f"{number}.":
                    cursor.movePosition(QTextCursor.StartOfBlock)
                    cursor.movePosition(
                        QTextCursor.Right,
                        QTextCursor.KeepAnchor,
                        len(num_match.group(0)),
                    )
                    cursor.removeSelectedText()
                    super().keyPressEvent(event)
                    return
                else:
                    super().keyPressEvent(event)
                    cursor = self.textCursor()
                    cursor.insertText(f"{number + 1}. ")
                    self.setTextCursor(cursor)
                    return
        super().keyPressEvent(event)

    def dropEvent(self, event: QDropEvent) -> None:
        if event.mimeData().hasUrls():
            for url in event.mimeData().urls():
                file_path = url.toLocalFile()
                if os.path.isfile(file_path):
                    main_window = self.window()
                    if hasattr(main_window, "attach_file_to_note_external"):
                        main_window.attach_file_to_note_external(file_path)
            event.acceptProposedAction()
        else:
            super().dropEvent(event)

    def _sanitize_windows_path(self, p: str) -> str:
        if not p:
            return p
        p = p.strip().strip('"')
        bad = "\ufeff\u200b\u200c\u200d\u200e\u200f\u202a\u202b\u202c\u202d\u202e\u2066\u2067\u2068\u2069"
        p = p.translate({ord(ch): None for ch in bad})
        p = p.replace("/", os.sep)
        parts = [re.sub(r"[ \.]+$", "", part) for part in p.split(os.sep)]
        p = os.sep.join(parts)
        return os.path.normpath(p)

    def _open_any_link(self, link: str) -> bool:
        try:
            s = html_lib.unescape((link or "").strip())
            url = QUrl.fromUserInput(s)
            is_file_like = (url.scheme().lower() == "file") or re.match(
                r"^[a-zA-Z]:[\\/]", s
            )
            if is_file_like:
                if url.scheme().lower() == "file":
                    local_path = url.toLocalFile()
                    if not local_path:
                        raw_path = (
                            QUrl(s).path() if s.lower().startswith("file://") else s
                        )
                        raw_path = raw_path or ""
                        raw_path = unquote(raw_path).lstrip("/").replace("/", os.sep)
                        local_path = raw_path
                else:
                    local_path = s
                local_path = self._sanitize_windows_path(local_path)

                def try_open(p: str) -> bool:
                    if sys.platform.startswith("win"):
                        try:
                            os.startfile(p)
                            return True
                        except Exception:
                            pass
                        try:
                            from PySide6.QtCore import QProcess

                            return QProcess.startDetached(
                                "cmd", ["/c", "start", "", f'"{p}"']
                            )
                        except Exception:
                            pass
                    return QDesktopServices.openUrl(QUrl.fromLocalFile(p))

                if os.path.exists(local_path):
                    return try_open(local_path)
                lower = local_path.replace("/", os.sep).replace("\\", os.sep).lower()
                sep_notes = f"{os.sep}notes{os.sep}"
                if sep_notes in lower:
                    suffix = local_path[
                        len(lower) - len(lower.split(sep_notes, 1)[1]) :
                    ]
                    candidate = os.path.join(NOTES_DIR, suffix)
                    if os.path.exists(candidate):
                        return try_open(candidate)
                mw = self.window()
                if hasattr(mw, "current_note") and mw.current_note:
                    folder = NotesApp.safe_folder_name(
                        mw.current_note.title,
                        mw.current_note.uuid,
                        mw.current_note.timestamp,
                    )
                    candidate = os.path.join(
                        NOTES_DIR, folder, os.path.basename(local_path)
                    )
                    if os.path.exists(candidate):
                        return try_open(candidate)
                return try_open(local_path)
            if not url.scheme():
                url = QUrl.fromUserInput("https://" + s)
            return QDesktopServices.openUrl(url)
        except Exception:
            return False

    def mousePressEvent(self, event: QMouseEvent) -> None:
        pos = event.position().toPoint()
        cursor = self.cursorForPosition(pos)
        char_format = cursor.charFormat()
        if event.button() == Qt.LeftButton:
            if char_format.isAnchor():
                link = char_format.anchorHref()
                if link.startswith("dropdown://"):
                    tail = link.split("://", 1)[1]
                    dd_id = tail.split("?", 1)[0].split("#", 1)[0]
                    main_window = self.window()
                    if hasattr(main_window, "show_dropdown_menu_for_token"):
                        rect = self.cursorRect(cursor)
                        global_pos = self.viewport().mapToGlobal(rect.bottomLeft())
                        main_window.show_dropdown_menu_for_token(dd_id, global_pos)
                    return
                if event.modifiers() & Qt.ControlModifier:
                    self._open_any_link(link)
                    return
                super().mousePressEvent(event)
                return
            if char_format.isImageFormat():
                if event.modifiers() & Qt.ControlModifier:
                    image_format = char_format.toImageFormat()
                    image_path = image_format.name()
                    if image_path.startswith("Data:image"):
                        header, b64data = image_path.split(",", 1)
                        suffix = ".png" if "png" in header else ".jpg"
                        with tempfile.NamedTemporaryFile(
                            delete=False, suffix=suffix
                        ) as tmpfile:
                            tmpfile.write(base64.b64decode(b64data))
                            tmpfile_path = tmpfile.name
                        QDesktopServices.openUrl(QUrl.fromLocalFile(tmpfile_path))
                    elif image_path.startswith("file://"):
                        file_path = QUrl(image_path).toLocalFile()
                        if os.path.exists(file_path):
                            self._open_any_link(image_path)
                else:
                    super().mousePressEvent(event)
                return
            if event.modifiers() & Qt.ControlModifier:
                block_cursor = self.cursorForPosition(pos)
                block_cursor.select(QTextCursor.SelectionType.BlockUnderCursor)
                block_html = block_cursor.selection().toHtml()
                m = re.search(r'href="((?:file|https?)://[^"]+)"', block_html)
                if m:
                    self._open_any_link(html_lib.unescape(m.group(1)))
                    return
            word_cursor = self.cursorForPosition(pos)
            word_cursor.select(QTextCursor.SelectionType.WordUnderCursor)
            word = word_cursor.selectedText()
            if word == "☐":
                word_cursor.insertText("☑")
                return
            elif word == "☑":
                word_cursor.insertText("☐")
                return
        super().mousePressEvent(event)

    def mouseDoubleClickEvent(self, event: QMouseEvent) -> None:
        pos = event.position().toPoint()
        cursor = self.cursorForPosition(pos)
        cf = cursor.charFormat()
        if cf.isAnchor():
            if event.modifiers() & Qt.ControlModifier:
                self._open_any_link(cf.anchorHref())
                return
            super().mouseDoubleClickEvent(event)
            return
        super().mouseDoubleClickEvent(event)

    def mouseMoveEvent(self, event: QMouseEvent) -> None:
        pos = event.position().toPoint()
        cf = self.cursorForPosition(pos).charFormat()
        if cf.isAnchor() and (event.modifiers() & Qt.ControlModifier):
            self.viewport().setCursor(Qt.PointingHandCursor)
        else:
            self.viewport().setCursor(Qt.IBeamCursor)
        super().mouseMoveEvent(event)

    def leaveEvent(self, event: QEvent) -> None:
        self.viewport().unsetCursor()
        super().leaveEvent(event)

    def create_drag_image(self, widget: QWidget) -> QPixmap:
        pixmap = widget.grab()
        scale_factor = 0.4
        scaled_pixmap = pixmap.scaled(
            pixmap.size() * scale_factor, Qt.KeepAspectRatio, Qt.SmoothTransformation
        )
        painter = QPainter(scaled_pixmap)
        path = QPainterPath()
        path.addRect(0, 0, scaled_pixmap.width() - 1, scaled_pixmap.height() - 1)
        pen = QPen(QColor(0, 0, 0), 1)
        painter.setPen(pen)
        painter.drawPath(path)
        return scaled_pixmap

    def delete_selection(self) -> None:
        cursor = self.textCursor()
        if cursor.hasSelection():
            cursor.removeSelectedText()

    def replace_word(self, cursor: QTextCursor, new_word: str) -> None:
        cursor.beginEditBlock()
        cursor.removeSelectedText()
        cursor.insertText(new_word)
        cursor.endEditBlock()

    def add_to_dictionary(self, word: str) -> None:
        highlighter = getattr(self.window(), "spell_highlighter", None)
        if highlighter:
            highlighter.add_to_dictionary(word)

    def contextMenuEvent(self, event: QContextMenuEvent) -> None:
        menu = QMenu(self)
        menu.setStyleSheet(CUSTOM_MENU_STYLE)
        undo_action = QAction("Undo", self)
        undo_action.setShortcut("Ctrl+Z")
        menu.addAction(undo_action)
        redo_action = QAction("Redo", self)
        redo_action.setShortcut("Ctrl+Y")
        undo_action.triggered.connect(lambda: self.window().undo() if hasattr(self.window(), "undo") else None)
        redo_action.triggered.connect(lambda: self.window().redo() if hasattr(self.window(), "redo") else None)
        menu.addAction(redo_action)
        menu.addSeparator()
        cut_action = QAction("Cut", self)
        cut_action.setShortcut("Ctrl+X")
        cut_action.triggered.connect(self.cut)
        menu.addAction(cut_action)
        copy_action = QAction("Copy", self)
        copy_action.setShortcut("Ctrl+C")
        copy_action.triggered.connect(self.copy)
        menu.addAction(copy_action)
        paste_action = QAction("Paste", self)
        paste_action.setShortcut("Ctrl+V")
        paste_action.triggered.connect(self.paste)
        menu.addAction(paste_action)
        delete_action = QAction("Delete", self)
        delete_action.triggered.connect(self.delete_selection)
        menu.addAction(delete_action)
        menu.addSeparator()
        self._inject_table_context_menu(menu, event.pos())
        menu.addSeparator()
        select_all_action = QAction("Select All", self)
        select_all_action.setShortcut("Ctrl+A")
        select_all_action.triggered.connect(self.selectAll)
        menu.addAction(select_all_action)
        menu.addSeparator()
        view_dict_action = QAction("View Dictionary", self)
        view_dict_action.triggered.connect(
            lambda checked=False: QDesktopServices.openUrl(
                QUrl.fromLocalFile(USER_DICT_PATH)
            )
        )
        menu.addAction(view_dict_action)
        word_cursor = self.cursorForPosition(event.pos())
        word_cursor.select(QTextCursor.WordUnderCursor)
        raw_word = word_cursor.selectedText()
        m = re.search(r"[A-Za-zА-Яа-яЁё']+", raw_word or "")
        if not m:
            menu.exec(self.mapToGlobal(event.pos()))
            return
        lw = m.group(0).lower()
        highlighter = getattr(self.window(), "spell_highlighter", None)
        spell = getattr(highlighter, "spell_checker", None)
        user_candidates = []
        if highlighter and getattr(highlighter, "user_words", None):
            try:
                user_candidates = difflib.get_close_matches(
                    lw, list(highlighter.user_words), n=7, cutoff=0.6
                )
            except Exception:
                user_candidates = []
        spell_candidates = []
        morph_candidates = []
        try:
            if spell and re.search(r"[А-Яа-яЁё]", lw):
                ref = ""
                word_core = lw
                for r in RU_REFLEXIVE:
                    if word_core.endswith(r) and len(word_core) - len(r) >= 3:
                        word_core = word_core[: -len(r)]
                        ref = r
                        break
                suf = ""
                for s in RU_SUFFIXES_RU:
                    if word_core.endswith(s) and len(word_core) - len(s) >= 3:
                        suf = s
                        word_core = word_core[: -len(s)]
                        break
                base_cands = set()
                try:
                    corr = spell.correction(word_core)
                    if corr:
                        base_cands.add(corr)
                except Exception:
                    pass
                try:
                    base_cands |= set(spell.candidates(word_core))
                except Exception:
                    pass
                for bc in base_cands:
                    cand = bc + suf + ref
                    morph_candidates.append(cand)
        except Exception:
            morph_candidates = []
        if spell:
            try:
                corr = spell.correction(lw)
                cands = spell.candidates(lw) if hasattr(spell, "candidates") else set()
                spell_candidates = [corr] + sorted(cands)
            except Exception:
                spell_candidates = []
        doc_candidates = []
        try:
            plain = self.window().text_edit.toPlainText()
            tokens = set(w.lower() for w in re.findall(r"[A-Za-zА-Яа-яЁё']+", plain))
            if lw in tokens:
                tokens.discard(lw)
            sample = list(tokens)[:5000]
            doc_candidates = difflib.get_close_matches(lw, sample, n=7, cutoff=0.65)
        except Exception:
            doc_candidates = []

        def _clean(seq):
            seen = set()
            out = []
            for s in seq:
                if not s:
                    continue
                sl = s.lower()
                if sl == lw or sl in seen:
                    continue
                out.append(s)
                seen.add(sl)
                if len(out) >= 7:
                    break
            return out

        spell_candidates = _clean(morph_candidates + spell_candidates)

        user_candidates = _clean(user_candidates)
        spell_candidates = _clean(spell_candidates)
        doc_candidates = _clean(doc_candidates)
        menu.addSeparator()
        copy_plain_action = QAction("Copy (without formatting)", self)
        copy_plain_action.triggered.connect(self.copy_without_formatting)
        copy_rich_action = QAction("Copy (with formatting)", self)
        copy_rich_action.triggered.connect(self.copy_with_formatting)
        menu.addAction(copy_plain_action)
        menu.addAction(copy_rich_action)
        menu.addSeparator()
        ignore_action = QAction("Ignore the word in the current note", self)
        ignore_action.triggered.connect(
            lambda checked=False, w=raw_word: self.ignore_in_this_note(w)
        )
        menu.addAction(ignore_action)
        add_dict_action = QAction("Add to Dictionary", self)
        add_dict_action.triggered.connect(
            lambda checked=False, w=raw_word: self.add_to_dictionary(w)
        )
        menu.addAction(add_dict_action)
        replace_menu = menu.addMenu("Replace")
        user_menu = replace_menu.addMenu("From your personal dictionary")
        if user_candidates:
            for s in user_candidates:
                act = user_menu.addAction(s)
                act.triggered.connect(
                    lambda _, s=s, c=QTextCursor(word_cursor): self.replace_word(c, s)
                )
        else:
            a = user_menu.addAction("— There are no options —")
            a.setEnabled(False)

        spell_menu = replace_menu.addMenu("Spelling options")
        if spell_candidates:
            for s in spell_candidates:
                act = spell_menu.addAction(s)
                act.triggered.connect(
                    lambda _, s=s, c=QTextCursor(word_cursor): self.replace_word(c, s)
                )
        else:
            a = spell_menu.addAction("— There are no options —")
            a.setEnabled(False)
        doc_menu = replace_menu.addMenu("From the current note")
        if doc_candidates:
            for s in doc_candidates:
                act = doc_menu.addAction(s)
                act.triggered.connect(
                    lambda _, s=s, c=QTextCursor(word_cursor): self.replace_word(c, s)
                )
        else:
            a = doc_menu.addAction("— There are no options —")
            a.setEnabled(False)
        menu.exec(self.mapToGlobal(event.pos()))

    def _current_table_pos(self):
        cur = self.textCursor()
        tbl = cur.currentTable()
        if not tbl:
            return None
        cell = tbl.cellAt(cur)
        if not cell.isValid():
            return None
        return tbl, cell.row(), cell.column()

    def _inject_table_context_menu(self, menu: QMenu, pos) -> None:
        cur = self.cursorForPosition(pos)
        tbl = cur.currentTable()
        if tbl:
            cell = tbl.cellAt(cur)
            if not cell.isValid():
                return
            r, c = cell.row(), cell.column()
            tmenu = menu.addMenu("Table")
            a_ins_row_above = tmenu.addAction("Вставить строку сверху")
            a_ins_row_below = tmenu.addAction("Вставить строку снизу")
            a_del_row = tmenu.addAction("Удалить текущую строку")
            tmenu.addSeparator()
            a_ins_col_left = tmenu.addAction("Вставить столбец слева")
            a_ins_col_right = tmenu.addAction("Вставить столбец справа")
            a_del_col = tmenu.addAction("Удалить текущий столбец")
            tmenu.addSeparator()
            a_split_h = tmenu.addAction("Разделить ячейку по горизонтали (2 строки)")
            a_split_v = tmenu.addAction("Разделить ячейку по вертикали (2 столбца)")
            a_merge = tmenu.addAction("Слить выделенные ячейки")
            tmenu.addSeparator()
            a_del_cell = tmenu.addAction("Удалить ячейку (под курсором)")
            a_del_table = tmenu.addAction("Удалить таблицу целиком")
            a_ins_row_above.triggered.connect(
                lambda: self._table_insert_row(r, above=True)
            )
            a_ins_row_below.triggered.connect(
                lambda: self._table_insert_row(r, above=False)
            )
            a_del_row.triggered.connect(lambda: self._table_delete_row(r))
            a_ins_col_left.triggered.connect(
                lambda: self._table_insert_col(c, left=True)
            )
            a_ins_col_right.triggered.connect(
                lambda: self._table_insert_col(c, left=False)
            )
            a_del_col.triggered.connect(lambda: self._table_delete_col(c))
            a_split_h.triggered.connect(lambda: self._table_split_cell(rows=2, cols=1))
            a_split_v.triggered.connect(lambda: self._table_split_cell(rows=1, cols=2))
            a_merge.triggered.connect(self._table_merge_selection)
            a_del_cell.triggered.connect(self._table_delete_cell)
            a_del_table.triggered.connect(self._table_delete_table)
        else:
            tmenu = menu.addMenu("Table")
            a_new = tmenu.addAction("Вставить таблицу 2×2")
            a_new.triggered.connect(lambda: self._insert_table(rows=2, cols=2))

    def _table_delete_cell(self) -> None:
        cur = self.textCursor()
        tbl = cur.currentTable()
        if not tbl:
            return
        cell = tbl.cellAt(cur)
        if not cell.isValid():
            return
        r, c = cell.row(), cell.column()
        rows, cols = tbl.rows(), tbl.columns()
        try:
            rs, cs = cell.rowSpan(), cell.columnSpan()
            if rs > 1 or cs > 1:
                tbl.splitCell(r, c, 1, 1)
                rows, cols = tbl.rows(), tbl.columns()
        except Exception:
            pass
        if cols > 1:
            tbl.removeColumns(c, 1)
            return
        if rows > 1:
            tbl.removeRows(r, 1)
            return
        self._table_delete_table()

    def _table_delete_table(self) -> None:
        cur = self.textCursor()
        tbl = cur.currentTable()
        if not tbl:
            return
        first = tbl.cellAt(0, 0).firstCursorPosition()
        last = tbl.cellAt(tbl.rows() - 1, tbl.columns() - 1).lastCursorPosition()
        c = QTextCursor(first)
        c.setPosition(last.position(), QTextCursor.KeepAnchor)
        c.beginEditBlock()
        c.removeSelectedText()
        c.endEditBlock()

    def _insert_table(self, rows=2, cols=2) -> None:
        cur = self.textCursor()
        table_fmt = QTextTableFormat()
        table_fmt.setCellPadding(4)
        table_fmt.setCellSpacing(0)
        border_qcolor = self.palette().text().color()
        table_fmt.setBorder(1)
        try:
            table_fmt.setBorderBrush(QBrush(border_qcolor))
        except Exception:
            pass
        table = cur.insertTable(rows, cols, table_fmt)
        cell_fmt = QTextTableCellFormat()
        cell_fmt.setBorder(1)
        try:
            cell_fmt.setBorderStyle(QTextFrameFormat.BorderStyle_Solid)
            cell_fmt.setBorderBrush(QBrush(border_qcolor))
        except Exception:
            pass
        for r in range(rows):
            for c in range(cols):
                try:
                    table.cellAt(r, c).setFormat(cell_fmt)
                except Exception:
                    pass

    def _table_insert_row(self, row: int, *, above: bool) -> None:
        cur = self.textCursor()
        tbl = cur.currentTable()
        if not tbl:
            return
        pos = max(0, min(row + (0 if above else 1), tbl.rows()))
        tbl.insertRows(pos, 1)

    def _table_delete_row(self, row: int) -> None:
        cur = self.textCursor()
        tbl = cur.currentTable()
        if not tbl:
            return
        if 0 <= row < tbl.rows():
            tbl.removeRows(row, 1)

    def _table_insert_col(self, col: int, *, left: bool) -> None:
        cur = self.textCursor()
        tbl = cur.currentTable()
        if not tbl:
            return
        pos = max(0, min(col + (0 if left else 1), tbl.columns()))
        tbl.insertColumns(pos, 1)

    def _table_delete_col(self, col: int) -> None:
        cur = self.textCursor()
        tbl = cur.currentTable()
        if not tbl:
            return
        if 0 <= col < tbl.columns():
            tbl.removeColumns(col, 1)

    def _table_split_cell(self, *, rows: int, cols: int) -> None:
        cur = self.textCursor()
        tbl = cur.currentTable()
        if not tbl:
            return
        cell = tbl.cellAt(cur)
        if not cell.isValid():
            return
        r, c = cell.row(), cell.column()
        rs, cs = cell.rowSpan(), cell.columnSpan()
        if (rows > 1 and rs > 1) or (cols > 1 and cs > 1):
            try:
                tbl.splitCell(r, c, rs, cs)
            except Exception:
                pass
            return
        if rows > 1 and cols == 1:
            tbl.insertRows(r + 1, 1)
            ncols = tbl.columns()
            for j in range(ncols):
                if j == c:
                    continue
                try:
                    tbl.mergeCells(r, j, 2, 1)
                except Exception:
                    pass
            return
        if cols > 1 and rows == 1:
            tbl.insertColumns(c + 1, 1)
            nrows = tbl.rows()
            for i in range(nrows):
                if i == r:
                    continue
                try:
                    tbl.mergeCells(i, c, 1, 2)
                except Exception:
                    pass
            return

    def _table_merge_selection(self) -> None:
        cur = self.textCursor()
        tbl = cur.currentTable()
        if not tbl:
            return
        if not cur.hasSelection():
            cell = tbl.cellAt(cur)
            if cell.isValid():
                tbl.mergeCells(cell.row(), cell.column(), 1, 1)
            return
        start = min(cur.selectionStart(), cur.selectionEnd() - 1)
        end = max(cur.selectionStart(), cur.selectionEnd() - 1)
        tl = tbl.cellAt(start)
        br = tbl.cellAt(end)
        if not (tl.isValid() and br.isValid()):
            return
        r1, c1 = tl.row(), tl.column()
        r2, c2 = br.row(), br.column()
        top, left = min(r1, r2), min(c1, c2)
        nrows = abs(r2 - r1) + 1
        ncols = abs(c2 - c1) + 1
        tbl.mergeCells(top, left, nrows, ncols)

    def _sanitize_plain_for_copy(self, s: str) -> str:
        s = s.replace("▾", "").replace("▼", "")
        s = re.sub(r"[\u200b\u200c\u200d\u200e\u200f\u202a-\u202e\u2066-\u2069]", "", s)
        s = re.sub(r"[ \t]+(\r?\n)", r"\1", s)
        return s

    def _sanitize_html_for_copy(self, html: str) -> str:
        if not html:
            return html
        html = html.replace("▾", "").replace("▼", "")
        html = re.sub(r"&(?:#9(?:660|662);|#9660;|#9662;)", "", html, flags=re.I)
        html = re.sub(
            r"[\u200b\u200c\u200d\u200e\u200f\u202a-\u202e\u2066-\u2069]", "", html
        )
        html = re.sub(
            r'<a\b[^>]*\bhref\s*=\s*["\']dropdown://[^"\']*["\'][^>]*>(.*?)</a>',
            r"\1",
            html,
            flags=re.I | re.S,
        )
        html = re.sub(
            r'\bhref\s*=\s*["\']dropdown://[^"\']*["\']', "", html, flags=re.I
        )
        return html

    def _build_rich_mime_from_selection(self) -> QMimeData:
        cur = self.textCursor()
        if cur.hasSelection():
            frag = cur.selection()
            html = frag.toHtml()
            plain = frag.toPlainText()
        else:
            html = self.document().toHtml()
            plain = self.toPlainText()
        html = self._sanitize_html_for_copy(html)
        plain = self._sanitize_plain_for_copy(plain)
        md = QMimeData()
        md.setHtml(html)
        md.setText(plain)
        return md

    def createMimeDataFromSelection(self):
        if self.textCursor().charFormat().isImageFormat():
            return super().createMimeDataFromSelection()
        return self._build_rich_mime_from_selection()

    def startDrag(self, supportedActions: Qt.DropActions) -> None:
        cursor = self.textCursor()
        if cursor.charFormat().isImageFormat():
            image_format = cursor.charFormat().toImageFormat()
            image_path = image_format.name()
            if os.path.exists(image_path):
                drag = QDrag(self)
                mimeData = self.createMimeDataFromSelection()
                drag.setMimeData(mimeData)
                drag.setPixmap(self.create_drag_image(self))
                drag.exec(supportedActions, Qt.CopyAction)
                return
        super().startDrag(supportedActions)


class Note:
    def __init__(
        self,
        title: str,
        content: str,
        tags: list[str],
        favorite: bool = False,
        history: list[str] | None = None,
        timestamp: str | None = None,
        reminder: str | None = None,
        uuid: str | None = None,
        reminder_repeat: str | None = None,
        pinned: bool = False,
        password_manager: str = "",
        rdp_1c8: str = "",
        custom_fields: list[dict] | None = None,
        ignored_words: list[str] | None = None,
    ) -> None:
        self.reminder_shown = False
        self.title = title
        self.content = content
        self.tags = tags
        self.favorite = favorite
        self.timestamp = timestamp or QDateTime.currentDateTime().toString(Qt.ISODate)
        self.reminder = reminder
        self.uuid = uuid or str(uuid.uuid4())
        self.history = history if history is not None else []
        self.history_index = len(self.history) - 1 if self.history else -1
        self.reminder_repeat = reminder_repeat
        self.pinned = pinned
        self.password_manager = password_manager
        self.rdp_1c8 = rdp_1c8
        self.custom_fields = custom_fields if custom_fields is not None else []
        self.password_manager_visible = False
        self.rdp_1c8_visible = False
        self.rdp_1c8_removed = False
        self.ignored_words = ignored_words or []

    def to_dict(self) -> dict:
        return {
            "pinned": self.pinned,
            "history": self.history,
            "history_index": self.history_index,
            "title": self.title,
            "content": self.content,
            "tags": self.tags,
            "favorite": self.favorite,
            "timestamp": self.timestamp,
            "reminder": self.reminder,
            "reminder_repeat": self.reminder_repeat,
            "uuid": self.uuid,
            "password_manager": self.password_manager,
            "rdp_1c8": self.rdp_1c8,
            "password_manager_visible": bool(self.password_manager_visible),
            "rdp_1c8_visible": bool(self.rdp_1c8_visible),
            "rdp_1c8_removed": bool(self.rdp_1c8_removed),
            "custom_fields": self.custom_fields,
            "ignored_words": self.ignored_words,
        }

    @staticmethod
    def from_dict(data: dict) -> "Note":
        note = Note(
            title=data.get("title", ""),
            content=data.get("content", ""),
            tags=list(data.get("tags", [])),
            favorite=data.get("favorite", False),
            history=list(data.get("history", [])),
            timestamp=data.get("timestamp"),
            reminder=data.get("reminder"),
            uuid=data.get("uuid"),
            reminder_repeat=data.get("reminder_repeat", None),
            ignored_words=data.get("ignored_words", []),
        )
        note.pinned = data.get("pinned", False)
        note.history_index = data.get("history_index", len(note.history) - 1)
        note.password_manager = data.get("password_manager", "")
        note.rdp_1c8 = data.get("rdp_1c8", "")
        note.password_manager_visible = bool(
            data.get("password_manager_visible", False)
        )
        note.rdp_1c8_visible = bool(data.get("rdp_1c8_visible", False))
        note.rdp_1c8_removed = bool(data.get("rdp_1c8_removed", False))
        note.custom_fields = list(data.get("custom_fields", []))
        return note


class NumberGeneratorDialog(QDialog):
    def __init__(self, parent=None):
        super().__init__(parent)
        self.setWindowTitle("Генератор чисел")
        self.setModal(True)
        self.resize(420, 260)

        form = QFormLayout()
        self.start_spin = QSpinBox(); self.start_spin.setRange(-10**9, 10**9); self.start_spin.setValue(1)
        self.end_spin   = QSpinBox(); self.end_spin.setRange(-10**9, 10**9);  self.end_spin.setValue(100)
        self.count_spin = QSpinBox(); self.count_spin.setRange(1, 10**6);     self.count_spin.setValue(10)

        self.base_combo = QComboBox(); self.base_combo.addItems(["10", "16", "2"])
        self.unique_cb  = QCheckBox("Только уникальные")
        self.sort_cb    = QCheckBox("Сортировать")
        self.desc_cb    = QCheckBox("По убыванию"); self.desc_cb.setEnabled(False)
        self.sort_cb.toggled.connect(lambda v: self.desc_cb.setEnabled(v))

        self.leading_zeros_cb = QCheckBox("Выровнять нулями (для 10/16/2)")
        self.sep_edit = QLineEdit(", "); self.sep_edit.setPlaceholderText("разделитель (по умолчанию: запятая+пробел)")

        form.addRow("Начало:", self.start_spin)
        form.addRow("Конец:", self.end_spin)
        form.addRow("Сколько:", self.count_spin)
        form.addRow("Система счисления:", self.base_combo)
        form.addRow(self.unique_cb)
        form.addRow(self.sort_cb)
        form.addRow(self.desc_cb)
        form.addRow(self.leading_zeros_cb)
        form.addRow("Разделитель:", self.sep_edit)

        self.preview = QPlainTextEdit(); self.preview.setReadOnly(True); self.preview.setPlaceholderText("Предпросмотр…")

        btns = QDialogButtonBox(QDialogButtonBox.Ok | QDialogButtonBox.Cancel)
        self.btn_generate = QPushButton("Сгенерировать")
        btns.addButton(self.btn_generate, QDialogButtonBox.ActionRole)
        self.btn_copy = QPushButton("Копировать"); self.btn_copy.setEnabled(False)
        btns.addButton(self.btn_copy, QDialogButtonBox.ActionRole)

        self.btn_generate.clicked.connect(self._on_generate)
        self.btn_copy.clicked.connect(self._on_copy)
        btns.accepted.connect(self.accept)
        btns.rejected.connect(self.reject)

        root = QVBoxLayout(self)
        root.addLayout(form)
        root.addWidget(self.preview, 1)
        root.addWidget(btns)

        # первоначальная генерация
        self._on_generate()

    def _format_num(self, n: int, base: int, width: int) -> str:
        if base == 10:
            s = str(n)
        elif base == 16:
            s = format(n if n >= 0 else (1<<64)+n, "x")  # простой хак для отрицательных
        else:  # base == 2
            s = format(n if n >= 0 else (1<<64)+n, "b")
        if self.leading_zeros_cb.isChecked() and width > 0:
            s = s.zfill(width)
        return s

    def _on_generate(self):
        a = self.start_spin.value()
        b = self.end_spin.value()
        if a > b:
            a, b = b, a
        k = self.count_spin.value()
        base = int(self.base_combo.currentText())
        sep = self.sep_edit.text() if self.sep_edit.text() != "" else ", "
        pool_size = b - a + 1

        # подберём ширину для нулей
        width = 0
        if self.leading_zeros_cb.isChecked():
            if base == 10:
                width = len(str(max(abs(a), abs(b))))
            elif base == 16:
                width = len(format(max(abs(a), abs(b)), "x"))
            else:
                width = len(format(max(abs(a), abs(b)), "b"))

        rng = range(a, b + 1)
        numbers = []
        if self.unique_cb.isChecked():
            if k > pool_size:
                k = pool_size
            numbers = random.sample(list(rng), k)
        else:
            numbers = [random.randint(a, b) for _ in range(k)]

        if self.sort_cb.isChecked():
            numbers.sort(reverse=self.desc_cb.isChecked())

        formatted = [self._format_num(n, base, width) for n in numbers]
        text = sep.join(formatted)
        self.preview.setPlainText(text)
        self.btn_copy.setEnabled(bool(text))

    def _on_copy(self):
        text = self.preview.toPlainText()
        if text:
            QApplication.clipboard().setText(text)
    
    def result_text(self) -> str:
        # отдаём готовый текст
        return self.preview.toPlainText()



class NotesLoaderThread(QThread):
    finished = Signal(list)

    def __init__(self, parent=None):
        super().__init__(parent)

    def run(self) -> None:
        try:
            notes = NotesApp.read_notes_snapshot()
        except Exception:
            notes = []
        self.finished.emit(notes)


class NotesApp(QMainWindow):
    TRASH_DIR = os.path.join(NOTES_DIR, "Trash")
    window_hidden = Signal()

    def __init__(self):
        super().__init__()
        self.current_note_path = None
        self.timer_btn = None
        self._live_toasts: list[QLabel] = []
        self.exiting = False
        self.notes = []
        self.audio_thread = None
        self.settings = QSettings(SETTINGS_PATH, QSettings.IniFormat)
        self._timer_mode = self.settings.value("timer/mode", "countdown")
        self._timer_total = self.settings.value("timer/total", 25 * 60, type=int)
        self._timer_left = int(self._timer_total)
        self._timer_running = False
        self._timer = QTimer(self)
        self._timer.setInterval(1000)
        self._timer.timeout.connect(self._tick_timer)
        self._global_hotkeys = []
        self.load_plugins_state()
        self.init_ui()
        self._timer_mode = self.settings.value("timer/mode", "countdown")
        self._timer_total = self.settings.value("timer/total", 25 * 60, type=int)
        self._timer_left = int(
            self._timer_total if self._timer_mode == "countdown" else 0
        )
        self._timer_running = False
        self._timer = QTimer(self)
        self._timer.setInterval(1000)
        self._timer.timeout.connect(self._tick_timer)
        self.debounce_ms = self.settings.value("autosave_debounce_ms", 1500, type=int)
        self.autosave_enabled = True
        self.debounce_timer = QTimer(self)
        self.debounce_timer.setSingleShot(True)
        self.debounce_timer.timeout.connect(self.autosave_current_note)
        self.text_edit.textChanged.connect(self._on_text_changed)
        self.current_note = None
        self.attachments_watcher = QFileSystemWatcher(self)
        self.attachments_watcher.directoryChanged.connect(self._refresh_attachments)
        self.attachments_watcher.fileChanged.connect(self._refresh_attachments)
        self.init_all_components()
        self.load_plugins()
        self.init_theme()
        self.load_settings()
        self._clipboard_watch_enabled = self.settings.value(
            "layout/clipboard_auto_convert", True, type=bool
        )
        self._clipboard_ignore = False
        self._last_clipboard_text = ""
        QApplication.clipboard().dataChanged.connect(self._on_clipboard_changed)
        self._background_loader = None
        QTimer.singleShot(0, self._start_background_preload)
        self.tray_icon = QSystemTrayIcon(QIcon(TRAY_ICON_PATH), self)
        self.tray_icon.activated.connect(self.on_tray_icon_activated)
        self.tray_icon.setToolTip("Мои Заметки")
        self.tray_icon.setVisible(True)
        menu = QMenu()
        restore_action = QAction("Открыть", self)
        restore_action.triggered.connect(self.show)
        open_launcher_action = QAction("Выбрать приложение…", self)
        open_launcher_action.triggered.connect(self.show_app_launcher)
        menu.addAction(open_launcher_action)
        translate_layout_action = QAction("Перевести буфер (раскладка)", self)
        translate_layout_action.triggered.connect(self.translate_layout_only)
        menu.addAction(translate_layout_action)
        toggle_clipboard_case_action = QAction("Сменить регистр буфера", self)
        toggle_clipboard_case_action.triggered.connect(self.translate_case_only)
        menu.addAction(toggle_clipboard_case_action)
        toggle_clipboard_layout_action = QAction(
            "Автоперекладка буфера", self, checkable=True
        )
        toggle_clipboard_layout_action.setChecked(self._clipboard_watch_enabled)
        toggle_clipboard_layout_action.triggered.connect(
            self._toggle_clipboard_layout_watch
        )
        menu.addAction(toggle_clipboard_layout_action)
        exit_action = QAction("Выход", self)
        exit_action.triggered.connect(self.exit_app)
        menu.addAction(restore_action)
        menu.addAction(exit_action)
        self.tray_icon.setContextMenu(menu)
        self.reminder_timer = QTimer(self)
        self.reminder_timer.timeout.connect(self.check_reminders)
        self.reminder_timer.start(60000)

    @staticmethod
    def read_notes_snapshot() -> list[Note]:
        if not os.path.exists(NOTES_DIR):
            os.makedirs(NOTES_DIR, exist_ok=True)
        loaded_notes = []
        for folder in os.listdir(NOTES_DIR):
            folder_path = os.path.join(NOTES_DIR, folder)
            if os.path.isdir(folder_path):
                file_path = os.path.join(folder_path, "note.json")
                if os.path.exists(file_path):
                    with open(file_path, "r", encoding="utf-8") as f:
                        data = json.load(f)
                    note = Note.from_dict(data)
                    if "content_txt" in data:
                        note.content_txt = data["content_txt"]
                    else:
                        doc = QTextDocument()
                        doc.setHtml(note.content)
                        note.content_txt = doc.toPlainText()
                    loaded_notes.append(note)
        unique: dict[str, Note] = {}
        for note in loaded_notes:
            unique[note.uuid] = note
        return list(unique.values())

    def _start_background_preload(self) -> None:
        if getattr(self, "_background_loader", None) is not None:
            try:
                self._background_loader.finished.disconnect()
            except Exception:
                pass
        self._background_loader = NotesLoaderThread(self)
        self._background_loader.finished.connect(self._apply_preloaded_notes)
        self._background_loader.finished.connect(
            lambda *_: getattr(self._background_loader, "deleteLater", lambda: None)()
        )
        self._background_loader.start()

    def _apply_preloaded_notes(self, notes: list[Note]) -> None:
        if not notes:
            return
        self.notes = notes
        self.deduplicate_notes()
        self.refresh_notes_list()

    def _install_dialog_probe():
        class _Probe(QDialog):
            def showEvent(self, e):
                try:
                    print("\n=== DIALOG OPENED ===", type(self).__name__)
                    traceback.print_stack(limit=25, file=sys.stdout)
                    print("===  END STACK   ===\n")
                except Exception:
                    pass
                super().showEvent(e)
        old_showEvent = QDialog.showEvent
        def patched(self, e):
            try:
                print(f"\n[Dialog] {type(self).__name__} shown")
            except Exception:
                pass
            return old_showEvent(self, e)
        QDialog.showEvent = patched

    _install_dialog_probe()

    def _on_text_changed(self) -> None:
        if not getattr(self, "current_note", None):
            return
        if getattr(self, "_is_switching_note", False):
            return
        if getattr(self, "_edit_session_note_uuid", None) and self._edit_session_note_uuid != self.current_note.uuid:
            return
        self._edit_session_note_uuid = self.current_note.uuid
        self.update_current_note_content()
        self.pending_save = True
        self._set_unsaved(True)
        if hasattr(self, "debounce_timer"):
            self.debounce_timer.start()

    def exit_app(self) -> None:
        self.tray_icon.hide()
        self._unregister_global_hotkeys()
        self.close()
        QApplication.instance().quit()

    def _set_unsaved(self, flag: bool) -> None:
        self.setWindowModified(bool(flag))
        if hasattr(self, "dock_editor"):
            self.dock_editor.setWindowTitle("Редактор *" if flag else "Редактор")

    @staticmethod
    def _log_startup():
        try:
            log_path = os.path.join(DATA_DIR, "last_run.log")
            os.makedirs(DATA_DIR, exist_ok=True)
            with open(log_path, "a", encoding="utf-8") as f:
                print("="*60, file=f)
                print("START", datetime.datetime.now().isoformat(), file=f)
                print("APPDIR:", APPDIR, file=f)
                print("CWD   :", os.getcwd(), file=f)
                print("Notes exists:", os.path.isdir(NOTES_DIR), file=f)
                print("Data  exists:", os.path.isdir(DATA_DIR), file=f)
        except Exception:
            pass

    def init_ui(self) -> None:
        self.setWindowTitle("Мои Заметки — [*]")
        self.setMinimumSize(1250, 800)
        self.setWindowIcon(QIcon(ICON_PATH))
        always_on_top = self.settings.value("ui/always_on_top", True, type=bool)
        self.setWindowFlag(Qt.WindowStaysOnTopHint, always_on_top)
        pm_label = self.settings.value("password_manager_label", "PasswordManager")
        rdp_label = self.settings.value("rdp_1c8_label", "1C8 RDP")
        self.new_note_button = QPushButton("Новая")
        self.save_note_button = QPushButton("Сохранить")
        self.delete_note_button = QPushButton("Удалить")
        self.undo_button = QPushButton("↩️")
        self.redo_button = QPushButton("↪️")
        self.undo_button.clicked.connect(self.undo)
        self.redo_button.clicked.connect(self.redo)
        self.search_field = QLineEdit()
        self.search_field.setPlaceholderText("Поиск по заметкам...")
        self.search_field.textChanged.connect(self.refresh_notes_list)
        self.topmost_checkbox = QCheckBox("Поверх всех окон")
        self.topmost_checkbox.setChecked(always_on_top)
        self.topmost_checkbox.toggled.connect(self._toggle_always_on_top)
        btns = QVBoxLayout()
        btns.addWidget(self.new_note_button)
        btns.addWidget(self.save_note_button)
        btns.addWidget(self.delete_note_button)
        btns.addWidget(self.undo_button)
        btns.addWidget(self.redo_button)
        btns.addWidget(self.topmost_checkbox)
        btns.addStretch()
        buttons_widget = QWidget()
        buttons_widget.setLayout(btns)
        self.notes_list = QListWidget()
        self.notes_list.setSelectionMode(QAbstractItemView.ExtendedSelection)
        self.notes_list.setDragEnabled(True)
        self.notes_list.setAcceptDrops(True)
        self.notes_list.setDropIndicatorShown(True)
        self.notes_list.setDefaultDropAction(Qt.MoveAction)
        self.notes_list.setDragDropMode(QAbstractItemView.InternalMove)
        self.notes_list.model().rowsMoved.connect(self.handle_note_reorder)
        self.notes_list.setContextMenuPolicy(Qt.CustomContextMenu)
        self.notes_list.customContextMenuRequested.connect(self.show_note_context_menu)
        self.history_label = QLabel("Версии:")
        self.history_list = QListWidget()
        self.history_list.setSelectionMode(QAbstractItemView.ExtendedSelection)
        self.history_list.itemClicked.connect(self.restore_version_from_history)
        self.delete_history_button = QPushButton("Удалить выбранные версии")
        self.delete_history_button.clicked.connect(self.delete_selected_history_entries)
        self._hist_delete_sc = QShortcut(QKeySequence.Delete, self.history_list)
        self._hist_delete_sc.setContext(Qt.WidgetWithChildrenShortcut)
        self._hist_delete_sc.activated.connect(self.delete_selected_history_entries)
        hist_layout = QVBoxLayout()
        hist_layout.addWidget(self.history_label)
        hist_layout.addWidget(self.history_list)
        hist_layout.addWidget(self.delete_history_button)
        self.history_widget = QWidget()
        self.history_widget.setLayout(hist_layout)
        self.history_widget.setMinimumWidth(200)
        self.history_widget.setSizePolicy(QSizePolicy.Expanding, QSizePolicy.Expanding)
        self.text_edit = CustomTextEdit(parent=self, paste_image_callback=self.insert_image_from_clipboard)
        self.autosave_delay_ms = 1000
        self.autosave_timer = QTimer(self)
        self.autosave_timer.setSingleShot(True)
        self.autosave_timer.timeout.connect(lambda: self.save_note_quiet(force=True))
        self.text_edit.textChanged.connect(self._schedule_autosave)
        self.text_edit.textChanged.connect(self.on_text_changed_autosave)
        self.text_edit.setUndoRedoEnabled(False)
        self.spell_highlighter = SpellCheckHighlighter(self.text_edit.document())
        self.spell_highlighter = SpellCheckHighlighter(self.text_edit.document())
        self.spell_online_enabled = self.settings.value("spellcheck/online_enabled", False, type=bool)
        self.spell_highlighter.set_online_enabled(self.spell_online_enabled)
        self._save_debouncer = QTimer(self)
        self._save_debouncer.setSingleShot(True)
        self._save_debouncer.setInterval(400)
        self.attachments_panel = QWidget()
        self.attachments_layout = QHBoxLayout(self.attachments_panel)
        self.attachments_layout.setAlignment(Qt.AlignLeft)
        self.attachments_layout.setContentsMargins(0, 0, 0, 0)
        self.attachments_scroll = QScrollArea()
        self.attachments_scroll.setWidgetResizable(True)
        self.attachments_scroll.setWidget(self.attachments_panel)
        self.attachments_scroll.setVisible(False)
        self.attachments_scroll.setSizePolicy(QSizePolicy.Expanding, QSizePolicy.Minimum)
        self._save_debouncer.timeout.connect(self.update_current_note_content)
        self.text_edit.textChanged.connect(self._save_debouncer.start)
        self.text_edit.cursorPositionChanged.connect(self.update_font_controls)
        self.text_edit.setReadOnly(True)
        self.text_edit.hide()
        self.text_edit.setFocusPolicy(Qt.StrongFocus)
        self.text_edit.setAlignment(Qt.AlignLeft)
        self.text_edit.setMinimumWidth(250)
        self.text_edit.setSizePolicy(QSizePolicy.Expanding, QSizePolicy.Expanding)
        self.text_edit.setStyleSheet("font-size: 14px; font-family: 'Times New Roman';")
        self.text_edit.installEventFilter(self)
        self.text_edit.textChanged.connect(lambda: self._set_unsaved(True))
        self.tags_label = QLabel("Теги: нет")
        self.tags_label.setAlignment(Qt.AlignLeft)
        self.attachments_panel = QWidget()
        self.attachments_layout = QHBoxLayout(self.attachments_panel)
        self.attachments_layout.setAlignment(Qt.AlignLeft)
        self.attachments_layout.setContentsMargins(0, 0, 0, 0)
        self.attachments_scroll = QScrollArea()
        self.attachments_scroll.setWidgetResizable(True)
        self.attachments_scroll.setWidget(self.attachments_panel)
        self.attachments_scroll.setVisible(False)
        self.attachments_scroll.setSizePolicy(
            QSizePolicy.Expanding, QSizePolicy.Minimum
        )
        editor_combined = QWidget()
        editor_layout_combined = QVBoxLayout(editor_combined)
        editor_layout_combined.addWidget(self.tags_label)
        editor_layout_combined.addWidget(self.text_edit)
        editor_layout_combined.addWidget(self.attachments_scroll)
        editor_layout_combined.setStretch(1, 1)
        editor_layout_combined.setStretch(2, 0)
        self.central_widget = QWidget()
        self.setCentralWidget(self.central_widget)
        self.init_toolbar()
        self.dock_notes_list = QDockWidget("Заметки", self)
        self.dock_notes_list.setObjectName("dock_notes_list")
        self.dock_notes_list.setWidget(self.notes_list)
        self.dock_notes_list.setAllowedAreas(
            Qt.LeftDockWidgetArea | Qt.RightDockWidgetArea
        )
        self.addDockWidget(Qt.LeftDockWidgetArea, self.dock_notes_list)
        self.dock_history = QDockWidget("История", self)
        self.dock_history.setObjectName("dock_history")
        self.dock_history.setWidget(self.history_widget)
        self.dock_history.setAllowedAreas(
            Qt.LeftDockWidgetArea | Qt.RightDockWidgetArea
        )
        self.addDockWidget(Qt.LeftDockWidgetArea, self.dock_history)
        self.dock_buttons = QDockWidget("Управление", self)
        self.dock_buttons.setObjectName("dock_buttons")
        self.dock_buttons.setWidget(buttons_widget)
        self.dock_buttons.setAllowedAreas(
            Qt.LeftDockWidgetArea | Qt.RightDockWidgetArea
        )
        self.addDockWidget(Qt.LeftDockWidgetArea, self.dock_buttons)
        self.dock_editor = QDockWidget("Редактор", self)
        self.dock_editor.setObjectName("dock_editor")
        self.dock_editor.setWidget(editor_combined)
        self.dock_editor.setAllowedAreas(Qt.AllDockWidgetAreas)
        self.setDockOptions(
            self.dockOptions()
            | QMainWindow.AllowNestedDocks
            | QMainWindow.AllowTabbedDocks
            | QMainWindow.AnimatedDocks
        )
        self.dock_editor.setSizePolicy(QSizePolicy.Expanding, QSizePolicy.Expanding)
        self.addDockWidget(Qt.RightDockWidgetArea, self.dock_editor)
        self.password_manager_field = QLineEdit()
        self.password_manager_field.setPlaceholderText(pm_label)
        self.rdp_1c8_field = QLineEdit()
        self.rdp_1c8_field.setPlaceholderText(rdp_label)
        self.password_manager_row = QWidget()
        _pm_row_layout = QHBoxLayout(self.password_manager_row)
        _pm_row_layout.setContentsMargins(0, 0, 0, 0)
        self.password_manager_label = QLineEdit()
        self.password_manager_label.setText(pm_label)
        self.password_manager_label.setMinimumWidth(140)
        self.password_manager_label.setFrame(False)
        self.password_manager_label.setStyleSheet("background: transparent;")
        self.password_manager_label.textChanged.connect(self.on_pm_label_changed)
        self.password_manager_label.editingFinished.connect(
            lambda: self.settings.setValue(
                "password_manager_label", self.password_manager_label.text()
            )
        )
        _pm_row_layout.addWidget(self.password_manager_label)
        _pm_row_layout.addWidget(self.password_manager_field, 1)
        self.password_manager_copy_btn = QPushButton("Копировать")
        self.password_manager_copy_btn.setFixedHeight(24)
        self.password_manager_copy_btn.setEnabled(False)
        self.password_manager_copy_btn.clicked.connect(
            self.copy_password_manager_to_clipboard
        )
        self.password_manager_field.textChanged.connect(
            lambda t: self.password_manager_copy_btn.setEnabled(bool(t))
        )
        _pm_row_layout.addWidget(self.password_manager_copy_btn)
        pm_up = QToolButton()
        pm_up.setText("▲")
        pm_up.setFixedSize(24, 24)
        pm_dn = QToolButton()
        pm_dn.setText("▼")
        pm_dn.setFixedSize(24, 24)
        pm_up.clicked.connect(
            lambda _, w=self.password_manager_row: self._move_pinned(w, -1)
        )
        pm_dn.clicked.connect(
            lambda _, w=self.password_manager_row: self._move_pinned(w, +1)
        )
        _pm_row_layout.addWidget(pm_up)
        _pm_row_layout.addWidget(pm_dn)
        self.rdp_1c8_row = QWidget()
        _rdp_row_layout = QHBoxLayout(self.rdp_1c8_row)
        _rdp_row_layout.setContentsMargins(0, 0, 0, 0)
        self.rdp_1c8_label = QLineEdit()
        self.rdp_1c8_label.setText(rdp_label)
        self.rdp_1c8_label.setMinimumWidth(140)
        self.rdp_1c8_label.setFrame(False)
        self.rdp_1c8_label.setStyleSheet("background: transparent;")
        self.rdp_1c8_label.textChanged.connect(self.on_rdp_label_changed)
        self.rdp_1c8_label.editingFinished.connect(
            lambda: self.settings.setValue("rdp_1c8_label", self.rdp_1c8_label.text())
        )
        _rdp_row_layout.addWidget(self.rdp_1c8_label)
        _rdp_row_layout.addWidget(self.rdp_1c8_field, 1)
        self.rdp_1c8_copy_btn = QPushButton("Копировать")
        self.rdp_1c8_copy_btn.setFixedHeight(24)
        self.rdp_1c8_copy_btn.setEnabled(False)
        self.rdp_1c8_copy_btn.clicked.connect(self.copy_rdp_1c8_to_clipboard)
        self.rdp_1c8_field.textChanged.connect(
            lambda t: self.rdp_1c8_copy_btn.setEnabled(bool(t))
        )
        self.rdp_1c8_delete_btn = QPushButton("✖")
        self.rdp_1c8_delete_btn.setFixedSize(24, 24)
        self.rdp_1c8_delete_btn.clicked.connect(self.delete_rdp_1c8_field)
        _rdp_row_layout.addWidget(self.rdp_1c8_copy_btn)
        _rdp_row_layout.addWidget(self.rdp_1c8_delete_btn)
        rdp_up = QToolButton()
        rdp_up.setText("▲")
        rdp_up.setFixedSize(24, 24)
        rdp_dn = QToolButton()
        rdp_dn.setText("▼")
        rdp_dn.setFixedSize(24, 24)
        rdp_up.clicked.connect(lambda _, w=self.rdp_1c8_row: self._move_pinned(w, -1))
        rdp_dn.clicked.connect(lambda _, w=self.rdp_1c8_row: self._move_pinned(w, +1))
        _rdp_row_layout.addWidget(rdp_up)
        _rdp_row_layout.addWidget(rdp_dn)
        self.pinned_container = QWidget()
        self.pinned_layout = QVBoxLayout(self.pinned_container)
        self.pinned_layout.setContentsMargins(0, 0, 0, 0)
        self.pinned_layout.setSpacing(6)
        self.pinned_layout.addWidget(self.password_manager_row)
        self._add_field_panel = QWidget()
        _add_l = QHBoxLayout(self._add_field_panel)
        _add_l.setContentsMargins(0, 0, 0, 0)
        self.add_field_btn = QPushButton("Добавить поле")
        self.add_field_btn.setFixedHeight(24)
        self.add_field_btn.setEnabled(False)
        self.add_field_btn.clicked.connect(self.add_custom_field)
        _add_l.addStretch(1)
        _add_l.addWidget(self.add_field_btn)
        self.pinned_layout.addWidget(self._add_field_panel)
        self.custom_fields_container = self.pinned_container
        self.custom_fields_layout = self.pinned_layout
        self.custom_fields_widgets: list[dict] = []
        editor_layout_combined.insertWidget(0, self.pinned_container)
        self.dock_toolbar = QDockWidget("Панель инструментов", self)
        self.dock_toolbar.setObjectName("dock_toolbar")
        self.dock_toolbar.setWidget(self.toolbar_scroll)
        self.dock_toolbar.setAllowedAreas(
            Qt.TopDockWidgetArea | Qt.BottomDockWidgetArea
        )
        self.addDockWidget(Qt.TopDockWidgetArea, self.dock_toolbar)
        self.resizeDocks(
            [self.dock_notes_list, self.dock_editor, self.dock_history],
            [280, 800, 240],
            Qt.Horizontal
        )
        self.resizeDocks(
            [self.dock_toolbar, self.dock_editor],
            [120, 1000],
            Qt.Vertical
        )
        for dock in (
            self.dock_notes_list,
            self.dock_history,
            self.dock_buttons,
            self.dock_editor,
            self.dock_toolbar,
        ):
            dock.setContextMenuPolicy(Qt.CustomContextMenu)
            dock.customContextMenuRequested.connect(
                lambda pos, d=dock: self.show_dock_context_menu(pos, d)
            )
        self.visibility_toolbar = QToolBar("Видимость полей", self)
        self.visibility_toolbar.setObjectName("visibility_toolbar")
        self.visibility_toolbar.setMovable(False)
        self.addToolBar(Qt.TopToolBarArea, self.visibility_toolbar)
        self.action_toggle_pm = QAction(
            f"🙈 {self.password_manager_label.text()}", self
        )
        self.action_toggle_pm.setCheckable(True)
        self.action_toggle_pm.toggled.connect(self.on_toggle_pm_visible)
        self.action_toggle_rdp = QAction(f"🙈 {self.rdp_1c8_label.text()}", self)
        self.action_toggle_rdp.setCheckable(True)
        self.action_toggle_rdp.toggled.connect(self.on_toggle_rdp_visible)
        self.visibility_toolbar.addAction(self.action_toggle_pm)
        self.visibility_toolbar.addAction(self.action_toggle_rdp)
        self.add_menu_bar()
        self.new_note_button.clicked.connect(self.new_note)
        self.save_note_button.clicked.connect(self.save_note)
        self.delete_note_button.clicked.connect(self.delete_note)
        self.notes_list.itemClicked.connect(self.load_note)
        self.current_note = None
        self._update_editor_visibility()
        sc_show_launcher = QShortcut(QKeySequence("Ctrl+Alt+L"), self)
        sc_show_launcher.setContext(Qt.ApplicationShortcut)
        sc_show_launcher.activated.connect(self.show_app_launcher)
        sc_undo_notes = QShortcut(QKeySequence.Undo, self)
        sc_undo_notes.setContext(Qt.ApplicationShortcut)
        sc_undo_notes.activated.connect(self.undo)
        sc_redo_notes = QShortcut(QKeySequence.Redo, self)
        sc_redo_notes.setContext(Qt.ApplicationShortcut)
        sc_redo_notes.activated.connect(self.redo)
        sc_translate_layout = QShortcut(QKeySequence("Ctrl+Shift+F12"), self)
        sc_translate_layout.setContext(Qt.ApplicationShortcut)
        sc_translate_layout.activated.connect(self.translate_layout_only)
        sc_toggle_case = QShortcut(QKeySequence("Ctrl+Alt+F12"), self)
        sc_toggle_case.setContext(Qt.ApplicationShortcut)
        sc_toggle_case.activated.connect(self.translate_case_only)
        self._register_global_hotkeys()
        app = QApplication.instance()
        if app is not None:
            app.aboutToQuit.connect(self._unregister_global_hotkeys)
        ratios = self.settings.value("ui/dock_ratios") or [0.22, 0.56, 0.22]
        try:
            ratios = [float(r) for r in ratios]
            total = sum(ratios)
            if total > 0:
                self._dock_ratios = [r / total for r in ratios[:3]]
            else:
                self._dock_ratios = [0.22, 0.56, 0.22]
        except Exception:
            ratios = self.settings.value("ui/dock_ratios") or [0.22, 0.56, 0.22]
        try:
            ratios = [float(r) for r in ratios]
            total = sum(ratios)
            if total > 0:
                self._dock_ratios = [r / total for r in ratios[:3]]
            else:
                self._dock_ratios = [0.22, 0.56, 0.22]
        except Exception:
            self._dock_ratios = [0.22, 0.56, 0.22]
        self._resizing_apply = False
        self._resize_deb = QTimer(self)

    def copy_password_manager_to_clipboard(self) -> None:
        text = self.password_manager_field.text().strip()
        if not text:
            QToolTip.showText(QCursor.pos(), "Поле пустое")
            return
        QApplication.clipboard().setText(text)
        self.show_toast(
            "Скопировано",
            boundary_widget=self.dock_editor.widget(),
            anchor_widget=self.password_manager_copy_btn,
        )

    def _schedule_autosave(self):
        if not getattr(self, "current_note", None):
            return
        if not getattr(self, "autosave_enabled", True):
            return
        self.autosave_timer.start(self.autosave_delay_ms)

    def _apply_dock_proportions(self):
        left_width = 360
        right_width = max(600, self.width() - left_width)
        self.resizeDocks(
            [self.dock_notes_list, self.dock_editor],
            [left_width, right_width],
            Qt.Horizontal
        )
        self.resizeDocks(
            [self.dock_notes_list, self.dock_history, self.dock_buttons],
            [1, 1, 1],
            Qt.Vertical
        )
        self._apply_dock_proportions()

    def show_number_generator(self):
        try:
            dlg = NumberGeneratorDialog(self)
            if dlg.exec() == QDialog.Accepted:
                text = dlg.result_text()
                if text and hasattr(self, "text_edit"):
                    # вставляем в позицию курсора
                    cur = self.text_edit.textCursor()
                    cur.insertText(text)
                    self.text_edit.setTextCursor(cur)
                    self._set_unsaved(True)
                    # автосейв, если включён
                    if getattr(self, "autosave_enabled", True):
                        if self.autosave_enabled:
                            self.save_note_quiet(force=True)
        except Exception as e:
            QMessageBox.warning(self, "Генератор чисел", f"Ошибка: {e}")

    def _apply_vertical_defaults(self):
        min_h = 36 
        h = self.settings.value("ui/dock_toolbar_height", min_h, type=int)
        try:
            h = int(h)
        except Exception:
            h = min_h
        h = max(min_h, min(h, 300))
        self.dock_toolbar.setMinimumHeight(min_h)
        self.resizeDocks(
            [self.dock_toolbar, self.dock_editor],
            [h, max(300, self.height() - h)],
            Qt.Vertical
        )

    def _current_dock_widths(self):
        l = max(1, self.dock_notes_list.width())
        c = max(1, self.dock_editor.width())
        r = max(1, self.dock_history.width())
        return [l, c, r]

    def _capture_dock_ratios(self):
        l, c, r = self._current_dock_widths()
        total = max(1, l + c + r)
        self._dock_ratios = [l/total, c/total, r/total]
        try:
            self.settings.setValue("ui/dock_ratios", self._dock_ratios)
        except Exception:
            pass

    def start_stopwatch(self):
        if not getattr(self, "stopwatch_running", False):
            self.stopwatch_running = True
            try:
                self.stopwatch_timer.start()
            except Exception:
                pass

    def resizeEvent(self, e):
        super().resizeEvent(e)
        if not hasattr(self, "_resize_deb"):
            self._resize_deb = QTimer(self)
            self._resize_deb.setSingleShot(True)
            self._resize_deb.setInterval(0)
            self._resize_deb.timeout.connect(self._apply_dock_ratios)
        self._resize_deb.start()

    def _apply_dock_ratios(self):
        if self._resizing_apply:
            return
        self._resizing_apply = True
        try:
            total = max(1, self.width())
            min_left = 180
            min_center = 400
            min_right = 180
            rl, rc, rr = self._dock_ratios
            left  = int(total * rl)
            center= int(total * rc)
            right = int(total * rr)
            if center < min_center:
                delta = min_center - center
                take = min(delta, max(0, left - min_left))
                left -= take
                delta -= take
                right = max(min_right, right - delta)
                center = min_center
            if left < min_left:
                need = min_left - left
                if center - need >= min_center:
                    center -= need
                else:
                    right = max(min_right, right - (need - (center - min_center)))
                    center = min_center
                left = min_left
            if right < min_right:
                need = min_right - right
                if center - need >= min_center:
                    center -= need
                else:
                    left = max(min_left, left - (need - (center - min_center)))
                    center = min_center
                right = min_right
            extra = (left + center + right) - total
            if extra != 0:
                shift = min(extra, center - min_center) if extra > 0 else -min(-extra, total)
                center -= shift
                extra -= shift
                right -= extra
            self.resizeDocks(
                [self.dock_notes_list, self.dock_editor, self.dock_history],
                [max(1,left), max(1,center), max(1,right)],
                Qt.Horizontal
            )
        finally:
            QTimer.singleShot(0, self._capture_dock_ratios)
            self._resizing_apply = False

    def stop_stopwatch(self):
        self.stopwatch_running = False
        self.stopwatch_timer.stop()

    def _set_html_preserve_view(self, html: str, preferred_pos: int | None = None) -> None:
        cur_before = self.text_edit.textCursor()
        block_num = cur_before.blockNumber()
        col_in_block = cur_before.positionInBlock()
        if preferred_pos is not None:
            tmp = QTextCursor(self.text_edit.document())
            tmp.setPosition(max(0, int(preferred_pos)))
            block_num = tmp.blockNumber()
            col_in_block = tmp.positionInBlock()
        vbar = self.text_edit.verticalScrollBar()
        old_scroll = vbar.value()
        self.text_edit.blockSignals(True)
        self.text_edit.setUpdatesEnabled(False)
        self.text_edit.setHtml(html)
        doc = self.text_edit.document()
        max_block = max(0, doc.blockCount() - 1)
        block_num = max(0, min(block_num, max_block))
        block = doc.findBlockByNumber(block_num)
        max_col = max(0, block.length() - 1)
        col_in_block = max(0, min(col_in_block, max_col))
        target_pos = block.position() + col_in_block
        cur_after = self.text_edit.textCursor()
        cur_after.setPosition(target_pos)
        self.text_edit.setTextCursor(cur_after)
        vbar.setValue(old_scroll)
        self.text_edit.setUpdatesEnabled(True)
        self.text_edit.blockSignals(False)

    def reset_stopwatch(self):
        self.stopwatch_running = False
        self.stopwatch_timer.stop()
        self.stopwatch_time = 0
        self.lbl_stopwatch.setText("00:00:00")

    def _update_stopwatch(self):
        if self.stopwatch_running:
            self.stopwatch_time += 1
            h = self.stopwatch_time // 3600
            m = (self.stopwatch_time % 3600) // 60
            s = self.stopwatch_time % 60
            self.lbl_stopwatch.setText(f"{h:02d}:{m:02d}:{s:02d}")

    def copy_rdp_1c8_to_clipboard(self) -> None:
        text = self.rdp_1c8_field.text().strip()
        if not text:
            QToolTip.showText(QCursor.pos(), "Поле пустое")
            return
        QApplication.clipboard().setText(text)
        self.show_toast(
            "Скопировано",
            boundary_widget=self.dock_editor.widget(),
            anchor_widget=self.rdp_1c8_copy_btn,
        )

    def add_word_to_note_ignore(self, word: str) -> None:
        if not getattr(self, "current_note", None):
            return
        w = (word or "").strip().lower()
        if not w:
            return
        if w not in self.current_note.ignored_words:
            self.current_note.ignored_words.append(w)
            if hasattr(self, "spell_highlighter"):
                self.spell_highlighter.set_local_ignored(
                    self.current_note.ignored_words
                )
            self.save_note_quiet(force=True)

    def _ensure_rdp_row_in_layout(self) -> None:
        if self.pinned_layout.indexOf(self.rdp_1c8_row) == -1:
            insert_at = self.pinned_layout.indexOf(self._add_field_panel)
            if insert_at == -1:
                insert_at = self.pinned_layout.count()
            self.pinned_layout.insertWidget(insert_at, self.rdp_1c8_row)

    def delete_rdp_1c8_field(self) -> None:
        self.rdp_1c8_field.clear()
        if getattr(self, "current_note", None):
            self.current_note.rdp_1c8 = ""
            self.current_note.rdp_1c8_visible = False
            self.current_note.rdp_1c8_removed = True
            self.save_note_to_file(self.current_note)
        self.rdp_1c8_row.setVisible(False)
        if hasattr(self, "action_toggle_rdp"):
            self.action_toggle_rdp.setEnabled(False)
            self.action_toggle_rdp.setVisible(False)
            self.action_toggle_rdp.blockSignals(True)
            self.action_toggle_rdp.setChecked(False)
            self._update_eye_action(
                self.action_toggle_rdp, False, self.rdp_1c8_label.text()
            )
            self.action_toggle_rdp.setEnabled(False)
            self.action_toggle_rdp.setVisible(False)
            self.action_toggle_rdp.blockSignals(False)

    def add_custom_field(self, data: dict | None = None) -> None:
        row = QWidget()
        layout = QHBoxLayout(row)
        layout.setContentsMargins(0, 0, 0, 0)
        label_edit = QLineEdit()
        label = data.get("label", "Новое поле") if data else "Новое поле"
        label_edit.setText(label)
        label_edit.setMinimumWidth(140)
        label_edit.setFrame(False)
        label_edit.setStyleSheet("background: transparent;")
        value_edit = QLineEdit()
        value = data.get("value", "") if data else ""
        value_edit.setText(value)
        value_edit.setPlaceholderText(label)
        remove_btn = QPushButton("✖")
        remove_btn.setFixedSize(24, 24)
        up_btn = QToolButton();  up_btn.setText("▲");  up_btn.setFixedSize(24, 24)
        down_btn = QToolButton(); down_btn.setText("▼"); down_btn.setFixedSize(24, 24)
        up_btn.clicked.connect(lambda _, w=row: self._move_pinned(w, -1))
        down_btn.clicked.connect(lambda _, w=row: self._move_pinned(w, +1))
        copy_btn = QPushButton("Копировать")
        copy_btn.setFixedHeight(24)
        copy_btn.setEnabled(bool(value))
        copy_btn.clicked.connect(
            lambda _, e=value_edit, b=copy_btn: self.copy_custom_field_to_clipboard(e, b)
        )
        value_edit.textChanged.connect(lambda _: self._schedule_autosave())
        value_edit.textChanged.connect(lambda t, b=copy_btn: b.setEnabled(bool(t)))
        value_edit.textChanged.connect(lambda _: self.update_current_note_custom_fields())
        value_edit.textChanged.connect(lambda _: self.save_note_quiet(force=True) if self.autosave_enabled else None)
        layout.addWidget(up_btn)
        layout.addWidget(down_btn)
        layout.addWidget(label_edit)
        layout.addWidget(value_edit, 1)
        layout.addWidget(copy_btn)
        layout.addWidget(remove_btn)
        insert_at = self.pinned_layout.indexOf(self._add_field_panel)
        self.pinned_layout.insertWidget(insert_at, row)
        fid = (data.get("id") if data else None) or f"cf::{uuid4().hex}"
        action = QAction(f"🙈 {label}", self)
        action.setCheckable(True)
        fid = label_edit.text().strip() or "field"
        visible = bool(int(self.settings.value(f"ui/customFieldVisible/{fid}", 1)))
        if data and "visible" in data:
            raw = data.get("visible", False)
            visible = (raw is True) or (str(raw).lower() in {"1", "true", "yes"})
        action.setChecked(visible)
        row.setVisible(visible)
        self._update_eye_action(action, visible, label)
        widget = {
            "row": row,
            "label_edit": label_edit,
            "value_edit": value_edit,
            "copy_btn": copy_btn,
            "remove_btn": remove_btn,
            "action": action,
            "id": fid,
        }
        if data and "visible" in data:
            raw = data["visible"]
            visible = (raw is True) or (str(raw).lower() in {"1", "true", "yes"})
        else:
            visible = bool(int(self.settings.value(f"ui/customFieldVisible/{fid}", 1)))
            action.blockSignals(True)
        action.setChecked(visible)
        action.blockSignals(False)
        row.setVisible(visible)
        self._update_eye_action(action, visible, label)
        action.toggled.connect(lambda checked, w=widget: self.on_toggle_custom_field(w, checked))
        label_edit.textChanged.connect(lambda text, w=widget: self.on_custom_field_label_changed(w, text))
        remove_btn.clicked.connect(lambda _, w=widget: self.remove_custom_field(w))
        self.visibility_toolbar.addAction(action)
        self.custom_fields_widgets.append(widget)
        self.update_current_note_custom_fields()

    def _pinned_widgets(self) -> list[QWidget]:
        out = []
        if not hasattr(self, "pinned_layout"):
            return out
        for i in range(self.pinned_layout.count()):
            w = self.pinned_layout.itemAt(i).widget()
            if not w or w is self._add_field_panel:
                continue
            out.append(w)
        return out

    def _sync_custom_widgets_order_from_layout(self) -> None:
        if not hasattr(self, "custom_fields_widgets"):
            return
        by_row = {w["row"]: w for w in self.custom_fields_widgets}
        new_list = []
        for w in self._pinned_widgets():
            if w in by_row:
                new_list.append(by_row[w])
        self.custom_fields_widgets = new_list

    def remove_custom_field(self, widget: dict) -> None:
        self.visibility_toolbar.removeAction(widget["action"])
        widget["row"].setParent(None)
        widget["action"].deleteLater()
        widget["row"].deleteLater()
        self.custom_fields_widgets.remove(widget)
        self.update_current_note_custom_fields()
        if self.current_note:
            self.save_note_to_file(self.current_note)

    def _move_pinned(self, widget: QWidget, delta: int) -> None:
        if not hasattr(self, "pinned_layout"):
            return
        cur = self.pinned_layout.indexOf(widget)
        if cur < 0:
            return
        last_index = self.pinned_layout.indexOf(self._add_field_panel) - 1
        if last_index < 0:
            last_index = self.pinned_layout.count() - 1
        target = max(0, min(cur + int(delta), last_index))
        if target == cur:
            return
        self.pinned_layout.removeWidget(widget)
        self.pinned_layout.insertWidget(target, widget)
        self._sync_custom_widgets_order_from_layout()
        self.update_current_note_custom_fields()
        if self.current_note:
            self.save_note_to_file(self.current_note)

    def clear_custom_fields(self) -> None:
        for w in list(self.custom_fields_widgets):
            self.visibility_toolbar.removeAction(w["action"])
            w["action"].deleteLater()
            w["row"].deleteLater()
        self.custom_fields_widgets.clear()
        self.update_current_note_custom_fields()

    def on_custom_field_label_changed(self, widget: dict, text: str) -> None:
        widget["value_edit"].setPlaceholderText(text)
        self._update_eye_action(widget["action"], widget["row"].isVisible(), text)
        self.update_current_note_custom_fields()
        if hasattr(self, "debounce_timer"):
            if self.autosave_enabled:
                self.save_note_quiet(force=True)

    def on_toggle_custom_field(self, widget: dict, checked: bool) -> None:
        if not self.current_note:
            widget["action"].blockSignals(True)
            widget["action"].setChecked(False)
            widget["action"].blockSignals(False)
            return
        widget["row"].setVisible(checked)
        self._update_eye_action(widget["action"], checked, widget["label_edit"].text())
        fid = widget.get("id")
        if fid:
            self.settings.setValue(f"ui/customFieldVisible/{fid}", int(checked))
            self.settings.sync()
        self.settings.setValue(f"ui/customFieldVisible/{fid}", int(checked))
        self.update_current_note_custom_fields()
        self.save_note_to_file(self.current_note)
        self.save_note_quiet(force=True)

    def copy_custom_field_to_clipboard(self, edit: QLineEdit, btn: QPushButton) -> None:
        text = edit.text().strip()
        if not text:
            QToolTip.showText(QCursor.pos(), "Поле пустое")
            return
        QApplication.clipboard().setText(text)
        self.show_toast(
            "Скопировано",
            boundary_widget=self.dock_editor.widget(),
            anchor_widget=btn,
        )

    def update_current_note_custom_fields(self) -> None:
        if not getattr(self, "current_note", None):
            return
        fields = []
        for w in self.custom_fields_widgets:
            fields.append(
                {
                    "id": w.get("id"),
                    "label": w["label_edit"].text(),
                    "value": w["value_edit"].text(),
                    "visible": w["row"].isVisible(),
                }
            )
        self.current_note.custom_fields = fields

    def on_pm_label_changed(self, text: str) -> None:
        self.password_manager_field.setPlaceholderText(text)
        if hasattr(self, "action_toggle_pm"):
            self._update_eye_action(
                self.action_toggle_pm,
                self.password_manager_row.isVisible(),
                text,
            )

    def on_rdp_label_changed(self, text: str) -> None:
        self.rdp_1c8_field.setPlaceholderText(text)
        if hasattr(self, "action_toggle_rdp"):
            self._update_eye_action(
                self.action_toggle_rdp,
                self.rdp_1c8_row.isVisible(),
                text,
            )
            self.action_toggle_rdp.setEnabled(False)
            self.action_toggle_rdp.setVisible(False)

    def add_toolbar(self) -> None:
        self.init_toolbar()
        layout_splitter = QSplitter(Qt.Vertical)
        layout_splitter.addWidget(self.toolbar_scroll)
        layout_splitter.addWidget(self.main_splitter)
        layout_splitter.setSizes([140, 1000])
        self.layout_splitter = layout_splitter
        self.main_layout = QHBoxLayout()
        self.main_layout.addWidget(self.layout_splitter)
        self.central_widget.setLayout(self.main_layout)

    def _update_eye_action(self, action: QAction, visible: bool, label: str) -> None:
        action.setText(("👁 " if visible else "🙈 ") + label)
        action.setToolTip(("Скрыть " if visible else "Показать ") + label)

    def on_toggle_pm_visible(self, checked: bool) -> None:
        if not self.current_note:
            return
        self.current_note.password_manager_visible = bool(checked)
        self.password_manager_row.setVisible(checked)
        self._update_eye_action(
            self.action_toggle_pm, checked, self.password_manager_label.text()
        )
        self.save_note_to_file(self.current_note)

    def on_toggle_rdp_visible(self, checked: bool) -> None:
        if checked:
            self._ensure_rdp_row_in_layout()
        if not self.current_note:
            return
        if getattr(self.current_note, "rdp_1c8_removed", False):
            if hasattr(self, "action_toggle_rdp"):
                self.action_toggle_rdp.blockSignals(True)
                self.action_toggle_rdp.setChecked(False)
                self.action_toggle_rdp.blockSignals(False)
                self.action_toggle_rdp.setEnabled(False)
                self.action_toggle_rdp.setVisible(False)
            return
        self.current_note.rdp_1c8_visible = bool(checked)
        self.rdp_1c8_row.setVisible(checked)
        self._update_eye_action(
            self.action_toggle_rdp, checked, self.rdp_1c8_label.text()
        )
        self.save_note_to_file(self.current_note)

    def delete_selected_history_entries(self) -> None:
        if not self.current_note:
            return
        selected_indexes = [item.row() for item in self.history_list.selectedIndexes()]
        if not selected_indexes:
            self.show_toast("Нет выбранных версий для удаления")
            return
        selected_indexes.sort(reverse=True)
        for index in selected_indexes:
            if 0 <= index < len(self.current_note.history):
                self.current_note.history.pop(index)
        if self.current_note.history_index >= len(self.current_note.history):
            self.current_note.history_index = len(self.current_note.history) - 1
        elif self.current_note.history_index in selected_indexes:
            self.current_note.history_index = max(
                0, self.current_note.history_index - 1
            )
        if not self.current_note.history:
            self.text_edit.blockSignals(True)
            self.text_edit.clear()
            self.text_edit.blockSignals(False)
            self.current_note.history_index = -1
        else:
            self.text_edit.blockSignals(True)
            self._set_html_preserve_view(
            self.current_note.history[self.current_note.history_index])
            self.text_edit.blockSignals(False)
        self.update_history_list()
        self.update_history_list_selection()
        self.save_note_to_file(self.current_note)
        self.show_toast("Выбранные версии удалены из истории")

    def insert_update_block(self) -> None:
        if not self.current_note:
            return

        now = QDateTime.currentDateTime()
        date_str = now.toString("dd.MM.yyyy")
        cursor = self.text_edit.textCursor()

        html = (
            f"<b>UPD [{date_str}]</b><br>"
            f"<b>Base:</b> <br>"
            f"<b>User:</b> <br>"
            f"<b>Result:</b> <br>"
            f"<b>Details:</b> <br><br>"
        )
        cursor.insertHtml(html)
        self.text_edit.setTextCursor(cursor)
        self.text_edit.setFocus()

    def restore_version_from_history(self, item: QListWidgetItem) -> None:
        if not self.current_note:
            return
        content = item.data(Qt.UserRole)
        if content:
            self._set_html_preserve_view(content)
            self.current_note.history_index = self.history_list.row(item)
            self.update_history_list_selection()

    def _detach_attachments_watcher(self, extra_paths: list[str] | None = None) -> None:
        try:
            if not hasattr(self, "attachments_watcher") or not self.attachments_watcher:
                return
            files = list(self.attachments_watcher.files() or [])
            dirs = list(self.attachments_watcher.directories() or [])
            to_remove = files + dirs
            if extra_paths:
                to_remove.extend(extra_paths)
            if to_remove:
                self.attachments_watcher.removePaths(list(set(to_remove)))
        except Exception:
            pass

    def move_note_to_trash(self, note: Note) -> None:
        self.save_note_to_specific_folder(note, self.TRASH_DIR)
        note_folder_name = NotesApp.safe_folder_name(
            note.title, note.uuid, note.timestamp
        )
        note_dir = os.path.join(NOTES_DIR, note_folder_name)
        trash_dir = os.path.join(self.TRASH_DIR, note_folder_name)
        self._detach_attachments_watcher([note_dir])
        if not os.path.exists(self.TRASH_DIR):
            os.makedirs(self.TRASH_DIR)
        if os.path.exists(trash_dir):
            shutil.rmtree(trash_dir)
        if os.path.exists(note_dir):
            shutil.move(note_dir, trash_dir)
        self.notes = [n for n in self.notes if n.uuid != note.uuid]

    def save_note_to_specific_folder(self, note: Note, folder: str) -> None:
        note_dir = os.path.join(
            folder, NotesApp.safe_folder_name(note.title, note.uuid, note.timestamp)
        )
        os.makedirs(note_dir, exist_ok=True)
        file_path = os.path.join(note_dir, "note.json")
        if self.current_note and note.uuid == self.current_note.uuid:
            note.content = self.text_edit.toHtml()
        note_dict = note.to_dict()
        try:
            if os.path.exists(file_path):
                with open(file_path, "r", encoding="utf-8") as bf:
                    prev_json = bf.read()
                safe_title = (
                    re.sub(r"[^a-zA-Zа-яА-Я0-9 _\-\.\(\)]", "", note.title)
                    .strip()
                    .replace(" ", "_")
                    or "note"
                )
                safe_title = safe_title[:100]
                backup_path = os.path.join(note_dir, f"backup({safe_title}).json")
                with open(backup_path, "w", encoding="utf-8") as out:
                    out.write(prev_json)
                    out.flush()
                    os.fsync(out.fileno())
        except Exception as e:
            print("Backup error:", e)
        if not note.reminder:
            note_dict.pop("reminder", None)
        with open(file_path, "w", encoding="utf-8") as f:
            doc = QTextDocument()
            doc.setHtml(note.content)
            plain_text = doc.toPlainText()
            note_dict["content_txt"] = plain_text
            json.dump(note_dict, f, ensure_ascii=False, indent=4)

    def bring_widget_to_front(self, w: QWidget) -> None:
        old = w.windowFlags()
        w.setWindowFlag(Qt.WindowStaysOnTopHint, True)
        w.show()
        w.raise_()
        w.activateWindow()

        def _restore():
            w.setWindowFlags(old)
            w.show()

        QTimer.singleShot(800, _restore)

    def _center_on_parent(self, dlg: QDialog) -> None:
        try:
            parent = dlg.parent() if isinstance(dlg.parent(), QWidget) else self
            if not isinstance(parent, QWidget):
                return
            dlg.adjustSize()
            rect = QStyle.alignedRect(
                Qt.LeftToRight, Qt.AlignCenter, dlg.sizeHint(), parent.frameGeometry()
            )
            dlg.setGeometry(rect)
        except Exception:
            pass

    def record_state_for_undo(self) -> None:
        note = self.current_note
        if not note:
            self.attachments_scroll.setVisible(False)
            return
        if not hasattr(note, "history") or note.history is None:
            note.history = []
        if note.history_index < len(note.history) - 1:
            note.history = note.history[: note.history_index + 1]
        current_html = self.text_edit.toHtml()
        current_plain = self.text_edit.toPlainText()
        if note.history:
            doc = QTextDocument()
            doc.setHtml(note.history[-1])
            if doc.toPlainText() == current_plain:
                return
        if len(note.history) >= MAX_HISTORY:
            note.history = [current_html]
            note.history_index = 0
        else:
            note.history.append(current_html)
            note.history_index = len(note.history) - 1
        self.update_history_list()
        self.update_history_list_selection()

    def on_tray_icon_activated(self, reason: QSystemTrayIcon.ActivationReason) -> None:
        if reason == QSystemTrayIcon.Trigger:
            self.show()
            self.raise_()
            self.activateWindow()

    def update_history_list(self) -> None:
        if not self.current_note:
            self.history_list.clear()
            return
        self.history_list.clear()
        for i, content in enumerate(self.current_note.history):
            item = QListWidgetItem(f"Версия {i+1}")
            item.setData(Qt.UserRole, content)
            self.history_list.addItem(item)
        self.history_list.scrollToBottom()

    def undo(self) -> None:
        if not self.current_note:
            return
        if self.current_note.history_index > 0:
            self.current_note.history_index -= 1
            self.text_edit.blockSignals(True)
            self._set_html_preserve_view(self.current_note.history[self.current_note.history_index])
            self.text_edit.blockSignals(False)
            self.update_history_list_selection()

    def redo(self) -> None:
        if not self.current_note:
            return
        if self.current_note.history_index < len(self.current_note.history) - 1:
            self.current_note.history_index += 1
            self.text_edit.blockSignals(True)
            self._set_html_preserve_view(self.current_note.history[self.current_note.history_index])
            self.text_edit.blockSignals(False)
            self.update_history_list_selection()

    def update_history_list_selection(self) -> None:
        if not self.current_note:
            return
        self.history_list.blockSignals(True)
        for i in range(self.history_list.count()):
            item = self.history_list.item(i)
            if i == self.current_note.history_index:
                item.setSelected(True)
            else:
                item.setSelected(False)
        self.history_list.blockSignals(False)

    def load_notes_from_disk(self) -> None:
        self.ensure_notes_directory()
        loaded_notes = []
        for folder in os.listdir(NOTES_DIR):
            folder_path = os.path.join(NOTES_DIR, folder)
            if os.path.isdir(folder_path):
                file_path = os.path.join(folder_path, "note.json")
                if os.path.exists(file_path):
                    with open(file_path, "r", encoding="utf-8") as f:
                        data = json.load(f)
                    note = Note.from_dict(data)
                    if "content_txt" in data:
                        note.content_txt = data["content_txt"]
                    else:
                        doc = QTextDocument()
                        doc.setHtml(note.content)
                        note.content_txt = doc.toPlainText()
                    loaded_notes.append(note)
        unique = {}
        for note in loaded_notes:
            unique[note.uuid] = note
        self.notes = list(unique.values())
        self.deduplicate_notes()

    def save_note_to_file(self, note: Note) -> None:
        note_dir = os.path.join(
            NOTES_DIR, NotesApp.safe_folder_name(note.title, note.uuid, note.timestamp)
        )
        os.makedirs(note_dir, exist_ok=True)
        file_path = os.path.join(note_dir, "note.json")
        try:
            if os.path.exists(file_path):
                with open(file_path, "r", encoding="utf-8") as bf:
                    prev_json = bf.read()
                safe_title = (
                    re.sub(r"[^a-zA-Zа-яА-Я0-9 _\-\.\(\)]", "", note.title)
                    .strip()
                    .replace(" ", "_")
                    or "note"
                )
                safe_title = safe_title[:100]
                backup_path = os.path.join(note_dir, f"backup({safe_title}).json")
                with open(backup_path, "w", encoding="utf-8") as out:
                    out.write(prev_json)
                    out.flush()
                    os.fsync(out.fileno())
        except Exception as e:
            print("Backup error:", e)
        if self.current_note and note.uuid == self.current_note.uuid:
            note.content = self.text_edit.toHtml()
        note_dict = note.to_dict()
        if not note.reminder:
            note_dict.pop("reminder", None)
        with open(file_path, "w", encoding="utf-8") as f:
            if self.current_note and note.uuid == self.current_note.uuid:
                doc = QTextDocument()
                doc.setHtml(self.text_edit.toHtml())
                plain_text = doc.toPlainText()
                note_dict["content_txt"] = plain_text
            else:
                doc = QTextDocument()
                doc.setHtml(note.content)
                plain_text = doc.toPlainText()
                note_dict["content_txt"] = plain_text
            json.dump(note_dict, f, ensure_ascii=False, indent=4)

    def save_all_notes_to_disk(self) -> None:
        self.ensure_notes_directory()
        unique_notes = {}
        for note in self.notes:
            unique_notes[note.uuid] = note
        self.notes = list(unique_notes.values())
        for note in self.notes:
            self.save_note_to_file(note)
        self.deduplicate_notes()

    def deduplicate_notes(self) -> None:
        unique = {}
        for note in self.notes:
            unique[note.uuid] = note
        self.notes = list(unique.values())

    def load_settings(self) -> None:
        last_note = self.settings.value("ui/lastOpenedNote", "", str)
        if last_note and os.path.exists(last_note):
            try:
                self.show_note_with_attachments(last_note)
                for row in range(self.notes_list.count()):
                    item = self.notes_list.item(row)
                    if item and getattr(item, "path", None) == last_note:
                        self.notes_list.setCurrentRow(row)
                        break
            except Exception as e:
                print(f"Не удалось открыть последнюю заметку: {e}")
        h = self.settings.value("ui/dock_toolbar_height", type=int)
        if not h:
            h = 36
        else:
            h = max(36, min(h, 300))
        self.resizeDocks([self.dock_toolbar, self.dock_editor],
                        [h, max(300, self.height() - h)], Qt.Vertical)
        self.dock_toolbar.setMinimumHeight(36)
        geometry = self.settings.value("geometry")
        if geometry:
            self.restoreGeometry(geometry)
        state = self.settings.value("windowState")
        if state:
            self.restoreState(state)
        try:
            self._apply_dock_ratios()
        except Exception:
            pass
        last_text = self.settings.value("lastNoteText")
        if last_text and not getattr(self, "current_note", None):
            self.text_edit.blockSignals(True)
            self.text_edit.setHtml(last_text)
            self.text_edit.blockSignals(False)
        else:
            self.settings.setValue("lastNoteText", "")
        self._apply_vertical_defaults()

    def _note_dir(self, note=None) -> str | None:
        n = note or getattr(self, "current_note", None)
        if not n:
            return None
        return os.path.join(
            NOTES_DIR,
            NotesApp.safe_folder_name(n.title, n.uuid, n.timestamp)
        )

    def save_settings(self) -> None:
        self.settings.setValue("geometry", self.saveGeometry())
        self.settings.setValue("ui/dock_toolbar_height", self.dock_toolbar.height())
        self.settings.setValue("windowState", self.saveState())
        self.settings.setValue("ui/dock_ratios", self._dock_ratios)
        self.settings.setValue("lastNoteText", "" if getattr(self, "current_note", None) else self.text_edit.toHtml(),)
        self.settings.sync()

    def confirm_delete_note(self, note: Note) -> None:
        reply = QMessageBox.question(
            self,
            "Удалить заметку",
            f"Вы уверены, что хотите удалить заметку '{note.title}'?",
            QMessageBox.Yes | QMessageBox.No,
        )
        if reply == QMessageBox.Yes:
            self.move_note_to_trash(note)
            note_dir = os.path.join(
                NOTES_DIR, NotesApp.safe_folder_name(note.title, note.uuid)
            )
            if self.current_note and self.current_note.uuid == note.uuid:
                self.current_note = None
                self.text_edit.clear()
            self.refresh_notes_list()
            self.save_all_notes_to_disk()

    def rename_note(self, note: Note) -> None:
        dlg = QInputDialog(self)
        dlg.setWindowTitle("Переименовать")
        dlg.setLabelText("Новое имя заметки:")
        dlg.setTextValue(note.title)
        edit = dlg.findChild(QLineEdit)
        if edit is not None:
            edit.selectAll()
            edit.setProperty("renameTitleEditor", True)
            edit.installEventFilter(self)
        if dlg.exec():
            new_title = dlg.textValue().strip()
            if new_title and new_title != note.title:
                old_dir = os.path.join(
                    NOTES_DIR,
                    NotesApp.safe_folder_name(note.title, note.uuid, note.timestamp),
                )
                note.title = new_title
                new_dir = os.path.join(
                    NOTES_DIR,
                    NotesApp.safe_folder_name(note.title, note.uuid, note.timestamp),
                )
                if os.path.exists(old_dir):
                    os.rename(old_dir, new_dir)
                self.save_note_to_file(note)
                self.refresh_notes_list()

    @staticmethod
    def safe_filename(title: str, ext: str) -> str:
        base = re.sub(r"[^a-zA-Zа-яА-Я0-9 _\-]", "", title)
        base = base.strip().replace(" ", "_")
        if len(base) > 100:
            base = base[:100]
        return f"{base or 'note'}.{ext}"

    def export_current_note_pdf(self) -> None:
        if not self.current_note:
            QMessageBox.warning(self, "Ошибка", "Нет выбранной заметки для экспорта.")
            return
        default_name = self.safe_filename(self.current_note.title, "pdf")
        file_path, _ = QFileDialog.getSaveFileName(
            self, "Сохранить как PDF", default_name, filter="PDF Files (*.pdf)"
        )
        if file_path:
            from PySide6.QtPrintSupport import QPrinter

            doc = QTextDocument()
            html = self.text_edit.toHtml()
            html = re.sub(r'font-size\s*:[^;"]*;?', "", html)
            doc.setHtml(html)
            font = QFont("Times New Roman", 14)
            doc.setDefaultFont(font)
            printer = QPrinter(QPrinter.HighResolution)
            printer.setOutputFormat(QPrinter.PdfFormat)
            printer.setOutputFileName(file_path)
            doc.print_(printer)
            self.show_toast("PDF сохранён")

    def export_current_note_txt(self) -> None:
        if not self.current_note:
            QMessageBox.warning(self, "Ошибка", "Нет выбранной заметки для экспорта.")
            return
        default_name = self.safe_filename(self.current_note.title, "txt")
        file_path, _ = QFileDialog.getSaveFileName(
            self, "Сохранить как TXT", default_name, filter="Text Files (*.txt)"
        )
        if file_path:
            text = self.text_edit.toPlainText()
            with open(file_path, "w", encoding="utf-8") as f:
                f.write(text)
            self.show_toast("TXT сохранён")

    def export_current_note_docx(self) -> None:
        if not self.current_note:
            QMessageBox.warning(self, "Ошибка", "Нет выбранной заметки для экспорта.")
            return
        default_name = self.safe_filename(self.current_note.title, "docx")
        file_path, _ = QFileDialog.getSaveFileName(
            self, "Сохранить как DOCX", default_name, filter="Word Documents (*.docx)"
        )
        if file_path:
            from docx import Document

            doc = Document()
            qt_doc = self.text_edit.document()
            block = qt_doc.firstBlock()
            while block.isValid():
                paragraph = doc.add_paragraph()
                it = block.begin()
                while not it.atEnd():
                    fragment = it.fragment()
                    if fragment.isValid():
                        char_format = fragment.charFormat()
                        if char_format.isImageFormat():
                            image_format = char_format.toImageFormat()
                            image_path = image_format.name()
                            run = paragraph.add_run()
                            if image_path.startswith("data:image"):
                                header, b64data = image_path.split(",", 1)
                                suffix = ".png" if "png" in header else ".jpg"
                                with tempfile.NamedTemporaryFile(
                                    delete=False, suffix=suffix
                                ) as tmp:
                                    tmp.write(base64.b64decode(b64data))
                                    tmp_path = tmp.name
                                run.add_picture(tmp_path)
                                os.unlink(tmp_path)
                            else:
                                if image_path.startswith("file://"):
                                    image_path = QUrl(image_path).toLocalFile()
                                if os.path.exists(image_path):
                                    run.add_picture(image_path)
                        else:
                            paragraph.add_run(fragment.text())
                    it += 1
                block = block.next()
            doc.save(file_path)
            self.show_toast("DOCX сохранён")

    def copy_note(self, note: Note) -> None:
        new_title = note.title + "(копия)"
        new_uuid = str(uuid.uuid4())
        new_note = Note(
            title=new_title,
            content=note.content,
            tags=note.tags.copy(),
            favorite=False,
            history=note.history.copy(),
            reminder=None,
            uuid=new_uuid,
            password_manager=note.password_manager,
            rdp_1c8=note.rdp_1c8,
            custom_fields=[dict(f) for f in getattr(note, "custom_fields", [])],
        )
        new_note.password_manager_visible = note.password_manager_visible
        new_note.rdp_1c8_visible = note.rdp_1c8_visible
        note_dir = os.path.join(
            NOTES_DIR, NotesApp.safe_folder_name(note.title, note.uuid, note.timestamp)
        )
        new_note_dir = os.path.join(
            NOTES_DIR,
            NotesApp.safe_folder_name(new_title, new_uuid, new_note.timestamp),
        )
        if os.path.exists(note_dir):
            shutil.copytree(note_dir, new_note_dir)
            self.save_note_to_file(new_note)
        self.load_notes_from_disk()
        self.refresh_notes_list()
        self.deduplicate_notes()

    def ensure_notes_directory(self) -> None:
        if not os.path.exists(NOTES_DIR):
            os.makedirs(NOTES_DIR)

    def create_new_note(self) -> None:
        self.new_note()

    def save_current_note(self) -> None:
        self.save_note()

    def delete_current_note(self) -> None:
        self.delete_note()

    def new_note(self) -> None:
        title, ok = QInputDialog.getText(self, "Новая заметка", "Введите название:")
        if ok and title:
            for note in self.notes:
                if note.title == title:
                    QMessageBox.warning(
                        self,
                        "Дубликат",
                        f"Заметка с названием '{title}' уже существует.",
                    )
                    return
            note_uuid = str(uuid.uuid4())
            note = Note(
                title=title,
                content="",
                tags=[],
                favorite=False,
                timestamp=QDateTime.currentDateTime().toString(Qt.ISODate),
                reminder=None,
                uuid=note_uuid,
            )
            note.history = [""]
            note.history_index = 0
            note_dir = os.path.join(
                NOTES_DIR, NotesApp.safe_folder_name(title, note_uuid)
            )
            os.makedirs(note_dir, exist_ok=True)
            self.notes.append(note)
            self.current_note = note
            if hasattr(self, "spell_highlighter"):
                self.spell_highlighter.set_local_ignored(
                    self.current_note.ignored_words
                )
            self.text_edit.blockSignals(True)
            self.text_edit.clear()
            self.text_edit.setFont(QFont("Times New Roman", 14))
            self.text_edit.setCurrentFont(QFont("Times New Roman"))
            self.text_edit.setFontPointSize(14)
            self.font_combo.setCurrentFont(QFont("Times New Roman"))
            self.font_size_spin.setValue(14)
            self.text_edit.blockSignals(False)
            self.tags_label.setText("Теги: нет")
            self.password_manager_field.clear()
            self.rdp_1c8_field.clear()
            self.refresh_notes_list()
            self.show_note_with_attachments(note)
            self.text_edit.setFocus()
            self.text_edit.setReadOnly(False)
            note.password_manager_visible = False
            note.rdp_1c8_visible = False
            note.rdp_1c8_removed = True
            self.password_manager_field.clear()
            self.rdp_1c8_field.clear()
            self.password_manager_row.setVisible(False)
            self.rdp_1c8_row.setVisible(False)
            self.clear_custom_fields()
            self.add_field_btn.setEnabled(True)
            if hasattr(self, "action_toggle_pm"):
                self.action_toggle_pm.blockSignals(True)
                self.action_toggle_pm.setChecked(False)
                self._update_eye_action(
                    self.action_toggle_pm, False, self.password_manager_label.text()
                )
                self.action_toggle_pm.blockSignals(False)
            if hasattr(self, "action_toggle_rdp"):
                self.action_toggle_rdp.blockSignals(True)
                self.action_toggle_rdp.setChecked(False)
                self._update_eye_action(
                    self.action_toggle_rdp, False, self.rdp_1c8_label.text()
                )
                self.action_toggle_rdp.blockSignals(False)
                self.action_toggle_rdp.setEnabled(False)
                self.action_toggle_rdp.setVisible(False)

    def save_note(self) -> None:
        if self.current_note:
            html = self.text_edit.toHtml()
            html = self._persist_dropdown_values_in_html(html)
            self.current_note.content = html
            self.current_note.password_manager = self.password_manager_field.text()
            self.current_note.rdp_1c8 = self.rdp_1c8_field.text()
            self.update_current_note_custom_fields()
            self.save_note_to_file(self.current_note)
            self.refresh_notes_list()
            self.record_state_for_undo()
            self.show_toast("Заметка сохранена")
            self._set_unsaved(False)

    def save_note_quiet(self, *, force: bool = False) -> None:
        if not force and not getattr(self, "autosave_enabled", True):
            return
        if self.current_note:
            html = self.text_edit.toHtml()
            html = self._persist_dropdown_values_in_html(html)
            self.current_note.content = html
            self.current_note.password_manager = self.password_manager_field.text()
            self.current_note.rdp_1c8 = self.rdp_1c8_field.text()
            self.update_current_note_custom_fields()
            self.save_note_to_file(self.current_note)
            self.record_state_for_undo()
            self._set_unsaved(False)

    def open_note_folder(self, note: Note) -> None:
        if not note:
            self.attachments_scroll.setVisible(False)
            return
        folder_name = NotesApp.safe_folder_name(note.title, note.uuid, note.timestamp)
        candidates = [
            os.path.join(NOTES_DIR, folder_name),
            os.path.join(self.TRASH_DIR, folder_name),
        ]
        for path in candidates:
            if os.path.isdir(path):
                QDesktopServices.openUrl(QUrl.fromLocalFile(os.path.abspath(path)))
                return
        self.save_note_to_file(note)
        path = os.path.join(NOTES_DIR, folder_name)
        os.makedirs(path, exist_ok=True)
        QDesktopServices.openUrl(QUrl.fromLocalFile(os.path.abspath(path)))

    def delete_note(self) -> None:
        selected_items = self.notes_list.selectedItems()
        if selected_items:
            reply = QMessageBox.question(
                self,
                "Удаление",
                f"Переместить выбранные заметки в корзину? ({len(selected_items)} шт.)",
                QMessageBox.Yes | QMessageBox.No,
            )
            if reply != QMessageBox.Yes:
                return
            for item in selected_items:
                note = item.data(Qt.UserRole)
                if note:
                    self.move_note_to_trash(note)
            self.refresh_notes_list()
            self.load_note(None)
            self.current_note = None
            self.show_toast("Перемещены в корзину")
            return
        if self.current_note:
            title = self.current_note.title
            reply = QMessageBox.question(
                self,
                "Удаление",
                f"Переместить заметку '{title}' в корзину?",
                QMessageBox.Yes | QMessageBox.No,
            )
            if reply == QMessageBox.Yes:
                self.move_note_to_trash(self.current_note)
                self.refresh_notes_list()
                self.load_note(None)
                self.current_note = None
                self.show_toast(f"«{title}» перемещена в корзину")
            return
        self.show_toast("Нет заметки для удаления")

    def load_templates(self) -> list[dict]:
        templates_path = os.path.join(DATA_DIR, "templates.json")
        base_templates = [
            {
                "name": "Список дел",
                "category": "Работа",
                "description": "Чекбокс-лист для задач",
                "content_html": "<b>Список дел:</b><br>☐ Первая задача<br>☐ Вторая задача<br>☐ Третья задача<br>",
            },
            {
                "name": "Встреча",
                "category": "Встречи",
                "description": "Заготовка для заметки о встрече",
                "content_html": "<b>Встреча</b><br>Дата: <br>Участники: <br>Цели: <br>Результаты: <br>",
            },
            {
                "name": "Дневник",
                "category": "Личное",
                "description": "Дневниковая запись",
                "content_html": "<b>Дневник</b><br>Дата: <br>Сегодня:<br><br>Настроение:<br>События:<br>",
            },
            {
                "name": "UPD блок",
                "category": "Работа",
                "description": "Блок обновления с датой, базой, пользователем, результатом и деталями",
                "content_html": "<b>UPD [{date}]</b><br><b>Base:</b> <br><b>User:</b> <br><b>Result:</b> <br><b>Details:</b> <br><br>",
            },
        ]
        auto_today_tpl = {
            "name": "Сегодня",
            "category": "Работа",
            "description": "Строка с текущей датой",
            "content_html": (
                '<table style="border-collapse:collapse;">'
                '<tr><td style="border:1px solid #bdbdbd; padding:4px 8px; border-radius:4px;">'
                'UPD [<span data-dyn="today">{today}</span>]'
                "</td></tr></table><br>"
            ),
        }
        correct_range5 = {
            "name": "Интервал 5 дней",
            "category": "Работа",
            "description": "Таблица 1×1 с диапазоном дат (N=5 дней, редактируемая начальная дата)",
            "content_html": (
                '<table style="border-collapse:collapse;">'
                '<tr><td style="border:1px solid #bdbdbd; padding:4px 8px; border-radius:4px;">'
                '<span style="font-weight:600;">{date}-{date+4}</span>'
                "</td></tr></table><br>"
            ),
        }
        if not os.path.exists(templates_path):
            templates = base_templates + [auto_today_tpl, correct_range5]
            os.makedirs(DATA_DIR, exist_ok=True)
            with open(templates_path, "w", encoding="utf-8") as f:
                json.dump(templates, f, ensure_ascii=False, indent=4)
            return templates
        try:
            with open(templates_path, "r", encoding="utf-8") as f:
                templates = json.load(f)
                if not isinstance(templates, list):
                    templates = []
        except Exception:
            templates = []
        by_name = {t.get("name"): t for t in templates}

        def upsert(tpl: dict) -> None:
            cur = by_name.get(tpl["name"])
            if cur is None:
                templates.append(tpl)
                by_name[tpl["name"]] = tpl
            else:
                if cur.get("content_html") != tpl.get("content_html"):
                    cur.update(tpl)

        for t in base_templates:
            upsert(t)
        upsert(auto_today_tpl)
        upsert(correct_range5)
        with open(templates_path, "w", encoding="utf-8") as f:
            json.dump(templates, f, ensure_ascii=False, indent=4)
        return templates

    def save_templates(self, templates: list[dict]) -> None:
        templates_path = os.path.join(DATA_DIR, "templates.json")
        os.makedirs(DATA_DIR, exist_ok=True)
        with open(templates_path, "w", encoding="utf-8") as f:
            json.dump(templates, f, ensure_ascii=False, indent=4)

    def save_current_as_template(self) -> None:
        if not self.current_note:
            QMessageBox.warning(self, "Нет заметки", "Сначала выбери заметку.")
            return
        dialog = QDialog(self)
        dialog.setWindowTitle("Сохранить как шаблон")
        layout = QFormLayout(dialog)
        name_edit = QLineEdit()
        category_edit = QLineEdit()
        description_edit = QLineEdit()
        only_selection = QCheckBox("Только выделенный фрагмент")
        layout.addRow("Название:", name_edit)
        layout.addRow("Категория:", category_edit)
        layout.addRow("Описание:", description_edit)
        layout.addRow("", only_selection)
        preview = QTextEdit()
        preview.setMinimumHeight(180)
        layout.addRow("Содержимое:", preview)

        def fill_preview():
            cur = self.text_edit.textCursor()
            if only_selection.isChecked() and cur.hasSelection():
                frag = cur.selection()
                preview.setHtml(frag.toHtml())
            else:
                preview.setHtml(self.text_edit.toHtml())

        only_selection.toggled.connect(fill_preview)
        fill_preview()
        buttons = QDialogButtonBox(QDialogButtonBox.Save | QDialogButtonBox.Cancel)
        layout.addRow(buttons)

        def on_accept():
            name = name_edit.text().strip()
            if not name:
                QMessageBox.warning(dialog, "Ошибка", "Укажи название шаблона.")
                return
            tpl = {
                "name": name,
                "category": (category_edit.text().strip() or "Без категории"),
                "description": description_edit.text().strip(),
                "content_html": preview.toHtml(),
            }
            templates = self.load_templates()
            exist_idx = next(
                (i for i, t in enumerate(templates) if t.get("name") == name), -1
            )
            if exist_idx >= 0:
                r = QMessageBox.question(
                    dialog,
                    "Перезаписать?",
                    f"Шаблон «{name}» уже существует. Перезаписать?",
                    QMessageBox.Yes | QMessageBox.No,
                )
                if r == QMessageBox.No:
                    return
                templates[exist_idx] = tpl
            else:
                templates.append(tpl)
            self.save_templates(templates)
            dialog.accept()
            self.show_toast("Шаблон сохранён")

        buttons.accepted.connect(on_accept)
        buttons.rejected.connect(dialog.reject)
        self._center_on_parent(dialog)
        dialog.exec()

    def _render_date_placeholders(self, html: str, base_date: QDate) -> str:
        pattern = re.compile(r"\{date(?:\s*(?P<sign>[+-])\s*(?P<days>\d+))?\}", re.I)

        def repl(m: re.Match) -> str:
            days = int(m.group("days") or 0)
            if m.group("sign") == "-":
                days = -days
            d = base_date.addDays(days)
            return d.toString("dd.MM.yyyy")

        return pattern.sub(repl, html)

    def _apply_dynamic_tokens(self, html: str) -> str:
        today = QDate.currentDate().toString("dd.MM.yyyy")
        html = re.sub(
            r'(<span\b[^>]*\bdata-dyn="today"[^>]*>)(.*?)(</span>)',
            lambda m: m.group(1) + today + m.group(3),
            html,
            flags=re.I | re.S,
        )
        html = re.sub(r"\{today\}", today, html, flags=re.I)
        return html

    def insert_template(self) -> None:
        if not self.current_note:
            QMessageBox.warning(self, "Нет заметки", "Создайте или выберите заметку.")
            return
        templates = self.load_templates()
        dialog = QDialog(self)
        dialog.setWindowFlag(Qt.WindowStaysOnTopHint, True)
        dialog.setWindowTitle("Вставить шаблон")
        layout = QVBoxLayout(dialog)
        combo = QComboBox(dialog)
        categories = sorted({tpl.get("category", "Без категории") for tpl in templates})
        category_combo = QComboBox(dialog)
        category_combo.addItem("Все категории")
        category_combo.addItems(categories)
        filtered_templates = templates.copy()

        def filter_templates():
            cat = category_combo.currentText()
            combo.clear()
            filtered = [
                tpl
                for tpl in templates
                if cat == "Все категории" or tpl.get("category", "Без категории") == cat
            ]
            filtered_templates.clear()
            filtered_templates.extend(filtered)
            combo.addItems(
                [f"{tpl['name']} ({tpl.get('description','')})" for tpl in filtered]
            )

        category_combo.currentIndexChanged.connect(filter_templates)
        layout.addWidget(category_combo)
        layout.addWidget(combo)
        preview = QTextEdit(dialog)
        preview.setReadOnly(True)
        preview.setMinimumHeight(100)
        layout.addWidget(preview)

        def update_preview(idx):
            if 0 <= idx < len(filtered_templates):
                content = filtered_templates[idx]["content_html"]
                today = QDate.currentDate()
                content = self._render_date_placeholders(content, today)
                preview.setHtml(content)
            else:
                preview.setHtml("")

        combo.currentIndexChanged.connect(update_preview)
        filter_templates()
        combo.setCurrentIndex(0)
        update_preview(0)
        buttons = QDialogButtonBox(QDialogButtonBox.Ok | QDialogButtonBox.Cancel)
        layout.addWidget(buttons)
        dialog.setLayout(layout)

        def _toggle_launcher_topmost(self, checked: bool) -> None:
            self.setWindowFlag(Qt.WindowStaysOnTopHint, checked)
            self.settings.setValue("ui/launcher_always_on_top", checked)
            self.show()

        def keyPressEvent(event):
            if event.key() == Qt.Key_Escape:
                dialog.reject()
            else:
                QDialog.keyPressEvent(dialog, event)

        dialog.keyPressEvent = keyPressEvent
        buttons.accepted.connect(dialog.accept)
        buttons.rejected.connect(dialog.reject)
        self._center_on_parent(dialog)
        if dialog.exec() == QDialog.Accepted:
            idx = combo.currentIndex()
            if 0 <= idx < len(filtered_templates):
                tpl = filtered_templates[idx]
                content_html = tpl["content_html"]
                date_pattern = re.compile(r"\{date(?:[+-]\s*\d+)?\}", re.I)
                base_date = QDate.currentDate()
                if date_pattern.search(content_html):
                    ask = QDialog(self)
                    ask.setWindowTitle("Начальная дата")
                    form = QFormLayout(ask)
                    de = QDateTimeEdit(ask)
                    de.setDisplayFormat("dd.MM.yyyy")
                    de.setCalendarPopup(True)
                    de.setDate(QDate.currentDate())
                    form.addRow("Начальная дата:", de)
                    buttons = QDialogButtonBox(
                        QDialogButtonBox.Ok | QDialogButtonBox.Cancel, ask
                    )
                    form.addRow(buttons)
                    buttons.accepted.connect(ask.accept)
                    buttons.rejected.connect(ask.reject)
                    self._center_on_parent(ask)
                    if ask.exec() != QDialog.Accepted:
                        return
                    base_date = de.date()
                content_html = self._render_date_placeholders(content_html, base_date)
                content_html = self._apply_dynamic_tokens(content_html)
                cursor = self.text_edit.textCursor()
                start = cursor.position()
                self.text_edit.insertHtml(content_html)
                end = self.text_edit.textCursor().position()
                if end > start:
                    cursor.setPosition(start)
                    cursor.setPosition(end, QTextCursor.KeepAnchor)
                    fmt = QTextCharFormat()
                    font = QFont("Times New Roman", 14)
                    fmt.setFont(font)
                    cursor.mergeCharFormat(fmt)
                    cursor.clearSelection()
                    self.text_edit.setTextCursor(cursor)
                self.record_state_for_undo()

    def manage_templates_dialog(self) -> None:
        templates = self.load_templates()
        dlg = QDialog(self)
        dlg.setWindowTitle("Шаблоны")
        v = QVBoxLayout(dlg)
        lst = QListWidget()

        def refresh():
            lst.clear()
            for t in templates:
                lst.addItem(f"{t.get('name','')} — {t.get('category','Без категории')}")

        refresh()
        v.addWidget(lst)
        form = QFormLayout()
        name_edit = QLineEdit()
        category_edit = QLineEdit()
        desc_edit = QLineEdit()
        form.addRow("Название:", name_edit)
        form.addRow("Категория:", category_edit)
        form.addRow("Описание:", desc_edit)
        v.addLayout(form)
        content = QTextEdit()
        content.setMinimumHeight(180)
        v.addWidget(content)
        btns_line = QHBoxLayout()
        btn_new = QPushButton("Новый")
        btn_save = QPushButton("Сохранить")
        btn_del = QPushButton("Удалить")
        btn_close = QPushButton("Закрыть")
        for b in (btn_new, btn_save, btn_del, btn_close):
            btns_line.addWidget(b)
        v.addLayout(btns_line)

        def load_current(i):
            if i < 0 or i >= len(templates):
                name_edit.clear()
                category_edit.clear()
                desc_edit.clear()
                content.clear()
                return
            t = templates[i]
            name_edit.setText(t.get("name", ""))
            category_edit.setText(t.get("category", ""))
            desc_edit.setText(t.get("description", ""))
            content.setHtml(t.get("content_html", ""))

        lst.currentRowChanged.connect(load_current)
        if templates:
            lst.setCurrentRow(0)

        def on_new():
            templates.append(
                {
                    "name": "Новый шаблон",
                    "category": "Без категории",
                    "description": "",
                    "content_html": "",
                }
            )
            refresh()
            lst.setCurrentRow(len(templates) - 1)

        def on_save():
            i = lst.currentRow()
            if i < 0:
                return
            templates[i] = {
                "name": name_edit.text().strip() or "Без имени",
                "category": category_edit.text().strip() or "Без категории",
                "description": desc_edit.text().strip(),
                "content_html": content.toHtml(),
            }
            self.save_templates(templates)
            refresh()
            lst.setCurrentRow(i)
            self.show_toast("Сохранено")

        def on_del():
            i = lst.currentRow()
            if i < 0:
                return
            r = QMessageBox.question(
                dlg,
                "Удалить",
                "Удалить выбранный шаблон?",
                QMessageBox.Yes | QMessageBox.No,
            )
            if r == QMessageBox.Yes:
                templates.pop(i)
                self.save_templates(templates)
                refresh()
                lst.setCurrentRow(min(i, len(templates) - 1))

        btn_new.clicked.connect(on_new)
        btn_save.clicked.connect(on_save)
        btn_del.clicked.connect(on_del)
        btn_close.clicked.connect(dlg.accept)
        self._center_on_parent(dlg)

        def _on_show(ev):
            QDialog.showEvent(dlg, ev)
            self._center_on_parent(dlg)

        dlg.showEvent = _on_show
        self._center_on_parent(dlg)
        dlg.exec()

    def show_trash(self) -> None:
        self.notes_list.clear()
        if not os.path.exists(self.TRASH_DIR):
            return
        for folder in os.listdir(self.TRASH_DIR):
            folder_path = os.path.join(self.TRASH_DIR, folder)
            if os.path.isdir(folder_path):
                file_path = os.path.join(folder_path, "note.json")
                if os.path.exists(file_path):
                    with open(file_path, "r", encoding="utf-8") as f:
                        data = json.load(f)
                        note = Note.from_dict(data)
                        doc = QTextDocument()
                        doc.setHtml(note.content)
                        note.content_txt = doc.toPlainText()
                        title = note.title
                        timestamp = QDateTime.fromString(note.timestamp, Qt.ISODate)
                        date_str = timestamp.toString("dd.MM.yyyy")
                        item_text = f"{title} — {date_str} 🗑"
                        item = QListWidgetItem(item_text)
                        item.setFont(QFont("Segoe UI Emoji", 10))
                        item.setData(Qt.UserRole, note)
                        item.setForeground(QColor("gray"))
                        self.notes_list.addItem(item)

    def restore_note_from_trash(self) -> None:
        selected_items = self.notes_list.selectedItems()
        if not selected_items:
            return
        item = selected_items[0]
        note = item.data(Qt.UserRole)
        trash_note_dir = os.path.join(
            self.TRASH_DIR,
            NotesApp.safe_folder_name(note.title, note.uuid, note.timestamp),
        )
        note_dir = os.path.join(
            NOTES_DIR, NotesApp.safe_folder_name(note.title, note.uuid, note.timestamp)
        )
        if os.path.exists(trash_note_dir):
            shutil.move(trash_note_dir, note_dir)
        self.load_notes_from_disk()
        self.refresh_notes_list()
        for i in range(self.notes_list.count()):
            item = self.notes_list.item(i)
            note_obj = item.data(Qt.UserRole)
            if note_obj.uuid == note.uuid:
                self.notes_list.setCurrentRow(i)
                self.select_note(note_obj)
                break
        self.show_toast("Заметка восстановлена")

    def delete_note_permanently(self) -> None:
        selected_items = self.notes_list.selectedItems()
        if not selected_items:
            return
        reply = QMessageBox.question(
            self,
            "Удалить навсегда",
            "Удалить выбранную заметку безвозвратно?",
            QMessageBox.Yes | QMessageBox.No,
        )
        if reply == QMessageBox.Yes:
            item = selected_items[0]
            note = item.data(Qt.UserRole)
            trash_note_dir = os.path.join(
                self.TRASH_DIR,
                NotesApp.safe_folder_name(note.title, note.uuid, note.timestamp),
            )
            if os.path.exists(trash_note_dir):
                self._detach_attachments_watcher([trash_note_dir])
                shutil.rmtree(trash_note_dir)
            self.refresh_notes_list()
            self.show_toast("Заметка удалена навсегда")

    def empty_trash(self) -> None:
        if os.path.exists(self.TRASH_DIR):
            reply = QMessageBox.question(
                self,
                "Очистить корзину",
                "Удалить все заметки из корзины безвозвратно?",
                QMessageBox.Yes | QMessageBox.No,
            )
            if reply == QMessageBox.Yes:
                for folder in os.listdir(self.TRASH_DIR):
                    folder_path = os.path.join(self.TRASH_DIR, folder)
                    shutil.rmtree(folder_path)
                self.refresh_notes_list()
                self.show_toast("Корзина очищена")

    def load_note(self, item: QListWidgetItem | None) -> None:
        if item is None:
            self.setWindowTitle("Мои Заметки — [*]")
            self.setWindowModified(False)
            self.text_edit.blockSignals(True)
            self.text_edit.clear()
            self.text_edit.blockSignals(False)
            self.text_edit.setReadOnly(True)
            self.text_edit.hide()
            self._detach_attachments_watcher()
            self.tags_label.setText(f"Теги: {', '.join(self.current_note.tags) if (self.current_note and self.current_note.tags) else 'нет'}")
            self.current_note = None
            if hasattr(self, "spell_highlighter"):
                self.spell_highlighter.set_local_ignored([])
            self.clear_custom_fields()
            self.add_field_btn.setEnabled(False)
            if hasattr(self, "history_list"):
                self.history_list.clear()
            self.attachments_scroll.setVisible(False)
            self._update_editor_visibility()
            return
        note = item.data(Qt.ItemDataRole.UserRole)
        self.select_note(note)
        self._update_editor_visibility()

    def select_note(self, note: Note | None) -> None:
        if note is None:
            return
        if getattr(self, "current_note", None) and self.current_note.uuid == note.uuid:
            return
        self._suspend_edit_signals()
        try:
            if getattr(self, "current_note", None):
                html = self.text_edit.toHtml()
                html = self._persist_dropdown_values_in_html(html)
                self.current_note.content = html
                self.save_note_to_file(self.current_note)
            self.current_note = note
            self.tags_label.setText(f"Теги: {', '.join(note.tags) if note and note.tags else 'нет'}")
            if hasattr(self, "settings"):
                self.settings.setValue("lastNoteText", "")
            if hasattr(self, "spell_highlighter"):
                self.spell_highlighter.set_local_ignored(note.ignored_words)
            self.text_edit.show()
            self.text_edit.setReadOnly(False)
            self.setWindowTitle(f"Мои Заметки — {note.title} [*]")
            self.setWindowModified(False)
            if hasattr(self, "dock_editor"):
                self.dock_editor.setWindowTitle("Редактор")
            pm_vis = bool(getattr(note, "password_manager_visible", False))
            rdp_vis = bool(getattr(note, "rdp_1c8_visible", False))
            rdp_removed = bool(getattr(note, "rdp_1c8_removed", False))
            if rdp_vis and not rdp_removed:
                self._ensure_rdp_row_in_layout()
            self.rdp_1c8_row.setVisible(False if rdp_removed else rdp_vis)
            self.password_manager_row.setVisible(pm_vis)
            self.rdp_1c8_row.setVisible(False if rdp_removed else rdp_vis)
            if hasattr(self, "action_toggle_pm"):
                self.action_toggle_pm.blockSignals(True)
                self.action_toggle_pm.setChecked(pm_vis)
                self._update_eye_action(
                    self.action_toggle_pm, pm_vis, self.password_manager_label.text()
                )
                self.action_toggle_pm.blockSignals(False)
            if hasattr(self, "action_toggle_rdp"):
                self.action_toggle_rdp.blockSignals(True)
                self.action_toggle_rdp.setVisible(not rdp_removed)
                self.action_toggle_rdp.setEnabled(not rdp_removed)
                self.action_toggle_rdp.setChecked(rdp_vis)
                self.action_toggle_rdp.setEnabled(False)
                self.action_toggle_rdp.setVisible(False)
                self._update_eye_action(
                    self.action_toggle_rdp, rdp_vis, self.rdp_1c8_label.text()
                )
                self.action_toggle_rdp.blockSignals(False)
            fields_data = list(getattr(note, "custom_fields", []))
            self.clear_custom_fields()
            for field in fields_data:
                self.add_custom_field(field)
            self.add_field_btn.setEnabled(True)

            if hasattr(self, "settings"):
                self.settings.setValue("lastNoteUUID", note.uuid)

            if not getattr(note, "history", None):
                note.history = [note.content]
                note.history_index = 0
            self.update_history_list()
            self.update_history_list_selection()
            self.show_note_with_attachments(note)

            self.tags_label.setText(f"Теги: {', '.join(note.tags) if note.tags else 'нет'}")
            self.password_manager_field.setText(note.password_manager)
            self.rdp_1c8_field.setText(note.rdp_1c8)

            for i in range(self.notes_list.count()):
                item = self.notes_list.item(i)
                n = item.data(Qt.UserRole)
                if n and n.uuid == note.uuid:
                    self.notes_list.setCurrentItem(item)
                    break
        finally:
            self._resume_edit_signals()

    def _html_escape(self, s: str) -> str:
        return s.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")

    def _load_dropdown_values(self) -> list[str]:
        try:
            vals = json.loads(self.settings.value("dropdown_values", "[]"))
            if not isinstance(vals, list):
                vals = []
        except Exception:
            vals = []
        if not vals:
            vals = ["Вариант 1", "Вариант 2", "Вариант 3"]
        return [str(v) for v in vals]

    def _dropdown_palette(self) -> tuple[str, str, str]:
        is_dark = self.settings.value("theme", "dark") == "dark"
        if is_dark:
            return ("#303030", "#ffffff", "#7a7a7a")
        return ("#f0f0f0", "#111111", "#b0b0b0")

    def _ensure_dd_map(self) -> None:
        if not hasattr(self, "_dropdown_tokens"):
            self._dropdown_tokens = {}

    def _get_dropdown_token_info(self, dd_id: str) -> dict | None:
        self._ensure_dd_map()
        if dd_id in self._dropdown_tokens:
            return self._dropdown_tokens[dd_id]

        html = self.text_edit.toHtml()
        href_prefix = f"dropdown://{re.escape(dd_id)}"
        m = re.search(
            rf'<a[^>]*href="{href_prefix}[^"]*"[^>]*>(.*?)</a>', html, flags=re.S
        )
        if not m:
            return None

        full_tag = m.group(0)
        inner = m.group(1)

        vals = None
        href_m = re.search(r'href="([^"]+)"', full_tag)
        if href_m:
            href_val = href_m.group(1)
            v_m = re.search(r'[?#]v=([^"&]*)', href_val)
            if v_m:
                try:
                    vals = json.loads(unquote(v_m.group(1)))
                except Exception:
                    vals = None

        if vals is None:
            title_m = re.search(r'title=(["\'])(.*?)\1', full_tag)
            if title_m:
                try:
                    vals = json.loads(unquote(title_m.group(2)))
                except Exception:
                    vals = None

        if vals is None:
            vals = self._load_dropdown_values()

        text = re.sub(r"<[^>]+>", "", inner)
        text = text.replace("▼", "").replace("&#9662;", "").strip()

        info = {"value": text, "values": vals}
        self._dropdown_tokens[dd_id] = info
        return info

    def _insert_dropdown_token(
        self, value: str, dd_id: str | None = None, values: list[str] | None = None
    ) -> str:
        self._ensure_dd_map()
        if dd_id is None:
            dd_id = uuid.uuid4().hex[:8]
        if values is None:
            values = self._load_dropdown_values()
        self._dropdown_tokens[dd_id] = {"value": value, "values": list(values)}
        bg, fg, br = self._dropdown_palette()
        inner = (
            f'<span style="display:inline-block; padding:2px 8px; '
            f"border-radius:6px; border:1px solid {br}; "
            f'background:{bg}; color:{fg};">{self._html_escape(value)} &#9662;</span>'
        )
        encoded = quote(json.dumps(values, ensure_ascii=False))
        html = (
            f'<a href="dropdown://{dd_id}" title="{encoded}" '
            f'style="text-decoration:none;">{inner}</a>&nbsp;'
        )
        c = self.text_edit.textCursor()
        c.insertHtml(html)
        self.text_edit.setTextCursor(c)
        self.record_state_for_undo()
        if hasattr(self, "debounce_timer"):
            if self.autosave_enabled:
                self.save_note_quiet(force=True)
        return dd_id

    def _update_dropdown_token(self, dd_id: str, new_value: str) -> None:
        info = self._get_dropdown_token_info(dd_id)
        if not info:
            return
        info["value"] = new_value
        values = info.get("values", [])
        bg, fg, br = self._dropdown_palette()
        inner = (
            f'<span style="display:inline-block; padding:2px 8px; '
            f"border-radius:6px; border:1px solid {br}; "
            f'background:{bg}; color:{fg};">{self._html_escape(new_value)} &#9662;</span>'
        )
        encoded = quote(json.dumps(values, ensure_ascii=False))
        new_anchor = (
            f'<a href="dropdown://{dd_id}?v={encoded}" '
            f'style="text-decoration:none;">{inner}</a>'
        )
        html = self.text_edit.toHtml()
        href_prefix = f"dropdown://{re.escape(dd_id)}"
        pattern = re.compile(rf'<a[^>]*href="{href_prefix}[^"]*"[^>]*>.*?</a>', re.S)
        new_html, count = pattern.subn(new_anchor, html, count=1)
        if count == 0:
            return
        self.text_edit.blockSignals(True)
        self.text_edit.setHtml(new_html)
        self.text_edit.blockSignals(False)
        self.record_state_for_undo()
        if hasattr(self, "debounce_timer"):
            if self.autosave_enabled:
                self.save_note_quiet(force=True)

    def _persist_dropdown_values_in_html(self, html: str) -> str:
        pattern = re.compile(
            r'<a[^>]*href="dropdown://([0-9a-fA-F]{8})[^"]*"[^>]*>(.*?)</a>', re.S
        )

        def repl(m):
            dd_id = m.group(1)
            _ = m.group(2)

            info = self._get_dropdown_token_info(dd_id)
            if not info:
                return m.group(0)
            value = info.get("value", "")
            values = info.get("values", []) or []
            bg, fg, br = self._dropdown_palette()
            inner = (
                f'<span style="display:inline-block; padding:2px 8px; '
                f"border-radius:6px; border:1px solid {br}; "
                f'background:{bg}; color:{fg};">{self._html_escape(value)} &#9662;</span>'
            )

            encoded = quote(json.dumps(values, ensure_ascii=False))
            return (
                f'<a href="dropdown://{dd_id}?v={encoded}" '
                f'style="text-decoration:none;">{inner}</a>'
            )

        return pattern.sub(repl, html)

    def _edit_values_and_reopen(self, dd_id: str) -> None:
        info = self._get_dropdown_token_info(dd_id)
        if not info:
            return
        vals = self._open_dropdown_values_editor(info.get("values"))
        if vals is not None:
            info["values"] = vals
            self._update_dropdown_token(dd_id, info.get("value", ""))
            self.show_dropdown_menu_for_token(dd_id)

    def show_dropdown_menu_for_token(self, dd_id: str, global_pos=None) -> None:
        info = self._get_dropdown_token_info(dd_id)
        values = (
            info.get("values", self._load_dropdown_values())
            if info
            else self._load_dropdown_values()
        )
        menu = QMenu(self)
        menu.setStyleSheet(CUSTOM_MENU_STYLE)
        for v in values:
            act = menu.addAction(v)
            act.triggered.connect(
                lambda _, val=v: self._update_dropdown_token(dd_id, val)
            )
        menu.addSeparator()
        add_act = menu.addAction("➕ Добавить значение…")
        add_act.triggered.connect(lambda: self._add_value_to_dropdown(dd_id))
        curr = (info or {}).get("value", "")
        del_act = menu.addAction(f"➖ Удалить «{curr}» из списка")
        del_act.setEnabled(bool(curr and curr in values))
        del_act.triggered.connect(lambda: self._remove_value_from_dropdown(dd_id, curr))
        menu.addSeparator()
        edit_act = menu.addAction("✏️ Редактировать список…")
        edit_act.triggered.connect(lambda: self._edit_values_and_reopen(dd_id))
        if global_pos is None:
            cr = self.text_edit.cursorRect(self.text_edit.textCursor())
            global_pos = self.text_edit.viewport().mapToGlobal(cr.bottomLeft())
        menu.exec(global_pos)

    def show_note_with_attachments(self, note: Note | None) -> None:
        self.current_note = note
        self.settings.setValue("lastNoteUuid", note.uuid)
        files = self.attachments_watcher.files()
        if files:
            self.attachments_watcher.removePaths(files)
        dirs = self.attachments_watcher.directories()
        if dirs:
            self.attachments_watcher.removePaths(dirs)
        if not note:
            self.attachments_scroll.setVisible(False)
            self.text_edit.blockSignals(True)
            self.text_edit.clear()
            self.text_edit.blockSignals(False)
            if hasattr(self, "attachments_layout"):
                for i in reversed(range(self.attachments_layout.count())):
                    widget = self.attachments_layout.itemAt(i).widget()
                    if widget:
                        widget.setParent(None)
                        widget.deleteLater()
            self.attachments_scroll.setVisible(False)
            return
        cursor_pos = None
        anchor_pos = None
        if note == self.current_note:
            cursor = self.text_edit.textCursor()
            cursor_pos = cursor.position()
            anchor_pos = cursor.anchor()
        self.text_edit.blockSignals(True)
        raw_html = (
            note.content_html
            if hasattr(note, "content_html") and note.content_html
            else note.content
        )
        self.settings.setValue("ui/lastOpenedNote", self.current_note_path)
        rendered = self._apply_dynamic_tokens(raw_html)
        self.text_edit.blockSignals(True)
        self.text_edit.setHtml(rendered)
        self.text_edit.blockSignals(False)
        if hasattr(self, "_dropdown_tokens"):
            self._dropdown_tokens.clear()
        self.text_edit.blockSignals(False)
        if cursor_pos is not None and anchor_pos is not None:
            self._safe_restore_cursor(anchor_pos, cursor_pos)
        if hasattr(self, "attachments_layout"):
            for i in reversed(range(self.attachments_layout.count())):
                widget = self.attachments_layout.itemAt(i).widget()
                if widget:
                    widget.setParent(None)
                    widget.deleteLater()
            note_dir = os.path.join(
                NOTES_DIR,
                NotesApp.safe_folder_name(note.title, note.uuid, note.timestamp),
            )
            attachments_found = False
            if os.path.isdir(note_dir):
                ignored_files = {"note.json", ".DS_Store", "Thumbs.db"}
                ignored_prefixes = ("~$", ".~")
                ignored_suffixes = ("~", ".tmp", ".temp")
                for filename in os.listdir(note_dir):
                    if (
                        filename in ignored_files
                        or filename.startswith(ignored_prefixes)
                        or filename.endswith(ignored_suffixes)
                        or (
                            filename.startswith("backup(")
                            and filename.endswith(".json")
                        )
                    ):
                        continue
                    attachments_found = True
                    file_path = os.path.join(note_dir, filename)
                    item_widget = QWidget()
                    layout = QHBoxLayout(item_widget)
                    if filename.lower().endswith(tuple(IMAGE_EXTENSIONS)):
                        pixmap = QPixmap(file_path)
                        icon_label = QLabel()
                        icon_label.setPixmap(
                            pixmap.scaled(
                                48, 48, Qt.KeepAspectRatio, Qt.SmoothTransformation
                            )
                        )
                        layout.addWidget(icon_label)
                    else:
                        icon_label = QLabel()
                        if os.path.exists(FILE_ICON_PATH):
                            icon_label.setPixmap(
                                QPixmap(FILE_ICON_PATH).scaled(
                                    48, 48, Qt.KeepAspectRatio, Qt.SmoothTransformation
                                )
                            )
                        else:
                            icon_label.setPixmap(
                                self.style()
                                .standardIcon(QStyle.SP_FileIcon)
                                .pixmap(32, 32)
                            )
                        layout.addWidget(icon_label)
                    label = QLabel(filename)
                    label.setTextInteractionFlags(Qt.TextSelectableByMouse)
                    layout.addWidget(label)
                    open_btn = QPushButton("Открыть")
                    open_btn.setToolTip("Открыть вложение")
                    open_btn.setFixedSize(60, 24)
                    open_btn.clicked.connect(
                        lambda _, path=file_path: self.text_edit._open_any_link(path)
                    )
                    layout.addWidget(open_btn)
                    del_btn = QPushButton("❌")
                    del_btn.setToolTip("Удалить вложение")
                    del_btn.setFixedSize(28, 24)
                    del_btn.clicked.connect(
                        lambda _, path=file_path: self.delete_attachment_from_panel(
                            path
                        )
                    )
                    layout.addWidget(del_btn)
                    item_widget.setSizePolicy(QSizePolicy.Maximum, QSizePolicy.Maximum)
                    if self.attachments_layout.count() > 0:
                        self.attachments_layout.addWidget(QLabel(" - "))
                    self.attachments_layout.addWidget(item_widget)
                if note_dir not in self.attachments_watcher.directories():
                    self.attachments_watcher.addPath(note_dir)
            self.attachments_scroll.setVisible(attachments_found)
        self.tags_label.setText(f"Теги: {', '.join(note.tags) if note.tags else 'нет'}")


    def _refresh_attachments(self, *_) -> None:
        if not self.current_note:
            return
        self._rebuild_attachments_panel(self.current_note)

    def delete_attachment_from_panel(self, file_path: str) -> None:
        reply = QMessageBox.question(
            self,
            "Удалить вложение",
            f"Удалить вложение '{os.path.basename(file_path)}'?\nФайл будет удалён безвозвратно.",
            QMessageBox.Yes | QMessageBox.No,
        )
        if reply == QMessageBox.Yes:
            b64_data = None
            if file_path.lower().endswith(tuple(IMAGE_EXTENSIONS)):
                try:
                    pixmap = QPixmap(file_path)
                    if not pixmap.isNull():
                        buffer = QBuffer()
                        buffer.open(QIODevice.WriteOnly)
                        pixmap.save(buffer, "PNG")
                        b64_data = base64.b64encode(buffer.data()).decode("utf-8")
                except Exception:
                    b64_data = None
            try:
                os.remove(file_path)
            except Exception as e:
                QMessageBox.warning(self, "Ошибка", f"Не удалось удалить файл:\n{e}")
            if self.current_note:
                file_path_abs = os.path.abspath(file_path)
                urls = {
                    QUrl.fromLocalFile(file_path_abs).toString(),
                    "file:///" + quote(file_path_abs.replace(os.sep, "/")),
                }
                html = self.text_edit.toHtml()
                for url in urls:
                    esc = re.escape(url)
                    html = re.sub(
                        rf"(<br>\s*)*(?:📄\s*)?<a[^>]+href=\"{esc}\"[^>]*>.*?</a>(<br>\s*)*",
                        "",
                        html,
                        flags=re.IGNORECASE,
                    )
                    html = re.sub(
                        rf"(<br>\s*)*<img[^>]+src=\"{esc}\"[^>]*>(<br>\s*)*",
                        "",
                        html,
                        flags=re.IGNORECASE,
                    )
                if b64_data:
                    esc_b64 = re.escape(b64_data)
                    html = re.sub(
                        rf"(<br>\s*)*<img[^>]+src=\"data:image/[^;]+;base64,{esc_b64}\"[^>]*>(<br>\s*)*",
                        "",
                        html,
                        flags=re.IGNORECASE,
                    )
                self.text_edit.blockSignals(True)
                self.text_edit.setHtml(html)
                self.text_edit.blockSignals(False)
                self.save_note_quiet(force=True)
                self._rebuild_attachments_panel(self.current_note)

    def _safe_restore_cursor(self, anchor_pos: int, cursor_pos: int) -> None:
        doc = self.text_edit.document()
        max_pos = max(0, doc.characterCount() - 1)
        a = max(0, min(anchor_pos, max_pos))
        p = max(0, min(cursor_pos, max_pos))
        cur = self.text_edit.textCursor()
        if a == p:
            cur.setPosition(p)
        else:
            cur.setPosition(a)
            cur.setPosition(p, QTextCursor.KeepAnchor)
        self.text_edit.setTextCursor(cur)

    def _atomic_json_dump(path: str, data: dict) -> None:
        tmp = path + ".tmp"
        with open(tmp, "w", encoding="utf-8") as f:
            json.dump(data, f, ensure_ascii=False, indent=4)
            f.flush()
            os.fsync(f.fileno())
        os.replace(tmp, path)

    def _update_editor_visibility(self):
        has = self.current_note is not None
        self.dock_editor.setVisible(has)
        self.dock_toolbar.setVisible(has)
        self.visibility_toolbar.setVisible(has)
        self.text_edit.setVisible(has)
        self.text_edit.setReadOnly(not has)
        self.tags_label.setVisible(has)
        if not has:
            self.tags_label.clear()
            self.add_field_btn.setEnabled(False)
            self.attachments_scroll.setVisible(False)
            self.password_manager_row.setVisible(False)
            self.rdp_1c8_row.setVisible(False)
            if hasattr(self, "action_toggle_pm"):
                self.action_toggle_pm.setEnabled(False)
            if hasattr(self, "action_toggle_rdp"):
                self.action_toggle_rdp.setEnabled(False)
                self.action_toggle_rdp.setVisible(False)
            for w in self.custom_fields_widgets:
                w["row"].setVisible(False)
                w["action"].setEnabled(False)
            return
        note = self.current_note
        pm_vis = bool(getattr(note, "password_manager_visible", False))
        rdp_vis = bool(getattr(note, "rdp_1c8_visible", False))
        rdp_removed = bool(getattr(note, "rdp_1c8_removed", False))
        if rdp_vis and not rdp_removed:
            self._ensure_rdp_row_in_layout()
        self.rdp_1c8_row.setVisible(False if rdp_removed else rdp_vis)
        self.password_manager_row.setVisible(pm_vis)
        if hasattr(self, "action_toggle_pm"):
            self.action_toggle_pm.setEnabled(True)
            self.action_toggle_pm.blockSignals(True)
            self.action_toggle_pm.setChecked(pm_vis)
            self._update_eye_action(
                self.action_toggle_pm, pm_vis, self.password_manager_label.text()
            )
            self.action_toggle_pm.blockSignals(False)
        if hasattr(self, "action_toggle_rdp"):
            self.action_toggle_rdp.setVisible(not rdp_removed)
            self.action_toggle_rdp.setEnabled(not rdp_removed)
            self.action_toggle_rdp.setEnabled(False)
            self.action_toggle_rdp.setVisible(False)
            self.action_toggle_rdp.blockSignals(True)
            self.action_toggle_rdp.setChecked(False if rdp_removed else rdp_vis)
            self._update_eye_action(
                self.action_toggle_rdp,
                (False if rdp_removed else rdp_vis),
                self.rdp_1c8_label.text(),
            )
            self.action_toggle_rdp.blockSignals(False)
        for w in self.custom_fields_widgets:
            w["action"].setEnabled(True)
            vis = bool(w["action"].isChecked())
            w["row"].setVisible(vis)
            self._update_eye_action(w["action"], vis, w["label_edit"].text())

    def attach_file_to_note(self) -> None:
        if not self.current_note:
            return
        file_path, _ = QFileDialog.getOpenFileName(
            self,
            "Прикрепить файл",
            "",
            ATTACH_FILE_FILTER,
        )
        if not file_path:
            return
        note_dir = os.path.join(
            NOTES_DIR,
            NotesApp.safe_folder_name(
                self.current_note.title,
                self.current_note.uuid,
                self.current_note.timestamp,
            ),
        )
        os.makedirs(note_dir, exist_ok=True)
        filename = os.path.basename(file_path)
        dest = os.path.join(note_dir, filename)
        try:
            shutil.copy(file_path, dest)
        except Exception as e:
            QMessageBox.critical(self, "Ошибка", f"Не удалось скопировать файл: {e}")
            return
        is_image = filename.lower().endswith(tuple(IMAGE_EXTENSIONS))
        cursor = self.text_edit.textCursor()
        self.text_edit.setTextCursor(cursor)
        if is_image:
            pixmap = QPixmap(dest)
            if not pixmap.isNull():
                buffer = QBuffer()
                buffer.open(QIODevice.WriteOnly)
                pixmap.save(buffer, "PNG")
                base64_data = base64.b64encode(buffer.data()).decode("utf-8")
                html_img = (
                    f'<img src="Data:image/png;base64,{base64_data}" width="200"><br>'
                )
                self.text_edit.insertHtml(html_img)
                self.record_state_for_undo()
            else:
                QMessageBox.warning(
                    self, "Ошибка", f"Не удалось открыть изображение: {filename}"
                )
        else:
            file_url = QUrl.fromLocalFile(os.path.abspath(dest)).toString()
            self.text_edit.insertHtml(f'📄 <a href="{file_url}">{filename}</a>')
        self.show_toast("Файл прикреплён к заметке.")
        self.save_note_quiet(force=True)
        self._rebuild_attachments_panel(self.current_note)

    def attach_file_to_note_external(self, file_path: str) -> None:
        if not self.current_note:
            return
        note_dir = os.path.join(
            NOTES_DIR,
            NotesApp.safe_folder_name(
                self.current_note.title,
                self.current_note.uuid,
                self.current_note.timestamp,
            ),
        )
        os.makedirs(note_dir, exist_ok=True)
        filename = os.path.basename(file_path)
        dest = os.path.join(note_dir, filename)
        try:
            shutil.copy(file_path, dest)
        except Exception as e:
            QMessageBox.critical(self, "Ошибка", f"Не удалось скопировать файл: {e}")
            return
        is_image = filename.lower().endswith(tuple(IMAGE_EXTENSIONS))
        cursor = self.text_edit.textCursor()
        self.text_edit.setTextCursor(cursor)
        if is_image:
            pixmap = QPixmap(dest)
            if not pixmap.isNull():
                buffer = QBuffer()
                buffer.open(QIODevice.WriteOnly)
                pixmap.save(buffer, "PNG")
                base64_data = base64.b64encode(buffer.data()).decode("utf-8")
                html_img = (
                    f'<img src="Data:image/png;base64,{base64_data}" width="200"><br>'
                )
                self.text_edit.insertHtml(html_img)
                self.record_state_for_undo()
            else:
                QMessageBox.warning(
                    self, "Ошибка", f"Не удалось открыть изображение: {filename}"
                )
        else:
            file_url = QUrl.fromLocalFile(os.path.abspath(dest)).toString()
            self.text_edit.insertHtml(f'📄 <a href="{file_url}">{filename}</a>')
        self.show_toast("Файл прикреплён к заметке.")
        self.save_note_quiet(force=True)
        self._rebuild_attachments_panel(self.current_note)

    def align_left(self) -> None:
        self.text_edit.setAlignment(Qt.AlignLeft)

    def align_center(self) -> None:
        self.text_edit.setAlignment(Qt.AlignCenter)

    def align_right(self) -> None:
        self.text_edit.setAlignment(Qt.AlignRight)

    def align_justify(self) -> None:
        self.text_edit.setAlignment(Qt.AlignJustify)

    def change_font(self, font: QFont) -> None:
        self.text_edit.setCurrentFont(font)

    def change_font_size(self, size: int) -> None:
        cursor = self.text_edit.textCursor()
        fmt = QTextCharFormat()
        fmt.setFontPointSize(size)
        if cursor.hasSelection():
            cursor.mergeCharFormat(fmt)
        else:
            self.text_edit.setCurrentCharFormat(fmt)

    def _current_font_size(self, cursor: QTextCursor) -> float:
        size = cursor.charFormat().fontPointSize()
        if not size:
            size = self.text_edit.fontPointSize() or 14
        return size

    def toggle_bold(self) -> None:
        cursor = self.text_edit.textCursor()
        if cursor.hasSelection():
            current_format = cursor.charFormat()
            fmt = QTextCharFormat()
            new_weight = (
                QFont.Normal
                if current_format.fontWeight() == QFont.Bold
                else QFont.Bold
            )
            fmt.setFontWeight(new_weight)
            cursor.mergeCharFormat(fmt)
        else:
            fmt = self.text_edit.currentCharFormat()
            new_weight = QFont.Normal if fmt.fontWeight() == QFont.Bold else QFont.Bold
            fmt.setFontWeight(new_weight)
            self.text_edit.setCurrentCharFormat(fmt)

    def toggle_italic(self) -> None:
        cursor = self.text_edit.textCursor()
        if cursor.hasSelection():
            current_format = cursor.charFormat()
            fmt = QTextCharFormat()
            fmt.setFontItalic(not current_format.fontItalic())
            cursor.mergeCharFormat(fmt)
        else:
            fmt = self.text_edit.currentCharFormat()
            fmt.setFontItalic(not fmt.fontItalic())
            self.text_edit.setCurrentCharFormat(fmt)

    def toggle_underline(self) -> None:
        cursor = self.text_edit.textCursor()
        if cursor.hasSelection():
            current_format = cursor.charFormat()
            fmt = QTextCharFormat()
            fmt.setFontUnderline(not current_format.fontUnderline())
            cursor.mergeCharFormat(fmt)
        else:
            fmt = self.text_edit.currentCharFormat()
            fmt.setFontUnderline(not fmt.fontUnderline())
            self.text_edit.setCurrentCharFormat(fmt)

    def _set_dropdown_values(self, dd_id: str, values: list[str]) -> None:
        info = self._get_dropdown_token_info(dd_id)
        if not info:
            return
        cleaned = []
        for v in values:
            v = (v or "").strip()
            if v and v not in cleaned:
                cleaned.append(v)
        info["values"] = cleaned
        self._update_dropdown_token(dd_id, info.get("value", ""))

    def _add_value_to_dropdown(self, dd_id: str) -> None:
        text, ok = QInputDialog.getText(self, "Добавить значение", "Текст:")
        if not ok:
            return
        val = (text or "").strip()
        if not val:
            return
        info = self._get_dropdown_token_info(dd_id)
        if not info:
            return
        values = list(info.get("values", []))
        if val not in values:
            values.append(val)
            self._set_dropdown_values(dd_id, values)
        self.show_dropdown_menu_for_token(dd_id)

    def _remove_value_from_dropdown(self, dd_id: str, value: str) -> None:
        info = self._get_dropdown_token_info(dd_id)
        if not info:
            return
        values = [v for v in info.get("values", []) if v != value]
        self._set_dropdown_values(dd_id, values)
        self.show_dropdown_menu_for_token(dd_id)

    def change_text_color(self) -> None:
        color = QColorDialog.getColor(Qt.black, self.text_edit, "Выберите цвет текста")
        if color.isValid():
            cursor = self.text_edit.textCursor()
            fmt = QTextCharFormat()
            fmt.setForeground(color)
            if cursor.hasSelection():
                cursor.mergeCharFormat(fmt)
            else:
                self.text_edit.setCurrentCharFormat(fmt)

    def change_background_color(self) -> None:
        color = QColorDialog.getColor(Qt.white, self.text_edit, "Выберите цвет фона")
        if color.isValid():
            cursor = self.text_edit.textCursor()
            fmt = QTextCharFormat()
            fmt.setBackground(color)
            if cursor.hasSelection():
                cursor.mergeCharFormat(fmt)
            else:
                self.text_edit.setCurrentCharFormat(fmt)

    def clear_formatting_complete(self) -> None:
        cursor = self.text_edit.textCursor()
        if cursor.hasSelection():
            selected_text = cursor.selectedText()
            fmt = QTextCharFormat()
            fmt.setFont(QFont("Times New Roman", 14))
            cursor.insertText(selected_text, fmt)
        else:
            fmt = QTextCharFormat()
            fmt.setFont(QFont("Times New Roman", 14))
            self.text_edit.setCurrentCharFormat(fmt)

    def apply_formatting(self, **kwargs) -> None:
        cursor = self.text_edit.textCursor()
        fmt = QTextCharFormat()
        for key, value in kwargs.items():
            if key == "bold":
                fmt.setFontWeight(QFont.Bold if value else QFont.Normal)
            elif key == "italic":
                fmt.setFontItalic(value)
            elif key == "underline":
                fmt.setFontUnderline(value)
            elif key == "color":
                fmt.setForeground(QColor(value))
            elif key == "background":
                fmt.setBackground(QColor(value))
            elif key == "font_size":
                fmt.setFontPointSize(value)
        if cursor.hasSelection():
            cursor.mergeCharFormat(fmt)
        else:
            self.text_edit.setCurrentCharFormat(fmt)

    def update_font_controls(self) -> None:
        cursor = self.text_edit.textCursor()
        fmt = cursor.charFormat()
        font = fmt.font()
        family = font.family() if font.family() else "Times New Roman"
        size = int(font.pointSize() if font.pointSize() > 0 else 14)
        self.font_combo.blockSignals(True)
        self.font_combo.setCurrentFont(QFont(family))
        self.font_combo.blockSignals(False)
        self.font_size_spin.blockSignals(True)
        self.font_size_spin.setValue(size)
        self.font_size_spin.blockSignals(False)

    def get_current_formatting(self):
        cursor = self.text_edit.textCursor()
        fmt = (
            cursor.charFormat()
            if cursor.hasSelection()
            else self.text_edit.currentCharFormat()
        )
        return {
            "bold": fmt.fontWeight() == QFont.Bold,
            "italic": fmt.fontItalic(),
            "underline": fmt.fontUnderline(),
            "color": fmt.foreground().color(),
            "background": fmt.background().color(),
            "font_size": fmt.fontPointSize(),
        }

    def set_font_family(self, font_family: str) -> None:
        cursor = self.text_edit.textCursor()
        fmt = QTextCharFormat()
        fmt.setFontFamily(font_family)
        if cursor.hasSelection():
            cursor.mergeCharFormat(fmt)
        else:
            self.text_edit.setCurrentCharFormat(fmt)

    def select_last_or_first_note(self) -> None:
        if not self.notes:
            return
        last_uuid = self.settings.value("lastNoteUuid", "")
        picked = next((n for n in self.notes if n.uuid == last_uuid), None) or self.notes[0]
        self.show_note_with_attachments(picked)
        items = self.notes_list.findItems(picked.title, Qt.MatchExactly)
        if items:
            self.notes_list.setCurrentItem(items[0])

    def toggle_strikethrough(self) -> None:
        cursor = self.text_edit.textCursor()
        if cursor.hasSelection():
            current_format = cursor.charFormat()
            fmt = QTextCharFormat()
            fmt.setFontStrikeOut(not current_format.fontStrikeOut())
            cursor.mergeCharFormat(fmt)
        else:
            fmt = self.text_edit.currentCharFormat()
            fmt.setFontStrikeOut(not fmt.fontStrikeOut())
            self.text_edit.setCurrentCharFormat(fmt)

    @staticmethod
    def _app_dir():
        return (
            os.path.dirname(sys.executable)
            if getattr(sys, "frozen", False)
            else os.path.abspath(os.path.dirname(__file__))
        )

    def open_readme(self):
        path = os.path.join(APPDIR, "README.md")
        if not os.path.exists(path):
            up = os.path.abspath(os.path.join(APPDIR, "..", "README.md"))
            if os.path.exists(up):
                path = up
        if not os.path.exists(path):
            QMessageBox.warning(self, "README", f"Файл не найден:\n{path}")
            return
        dlg = QDialog(self)
        dlg.setWindowTitle("README.md")
        dlg.resize(900, 700)
        v = QVBoxLayout(dlg)
        view = QTextBrowser(dlg)
        view.setOpenLinks(False)
        view.setOpenExternalLinks(False)
        view.setTextInteractionFlags(Qt.TextBrowserInteraction)
        v.addWidget(view)
        btns = QDialogButtonBox(QDialogButtonBox.Close)
        btns.addButton("Открыть в редакторе", QDialogButtonBox.ActionRole)
        v.addWidget(btns)
        with open(path, "r", encoding="utf-8") as f:
            md = f.read()
        try:
            view.setMarkdown(md)
        except AttributeError:
            view.setPlainText(md)
        base = QUrl.fromLocalFile(os.path.abspath(os.path.dirname(path)) + os.sep)
        view.document().setBaseUrl(base)
        self._ensure_heading_ids(view)
        if hasattr(self, "_linkify_urls_in_browser"):
            self._linkify_urls_in_browser(view)
            view.document().setBaseUrl(base)
            view.setOpenLinks(False)
            view.setOpenExternalLinks(False)

        def _on_anchor_clicked(url: QUrl):
            if (not url.scheme()) and url.fragment():
                frag = url.fragment()
                if not view.scrollToAnchor(frag):
                    view.scrollToAnchor(frag.lower())
                return
            if url.isRelative():
                QDesktopServices.openUrl(view.document().baseUrl().resolved(url))
                return
            QDesktopServices.openUrl(url)

        view.anchorClicked.connect(_on_anchor_clicked)

        def _open_external():
            QDesktopServices.openUrl(QUrl.fromLocalFile(os.path.abspath(path)))

        btns.rejected.connect(dlg.reject)
        for b in btns.buttons():
            if btns.buttonRole(b) == QDialogButtonBox.ActionRole:
                b.clicked.connect(_open_external)
                break
        fg = dlg.frameGeometry()
        fg.moveCenter(self.frameGeometry().center())
        dlg.move(fg.topLeft())
        dlg.exec()

    def _ensure_heading_ids(self, view: QTextBrowser):
        html = view.toHtml()

        def slugify(text: str) -> str:
            s = text.strip().lower()
            s = re.sub(r"[^0-9a-zа-яё \-]+", "", s, flags=re.IGNORECASE)
            s = re.sub(r"\s+", "-", s)
            s = re.sub(r"-{2,}", "-", s)
            return s.strip("-")

        def repl(m):
            tag, attrs, inner = m.group(1), m.group(2), m.group(3)
            if re.search(r"\b(id|name)\s*=", attrs, flags=re.I):
                return m.group(0)
            text_only = re.sub(r"<[^>]*>", "", inner)
            anchor = slugify(text_only)
            if not anchor:
                return m.group(0)
            return f'<{tag}{attrs} id="{anchor}" name="{anchor}">{inner}</{tag}>'

        html = re.sub(r"<(h[1-6])([^>]*)>(.*?)</\1>", repl, html, flags=re.I | re.S)
        view.setHtml(html)
        base = view.document().baseUrl()
        view.document().setBaseUrl(base)
        view.setOpenLinks(False)
        view.setOpenExternalLinks(False)

    def _linkify_urls_in_browser(self, view: QTextBrowser):
        html = view.toHtml()
        pattern = re.compile(r'(?<!href=")(?<!src=")((?:https?://|www\.)[^\s"<>]+)')

        def repl(m):
            url = m.group(1)
            href = url if url.startswith(("http://", "https://")) else f"http://{url}"
            return f'<a href="{href}">{url}</a>'

        html = pattern.sub(repl, html)
        view.setHtml(html)
        view.setOpenExternalLinks(True)
        view.setOpenLinks(False)

    def _replace_selection_with_text(
        self, cursor: QTextCursor, new_text: str
    ) -> None:
        selection_start = cursor.selectionStart()
        cursor.insertText(new_text)
        cursor.setPosition(selection_start)
        cursor.setPosition(selection_start + len(new_text), QTextCursor.KeepAnchor)
        self.text_edit.setTextCursor(cursor)

    def toggle_case(self) -> None:
        cursor = self.text_edit.textCursor()
        if cursor.hasSelection():
            text = cursor.selectedText()
            text = text.swapcase()
            self._replace_selection_with_text(cursor, text)
            self.record_state_for_undo()

    def _convert_layout_text(self, text: str) -> str:
        ru_count = sum(1 for ch in text if ch.lower() in RU_LAYOUT)
        en_count = sum(1 for ch in text if ch.lower() in EN_LAYOUT)
        if ru_count == 0 and en_count == 0:
            return text
        table = EN_TO_RU if en_count >= ru_count else RU_TO_EN
        return text.translate(table)

    def _swap_case_text(self, text: str) -> str:
        return text.swapcase()

    def _apply_text_transform(self, transform: Callable[[str], str], message: str) -> None:
        cursor = self.text_edit.textCursor()
        used_selection = False
        if cursor.hasSelection():
            text = cursor.selectedText()
            used_selection = True
        else:
            clipboard = QApplication.clipboard()
            text = clipboard.text()
        if not text:
            return
        converted = transform(text)
        if used_selection:
            self._replace_selection_with_text(cursor, converted)
            self.record_state_for_undo()
        else:
            clipboard.setText(converted)
            if self.isVisible():
                self.show_toast(message)
            else:
                try:
                    self.tray_icon.showMessage(
                        "Мои Заметки",
                        message,
                        QSystemTrayIcon.Information,
                        3000,
                    )
                except Exception:
                    pass

    def translate_layout_only(self) -> None:
        self._apply_text_transform(
            self._convert_layout_text, "Текст в буфере переведен (раскладка)"
        )

    def translate_case_only(self) -> None:
        self._apply_text_transform(
            self._swap_case_text, "Регистр текста в буфере изменен"
        )

    def _register_global_hotkeys(self) -> None:
        if keyboard is None:
            return
        self._unregister_global_hotkeys()
        try:
            hk_layout = keyboard.add_hotkey(
                "ctrl+shift+f12", self.translate_layout_only, suppress=False
            )
            hk_case = keyboard.add_hotkey(
                "ctrl+alt+f12", self.translate_case_only, suppress=False
            )
            self._global_hotkeys = [hk_layout, hk_case]
        except Exception:
            self._global_hotkeys = []

    def _unregister_global_hotkeys(self) -> None:
        if keyboard is None:
            return
        for hk in getattr(self, "_global_hotkeys", []):
            try:
                keyboard.remove_hotkey(hk)
            except Exception:
                pass
        self._global_hotkeys = []

    def _toggle_clipboard_layout_watch(self, enabled: bool) -> None:
        self._clipboard_watch_enabled = bool(enabled)
        self.settings.setValue("layout/clipboard_auto_convert", self._clipboard_watch_enabled)
        if not self._clipboard_watch_enabled:
            self._last_clipboard_text = ""

    def _on_clipboard_changed(self) -> None:
        if not self._clipboard_watch_enabled or self._clipboard_ignore:
            return
        clipboard = QApplication.clipboard()
        text = clipboard.text()
        if not text:
            self._last_clipboard_text = ""
            return
        if text == self._last_clipboard_text:
            return
        converted = self._convert_layout_text(text)
        self._last_clipboard_text = text
        if converted == text:
            return
        self._clipboard_ignore = True
        try:
            clipboard.setText(converted)
            self._last_clipboard_text = converted
            if self.isVisible():
                self.show_toast("Текст в буфере переведен (раскладка)")
            elif getattr(self, "tray_icon", None) is not None:
                try:
                    self.tray_icon.showMessage(
                        "Мои Заметки",
                        "Текст в буфере переведен (раскладка)",
                        QSystemTrayIcon.Information,
                        3000,
                    )
                except Exception:
                    pass
        finally:
            self._clipboard_ignore = False

    def apply_heading(self, level: int) -> None:
        cursor = self.text_edit.textCursor()
        if not cursor.hasSelection():
            cursor.select(QTextCursor.WordUnderCursor)
        fmt = QTextCharFormat()
        size_map = {1: 24, 2: 18, 3: 14}
        fmt.setFontPointSize(size_map.get(level, 12))
        fmt.setFontWeight(QFont.Bold)
        cursor.mergeCharFormat(fmt)

    def insert_list(self, style: QTextListFormat.Style) -> None:
        cursor = self.text_edit.textCursor()
        if cursor.hasSelection():
            selection_start = cursor.selectionStart()
            selection_end = cursor.selectionEnd()
            cursor.setPosition(selection_start)
            cursor.setPosition(selection_end, QTextCursor.KeepAnchor)
            block_format = QTextListFormat()
            block_format.setStyle(style)
            cursor.createList(block_format)
        else:
            if style == QTextListFormat.ListDisc:
                cursor.insertText("• ")
            elif style == QTextListFormat.ListDecimal:
                cursor.insertText("1. ")

    def insert_bullet_list(self) -> None:
        self.insert_list(QTextListFormat.ListDisc)

    def insert_numbered_list(self) -> None:
        self.insert_list(QTextListFormat.ListDecimal)

    def insert_checkbox(self) -> None:
        cursor = self.text_edit.textCursor()
        if cursor.hasSelection():
            selection_start = cursor.selectionStart()
            selection_end = cursor.selectionEnd()

            doc = self.text_edit.document()
            start_block = doc.findBlock(selection_start)
            end_block = doc.findBlock(selection_end)

            block = start_block
            while True:
                block_cursor = QTextCursor(block)
                block_cursor.movePosition(QTextCursor.StartOfBlock)
                block_cursor.movePosition(QTextCursor.Right, QTextCursor.MoveAnchor, 0)
                text = block.text()
                if text.strip() and not (
                    text.startswith("☐ ") or text.startswith("☑ ")
                ):
                    block_cursor.insertText("☐ ")
                if block == end_block:
                    break
                block = block.next()
        else:
            cursor.movePosition(QTextCursor.StartOfLine)
            text = cursor.block().text()
            if not (text.startswith("☐ ") or text.startswith("☑ ")):
                cursor.insertText("☐ ")
                cursor.movePosition(QTextCursor.Right, QTextCursor.MoveAnchor, 2)
                self.text_edit.setTextCursor(cursor)
                self.text_edit.setFocus()

    def insert_horizontal_line(self) -> None:
        cursor = self.text_edit.textCursor()
        cursor.beginEditBlock()
        cursor.insertHtml("<hr style='border:1px solid #888; margin:0;'>")
        cursor.endEditBlock()
        self.text_edit.setTextCursor(cursor)

    def insert_image_html(
        self, html_img: str, replace_current: bool = False, orig_path: str | None = None
    ) -> None:
        doc = self.text_edit.document()
        cursor = QTextCursor(doc)
        replaced = False
        while not cursor.isNull() and not cursor.atEnd():
            if cursor.charFormat().isImageFormat():
                img_format = cursor.charFormat().toImageFormat()
                if orig_path and (
                    img_format.name() == orig_path
                    or (
                        img_format.name().startswith("Data:image")
                        and orig_path.startswith("Data:image")
                    )
                ):
                    cursor.movePosition(QTextCursor.Right, QTextCursor.KeepAnchor, 1)
                    cursor.insertHtml(html_img)
                    replaced = True
                    break
            cursor.movePosition(QTextCursor.NextCharacter)
        if not replaced:
            insert_cursor = self.text_edit.textCursor()
            insert_cursor.movePosition(QTextCursor.End)
            self.text_edit.setTextCursor(insert_cursor)
            self.text_edit.insertHtml(html_img)
            self.record_state_for_undo()

    def insert_image_into_note(self, image_path: str) -> None:
        pixmap = QPixmap(image_path)
        if not pixmap.isNull():
            buffer = QBuffer()
            buffer.open(QIODevice.OpenModeFlag.WriteOnly)
            pixmap.save(buffer, "PNG")
            base64_data = base64.b64encode(buffer.data()).decode("utf-8")
            html_img = (
                f'<img src="Data:image/png;base64,{base64_data}" width="300"><br>'
            )
            cursor = self.text_edit.textCursor()
            cursor.movePosition(QTextCursor.End)
            self.text_edit.setTextCursor(cursor)
            self.text_edit.insertHtml(html_img)
            self.record_state_for_undo()

    def _insert_image_at_cursor(self, image_path: str, width: int = 400) -> None:
        pixmap = QPixmap(image_path)
        if pixmap.isNull():
            QMessageBox.warning(self, "Ошибка", "Не удалось загрузить изображение.")
            return
        buffer = QBuffer()
        buffer.open(QIODevice.OpenModeFlag.WriteOnly)
        pixmap.save(buffer, "PNG")
        base64_data = base64.b64encode(buffer.data()).decode("utf-8")
        html_img = (
            f'<img src="Data:image/png;base64,{base64_data}" width="{width}"><br>'
        )
        cursor = self.text_edit.textCursor()
        self.text_edit.setTextCursor(cursor)
        self.text_edit.insertHtml(html_img)
        self.record_state_for_undo()

    def open_drawing_dialog(self) -> None:
        if not self.current_note:
            QMessageBox.warning(
                self, "Нет заметки", "Пожалуйста, выбери или создай заметку."
            )
            return

        dlg = DrawingDialog(self)
        if dlg.exec() == QDialog.DialogCode.Accepted:
            img = dlg.get_image()

            note_dir = os.path.join(
                NOTES_DIR,
                NotesApp.safe_folder_name(
                    self.current_note.title,
                    self.current_note.uuid,
                    self.current_note.timestamp,
                ),
            )
            os.makedirs(note_dir, exist_ok=True)

            default_base = (
                f"рисунок_{QDateTime.currentDateTime().toString('yyyyMMdd_HHmmss')}"
            )
            name, ok = QInputDialog.getText(
                self,
                "Имя рисунка",
                "Имя файла (без расширения):",
                QLineEdit.EchoMode.Normal,
                default_base,
            )
            if not ok:
                return

            base = (name or "").strip()
            if not base:
                base = default_base

            base = re.sub(r"[^a-zA-Zа-яА-Я0-9 _\-\.\(\)]", "_", base).strip(" ._")
            if not base:
                base = default_base

            filename = base if base.lower().endswith(".png") else base + ".png"
            filepath = os.path.join(note_dir, filename)

            if os.path.exists(filepath):
                stem, ext = os.path.splitext(filename)
                i = 1
                while os.path.exists(os.path.join(note_dir, f"{stem}_{i}{ext}")):
                    i += 1
                filepath = os.path.join(note_dir, f"{stem}_{i}{ext}")

            try:
                img.save(filepath, "PNG")
            except Exception as e:
                QMessageBox.critical(
                    self, "Ошибка", f"Не удалось сохранить рисунок:\n{e}"
                )
                return

            self._insert_image_at_cursor(filepath, width=400)
            self.save_note()
            self._rebuild_attachments_panel(self.current_note)

    def insert_image_from_clipboard(self, image: QImage) -> None:
        if not self.current_note:
            QMessageBox.warning(
                self, "Нет заметки", "Пожалуйста, выберите или создайте заметку."
            )
            return
        note_dir = os.path.join(
            NOTES_DIR,
            NotesApp.safe_folder_name(
                self.current_note.title,
                self.current_note.uuid,
                self.current_note.timestamp,
            ),
        )
        os.makedirs(note_dir, exist_ok=True)
        filename = f"clipboard_{uuid.uuid4().hex}.png"
        filepath = os.path.join(note_dir, filename)
        try:
            image.save(filepath)
        except Exception as e:
            QMessageBox.critical(
                self, "Ошибка", f"Не удалось сохранить изображение: {e}"
            )
            return
        pixmap = QPixmap(filepath)
        if not pixmap.isNull():
            buffer = QBuffer()
            buffer.open(QIODevice.WriteOnly)
            pixmap.save(buffer, "PNG")
            base64_data = base64.b64encode(buffer.data()).decode("utf-8")
            html_img = (
                f'<img src="Data:image/png;base64,{base64_data}" width="200"><br>'
            )
            cursor = self.text_edit.textCursor()
            cursor.movePosition(QTextCursor.End)
            self.text_edit.setTextCursor(cursor)
            self.text_edit.insertHtml(html_img)
            self.record_state_for_undo()

        self.save_note()

    def insert_audio_link(self, filepath: str) -> None:
        filename = os.path.basename(filepath)
        file_url = QUrl.fromLocalFile(os.path.abspath(filepath)).toString()
        self.text_edit.insertHtml(f'📄 <a href="{file_url}">{filename}</a>')
        self.save_note()

    def toggle_audio_recording(self) -> None:
        if self.audio_thread and self.audio_thread.isRunning():
            self.audio_thread.stop()
            self.audio_thread.wait()
            self.audio_thread = None
            self.audio_button.setText("🎤")
        else:
            if not self.current_note:
                QMessageBox.warning(
                    self, "Нет заметки", "Пожалуйста, выбери или создай заметку."
                )
                return
            note_dir = os.path.join(
                NOTES_DIR,
                NotesApp.safe_folder_name(
                    self.current_note.title,
                    self.current_note.uuid,
                    self.current_note.timestamp,
                ),
            )
            os.makedirs(note_dir, exist_ok=True)
            filename = (
                f"audio_{QDateTime.currentDateTime().toString('yyyyMMdd_HHmmss')}.wav"
            )
            full_path = os.path.join(note_dir, filename)

            self.audio_thread = AudioRecorderThread(full_path)
            self.audio_thread.recording_finished.connect(self.insert_audio_link)
            self.audio_thread.start()
            self.audio_button.setText("⏹")

    def apply_sorting(self) -> None:
        index = self.sort_combo.currentIndex()
        reverse = self.sort_order_combo.currentIndex() == 1

        if index == 0:
            self.notes.sort(key=lambda note: note.title.lower(), reverse=reverse)
            self.refresh_notes_list()
        elif index == 1:
            self.notes.sort(key=lambda note: note.timestamp, reverse=reverse)
            self.refresh_notes_list()
        elif index == 2:
            self.show_favorites_only()

    def filter_notes(self):
        if self.current_note:
            self.current_note.content = self.text_edit.toHtml()
        query = self.search_field.text().strip().lower()
        if not query:
            return self.notes
        filtered = []
        for note in self.notes:
            if self.current_note and note.uuid == self.current_note.uuid:
                content_plain = self.text_edit.toPlainText()
            else:
                doc = QTextDocument()
                doc.setHtml(note.content)
                content_plain = doc.toPlainText()
            reminder_raw = str(note.reminder or "").lower()
            reminder_human = ""
            if note.reminder:
                try:
                    dt = QDateTime.fromString(note.reminder, Qt.ISODate)
                    if dt.isValid():
                        reminder_human = dt.toString("dd.MM.yyyy HH:mm")
                except Exception:
                    reminder_human = ""
            if (
                query in note.title.lower()
                or query in content_plain.lower()
                or query in reminder_raw
                or (reminder_human and query in reminder_human.lower())
            ):
                filtered.append(note)
        return filtered

    def refresh_notes_list(self) -> None:
        self.notes_list.clear()
        filtered_notes = self.filter_notes()
        pinned_notes = [note for note in filtered_notes if note.pinned]
        unpinned_notes = [note for note in filtered_notes if not note.pinned]
        fav_color = self.get_contrast_favorite_color()
        current_uuid = self.current_note.uuid if self.current_note else None

        def _add_note(note):
            title = note.title
            timestamp = QDateTime.fromString(note.timestamp, Qt.ISODate)
            date_str = timestamp.toString("dd.MM.yyyy")
            reminder_symbol = " 🔔" if note.reminder else ""
            item_text = f"{title} — {date_str}{reminder_symbol}"
            item = QListWidgetItem(item_text)
            item.setData(Qt.UserRole, note)
            item.setFont(QFont("Segoe UI Emoji", 10))
            if note.favorite:
                item.setForeground(fav_color)
            self.notes_list.addItem(item)
            return item

        for note in pinned_notes:
            _add_note(note)

        if pinned_notes and unpinned_notes:
            sep_item = QListWidgetItem()
            sep_item.setFlags(Qt.NoItemFlags)
            line = QFrame()
            line.setFrameShape(QFrame.HLine)
            line.setFrameShadow(QFrame.Sunken)
            sep_item.setSizeHint(line.sizeHint())
            self.notes_list.addItem(sep_item)
            self.notes_list.setItemWidget(sep_item, line)

        for note in unpinned_notes:
            _add_note(note)

        if current_uuid:
            for i in range(self.notes_list.count()):
                it = self.notes_list.item(i)
                n = it.data(Qt.UserRole)
                if n and n.uuid == current_uuid:
                    self.notes_list.setCurrentItem(it)
                    n = it.data(Qt.UserRole)
                    if n:
                        self.select_note(n)
                    break

    def get_contrast_favorite_color(self) -> QColor:
        bg = self.notes_list.palette().color(self.notes_list.backgroundRole())
        avg = (bg.red() + bg.green() + bg.blue()) / 3
        if avg > 200:
            return QColor(180, 0, 0)
        elif avg > 127:
            return QColor(220, 25, 25)
        else:
            return QColor(255, 60, 60)

    def _show_refresh_busy(self, busy: bool) -> None:
        if busy:
            if not hasattr(self, "_busy_box") or self._busy_box is None:
                self._busy_box = QWidget(self)
                self._busy_box.setObjectName("BusyBox")
                lay = QVBoxLayout(self._busy_box)
                lay.setContentsMargins(0, 0, 0, 0)
                lay.setSpacing(0)
                self._busy_label = QLabel("0%", self._busy_box)
                self._busy_label.setAlignment(Qt.AlignHCenter | Qt.AlignBottom)
                self._busy_label.setStyleSheet(
                    "font-size:10px; font-weight:600; margin-bottom:1px;"
                )
                self._busy_progress = QProgressBar(self._busy_box)
                self._busy_progress.setRange(0, 100)
                self._busy_progress.setValue(0)
                self._busy_progress.setFixedHeight(8)
                self._busy_progress.setTextVisible(False)
                lay.addWidget(self._busy_label)
                lay.addWidget(self._busy_progress)
                self.statusBar().addPermanentWidget(self._busy_box, 1)
            if not hasattr(self, "_busy_timer") or self._busy_timer is None:
                self._busy_timer = QTimer(self)
                self._busy_timer.setInterval(60)
                self._busy_timer.timeout.connect(self._tick_busy_progress)
            self._busy_target = 95
            self._busy_step = 2
            self._busy_timer.start()
            self.statusBar().showMessage("Обновляю…")
            QApplication.setOverrideCursor(Qt.BusyCursor)
            if getattr(self, "refresh_button", None):
                self.refresh_button.setEnabled(False)

        else:
            try:
                if getattr(self, "_busy_timer", None):
                    self._busy_timer.stop()
            except Exception:
                pass
            if getattr(self, "_busy_progress", None):

                def _finish():
                    if not getattr(self, "_busy_progress", None):
                        return
                    cur = self._busy_progress.value()
                    if cur < 100:
                        new_v = min(100, cur + 4)
                        self._busy_progress.setValue(new_v)
                        if getattr(self, "_busy_label", None):
                            self._busy_label.setText(f"{new_v}%")
                        QTimer.singleShot(30, _finish)
                    else:
                        if getattr(self, "_busy_label", None):
                            self._busy_label.setText("100%")
                        QTimer.singleShot(300, self._remove_busy_progress)

                _finish()
            else:
                self._remove_busy_progress()

    def _tick_busy_progress(self):
        if not getattr(self, "_busy_progress", None):
            return
        v = self._busy_progress.value()
        if v < getattr(self, "_busy_target", 95):
            new_v = min(self._busy_target, v + self._busy_step)
            self._busy_progress.setValue(new_v)
            if getattr(self, "_busy_label", None):
                self._busy_label.setText(f"{new_v}%")

    def _remove_busy_progress(self):
        self.statusBar().clearMessage()
        QApplication.restoreOverrideCursor()
        if getattr(self, "refresh_button", None):
            self.refresh_button.setEnabled(True)
        if getattr(self, "_busy_box", None):
            try:
                self.statusBar().removeWidget(self._busy_box)
                self._busy_box.deleteLater()
            except Exception:
                pass
            self._busy_box = None
        if getattr(self, "_busy_progress", None):
            try:
                self._busy_progress.deleteLater()
            except Exception:
                pass
            self._busy_progress = None
        if getattr(self, "_busy_label", None):
            try:
                self._busy_label.deleteLater()
            except Exception:
                pass
            self._busy_label = None
        if getattr(self, "_busy_timer", None):
            try:
                self._busy_timer.stop()
            except Exception:
                pass
            self._busy_timer = None

    def on_refresh_clicked(self) -> None:
        self._show_refresh_busy(True)
        try:
            self.show_toast(
                "Обновляю…",
                timeout_ms=1600,
                anchor_widget=getattr(self, "refresh_button", None),
            )
        except Exception:
            pass
        QApplication.processEvents()

        try:
            if hasattr(self, "refresh_all"):
                self.refresh_all()
            else:
                self.refresh_notes_list()
                if getattr(self, "current_note", None):
                    self.show_note_with_attachments(self.current_note)
        except Exception as e:
            QMessageBox.warning(self, "Обновление", f"Не удалось обновить: {e}")
        finally:
            self._show_refresh_busy(False)
            try:
                self.show_toast(
                    "Готово",
                    timeout_ms=1400,
                    anchor_widget=getattr(self, "refresh_button", None),
                )
            except Exception:
                pass

    def sort_notes_by_title(self) -> None:
        self.notes.sort(key=lambda note: note.title.lower())
        self.refresh_notes_list()

    def sort_notes_by_date(self) -> None:
        self.notes.sort(key=lambda note: note.timestamp, reverse=True)
        self.refresh_notes_list()

    def toggle_favorite(self):
        if self.current_note:
            self.current_note.favorite = not self.current_note.favorite
            self.refresh_notes_list()

    def show_favorites_only(self) -> None:
        self.notes_list.clear()
        for note in self.notes:
            if note.favorite:
                timestamp = QDateTime.fromString(note.timestamp, Qt.ISODate)
                date_str = timestamp.toString("dd.MM.yyyy")
                reminder_symbol = " 🔔" if note.reminder else ""
                item_text = f"{note.title} — {date_str}{reminder_symbol}"
                item = QListWidgetItem(item_text)
                item.setData(Qt.UserRole, note)
                item.setFont(QFont("Segoe UI Emoji", 10))
                item.setForeground(QColor("gold"))
                self.notes_list.addItem(item)

    def add_tag_to_note(self) -> None:
        if not self.current_note:
            return
        dialog = QDialog(self)
        dialog.setWindowTitle("Добавить тег")
        layout = QVBoxLayout(dialog)
        combo = QComboBox(dialog)
        combo.setEditable(True)
        all_tags = self.get_all_tags()
        combo.addItems(all_tags)
        layout.addWidget(combo)
        buttons = QDialogButtonBox(QDialogButtonBox.Ok | QDialogButtonBox.Cancel)
        layout.addWidget(buttons)
        dialog.setLayout(layout)
        buttons.accepted.connect(dialog.accept)
        buttons.rejected.connect(dialog.reject)
        if dialog.exec() == QDialog.Accepted:
            tag_line = combo.currentText()
            tags = [t.strip() for t in tag_line.split(",") if t.strip()]
            added = []
            for tag in tags:
                if tag and tag not in self.current_note.tags:
                    self.current_note.tags.append(tag)
                    added.append(tag)
            if added:
                self.update_tag_filter_items()
                self.show_toast("Теги добавлены", "Добавлены теги: " + ", ".join(added))
                self.tags_label.setText(f"Теги: {', '.join(self.current_note.tags)}")
            else:
                self.show_toast(self, "Нет новых тегов", "Все введённые теги уже есть у заметки.")

    def manage_tags_dialog(self) -> None:
        dialog = QDialog(self)
        dialog.setWindowFlag(Qt.WindowStaysOnTopHint, True)
        dialog.setWindowTitle("Управление тегами")
        layout = QVBoxLayout(dialog)
        all_tags = sorted(self.get_all_tags())
        combo = QComboBox(dialog)
        combo.addItems(all_tags)
        layout.addWidget(combo)
        rename_btn = QPushButton("Переименовать тег")
        delete_btn = QPushButton("Удалить тег")
        layout.addWidget(rename_btn)
        layout.addWidget(delete_btn)

        def rename_tag():
            old_tag = combo.currentText()
            new_tag, ok = QInputDialog.getText(
                dialog, "Переименовать тег", f"Новый тег для '{old_tag}':"
            )
            if ok and new_tag and new_tag != old_tag:
                for note in self.notes:
                    note.tags = [new_tag if t == old_tag else t for t in note.tags]
                self.save_all_notes_to_disk()
                self.update_tag_filter_items()
                self.refresh_notes_list()
                combo.clear()
                combo.addItems(sorted(self.get_all_tags()))
                if self.current_note:
                    self.tags_label.setText(
                        f"Теги: {', '.join(self.current_note.tags) if self.current_note.tags else 'нет'}"
                    )

        def delete_tag():
            tag_to_delete = combo.currentText()
            if not tag_to_delete:
                return
            reply = QMessageBox.question(
                dialog,
                "Удалить тег?",
                f"Удалить тег '{tag_to_delete}' из всех заметок?",
                QMessageBox.Yes | QMessageBox.No,
            )
            if reply == QMessageBox.Yes:
                for note in self.notes:
                    if tag_to_delete in note.tags:
                        note.tags.remove(tag_to_delete)
                self.save_all_notes_to_disk()
                self.update_tag_filter_items()
                self.refresh_notes_list()
                combo.clear()
                combo.addItems(sorted(self.get_all_tags()))
                if self.current_note:
                    self.tags_label.setText(
                        f"Теги: {', '.join(self.current_note.tags) if self.current_note.tags else 'нет'}"
                    )

        rename_btn.clicked.connect(rename_tag)
        delete_btn.clicked.connect(delete_tag)
        buttons = QDialogButtonBox(QDialogButtonBox.Close)
        layout.addWidget(buttons)
        buttons.rejected.connect(dialog.reject)
        dialog.setLayout(layout)
        dialog.exec()

    def get_all_tags(self) -> list[str]:
        tags: set[str] = set()
        for note in self.notes:
            tags.update(note.tags)
        return sorted(tags)

    def show_all_notes(self) -> None:
        self.refresh_notes_list()

    def show_notes_by_tag(self, tag: str) -> None:
        self.notes_list.clear()
        for note in self.notes:
            if tag in note.tags:
                timestamp = QDateTime.fromString(note.timestamp, Qt.ISODate)
                date_str = timestamp.toString("dd.MM.yyyy")
                reminder_symbol = " 🔔" if note.reminder else ""
                item_text = f"{note.title} — {date_str}{reminder_symbol}"
                item = QListWidgetItem(item_text)
                item.setData(Qt.ItemDataRole.UserRole, note)
                item.setFont(QFont("Segoe UI Emoji", 10))
                if note.favorite:
                    item.setForeground(QColor("gold"))
                self.notes_list.addItem(item)

    def apply_tag_filter(self) -> None:
        selected_tag = self.tag_filter.currentText()
        if selected_tag == "Все теги" or not selected_tag:
            self.show_all_notes()
        else:
            self.show_notes_by_tag(selected_tag)

    def _rebuild_attachments_panel(self, note: Note | None) -> None:
        if hasattr(self, "attachments_layout"):
            for i in reversed(range(self.attachments_layout.count())):
                w = self.attachments_layout.itemAt(i).widget()
                if w:
                    w.setParent(None)
                    w.deleteLater()
        if not note:
            self.attachments_scroll.setVisible(False)
            return
        note_dir = os.path.join(
            NOTES_DIR, NotesApp.safe_folder_name(note.title, note.uuid, note.timestamp)
        )
        attachments_found = False
        if os.path.isdir(note_dir):
            ignored_files = {"note.json", ".DS_Store", "Thumbs.db"}
            ignored_prefixes = ("~$", ".~")
            ignored_suffixes = ("~", ".tmp", ".temp")
            for filename in os.listdir(note_dir):
                if (
                    filename in ignored_files
                    or filename.startswith(ignored_prefixes)
                    or filename.endswith(ignored_suffixes)
                    or (filename.startswith("backup(") and filename.endswith(".json"))
                ):
                    continue
                attachments_found = True
                file_path = os.path.join(note_dir, filename)
                item_widget = QWidget()
                layout = QHBoxLayout(item_widget)
                icon_label = QLabel()
                if filename.lower().endswith(tuple(IMAGE_EXTENSIONS)):
                    pixmap = QPixmap(file_path)
                    icon_label.setPixmap(
                        pixmap.scaled(
                            48, 48, Qt.KeepAspectRatio, Qt.SmoothTransformation
                        )
                    )
                else:
                    if os.path.exists(FILE_ICON_PATH):
                        icon_label.setPixmap(
                            QPixmap(FILE_ICON_PATH).scaled(
                                48, 48, Qt.KeepAspectRatio, Qt.SmoothTransformation
                            )
                        )
                    else:
                        icon_label.setPixmap(
                            self.style().standardIcon(QStyle.SP_FileIcon).pixmap(32, 32)
                        )
                layout.addWidget(icon_label)
                label = QLabel(filename)
                label.setTextInteractionFlags(Qt.TextSelectableByMouse)
                layout.addWidget(label)
                open_btn = QPushButton("Открыть")
                open_btn.setToolTip("Открыть вложение")
                open_btn.setFixedSize(60, 24)
                open_btn.clicked.connect(
                    lambda _, path=file_path: self.text_edit._open_any_link(path)
                )
                layout.addWidget(open_btn)
                del_btn = QPushButton("❌")
                del_btn.setToolTip("Удалить вложение")
                del_btn.setFixedSize(28, 24)
                del_btn.clicked.connect(
                    lambda _, path=file_path: self.delete_attachment_from_panel(path)
                )
                layout.addWidget(del_btn)
                self.attachments_layout.addWidget(item_widget)
        self.attachments_scroll.setVisible(attachments_found)

    def insert_link(self) -> None:
        cursor = self.text_edit.textCursor()
        url, ok = QInputDialog.getText(self, "Вставить ссылку", "URL:")
        if not (ok and url):
            return
        url = url.strip()
        href = url
        if os.path.isabs(url) or os.path.exists(url):
            href = QUrl.fromLocalFile(os.path.abspath(url)).toString()
        else:
            if not re.match(r"^[a-zA-Z][a-zA-Z0-9+\-.]*://", url):
                href = "https://" + url
        fmt = QTextCharFormat()
        fmt.setAnchor(True)
        fmt.setAnchorHref(href)
        fmt.setForeground(Qt.blue)
        fmt.setFontUnderline(True)
        fmt.setFont(QFont("Times New Roman", 14))

        if cursor.hasSelection() and cursor.selectedText().strip():
            cursor.mergeCharFormat(fmt)
        else:
            fmt.setFont(QFont("Times New Roman", 14))
            cursor.insertText(url, fmt)

        def _select_entire_anchor(self, cursor: QTextCursor) -> QTextCursor | None:
            fmt = cursor.charFormat()
            if not fmt.isAnchor():
                return None
            href = fmt.anchorHref()
            c = QTextCursor(cursor)
            max_steps = cursor.document().characterCount() + 5
            steps = 0
            while steps < max_steps:
                steps += 1
                left = QTextCursor(c)
                moved = left.movePosition(QTextCursor.Left, QTextCursor.KeepAnchor, 1)
                if (not moved) or left.atStart():
                    break
                f = left.charFormat()
                if not (f.isAnchor() and f.anchorHref() == href):
                    break
                if left.position() == c.position() and left.anchor() == c.anchor():
                    break
                c = left
            while steps < max_steps:
                steps += 1
                right = QTextCursor(c)
                moved = right.movePosition(QTextCursor.Right, QTextCursor.KeepAnchor, 1)
                if (not moved) or right.atEnd():
                    break
                f = right.charFormat()
                if not (f.isAnchor() and f.anchorHref() == href):
                    break
                if right.position() == c.position() and right.anchor() == c.anchor():
                    break
                c = right
            return c

    def edit_link(self) -> None:
        cursor = self.text_edit.textCursor()
        if not cursor.hasSelection():
            cursor.select(QTextCursor.WordUnderCursor)
        if not cursor.charFormat().isAnchor():
            QMessageBox.information(self, "Ссылка", "Курсор не находится на ссылке.")
            return
        full = self._select_entire_anchor(cursor)
        if not full:
            return
        fmt = full.charFormat()
        old_url = fmt.anchorHref()
        new_url, ok = QInputDialog.getText(
            self, "Изменить ссылку", "URL:", text=old_url
        )
        if not (ok and new_url):
            return
        visible = full.selectedText()
        if visible.strip() == old_url.strip() or re.match(
            r"^\s*(https?://|www\.)", visible, flags=re.I
        ):
            new_visible = new_url
        else:
            new_visible = visible
        fmt.setAnchor(True)
        fmt.setAnchorHref(new_url)
        full.insertText(new_visible, fmt)

    def remove_link(self) -> None:
        cursor = self.text_edit.textCursor()
        if not cursor.hasSelection():
            cursor.select(QTextCursor.WordUnderCursor)
        fmt = QTextCharFormat(cursor.charFormat())
        fmt.setAnchor(False)
        fmt.setAnchorHref("")
        fmt.setForeground(self.text_edit.palette().text())
        fmt.setFontUnderline(False)
        cursor.mergeCharFormat(fmt)

    def insert_table(self) -> None:
        if not self.current_note:
            QMessageBox.warning(
                self, "Нет заметки", "Сначала выбери или создай заметку."
            )
            return
        cursor = self.text_edit.textCursor()
        has_sel = cursor.hasSelection()
        default_rows = 1 if has_sel else 2
        default_cols = 1 if has_sel else 2
        rows, ok1 = QInputDialog.getInt(
            self, "Вставить таблицу", "Количество строк:", default_rows, 1, 100
        )
        cols, ok2 = QInputDialog.getInt(
            self, "Вставить таблицу", "Количество столбцов:", default_cols, 1, 100
        )
        if not (ok1 and ok2):
            return
        is_dark = self.settings.value("theme", "dark") == "dark"
        border_color = "white" if is_dark else "black"
        selected_html = None
        if has_sel:
            sel_text = cursor.selectedText()
            selected_html = html_lib.escape(sel_text).replace("\u2029", "<br/>")
        rows_html = []
        for r in range(rows):
            cells_html = []
            for c in range(cols):
                if r == 0 and c == 0 and selected_html is not None:
                    cell_content = selected_html
                else:
                    cell_content = "&nbsp;"
                cells_html.append(
                    f"<td style='min-width:1em; padding:3px; border:1px solid {border_color};'>{cell_content}</td>"
                )
            rows_html.append("<tr>" + "".join(cells_html) + "</tr>")
        table_html = (
            f"<table border='1' cellspacing='0' cellpadding='3' "
            f"style='border-collapse:collapse; border:1px solid {border_color};'>"
            + "".join(rows_html)
            + "</table>"
        )
        cursor.beginEditBlock()
        if has_sel:
            cursor.removeSelectedText()
        cursor.insertHtml(table_html)
        cursor.endEditBlock()

    def insert_dropdown(self) -> None:
        if not self.current_note:
            QMessageBox.warning(
                self, "Нет заметки", "Сначала выбери или создай заметку."
            )
            return
        values = self._load_dropdown_values()
        initial = values[0] if values else "Выбрать…"
        dd_id = self._insert_dropdown_token(initial, values=values)
        cr = self.text_edit.cursorRect(self.text_edit.textCursor())
        pos = self.text_edit.viewport().mapToGlobal(cr.bottomLeft())
        self.show_dropdown_menu_for_token(dd_id, pos)

    def _open_dropdown_values_editor(
        self, initial_values: list[str] | None = None
    ) -> list[str] | None:
        dlg = QDialog(self)
        dlg.setWindowTitle("Значения выпадающего списка")
        v = QVBoxLayout(dlg)
        lst = QListWidget(dlg)
        v.addWidget(lst)
        if initial_values is None:
            try:
                last = json.loads(self.settings.value("dropdown_values", "[]"))
                if not isinstance(last, list):
                    last = []
            except Exception:
                last = []
            if not last:
                last = ["Вариант 1", "Вариант 2", "Вариант 3"]
        else:
            last = (
                list(initial_values)
                if initial_values
                else ["Вариант 1", "Вариант 2", "Вариант 3"]
            )
        for s in last:
            lst.addItem(str(s))
        btns_line = QHBoxLayout()
        btn_add = QPushButton("Добавить")
        btn_ren = QPushButton("Переименовать")
        btn_del = QPushButton("Удалить")
        btn_up = QPushButton("↑")
        btn_dn = QPushButton("↓")
        btns_line.addWidget(btn_add)
        btns_line.addWidget(btn_ren)
        btns_line.addWidget(btn_del)
        btns_line.addStretch(1)
        btns_line.addWidget(btn_up)
        btns_line.addWidget(btn_dn)
        v.addLayout(btns_line)

        def add_item():
            text, ok = QInputDialog.getText(dlg, "Добавить значение", "Текст:")
            if ok and text.strip():
                lst.addItem(text.strip())

        def rename_item():
            it = lst.currentItem()
            if not it:
                return
            text, ok = QInputDialog.getText(
                dlg, "Переименовать", "Текст:", text=it.text()
            )
            if ok and text.strip():
                it.setText(text.strip())

        def delete_item():
            row = lst.currentRow()
            if row >= 0:
                lst.takeItem(row)

        def move_up():
            row = lst.currentRow()
            if row > 0:
                it = lst.takeItem(row)
                lst.insertItem(row - 1, it)
                lst.setCurrentRow(row - 1)

        def move_dn():
            row = lst.currentRow()
            if 0 <= row < lst.count() - 1:
                it = lst.takeItem(row)
                lst.insertItem(row + 1, it)
                lst.setCurrentRow(row + 1)

        btn_add.clicked.connect(add_item)
        btn_ren.clicked.connect(rename_item)
        btn_del.clicked.connect(delete_item)
        btn_up.clicked.connect(move_up)
        btn_dn.clicked.connect(move_dn)
        buttons = QDialogButtonBox(QDialogButtonBox.Ok | QDialogButtonBox.Cancel, dlg)
        v.addWidget(buttons)
        buttons.accepted.connect(dlg.accept)
        buttons.rejected.connect(dlg.reject)
        fg = dlg.frameGeometry()
        fg.moveCenter(self.frameGeometry().center())
        dlg.move(fg.topLeft())
        if dlg.exec() == QDialog.Accepted:
            values = [
                lst.item(i).text().strip()
                for i in range(lst.count())
                if lst.item(i).text().strip()
            ]
            if initial_values is None:
                self.settings.setValue(
                    "dropdown_values", json.dumps(values, ensure_ascii=False)
                )
            return values
        return None

    def _show_combo_popup(self, values: list[str]) -> None:
        menu = QMenu(self)
        fm = self.text_edit.fontMetrics()
        w = (
            max([180] + [fm.horizontalAdvance(v) for v in values])
            + fm.averageCharWidth() * 3
        )
        menu.setFixedWidth(int(w))
        for val in values:
            act = menu.addAction(val)
            act.triggered.connect(lambda _, v=val: self._insert_dropdown_plain(v))
        cr = self.text_edit.cursorRect(self.text_edit.textCursor())
        pos = self.text_edit.viewport().mapToGlobal(cr.bottomLeft())
        menu.exec(pos)

    def _insert_dropdown_plain(self, value: str) -> None:
        fmt = self.text_edit.currentCharFormat()
        c = self.text_edit.textCursor()
        c.insertText(value, fmt)
        self.text_edit.setTextCursor(c)
        self.record_state_for_undo()
        if hasattr(self, "debounce_timer"):
            if self.autosave_enabled:
                self.save_note_quiet(force=True)

    def _commit_dropdown_value(self, combo: QComboBox, value: str) -> None:
        fmt = self.text_edit.currentCharFormat()
        cursor = self.text_edit.textCursor()
        cursor.insertText(value, fmt)
        self.text_edit.setTextCursor(cursor)
        combo.hide()
        self.record_state_for_undo()
        if hasattr(self, "debounce_timer"):
            if self.autosave_enabled:
                self.save_note_quiet(force=True)

    def _insert_dropdown_plain(self, value: str) -> None:
        fmt = self.text_edit.currentCharFormat()
        c = self.text_edit.textCursor()
        c.insertText(value, fmt)
        self.text_edit.setTextCursor(c)
        self.record_state_for_undo()
        if hasattr(self, "debounce_timer"):
            if self.autosave_enabled:
                self.save_note_quiet(force=True)

    def init_toolbar(self):
        full_toolbar_widget = QWidget()
        full_layout = QVBoxLayout(full_toolbar_widget)
        full_layout.setContentsMargins(0, 0, 0, 0)
        full_layout.setSpacing(0)
        top_toolbar_widget = QWidget()
        flow_layout = FlowLayout(top_toolbar_widget)

        def add_tool_button(icon_path, tooltip, callback):
            button = QPushButton()
            button.setText(tooltip.split()[0])
            button.setToolTip(tooltip)
            button.clicked.connect(callback)
            button.setFixedSize(32, 32)
            flow_layout.addWidget(button)
            return button

        def flow_break():
            br = QWidget()
            br.setProperty("flow_break", True)
            br.setFixedSize(0, 0)
            flow_layout.addWidget(br)

        toggle_fav_button = QPushButton("⭐")
        toggle_fav_button.setToolTip("⭐ - Добавить/удалить из избранного")
        toggle_fav_button.setFixedSize(32, 32)
        toggle_fav_button.clicked.connect(self.toggle_favorite)
        flow_layout.addWidget(toggle_fav_button)
        add_tool_button("", "🗑 - Корзина", self.show_trash)
        add_tool_button("", "🎲 - Генератор чисел", self.open_number_generator)
        add_tool_button("", "📒 - Заметки", self.show_all_notes)
        self.refresh_button = add_tool_button(
            "", "🔄 - Обновить", self.on_refresh_clicked
        )
        add_tool_button("", "➕ - Новая", self.create_new_note)
        add_tool_button("", "💾 - Сохранить", self.save_note)
        add_tool_button("📎", "📎 - Приерепить файл", self.attach_file_to_note)
        add_tool_button("", "🖌 - Рисовать", self.open_drawing_dialog)
        self.audio_button = QPushButton("🎤")
        self.audio_button.setToolTip("🎤 - Записать аудио")
        self.audio_button.setFixedSize(32, 32)
        self.audio_button.clicked.connect(self.toggle_audio_recording)
        flow_layout.addWidget(self.audio_button)
        add_tool_button("", "𝐁 - Жирный", self.toggle_bold)
        add_tool_button("", "𝐼 - Курсив", self.toggle_italic)
        add_tool_button("", "U̲ - Подчёркнутый", self.toggle_underline)
        add_tool_button("", "̶̶̶Z - Зачеркнуть", self.toggle_strikethrough)
        add_tool_button("", "🔽 - Выпадающий список", self.insert_dropdown)
        add_tool_button("", "🧹 - Сбросить формат", self.clear_formatting)
        add_tool_button("", "🌈 - Цвет текста", self.change_text_color)
        add_tool_button("", "🅰️ - Фон текста", self.change_background_color)
        add_tool_button("", "← - Расположить слева", self.align_left)
        add_tool_button("", "→← - Расположить по центру", self.align_center)
        add_tool_button("", "→ - Расположить справа", self.align_right)
        add_tool_button("", "≡ - По ширине", self.align_justify)
        add_tool_button("", "H1 - Заголовок 1", lambda: self.apply_heading(1))
        add_tool_button("", "H2 - Заголовок 2", lambda: self.apply_heading(2))
        add_tool_button("", "H3 - Заголовок 3", lambda: self.apply_heading(3))
        add_tool_button("", "Aa - Регистр (выделение)", self.translate_case_only)
        add_tool_button("", "⌨ - Раскладка (выделение)", self.translate_layout_only)
        add_tool_button("", "• - Маркированный  список", self.insert_bullet_list)
        add_tool_button("", "1. - Нумерованный список", self.insert_numbered_list)
        add_tool_button("", "☑ - Чекбокс", self.insert_checkbox)
        add_tool_button("", "📅 - Таблица", self.insert_table)
        add_tool_button("", "🔗 - Ссылка", self.insert_link)
        add_tool_button("", "❌ - Удалить ссылку", self.remove_link)
        add_tool_button("", "▁ - Горизонтальная линия", self.insert_horizontal_line)
        add_tool_button("", "+🏷 - Добавить тег", self.add_tag_to_note)
        self.tag_filter = QComboBox()
        self.tag_filter.setEditable(False)
        self.tag_filter.setFixedWidth(180)
        self.update_tag_filter_items()
        flow_layout.addWidget(self.tag_filter)
        manage_tags_button = QPushButton("🏷")
        manage_tags_button.setToolTip("Управление тегами")
        manage_tags_button.clicked.connect(self.manage_tags_dialog)
        flow_layout.addWidget(manage_tags_button)
        favorites_button = QPushButton("★Избранные")
        favorites_button.clicked.connect(self.show_favorites_only)
        flow_layout.addWidget(favorites_button)
        reset_button = QPushButton("Сбросить фильтр")
        reset_button.clicked.connect(self.refresh_notes_list)
        flow_layout.addWidget(reset_button)
        self.sort_combo = QComboBox()
        self.sort_combo.addItems(["По заголовку", "По дате", "Избранные"])
        self.sort_combo.setToolTip("Сортировка")
        self.sort_combo.currentIndexChanged.connect(self.apply_sorting)
        flow_layout.addWidget(self.sort_combo)
        self.sort_order_combo = QComboBox()
        self.sort_order_combo.addItems(["↑", "↓"])
        self.sort_order_combo.setToolTip("Порядок сортировки")
        self.sort_order_combo.currentIndexChanged.connect(self.apply_sorting)
        flow_layout.addWidget(self.sort_order_combo)
        self.search_mode_combo = QComboBox()
        self.search_mode_combo.addItems(["Заголовок", "Содержимое", "Напоминание"])
        flow_layout.addWidget(self.search_mode_combo)
        self.search_bar = QLineEdit()
        self.search_bar.setPlaceholderText("Поиск...")
        flow_layout.addWidget(self.search_bar)
        search_button = QPushButton("🔍 - Поиск")
        search_button.clicked.connect(self.trigger_search)
        flow_layout.addWidget(search_button)
        self.font_combo = QFontComboBox()
        self.font_combo.setCurrentFont(QFont("Times New Roman"))
        self.font_combo.currentFontChanged.connect(self.change_font)
        flow_layout.addWidget(self.font_combo)
        self.font_size_spin = QSpinBox()
        self.font_size_spin.setRange(8, 48)
        self.font_size_spin.setValue(14)
        self.font_size_spin.valueChanged.connect(self.change_font_size)
        flow_layout.addWidget(self.font_size_spin)
        reminder_layout = QHBoxLayout()
        reminder_layout.setContentsMargins(0, 0, 0, 0)
        reminder_layout.setSpacing(10)
        reminder_button = QPushButton("Установить напоминание")
        reminder_button.clicked.connect(self.set_reminder_for_note)
        edit_reminder_button = QPushButton("Изменить напоминание")
        edit_reminder_button.clicked.connect(self.edit_reminder_for_note)
        remove_reminder_button = QPushButton("Удалить напоминание")
        remove_reminder_button.clicked.connect(self.remove_reminder_from_note)
        reminder_layout.addWidget(reminder_button)
        reminder_layout.addWidget(edit_reminder_button)
        reminder_layout.addWidget(remove_reminder_button)
        bottom_toolbar_widget = QWidget()
        bottom_toolbar_widget.setLayout(reminder_layout)
        full_layout.addWidget(top_toolbar_widget)
        self.toolbar_scroll = QScrollArea()
        self.toolbar_scroll.setWidget(full_toolbar_widget)
        self.toolbar_scroll.setWidgetResizable(True)
        self.toolbar_scroll.setMaximumHeight(140)
        self.toolbar_scroll.setHorizontalScrollBarPolicy(Qt.ScrollBarAlwaysOff)
        self.toolbar_scroll.setVerticalScrollBarPolicy(Qt.ScrollBarAlwaysOff)
        self.search_bar.setFocusPolicy(Qt.FocusPolicy.ClickFocus)
        self.timer_btn = QPushButton()
        self.timer_btn.setFixedHeight(32)
        self.timer_btn.setContextMenuPolicy(Qt.CustomContextMenu)
        self.timer_btn.clicked.connect(self._on_timer_clicked)
        self.timer_btn.customContextMenuRequested.connect(self._timer_context_menu)
        flow_layout.addWidget(self.timer_btn)
        self._update_timer_ui()
        self.stopwatch_time = 0
        self.stopwatch_running = False
        self.stopwatch_timer = QTimer(self)
        self.stopwatch_timer.setInterval(1000)
        self.stopwatch_timer.timeout.connect(self._update_stopwatch)
        if not hasattr(self, "lbl_stopwatch"):
            self.lbl_stopwatch = QLabel("00:00:00", self)
        self.btn_sw_start = QPushButton("▶️")
        self.btn_sw_stop = QPushButton("⏸️")
        self.btn_sw_reset = QPushButton("⏹️")
        self.btn_sw_start.clicked.connect(self.start_stopwatch)
        self.btn_sw_stop.clicked.connect(self.stop_stopwatch)
        self.btn_sw_reset.clicked.connect(self.reset_stopwatch)
        sw_layout = QHBoxLayout()
        sw_layout.setContentsMargins(0, 0, 0, 0)
        sw_layout.setSpacing(4)
        sw_layout.addWidget(self.lbl_stopwatch)
        sw_layout.addWidget(self.btn_sw_start)
        sw_layout.addWidget(self.btn_sw_stop)
        sw_layout.addWidget(self.btn_sw_reset)
        sw_widget = QWidget()
        sw_widget.setLayout(sw_layout)
        flow_layout.addWidget(sw_widget)

    def _timer_context_menu(self, pos):
        menu = QMenu(self)
        act_reset = menu.addAction("Сброс")
        act_mode = menu.addAction(
            "Режим: "
            + ("Секундомер" if self._timer_mode == "countdown" else "Обратный отсчёт")
        )
        sender = self.sender()
        try:
            global_pos = (
                sender.mapToGlobal(pos) if sender is not None else QCursor.pos()
            )
        except Exception:
            global_pos = QCursor.pos()
        act = menu.exec(global_pos)
        if act == act_reset:
            self._reset_timer()
        elif act == act_mode:
            self.toggle_timer_mode()

    def open_mass_reminders_dialog(self):
        dialog = QDialog(self)
        dialog.setWindowFlag(Qt.WindowStaysOnTopHint, True)
        dialog.setWindowTitle("Массовое управление напоминаниями")
        layout = QVBoxLayout(dialog)
        list_widget = QListWidget()
        notes_with_reminder = [n for n in self.notes if n.reminder]
        for note in notes_with_reminder:
            item = QListWidgetItem(f"{note.title} [{note.reminder}]")
            item.setData(Qt.UserRole, note)
            item.setCheckState(Qt.Unchecked)
            list_widget.addItem(item)
        layout.addWidget(list_widget)
        btn_remove = QPushButton("Удалить напоминание у выбранных")
        btn_remove.clicked.connect(
            lambda _=False, lw=list_widget, dlg=dialog: self.mass_remove_reminders(
                lw, dlg
            )
        )
        layout.addWidget(btn_remove)
        dialog.setLayout(layout)
        dialog.exec()

    def mass_remove_reminders(self, list_widget, dialog):
        changed = False
        for i in range(list_widget.count()):
            item = list_widget.item(i)
            note = item.data(Qt.UserRole)
            if item.checkState() == Qt.Checked:
                note.reminder = None
                note.reminder_repeat = None
                self.save_note_to_file(note)
                changed = True
        if changed:
            self.refresh_notes_list()
            dialog.accept()

    def open_number_generator(self):
        dlg = NumberGeneratorDialog(self)
        dlg.exec()

    def refresh_all(self) -> None:
        try:
            current_uuid = (
                self.current_note.uuid if getattr(self, "current_note", None) else None
            )

            self.load_notes_from_disk()
            self.refresh_notes_list()

            if current_uuid:
                note = next((n for n in self.notes if n.uuid == current_uuid), None)
                if not note:
                    self.show_note_with_attachments(None)
                    return

                note_dir = os.path.join(
                    NOTES_DIR,
                    NotesApp.safe_folder_name(note.title, note.uuid, note.timestamp),
                )
                json_path = os.path.join(note_dir, "note.json")
                if os.path.exists(json_path):
                    with open(json_path, "r", encoding="utf-8") as f:
                        data = json.load(f)
                    fresh = Note.from_dict(data)

                    note.title = fresh.title
                    note.content = fresh.content
                    note.tags = list(getattr(fresh, "tags", []))
                    note.password_manager = getattr(fresh, "password_manager", "")
                    note.rdp_1c8 = getattr(fresh, "rdp_1c8", "")
                    note.custom_fields = list(getattr(fresh, "custom_fields", []))
                    note.favorite = getattr(
                        fresh, "favorite", getattr(note, "favorite", False)
                    )

                self.current_note = note
                self.text_edit.blockSignals(True)
                self.text_edit.setHtml(
                    getattr(note, "content_html", "") or note.content
                )
                self.text_edit.blockSignals(False)

                self.tags_label.setText(
                    f"Теги: {', '.join(note.tags) if note.tags else 'нет'}"
                )
                self.password_manager_field.setText(
                    getattr(note, "password_manager", "")
                )
                self.rdp_1c8_field.setText(getattr(note, "rdp_1c8", ""))

                self._rebuild_attachments_panel(note)

        except Exception as e:
            QMessageBox.warning(self, "Обновление", f"Не удалось обновить: {e}")

    def rebuild_toolbar(self):
        dock = getattr(self, "dock_toolbar", None)
        was_visible = bool(dock and dock.isVisible())
        area = self.dockWidgetArea(dock) if dock else Qt.TopDockWidgetArea
        old_scroll = getattr(self, "toolbar_scroll", None)
        if old_scroll is not None:
            try:
                old_scroll.setParent(None)
                old_scroll.deleteLater()
            except Exception:
                pass
        self.init_toolbar()
        if dock is None:
            self.dock_toolbar = QDockWidget("Панель инструментов", self)
            self.dock_toolbar.setObjectName("dock_toolbar")
            self.dock_toolbar.setAllowedAreas(
                Qt.TopDockWidgetArea | Qt.BottomDockWidgetArea
            )
            self.addDockWidget(area, self.dock_toolbar)
            dock = self.dock_toolbar
        dock.setWidget(self.toolbar_scroll)
        dock.setVisible(was_visible)

    def show_random_number(self):
        gen = NumberGenerator(1, 1000)  # диапазон можно изменить
        num = gen.generate()
        QMessageBox.information(self, "Случайное число", f"Результат: {num}")


    def show_reminder_dialog(self, note):
        time_str = ""
        if note.reminder:
            dt = QDateTime.fromString(note.reminder, Qt.ISODate)
            if dt.isValid():
                time_str = dt.toString("dd.MM.yyyy HH:mm")
        text = f"<b>Напоминание по заметке:</b><br><b>{note.title}</b>"
        if time_str:
            text += f"<br><b>Время:</b> {time_str}"
        msg = QMessageBox(self)
        msg.setIcon(QMessageBox.Information)
        msg.setWindowTitle("Напоминание")
        msg.setTextFormat(Qt.RichText)
        msg.setText(text)
        msg.setStandardButtons(QMessageBox.Ok)
        msg.setWindowFlag(Qt.WindowStaysOnTopHint, True)
        msg.exec()
        note.reminder_shown = True

    def show_note_context_menu(self, position):
        item = self.notes_list.itemAt(position)
        if not item:
            menu = QMenu(self)
            create_action = QAction("Создать", self)
            create_action.triggered.connect(self.create_new_note)
            menu.addAction(create_action)
            menu.exec(self.notes_list.viewport().mapToGlobal(position))
            return
        note = item.data(Qt.UserRole)
        menu = QMenu()
        menu.setStyleSheet(
            """
            QMenu {
                background-color: #282828;
                color: #fff;
                border-radius: 10px;
                font-size: 16px;
                padding: 6px;
            }
            QMenu::item {
                background-color: transparent;
                color: #fff;
                padding: 8px 32px 8px 24px;
                border-radius: 6px;
            }
            QMenu::item:selected {
                background-color: #fff;
                color: #111;
                border-radius: 6px;
                border: 2px solid #7a7a7a;
            }
            QMenu::separator {
                height: 1px;
                background: #444;
                margin: 8px 0;
            }
            """
        )
        toggle_pin_action = QAction(
            "Закрепить" if not note.pinned else "Открепить", self
        )
        toggle_pin_action.triggered.connect(lambda: self.toggle_pin(note))
        menu.addAction(toggle_pin_action)
        open_action = QAction("Открыть", self)
        open_action.triggered.connect(lambda: self.select_note(note))
        menu.addAction(open_action)
        open_in_fs_action = QAction("Открыть в проводнике", self)
        open_in_fs_action.triggered.connect(lambda: self.open_note_folder(note))
        menu.addAction(open_in_fs_action)
        rename_action = QAction("Переименовать", self)
        rename_action.triggered.connect(lambda: self.rename_note(note))
        menu.addAction(rename_action)
        copy_action = QAction("Копировать", self)
        copy_action.triggered.connect(lambda: self.copy_note(note))
        menu.addAction(copy_action)
        menu.addSeparator()
        delete_perm_action = QAction("Удалить безвозвратно", self)
        delete_perm_action.triggered.connect(lambda: self.delete_note_completely(note))
        menu.addAction(delete_perm_action)
        note_folder_name = NotesApp.safe_folder_name(
            note.title, note.uuid, note.timestamp
        )
        note_in_trash = os.path.exists(os.path.join(self.TRASH_DIR, note_folder_name))
        if note_in_trash:
            delete_action = QAction("Удалить навсегда", self)

            def delete_permanently():
                reply = QMessageBox.question(
                    self,
                    "Удалить навсегда",
                    f"Удалить заметку '{note.title}' безвозвратно?",
                    QMessageBox.Yes | QMessageBox.No,
                )
                if reply == QMessageBox.Yes:
                    trash_note_dir = os.path.join(self.TRASH_DIR, note_folder_name)
                    if os.path.exists(trash_note_dir):
                        shutil.rmtree(trash_note_dir)
                    self.refresh_notes_list()

            delete_action.triggered.connect(delete_permanently)
            menu.addAction(delete_action)
        else:
            delete_action = QAction("Удалить", self)

            def move_to_trash():
                reply = QMessageBox.question(
                    self,
                    "Удаление",
                    f"Переместить заметку '{note.title}' в корзину?",
                    QMessageBox.Yes | QMessageBox.No,
                )
                if reply == QMessageBox.Yes:
                    self.move_note_to_trash(note)
                    self.refresh_notes_list()
                    self.load_note(None)
                    self.current_note = None

            delete_action.triggered.connect(move_to_trash)
            menu.addAction(delete_action)
        menu.exec(self.notes_list.viewport().mapToGlobal(position))

    def autosave(self) -> None:
        if self.current_note and not self.text_edit.isReadOnly():
            current_html = self.text_edit.toHtml()
            last_html = (
                self.current_note.history[-1] if self.current_note.history else ""
            )
            if current_html != last_html:
                if len(self.current_note.history) >= MAX_HISTORY:
                    self.current_note.history.pop(0)
                self.current_note.history.append(current_html)
                self.current_note.history_index = len(self.current_note.history) - 1
                self.update_history_list()
                self.update_history_list_selection()
            self.current_note.content = current_html
            self.current_note.password_manager = self.password_manager_field.text()
            self.current_note.rdp_1c8 = self.rdp_1c8_field.text()
            self.save_note_to_file(self.current_note)

    def attach_docks(self, dock1, dock2):
        area = self.dockWidgetArea(dock1)
        self.addDockWidget(area, dock2)
        self.tabifyDockWidget(dock1, dock2)
        dock2.show()
        dock1.raise_()

    def detach_dock(self, dock):
        dock.setFloating(True)
        dock.show()

    def move_dock(self, dock, area):
        self.addDockWidget(area, dock)
        dock.show()
        dock.raise_()

    def show_dock_context_menu(self, position, dock):
        menu = QMenu()
        attach_menu = menu.addMenu("Приклеить к...")
        docks = [
            (self.dock_notes_list, "Заметки"),
            (self.dock_history, "История"),
            (self.dock_buttons, "Кнопки"),
            (self.dock_editor, "Редактор"),
            (self.dock_toolbar, "Панель инструментов"),
        ]
        for target_dock, name in docks:
            if target_dock != dock:
                action = attach_menu.addAction(name)
                action.triggered.connect(
                    lambda checked, d1=target_dock, d2=dock: self.attach_docks(d1, d2)
                )

        move_menu = menu.addMenu("Переместить к краю")
        areas = [
            (Qt.LeftDockWidgetArea, "Слева"),
            (Qt.RightDockWidgetArea, "Справа"),
            (Qt.TopDockWidgetArea, "Сверху"),
            (Qt.BottomDockWidgetArea, "Снизу"),
        ]
        for area, name in areas:
            move_action = move_menu.addAction(name)
            move_action.triggered.connect(
                lambda checked, a=area, d=dock: self.move_dock(d, a)
            )
        detach_action = menu.addAction("Открепить")
        detach_action.triggered.connect(lambda: self.detach_dock(dock))
        menu.exec(dock.mapToGlobal(position))

    def clear_formatting(self):
        te = self.text_edit
        c = te.textCursor()
        ch = QTextCharFormat()
        f = QFont("Times New Roman", 14)
        ch.setFont(f)
        ch.setFontWeight(QFont.Normal)
        ch.setFontItalic(False)
        ch.setFontUnderline(False)
        ch.setFontStrikeOut(False)
        ch.setForeground(QBrush())
        ch.setBackground(QBrush())
        bf = QTextBlockFormat()
        bf.setHeadingLevel(0)
        if c.hasSelection():
            c.mergeBlockFormat(bf)
            c.mergeCharFormat(ch)
        else:
            te.setCurrentCharFormat(ch)
            c.mergeBlockFormat(bf)
        te.setTextCursor(c)

    def toggle_pin(self, note):
        note.pinned = not note.pinned
        self.save_note_to_file(note)
        self.refresh_notes_list()

    def add_menu_bar(self):
        menu_bar = self.menuBar()
        apps_menu = menu_bar.addMenu("Приложения")
        act_switch = QAction("Выбрать приложение…  (Ctrl+Alt+L)", self)
        act_switch.triggered.connect(self.show_app_launcher)
        apps_menu.addAction(act_switch)
        plugins_menu = menu_bar.addMenu("Плагины")
        act = QAction("Управление плагинами", self)
        act.triggered.connect(self.manage_plugins_dialog)
        plugins_menu.addAction(act)
        act = QAction("Перезагрузить плагины", self)
        act.triggered.connect(self.load_plugins)
        plugins_menu.addAction(act)
        tools_menu = menu_bar.addMenu("Инструменты")
        act_rng = QAction("🎲 Генератор чисел", self)
        act_rng.triggered.connect(self.open_number_generator)
        tools_menu.addAction(act_rng)
        template_menu = self.menuBar().addMenu("Шаблоны")
        act = QAction("Вставить шаблон", self)
        act.triggered.connect(self.insert_template)
        template_menu.addAction(act)
        act = QAction("Сохранить как шаблон", self)
        act.setShortcut("Ctrl+Shift+S")
        act.triggered.connect(self.save_current_as_template)
        template_menu.addAction(act)
        act = QAction("Управление шаблонами", self)
        act.triggered.connect(self.manage_templates_dialog)
        template_menu.addAction(act)
        file_menu = menu_bar.addMenu("Файл")
        act = QAction("Импорт", self)
        act.triggered.connect(self.import_note)
        file_menu.addAction(act)
        act = QAction("Экспорт в PDF", self)
        act.triggered.connect(self.export_current_note_pdf)
        file_menu.addAction(act)
        act = QAction("Экспорт в TXT", self)
        act.triggered.connect(self.export_current_note_txt)
        file_menu.addAction(act)
        act = QAction("Экспорт в DOCX", self)
        act.triggered.connect(self.export_current_note_docx)
        file_menu.addAction(act)
        act = QAction("Экспорт в JSON", self)
        act.triggered.connect(self.export_note)
        file_menu.addAction(act)
        help_menu = menu_bar.addMenu("Справка")
        act = QAction("Открыть README.md", self)
        act.triggered.connect(self.open_readme)
        help_menu.addAction(act)
        settings_menu = menu_bar.addMenu("Настройки")
        settings_action = QAction("Настройки:", self)
        settings_action.setShortcut("Ctrl+,")
        settings_action.triggered.connect(self.show_settings_window)
        settings_menu.addAction(settings_action)
        self.view_menu = menu_bar.addMenu("Вид")
        for dock in [
            self.dock_notes_list,
            self.dock_history,
            self.dock_buttons,
            self.dock_editor,
        ]:
            action = dock.toggleViewAction()
            action.setText(dock.windowTitle())
            self.view_menu.addAction(action)
        self.view_menu.addSeparator()
        self.action_always_on_top = QAction("Поверх всех окон", self)
        self.action_always_on_top.setCheckable(True)
        self.action_always_on_top.setChecked(
            self.settings.value("ui/always_on_top", True, type=bool)
        )
        self.action_always_on_top.toggled.connect(self._toggle_always_on_top)
        self.view_menu.addAction(self.action_always_on_top)
        self.toolbars_menu = self.view_menu.addMenu("Панели инструментов")
        dock_action = self.dock_toolbar.toggleViewAction()
        dock_action.setText(self.dock_toolbar.windowTitle() or "Панель инструментов")
        self.toolbars_menu.addAction(dock_action)
        if hasattr(self, "visibility_toolbar") and isinstance(
            self.visibility_toolbar, QToolBar
        ):
            tb_action = self.visibility_toolbar.toggleViewAction()
            tb_action.setText(
                self.visibility_toolbar.windowTitle() or "Видимость полей"
            )
            self.toolbars_menu.addAction(tb_action)
        trash_menu = self.menuBar().addMenu("Корзина")
        act = QAction("Показать корзину", self)
        act.triggered.connect(self.show_trash)
        trash_menu.addAction(act)
        act = QAction("Восстановить заметку", self)
        act.triggered.connect(self.restore_note_from_trash)
        trash_menu.addAction(act)
        act = QAction("Удалить безвозвратно", self)
        act.triggered.connect(self.delete_note_permanently)
        trash_menu.addAction(act)
        act = QAction("Очистить корзину", self)
        act.triggered.connect(self.empty_trash)
        trash_menu.addAction(act)
        reminders_menu = menu_bar.addMenu("Управление напоминаниями")
        act = QAction("Добавить напоминание к текущей заметки", self)
        act.triggered.connect(self.set_reminder_for_note)
        reminders_menu.addAction(act)
        act = QAction("Изменить напоминание у текущей заметки", self)
        act.triggered.connect(self.edit_reminder_for_note)
        reminders_menu.addAction(act)
        act = QAction("Удалить напоминание у текущей заметки", self)
        act.triggered.connect(self.remove_reminder_from_note)
        reminders_menu.addAction(act)
        reminders_menu.addSeparator()
        act = QAction("Массовое удаление", self)
        act.triggered.connect(self.open_mass_reminders_dialog)
        reminders_menu.addAction(act)

    def _toggle_always_on_top(self, checked: bool) -> None:
        self.setWindowFlag(Qt.WindowStaysOnTopHint, checked)
        self.settings.setValue("ui/always_on_top", checked)
        if (
            hasattr(self, "topmost_checkbox")
            and self.topmost_checkbox.isChecked() != checked
        ):
            self.topmost_checkbox.blockSignals(True)
            self.topmost_checkbox.setChecked(checked)
            self.topmost_checkbox.blockSignals(False)
        if (
            hasattr(self, "action_always_on_top")
            and self.action_always_on_top.isChecked() != checked
        ):
            self.action_always_on_top.blockSignals(True)
            self.action_always_on_top.setChecked(checked)
            self.action_always_on_top.blockSignals(False)
        self.show()

    def show_settings_window(self):
        dialog = QDialog(self)
        dialog.setWindowFlag(Qt.WindowStaysOnTopHint, True)
        dialog.setWindowTitle("Настройки")
        layout = QFormLayout(dialog)
        theme_combo = QComboBox()
        theme_combo.addItems(["Светлая", "Тёмная"])
        theme_combo.setCurrentIndex(
            0 if self.settings.value("theme", "dark") == "light" else 1
        )
        layout.addRow("Тема:", theme_combo)
               
        online_spell_cb = QCheckBox()
        curr_online = self.settings.value("spellcheck/online_enabled", True, type=bool)
        online_spell_cb.setChecked(curr_online)
        online_spell_cb.setToolTip("Если выключено — подчёркивание строится только по локальному словарю.")
        layout.addRow("Онлайн-проверка грамматики:", online_spell_cb)
        buttons = QDialogButtonBox(QDialogButtonBox.Ok | QDialogButtonBox.Cancel)
        buttons.accepted.connect(dialog.accept)
        buttons.rejected.connect(dialog.reject)
        layout.addWidget(buttons)
        dialog.setLayout(layout)
        if dialog.exec() == QDialog.DialogCode.Accepted:
            new_theme = "light" if theme_combo.currentIndex() == 0 else "dark"
            self.settings.setValue("theme", new_theme)
            self.spell_online_enabled = online_spell_cb.isChecked()
            self.settings.setValue("spellcheck/online_enabled", self.spell_online_enabled)
            if hasattr(self, "spell_highlighter"):
                self.spell_highlighter.set_online_enabled(self.spell_online_enabled)
            else:
                self.autosave_timer.stop()
                if hasattr(self, "debounce_timer"):
                    self.debounce_timer.stop()
            self.init_theme()

    def _fmt_mmss(self, secs: int) -> str:
        secs = max(0, int(secs))
        return f"{secs // 60:02d}:{secs % 60:02d}"

    def _update_timer_ui(self) -> None:
        if hasattr(self, "timer_btn"):
            txt = f"⏱ {self._fmt_mmss(self._timer_left if self._timer_mode=='countdown' else self._timer_total if not self._timer_running else self._timer_left)}"
            txt = f"⏱ {self._fmt_mmss(self._timer_left)}"
            self.timer_btn.setText(txt)
            self.timer_btn.setToolTip(
                f"⏱ {txt[2:].strip()} — {('идёт' if self._timer_running else 'пауза')}; "
                "ЛКМ — старт/пауза, ПКМ — сброс, Shift+ЛКМ — задать длительность"
            )

    def _on_timer_clicked(self) -> None:
        if QApplication.keyboardModifiers() & Qt.ShiftModifier:
            self._set_timer_duration_dialog()
            return
        if not self._timer_running:
            self._timer_running = True
            self._timer.start()
        else:
            self._timer_running = False
            self._timer.stop()
        self._update_timer_ui()

    def _reset_timer(self) -> None:
        if self._timer_mode == "countdown":
            self._timer_left = int(self._timer_total)
        else:
            self._timer_left = 0
        self._timer_running = False
        self._timer.stop()
        self._update_timer_ui()

    def _tick_timer(self) -> None:
        if self._timer_mode == "countdown":
            self._timer_left -= 1
            if self._timer_left <= 0:
                self._timer_left = 0
                self._timer_running = False
                self._timer.stop()
                try:
                    QApplication.beep()
                    self.show_toast(
                        "⏱ Таймер: время вышло!",
                        timeout_ms=3000,
                        anchor_widget=getattr(self, "timer_btn", None),
                    )
                except Exception:
                    pass
        else:
            self._timer_left += 1
        self._update_timer_ui()

    def _set_timer_duration_dialog(self) -> None:
        minutes, ok = QInputDialog.getInt(
            self, "⏱ Длительность", "Минут:", max(1, self._timer_total // 60), 1, 240, 1
        )
        if not ok:
            return
        self._timer_total = minutes * 60
        self.settings.setValue("timer/total", int(self._timer_total))
        if self._timer_mode == "countdown":
            self._timer_left = int(self._timer_total)
        self._update_timer_ui()

    def toggle_timer_mode(self) -> None:
        self._timer_mode = (
            "stopwatch" if self._timer_mode == "countdown" else "countdown"
        )
        self.settings.setValue("timer/mode", self._timer_mode)
        self._reset_timer()

    def show_help_window(self):
        dialog = QDialog(self)
        dialog.setWindowFlag(Qt.WindowStaysOnTopHint, True)
        dialog.setWindowTitle("Справка — Мои Заметки")
        dialog.setMinimumSize(600, 620)
        layout = QVBoxLayout(dialog)
        browser = QTextBrowser(dialog)
        browser.setOpenExternalLinks(True)
        browser.setHtml(
            """
                        <h2 style="text-align:center;">🗒️ Справка по приложению "Мои Заметки"</h2>

                        <h3>Основные возможности</h3>
                        <ul>
                        <li>Создание, редактирование, удаление и восстановление заметок</li>
                        <li>Форматирование текста (жирный, курсив, подчёркнутый, зачёркнутый, цвет, списки, чекбоксы, линии)</li>
                        <li>Прикрепление файлов и изображений, просмотр и удаление вложений</li>
                        <li>Вставка ссылок, аудиозаписей</li>
                        <li>Работа с тегами (создание, добавление, фильтрация, переименование, удаление)</li>
                        <li>История изменений и версий заметки</li>
                        <li>Корзина: перемещение, восстановление, удаление навсегда, очистка</li>
                        <li>Избранные и закреплённые заметки</li>
                        <li>Поиск, сортировка и фильтрация заметок</li>
                        <li>Напоминания с возможностью повторения</li>
                        <li>Экспорт заметок в PDF, TXT, DOCX</li>
                        </ul>

                        <h3>Горячие клавиши</h3>
                        <ul>
                        <li><b>Ctrl+N</b> — новая заметка</li>
                        <li><b>Ctrl+S</b> — сохранить заметку</li>
                        <li><b>Ctrl+B</b> — сделать текст жирным</li>
                        <li><b>Ctrl+I</b> — сделать текст курсивом</li>
                        <li><b>Ctrl+U</b> — подчеркнуть текст</li>
                        <li><b>Ctrl+Z</b> — отменить действие (Undo)</li>
                        <li><b>Ctrl+Y</b> — повторить действие (Redo)</li>
                        <li><b>Ctrl+Space</b> — сбросить форматирование выделенного текста</li>
                        <li><b>Ctrl+V</b> — вставить (текст, изображение из буфера обмена)</li>
                        <li><b>Delete</b> — удалить выделенную заметку (в корзину)</li>
                        <li><b>Esc</b> — выйти из режима редактирования заметки</li>
                        <li><b>Ctrl + Клик по картинке</b> — открыть изображение во внешней программе</li>
                        <li><b>ПКМ по заметке</b> — контекстное меню (открыть, удалить, переименовать, копировать, закрепить)</li>
                        </ul>

                        <h3>Форматирование и работа с текстом</h3>
                        <ul>
                        <li><b>𝐁</b> — жирный текст</li>
                        <li><b>𝐼</b> — курсив</li>
                        <li><b>U̲</b> — подчёркнутый</li>
                        <li><b>S̶</b> — зачёркнутый</li>
                        <li><b>🌈</b> — изменить цвет текста</li>
                        <li><b>🅰️</b> — изменить цвет фона текста</li>
                        <li><b>🧹</b> — сбросить всё форматирование</li>
                        <li><b>🔗</b> — вставить гиперссылку</li>
                        <li><b>•</b> — маркированный список</li>
                        <li><b>1.</b> — нумерованный список</li>
                        <li><b>☑</b> — чекбокс (флажок)</li>
                        <li><b>▁</b> — горизонтальная линия</li>
                        <li>Изменение размера и семейства шрифта — через панель инструментов сверху</li>
                        </ul>

                        <h3>Вложения, изображения и аудио</h3>
                        <ul>
                        <li><b>📎</b> — прикрепить файл к заметке</li>
                        <li><b>🖼</b> — вставить изображение (через меню или Ctrl+V)</li>
                        <li>Вложенные изображения и файлы отображаются под текстом заметки с кнопками "Открыть" и "Удалить"</li>
                        <li><b>🎤</b> — записать аудиофрагмент (отдельный файл, появится как вложение)</li>
                        </ul>

                        <h3>Работа с заметками</h3>
                        <ul>
                        <li>Новая, сохранить, удалить — через кнопки и горячие клавиши</li>
                        <li>Переименовать, копировать, закрепить — через контекстное меню</li>
                        <li>Закреплённые заметки отображаются выше</li>
                        <li>Избранное (⭐) — отметка заметки, быстрая фильтрация</li>
                        <li>Восстановление из корзины — через меню корзины</li>
                        </ul>

                        <h3>Корзина</h3>
                        <ul>
                        <li>Удалённые заметки перемещаются в корзину</li>
                        <li>В корзине можно восстановить заметку или удалить навсегда</li>
                        <li>Кнопка "Очистить корзину" — удаляет все заметки безвозвратно</li>
                        </ul>

                        <h3>Теги, фильтрация и сортировка</h3>
                        <ul>
                        <li>Теги отображаются под заголовком заметки</li>
                        <li>Добавление, удаление, переименование, фильтрация заметок по тегам</li>
                        <li>Строка поиска — фильтрация по заголовку, содержимому, тегам</li>
                        <li>Сортировка — по имени, дате, избранному (меню сортировки)</li>
                        <li>Быстрая фильтрация по тегам</li>
                        </ul>

                        <h3>История изменений</h3>
                        <ul>
                        <li>Для каждой заметки ведётся история версий (до 20 последних)</li>
                        <li>Можно откатываться к любой версии, удалять отдельные версии</li>
                        <li>История отображается в отдельной панели "История"</li>
                        </ul>

                        <h3>Напоминания</h3>
                        <ul>
                        <li>Можно задать напоминание для заметки (одноразовое или повторяющееся)</li>
                        <li>Время и периодичность задаются через форму напоминания</li>
                        <li>Сработавшее напоминание покажет уведомление, если приложение запущено</li>
                        </ul>

                        <h3>Экспорт</h3>
                        <ul>
                        <li>Заметку можно экспортировать в PDF, TXT или DOCX</li>
                        <li>Экспорт доступен через меню или кнопку экспорта</li>
                        </ul>

                        <h3>Прочее</h3>
                        <ul>
                        <li>Автосохранение изменений заметки каждые 60 секунд</li>
                        <li>Поддержка drag&amp;drop для вложений и заметок</li>
                        <li>Все данные хранятся локально в папке "Notes" рядом с программой</li>
                        <li>Удаление заметок и вложений из корзины — без возможности восстановления</li>
                        </ul>

                        <p style="font-size:11px; color:#888;">Обновлено: Август 2025</p>
                        """
        )
        close_btn = QPushButton("Закрыть")
        close_btn.clicked.connect(dialog.accept)
        layout.addWidget(browser)
        layout.addWidget(close_btn)
        dialog.exec()

    def export_note(self):
        if not self.current_note:
            QMessageBox.warning(
                self, "Экспорт", "Сначала выберите заметку для экспорта."
            )
            return
        file_path, _ = QFileDialog.getSaveFileName(
            self,
            "Экспорт заметки",
            f"{self.current_note.title}.json",
            "JSON Files (*.json)",
        )
        if file_path:
            note_dict = self.current_note.to_dict()
            try:
                with open(file_path, "w", encoding="utf-8") as f:
                    json.dump(note_dict, f, ensure_ascii=False, indent=4)
                self.show_toast("Экспорт завершён")
            except Exception as e:
                QMessageBox.critical(self, "Ошибка", f"Ошибка экспорта: {e}")

    def import_note(self):
        file_path, _ = QFileDialog.getOpenFileName(
            self, "Импорт заметки", "", "JSON Files (*.json)"
        )
        if file_path:
            try:
                with open(file_path, "r", encoding="utf-8") as f:
                    note_data = json.load(f)
                note_data["uuid"] = str(uuid.uuid4())
                note_data["title"] += " (импорт)"
                imported_note = Note.from_dict(note_data)
                self.save_note_to_file(imported_note)
                self.load_notes_from_disk()
                self.refresh_notes_list()
                self.show_toast("Заметка импортирована")
            except Exception as e:
                QMessageBox.critical(self, "Ошибка", f"Ошибка импорта: {e}")
                self.deduplicate_notes()

    def delete_note_completely(self, note: Note) -> None:
        reply = QMessageBox.question(
            self,
            "Удалить навсегда",
            f"Безвозвратно удалить заметку '{note.title}'? Данные будут удалены полностью.",
            QMessageBox.Yes | QMessageBox.No,
        )
        if reply != QMessageBox.Yes:
            return
        candidates = []
        try:
            name_with_ts = NotesApp.safe_folder_name(
                note.title, note.uuid, note.timestamp
            )
        except Exception:
            name_with_ts = NotesApp.safe_folder_name(note.title, note.uuid)
        name_no_ts = NotesApp.safe_folder_name(note.title, note.uuid)
        for base in (NOTES_DIR, self.TRASH_DIR):
            for fname in (name_with_ts, name_no_ts):
                p = os.path.join(base, fname)
                if os.path.exists(p):
                    candidates.append(p)
        try:
            if hasattr(self, "attachments_watcher") and self.attachments_watcher:
                watched = set(self.attachments_watcher.directories() or [])
                to_remove = [p for p in candidates if p in watched]
                if to_remove:
                    self.attachments_watcher.removePaths(list(set(to_remove)))
        except Exception:
            pass
        QApplication.processEvents()
        QThread.msleep(50)
        for p in candidates:
            if os.path.exists(p):
                shutil.rmtree(p)
        self.notes = [n for n in self.notes if n.uuid != note.uuid]
        if self.current_note and self.current_note.uuid == note.uuid:
            self.current_note = None
            self.text_edit.clear()
            if hasattr(self, "history_list"):
                self.history_list.clear()
            self.attachments_scroll.setVisible(False)

        self.refresh_notes_list()
        self.save_all_notes_to_disk()
        self.show_toast(f"«{note.title}» удалена безвозвратно")

    def trigger_search(self):
        self.handle_combined_search()

    def check_reminders(self):
        now = QDateTime.currentDateTime()
        changed = False
        for note in self.notes:
            if note.reminder and not getattr(note, "reminder_shown", False):
                reminder_time = QDateTime.fromString(note.reminder, Qt.ISODate)
                if not reminder_time.isValid():
                    continue
                if reminder_time <= now:
                    self.show_reminder_dialog(note)
                    changed = True
                    repeat = (note.reminder_repeat or "").lower()
                    if repeat in ("once", "", None):
                        note.reminder = None
                        note.reminder_repeat = None
                    elif repeat == "workdays":
                        next_dt = reminder_time.addDays(1)
                        while next_dt.date().dayOfWeek() > 5:
                            next_dt = next_dt.addDays(1)
                        note.reminder = next_dt.toString(Qt.ISODate)
                        note.reminder_shown = False
                    elif repeat == "weekends":
                        next_dt = reminder_time.addDays(1)
                        while next_dt.date().dayOfWeek() < 6:
                            next_dt = next_dt.addDays(1)
                        note.reminder = next_dt.toString(Qt.ISODate)
                        note.reminder_shown = False
                    elif repeat == "daily":
                        next_dt = reminder_time.addDays(1)
                        note.reminder = next_dt.toString(Qt.ISODate)
                        note.reminder_shown = False
                    elif repeat == "weekly":
                        next_dt = reminder_time.addDays(7)
                        note.reminder = next_dt.toString(Qt.ISODate)
                        note.reminder_shown = False
                    elif repeat == "minutely":
                        next_dt = reminder_time.addSecs(60)
                        note.reminder = next_dt.toString(Qt.ISODate)
                        note.reminder_shown = False
                    elif repeat == "hourly":
                        next_dt = reminder_time.addSecs(3600)
                        note.reminder = next_dt.toString(Qt.ISODate)
                        note.reminder_shown = False
                    else:
                        note.reminder = None
                        note.reminder_repeat = None

        if changed:
            self.save_all_notes_to_disk()
            self.refresh_notes_list()

    def show_reminder_dialog(self, note):
        msg = QMessageBox(self)
        msg.setIcon(QMessageBox.Information)
        msg.setWindowTitle("Напоминание")
        note_time = note.reminder or ""
        msg.setText(
            f"Напоминание по заметке:\n\n{note.title}\n\n"
            f"{QDateTime.fromString(note_time, Qt.ISODate).toString('dd.MM.yyyy HH:mm')}"
        )
        msg.setStandardButtons(QMessageBox.Ok)
        msg.setWindowFlag(Qt.WindowStaysOnTopHint, True)
        msg.exec()
        note.reminder_shown = True

    def handle_combined_search(self):
        tag = self.tag_filter.currentText().strip().lower()
        if tag == "все теги":
            tag = ""
        text = self.search_bar.text().strip().lower()
        mode = self.search_mode_combo.currentText()
        if not text and not tag:
            self.refresh_notes_list()
            return

        def _matches(note):
            matches_tag = tag in [t.lower() for t in note.tags] if tag else True
            if not text:
                return matches_tag
            if mode == "Заголовок":
                matches_search = text in note.title.lower()
            elif mode == "Напоминание":
                matches_search = (
                    bool(note.reminder) and text in str(note.reminder).lower()
                )
            else:
                doc = QTextDocument()
                doc.setHtml(note.content)
                matches_search = text in doc.toPlainText().lower()
            return matches_tag and matches_search

        filtered = [n for n in self.notes if _matches(n)]
        fav_color = self.get_contrast_favorite_color()
        current_uuid = self.current_note.uuid if self.current_note else None
        self.notes_list.clear()

        def _add_note(note):
            timestamp = QDateTime.fromString(note.timestamp, Qt.ISODate)
            date_str = timestamp.toString("dd.MM.yyyy")
            reminder_symbol = " 🔔" if note.reminder else ""
            item_text = f"{note.title} — {date_str}{reminder_symbol}"
            item = QListWidgetItem(item_text)
            item.setData(Qt.UserRole, note)
            item.setFont(QFont("Segoe UI Emoji", 10))
            if note.favorite:
                item.setForeground(fav_color)
            self.notes_list.addItem(item)

        pinned = [n for n in filtered if n.pinned]
        unpinned = [n for n in filtered if not n.pinned]

        for n in pinned:
            _add_note(n)
        if pinned and unpinned:
            sep_item = QListWidgetItem()
            sep_item.setFlags(Qt.NoItemFlags)
            line = QFrame()
            line.setFrameShape(QFrame.HLine)
            line.setFrameShadow(QFrame.Sunken)
            sep_item.setSizeHint(line.sizeHint())
            self.notes_list.addItem(sep_item)
            self.notes_list.setItemWidget(sep_item, line)
        for n in unpinned:
            _add_note(n)

        if current_uuid:
            for i in range(self.notes_list.count()):
                it = self.notes_list.item(i)
                n = it.data(Qt.UserRole)
                if n and n.uuid == current_uuid:
                    self.notes_list.setCurrentItem(it)
                    break

    def create_history_toolbar(self):
        toolbar = QToolBar("История версий", self)
        toolbar.setObjectName("history_toolbar")
        toolbar.addWidget(self.history_label)
        toolbar.addWidget(self.history_list)
        self.addToolBar(Qt.ToolBarArea.LeftToolBarArea, toolbar)

    def list_attachments_for_current_note(self):
        if not self.current_note:
            return
        note_dir = os.path.join(NOTES_DIR, self.current_note.uuid)
        if not os.path.exists(note_dir):
            return
        attachments = []
        for file_name in os.listdir(note_dir):
            if file_name != "note.json":
                attachments.append(file_name)
        attachment_links = "\n".join(
            f'<a href="file://{os.path.abspath(os.path.join(note_dir, f))}">{f}</a>'
            for f in attachments
        )
        if attachment_links:
            cursor = self.text_edit.textCursor()
            html = "<br>--- Attachments ---<br>" + attachment_links.replace(
                "\n", "<br>"
            )
            cursor.insertHtml(html)
            self.text_edit.setTextCursor(cursor)

    def _resize_editor_dock(self, delta_px: int) -> None:
        if not hasattr(self, "dock_editor") or self.dock_editor is None:
            return
        left_dock = getattr(self, "dock_notes_list", None)
        if left_dock is None:
            return
        editor_w = max(0, self.dock_editor.width())
        left_w = max(0, left_dock.width())
        total_w = max(1, editor_w + left_w)
        min_editor = 300
        max_editor = int(self.width() * 0.85)
        new_editor = max(min_editor, min(max_editor, editor_w + delta_px))
        new_left = max(80, total_w - new_editor)
        try:
            self.resizeDocks(
                [left_dock, self.dock_editor], [new_left, new_editor], Qt.Horizontal
            )
        except Exception:
            self.dock_editor.setMinimumWidth(new_editor)
            self.dock_editor.resize(new_editor, self.dock_editor.height())

    def update_tag_filter_items(self):
        all_tags = sorted(self.get_all_tags())
        self.tag_filter.clear()
        self.tag_filter.addItem("Все теги")
        self.tag_filter.addItems(all_tags)

    def update_search_filter_items(self) -> None:
        if not hasattr(self, "search_mode_combo") or self.search_mode_combo is None:
            return
        modes = ["Заголовок", "Содержимое", "Напоминание"]
        self.search_mode_combo.blockSignals(True)
        self.search_mode_combo.clear()
        self.search_mode_combo.addItems(modes)
        self.search_mode_combo.blockSignals(False)

    def set_reminder_for_note(self):
        if not self.current_note:
            return
        dialog = QDialog(self)
        dialog.setWindowFlag(Qt.WindowStaysOnTopHint, True)
        dialog.setWindowTitle("Установить напоминание")
        layout = QFormLayout(dialog)
        datetime_edit = QDateTimeEdit(QDateTime.currentDateTime())
        datetime_edit.setCalendarPopup(True)
        layout.addRow("Напоминание:", datetime_edit)
        repeat_combo = QComboBox()
        repeat_combo.addItems(
            ["Однократно", "По будням", "Каждый день", "Каждую неделю", "Каждый месяц"]
        )
        layout.addRow("Повторять:", repeat_combo)

        buttons = QDialogButtonBox(
            QDialogButtonBox.StandardButton.Ok | QDialogButtonBox.StandardButton.Cancel
        )
        buttons.accepted.connect(dialog.accept)
        buttons.rejected.connect(dialog.reject)
        layout.addWidget(buttons)
        if dialog.exec() == QDialog.DialogCode.Accepted:
            dt = datetime_edit.dateTime()
            self.current_note.reminder = dt.toString("yyyy-MM-dd HH:mm")
            self.save_note_to_file(self.current_note)
            self.refresh_notes_list()
            self.show_toast(f"Напоминание: {self.current_note.reminder}")
            repeat_option = repeat_combo.currentText()
            if repeat_option == "Однократно":
                self.current_note.reminder_repeat = None
            else:
                self.current_note.reminder_repeat = repeat_option

    def edit_reminder_for_note(self):
        if not self.current_note:
            QMessageBox.warning(self, "Нет заметки", "Сначала выберите заметку.")
            return
        dialog = QDialog(self)
        dialog.setWindowFlag(Qt.WindowStaysOnTopHint, True)
        dialog.setWindowTitle("Изменить напоминание")
        layout = QFormLayout(dialog)
        if self.current_note.reminder:
            dt = QDateTime.fromString(self.current_note.reminder, "yyyy-MM-dd HH:mm")
            if not dt.isValid():
                dt = QDateTime.fromString(self.current_note.reminder, Qt.ISODate)
            if not dt.isValid():
                dt = QDateTime.currentDateTime()
        else:
            dt = QDateTime.currentDateTime()
        datetime_edit = QDateTimeEdit(dt)
        datetime_edit.setCalendarPopup(True)
        layout.addRow("Напоминание:", datetime_edit)
        repeat_combo = QComboBox()
        options = [
            "Однократно",
            "По будням",
            "Каждый день",
            "Каждую неделю",
            "Каждый месяц",
        ]
        repeat_combo.addItems(options)
        if getattr(self.current_note, "reminder_repeat", None) in options:
            repeat_combo.setCurrentText(
                self.current_note.reminder_repeat or "Однократно"
            )
        else:
            repeat_combo.setCurrentIndex(0)
        layout.addRow("Повторять:", repeat_combo)
        buttons = QDialogButtonBox(QDialogButtonBox.Ok | QDialogButtonBox.Cancel)
        buttons.accepted.connect(dialog.accept)
        buttons.rejected.connect(dialog.reject)
        layout.addWidget(buttons)

        if dialog.exec() == QDialog.Accepted:
            dt = datetime_edit.dateTime()
            self.current_note.reminder = dt.toString("yyyy-MM-dd HH:mm")
            selected = repeat_combo.currentText()
            self.current_note.reminder_repeat = (
                None if selected == "Однократно" else selected
            )
            self.save_note_to_file(self.current_note)
            self.refresh_notes_list()
            self.show_toast("Напоминание обновлено")

    def remove_reminder_from_note(self):
        if self.current_note:
            self.current_note.reminder = None
            self.save_note_to_file(self.current_note)
            self.refresh_notes_list()
            self.show_toast("Напоминание удалено")

    def check_upcoming_reminders(self):
        now = QDateTime.currentDateTime()
        for note in self.notes:
            if note.reminder:
                reminder_dt = QDateTime.fromString(note.reminder, "yyyy-MM-dd HH:mm")
                if -60 <= now.secsTo(reminder_dt) <= 60:
                    self.tray_icon.showMessage(
                        "Напоминание",
                        f"Заметка: {note.title}",
                        QSystemTrayIcon.Information,
                        10000,
                    )
                    self.show_reminder_popup(note)
                    if note.reminder_repeat == "Каждый день":
                        note.reminder = reminder_dt.addDays(1).toString(
                            "yyyy-MM-dd HH:mm"
                        )
                        self.save_note_to_file(note)
                    elif note.reminder_repeat == "По будням":
                        if reminder_dt.date().dayOfWeek() < 6:
                            note.reminder = reminder_dt.addDays(1).toString(
                                "yyyy-MM-dd HH:mm"
                            )
                        else:
                            days_to_monday = 8 - reminder_dt.date().dayOfWeek()
                            note.reminder = reminder_dt.addDays(
                                days_to_monday
                            ).toString("yyyy-MM-dd HH:mm")
                        self.save_note_to_file(note)
                    elif note.reminder_repeat == "Каждую неделю":
                        note.reminder = reminder_dt.addDays(7).toString(
                            "yyyy-MM-dd HH:mm"
                        )
                        self.save_note_to_file(note)
                    elif note.reminder_repeat == "Каждый месяц":
                        note.reminder = reminder_dt.addMonths(1).toString(
                            "yyyy-MM-dd HH:mm"
                        )
                        self.save_note_to_file(note)
                    else:
                        note.reminder = None
                        self.save_note_to_file(note)
                elif now > reminder_dt.addSecs(60):
                    note.reminder = None
                    self.save_note_to_file(note)
                    self.refresh_notes_list()

    def show_reminder_popup(self, note):
        dlg = QDialog(self, Qt.WindowStaysOnTopHint)
        dlg.setWindowTitle("Напоминание")
        dlg.setWindowModality(Qt.NonModal)
        dlg.setMinimumWidth(300)
        layout = QVBoxLayout(dlg)
        layout.addWidget(QLabel(f"Напоминание по заметке: <b>{note.title}</b>"))
        layout.addWidget(QLabel(f"Время: {note.reminder}"))
        btn = QPushButton("ОК")
        btn.clicked.connect(dlg.accept)
        layout.addWidget(btn)
        dlg.show()

    def setup_reminder_timer(self):
        self.reminder_timer = QTimer(self)
        self.reminder_timer.timeout.connect(self.check_upcoming_reminders)
        self.reminder_timer.start(60000)

    def handle_note_reorder(self):
        new_order = []
        for i in range(self.notes_list.count()):
            item = self.notes_list.item(i)
            note = item.data(Qt.UserRole)
            new_order.append(note)
        self.notes = new_order
        self.save_all_notes_to_disk()

    def init_all_components(self):
        QShortcut(QKeySequence("Escape"), self.text_edit, self.exit_note)
        QShortcut(QKeySequence("Ctrl+B"), self.text_edit, self.toggle_bold)
        QShortcut(QKeySequence("Ctrl+I"), self.text_edit, self.toggle_italic)
        QShortcut(QKeySequence("Ctrl+U"), self.text_edit, self.toggle_underline)
        QShortcut(QKeySequence("Ctrl+N"), self, self.create_new_note)
        QShortcut(QKeySequence("Ctrl+S"), self, self.save_note)
        QShortcut(QKeySequence("Ctrl+Space"), self.text_edit, self.clear_formatting)
        QShortcut(
            QKeySequence("Ctrl+Shift+D"), self, activated=self.open_drawing_dialog
        )
        self._notes_delete_sc = QShortcut(QKeySequence.Delete, self.notes_list)
        self._notes_delete_sc.setContext(Qt.WidgetWithChildrenShortcut)
        self._notes_delete_sc.activated.connect(self.delete_note)
        self.load_notes_from_disk()
        self.refresh_notes_list()
        last_uuid = self.settings.value("lastNoteUUID", "")
        if last_uuid:
            for i in range(self.notes_list.count()):
                it = self.notes_list.item(i)
                n = it.data(Qt.UserRole)
                if n and getattr(n, "uuid", None) == last_uuid:
                    self.notes_list.setCurrentRow(i)
                    self.select_note(n)
                    break
        self.update_tag_filter_items()
        self.update_search_filter_items()
        self.setup_reminder_timer()
        self.setAcceptDrops(True)
        self.autosave_enabled = self.settings.value("autosave_enabled", True, type=bool)
        self.autosave_interval = self.settings.value(
            "autosave_interval", 60000, type=int
        )
        self.autosave_timer = QTimer(self)
        self.autosave_timer.timeout.connect(self.autosave_current_note)
        self._edit_session_note_uuid = None

    def _suspend_edit_signals(self):
        try:
            self._is_switching_note = True
            if hasattr(self, "debounce_timer"):
                self.debounce_timer.stop()
            self.text_edit.blockSignals(True)
        except Exception:
            pass

    def _resume_edit_signals(self):
        try:
            self.text_edit.blockSignals(False)
        finally:
            self._is_switching_note = False

    def init_theme(self):
        theme = self.settings.value("theme", "dark")
        if theme == "dark":
            self.apply_dark_theme()
        else:
            self.apply_light_theme()

    def apply_dark_theme(self):
        self.setStyle(QStyleFactory.create("Fusion"))
        dark_palette = QPalette()
        dark_palette.setColor(QPalette.ColorRole.Window, QColor(53, 53, 53))
        dark_palette.setColor(QPalette.ColorRole.WindowText, Qt.GlobalColor.white)
        dark_palette.setColor(QPalette.ColorRole.Base, QColor(42, 42, 42))
        dark_palette.setColor(QPalette.ColorRole.AlternateBase, QColor(66, 66, 66))
        dark_palette.setColor(QPalette.ColorRole.ToolTipBase, Qt.GlobalColor.white)
        dark_palette.setColor(QPalette.ColorRole.ToolTipText, Qt.GlobalColor.white)
        dark_palette.setColor(QPalette.ColorRole.Text, Qt.GlobalColor.white)
        dark_palette.setColor(QPalette.ColorRole.Button, QColor(53, 53, 53))
        dark_palette.setColor(QPalette.ColorRole.ButtonText, Qt.GlobalColor.white)
        dark_palette.setColor(QPalette.ColorRole.BrightText, Qt.GlobalColor.red)
        dark_palette.setColor(
            QPalette.ColorRole.Highlight, QColor(142, 45, 197).lighter()
        )
        dark_palette.setColor(QPalette.ColorRole.HighlightedText, Qt.GlobalColor.black)
        dark_palette.setColor(QPalette.ColorRole.Link, Qt.GlobalColor.white)
        dark_palette.setColor(QPalette.ColorRole.LinkVisited, Qt.GlobalColor.white)
        self.notes_list.setStyleSheet("color: white; background-color: #2b2b2b;")
        self.text_edit.setFont(QFont("Times New Roman", 14))
        self.text_edit.setStyleSheet(
            "font-size: 14px; font-family: 'Times New Roman'; color: white; background-color: #2b2b2b;"
        )
        self.text_edit.document().setDefaultStyleSheet("a { color: white; }")
        QApplication.instance().setPalette(dark_palette)
        self.setStyleSheet(
            """
            QToolTip {
                background-color: #2a2a2a;
                color: white;
                border: 1px solid white;
                padding: 5px;
                font-size: 12px;
            }
            QCheckBox { color: white; }
        """
        )
        self.rebuild_toolbar()

    def show_toast(
        self,
        text: str,
        ms: int = 1200,
        *,
        boundary_widget: QWidget | None = None,
        anchor_widget: QWidget | None = None,
    ) -> None:
        overlay_parent = boundary_widget or self
        lbl = QLabel(text, overlay_parent)
        lbl.setObjectName("toast")
        lbl.setAttribute(Qt.WA_TransparentForMouseEvents)
        lbl.setStyleSheet(
            """
            QLabel {
                background: #ffffff;
                color: #000000;
                border: 1px solid #9a9a9a;
                border-radius: 6px;
                padding: 4px 8px;
                font-size: 12px;
            }
        """
        )
        lbl.setWordWrap(False)
        lbl.adjustSize()
        if anchor_widget is not None:
            anchor_center = anchor_widget.mapTo(
                overlay_parent, anchor_widget.rect().center()
            )
            x = anchor_center.x() - lbl.width() - 8
            y = anchor_center.y() - lbl.height() // 2
        else:
            x = overlay_parent.width() - lbl.width() - 8
            y = overlay_parent.height() - lbl.height() - 8
        stack_idx = sum(1 for t in self._live_toasts if t.parent() is overlay_parent)
        y -= stack_idx * (lbl.height() + 6)
        x = max(6, min(x, overlay_parent.width() - lbl.width() - 6))
        y = max(6, min(y, overlay_parent.height() - lbl.height() - 6))
        lbl.move(x, y)
        lbl.raise_()
        lbl.show()
        self._live_toasts.append(lbl)

        def _close():
            try:
                lbl.close()
            finally:
                if lbl in self._live_toasts:
                    self._live_toasts.remove(lbl)
                lbl.deleteLater()

        QTimer.singleShot(ms, _close)

    def exit_note(self):
        self.load_note(None)
        self.current_note = None
        self.notes_list.clearSelection()
        self._update_editor_visibility()

    def apply_light_theme(self):
        self.setStyle(QStyleFactory.create("Fusion"))
        default_palette = QApplication.style().standardPalette()
        default_palette.setColor(QPalette.ColorRole.Link, Qt.GlobalColor.black)
        default_palette.setColor(QPalette.ColorRole.LinkVisited, Qt.GlobalColor.black)
        QApplication.instance().setPalette(default_palette)
        self.notes_list.setStyleSheet("color: black; background-color: white;")
        self.text_edit.setFont(QFont("Times New Roman", 14))
        self.text_edit.setStyleSheet(
            "font-size: 14px; font-family: 'Times New Roman'; color: black; background-color: white;"
        )
        self.text_edit.document().setDefaultStyleSheet("a { color: black; }")
        self.new_note_button.setStyleSheet("")
        self.save_note_button.setStyleSheet("")
        self.setStyleSheet(
            """
            QToolTip {
                background-color: #ffffff;
                color: black;
                border: 1px solid black;
                padding: 5px;
                font-size: 12px;
            }
            QCheckBox { color: black; }
        """
        )
        self.delete_note_button.setStyleSheet("")
        self.audio_button.setStyleSheet("")
        self.menuBar().setStyleSheet("")
        self.rebuild_toolbar()

    def autosave_current_note(self):
        if not getattr(self, "autosave_enabled", True):
            return
        if not getattr(self, "current_note", None):
            return
        if getattr(self, "_edit_session_note_uuid", None) and \
        self.current_note and self.current_note.uuid != self._edit_session_note_uuid:
            return
        self.current_note.content = self.text_edit.toHtml()
        self.current_note.password_manager = self.password_manager_field.text()
        self.current_note.rdp_1c8 = self.rdp_1c8_field.text()
        self.update_current_note_custom_fields()
        self.save_note_to_file(self.current_note)
        self._set_unsaved(False)

    def eventFilter(self, source, event):
        if (
            isinstance(source, QLineEdit)
            and source.property("renameTitleEditor")
            and event.type() == QEvent.KeyPress
        ):
            if source.hasSelectedText() and len(source.selectedText()) == len(
                source.text()
            ):
                if event.key() == Qt.Key_Left:
                    source.setCursorPosition(0)
                    source.setSelection(0, 0)
                    return True
                if event.key() == Qt.Key_Right:
                    source.setCursorPosition(len(source.text()))
                    source.setSelection(0, 0)
                    return True
        if source is self.text_edit and event.type() == QEvent.KeyPress:
            if event.matches(QKeySequence.Undo):
                self.undo()
                return True
            elif event.matches(QKeySequence.Redo):
                self.redo()
                return True
        return super().eventFilter(source, event)

    def dragEnterEvent(self, event):
        if event.mimeData().hasUrls():
            event.acceptProposedAction()

    def dropEvent(self, event):
        if not self.current_note:
            return
        for url in event.mimeData().urls():
            file_path = url.toLocalFile()
            if os.path.isfile(file_path):
                note_dir = os.path.join(
                    NOTES_DIR,
                    NotesApp.safe_folder_name(
                        self.current_note.title,
                        self.current_note.uuid,
                        self.current_note.timestamp,
                    ),
                )
                if not os.path.exists(note_dir):
                    os.makedirs(note_dir)
                filename = os.path.basename(file_path)
                dest = os.path.join(note_dir, filename)
                shutil.copy(file_path, dest)
                if filename.lower().endswith(tuple(IMAGE_EXTENSIONS)):
                    pixmap = QPixmap(dest)
                    if not pixmap.isNull():
                        buffer = QBuffer()
                        buffer.open(QIODevice.WriteOnly)
                        pixmap.save(buffer, "PNG")
                        base64_data = base64.b64encode(buffer.data()).decode("utf-8")
                        html_img = f'<img src="Data:image/png;base64,{base64_data}" width="200"><br>'
                        cursor = self.text_edit.textCursor()
                        cursor.movePosition(QTextCursor.End)
                        self.text_edit.setTextCursor(cursor)
                        self.text_edit.insertHtml(html_img)
                        self.record_state_for_undo()
                else:
                    file_url = QUrl.fromLocalFile(os.path.abspath(dest)).toString()
                    cursor = self.text_edit.textCursor()
                    cursor.movePosition(QTextCursor.End)
                    self.text_edit.setTextCursor(cursor)
                    self.text_edit.insertHtml(f'📄 <a href="{file_url}">{filename}</a>')
        try:
            self.show_toast("Файлы прикреплены", timeout_ms=1400,
                            anchor_widget=getattr(self, "attachments_scroll", None))
        except Exception:
            pass
        self.save_note()

    def closeEvent(self, event: QCloseEvent):
        self.save_settings()
        try:
            self.save_note_quiet(force=True)
            if getattr(self, "current_note", None):
                self.save_note_quiet(force=True)
            self.save_all_notes_to_disk()
        except Exception as e:
            print("Save on close failed:", e)
        event.ignore()
        self.hide()
        self.tray_icon.show()
        self.window_hidden.emit()

    @staticmethod
    def safe_folder_name(title, note_uuid, timestamp=None):
        safe_title = re.sub(r"[^\w\-_ ]", "_", title)[:40].strip()
        if timestamp is None:
            date_str = QDateTime.currentDateTime().toString("dd.MM.yyyy")
        else:
            if isinstance(timestamp, str):
                dt = QDateTime.fromString(timestamp, Qt.ISODate)
            else:
                dt = timestamp
            date_str = dt.toString("dd.MM.yyyy")
        uuid_short = (
            note_uuid if note_uuid.startswith("UUID_") else f"UUID({note_uuid[:8]})"
        )
        return f"{safe_title} от {date_str}, {uuid_short}"

    def manage_plugins_dialog(self):
        plugins_folder = os.path.join(APPDIR, "Plugins")
        os.makedirs(plugins_folder, exist_ok=True)
        plugins_state_path = os.path.join(DATA_DIR, "plugins_state.json")
        plugins = []
        if os.path.exists(plugins_folder):
            for fname in os.listdir(plugins_folder):
                if fname.endswith(".py"):
                    plugins.append(fname[:-3])
        if os.path.exists(plugins_state_path):
            with open(plugins_state_path, "r", encoding="utf-8") as f:
                plugins_state = json.load(f)
        else:
            plugins_state = {}

        dialog = QDialog(self)
        dialog.setWindowTitle("Управление плагинами")
        dialog.setWindowFlag(Qt.WindowStaysOnTopHint, True)
        layout = QVBoxLayout(dialog)
        checkboxes = {}
        for plugin in sorted(plugins):
            cb = QCheckBox(plugin)
            cb.setChecked(plugins_state.get(plugin, False))
            layout.addWidget(cb)
            checkboxes[plugin] = cb

        btn_add = QPushButton("Добавить плагин")
        btn_del = QPushButton("Удалить плагин")
        btn_rename = QPushButton("Переименовать плагин")
        btns_h = QHBoxLayout()
        btns_h.addWidget(btn_add)
        btns_h.addWidget(btn_del)
        btns_h.addWidget(btn_rename)
        layout.addLayout(btns_h)

        def add_plugin():
            file_path, _ = QFileDialog.getOpenFileName(
                dialog, "Выберите .py файл плагина", "", "Python Files (*.py)"
            )
            if file_path:
                dest_path = os.path.join(plugins_folder, os.path.basename(file_path))
                if os.path.exists(dest_path):
                    QMessageBox.warning(
                        dialog, "Ошибка", "Плагин с таким именем уже существует."
                    )
                    return
                shutil.copy(file_path, dest_path)
                dialog.close()
                self.manage_plugins_dialog()

        def del_plugin():
            names = [name for name, cb in checkboxes.items() if cb.isChecked()]
            if not names:
                QMessageBox.warning(
                    dialog,
                    "Нет выбора",
                    "Отметьте чекбокс напротив плагина для удаления.",
                )
                return
            for fname in names:
                path = os.path.join(plugins_folder, fname + ".py")
                if os.path.exists(path):
                    reply = QMessageBox.question(
                        dialog,
                        "Удалить плагин",
                        f"Удалить плагин {fname}?",
                        QMessageBox.Yes | QMessageBox.No,
                    )
                    if reply == QMessageBox.Yes:
                        os.remove(path)
                        plugins_state.pop(fname, None)
            dialog.close()
            self.manage_plugins_dialog()

        def rename_plugin():
            names = [name for name, cb in checkboxes.items() if cb.isChecked()]
            if not names:
                QMessageBox.warning(
                    dialog,
                    "Нет выбора",
                    "Отметьте чекбокс напротив плагина для переименования.",
                )
                return
            old_name = names[0]
            new_name, ok = QInputDialog.getText(
                dialog,
                "Переименовать плагин",
                "Новое имя файла (без .py):",
                text=old_name,
            )
            if ok and new_name and new_name != old_name:
                old_path = os.path.join(plugins_folder, old_name + ".py")
                new_path = os.path.join(plugins_folder, new_name + ".py")
                if os.path.exists(new_path):
                    QMessageBox.warning(
                        dialog, "Ошибка", "Файл с таким именем уже существует."
                    )
                    return
                os.rename(old_path, new_path)
                plugins_state[new_name] = plugins_state.pop(old_name)
                dialog.close()
                self.manage_plugins_dialog()

        btn_add.clicked.connect(add_plugin)
        btn_del.clicked.connect(del_plugin)
        btn_rename.clicked.connect(rename_plugin)
        buttons = QDialogButtonBox(QDialogButtonBox.Ok | QDialogButtonBox.Cancel)
        layout.addWidget(buttons)

        def on_accept():
            for plugin, cb in checkboxes.items():
                plugins_state[plugin] = cb.isChecked()
            with open(plugins_state_path, "w", encoding="utf-8") as f:
                json.dump(plugins_state, f, ensure_ascii=False, indent=2)
            dialog.accept()
            self.load_plugins()

        buttons.accepted.connect(on_accept)
        buttons.rejected.connect(dialog.reject)
        dialog.setLayout(layout)
        dialog.exec()

    def clear_plugin_menu_actions(self):
        menu_bar = self.menuBar()
        plugins_menu = None
        for menu in menu_bar.children():
            try:
                if menu.title() == "Плагины":
                    plugins_menu = menu
                    break
            except Exception:
                continue
        if plugins_menu:
            static_actions = {"Управление плагинами", "Перезагрузить плагины"}
            actions_to_remove = []
            for action in plugins_menu.actions():
                if action.text() not in static_actions:
                    actions_to_remove.append(action)
            for action in actions_to_remove:
                plugins_menu.removeAction(action)

    def load_plugins(self):
        import importlib

        importlib.invalidate_caches()
        self.clear_plugin_menu_actions()
        for module in getattr(self, "active_plugins", {}).values():
            try:
                if hasattr(module, "on_disable"):
                    module.on_disable(parent=self)
            except Exception as e:
                print(
                    f"Ошибка при деинициализации плагина {getattr(module, '__name__', '')}: {e}"
                )
        plugins_folder = os.path.join(APPDIR, "Plugins")
        plugins_state_path = os.path.join(DATA_DIR, "plugins_state.json")
        if os.path.exists(plugins_state_path):
            with open(plugins_state_path, "r", encoding="utf-8") as f:
                plugins_state = json.load(f)
        else:
            plugins_state = {}
        self.active_plugins = {}
        if not os.path.exists(plugins_folder):
            return
        for fname in os.listdir(plugins_folder):
            if fname.endswith(".py"):
                plugin_name = fname[:-3]
                plugin_path = os.path.join(plugins_folder, fname)
                try:
                    spec = importlib.util.spec_from_file_location(
                        plugin_name, plugin_path
                    )
                    module = importlib.util.module_from_spec(spec)
                    spec.loader.exec_module(module)
                    self.active_plugins[plugin_name] = module
                    if plugins_state.get(plugin_name, False):
                        if hasattr(module, "on_enable"):
                            module.on_enable(parent=self)
                    else:
                        if hasattr(module, "on_disable"):
                            module.on_disable(parent=self)
                except Exception as e:
                    print(f"Ошибка при загрузке плагина {fname}:", e)

    def detect_plugins(self):
        plugins = []
        plugins_folder = os.path.join(APPDIR, "Plugins")
        if os.path.exists(plugins_folder):
            for file in os.listdir(plugins_folder):
                if file.endswith(".py"):
                    plugins.append(file[:-3])
        return plugins

    def load_plugins_state(self):
        plugins_path = os.path.join(DATA_DIR, "plugins_state.json")
        if os.path.exists(plugins_path):
            with open(plugins_path, "r", encoding="utf-8") as f:
                self.plugins_state = json.load(f)
        else:
            self.plugins_state = {}

    def save_plugins_state(self):
        plugins_path = os.path.join(DATA_DIR, "plugins_state.json")
        with open(plugins_path, "w", encoding="utf-8") as f:
            json.dump(self.plugins_state, f, ensure_ascii=False, indent=2)

    def update_current_note_content(self):
        if not hasattr(self, "current_note") or self.current_note is None:
            return
        if not hasattr(self, "current_note") or self.current_note is None:
            return
        html = self.text_edit.toHtml()
        html = self._persist_dropdown_values_in_html(html)
        self.current_note.content = html
        self.record_state_for_undo()

    def hideEvent(self, event):
        self.save_settings()
        super().hideEvent(event)

    def show_from_tray(self):
        self.bring_widget_to_front(self)

    def show_app_launcher(self):
        app = QApplication.instance()
        launcher = getattr(app, "launcher_window", None)
        if launcher is None:
            launcher = LauncherWindow()
            app.launcher_window = launcher
        launcher.show()
        launcher.raise_()
        launcher.activateWindow()

    def on_tray_icon_activated(self, reason: QSystemTrayIcon.ActivationReason) -> None:
        if reason == QSystemTrayIcon.Trigger:
            self.show_from_tray()

    def resizeEvent(self, event):
        try:
            self._apply_dock_proportions()
        except Exception:
            pass
        self.save_settings()
        super().resizeEvent(event)


    def moveEvent(self, event):
        self.save_settings()
        super().moveEvent(event)

    def toggle_autosave(self, state: int):
        self.autosave_enabled = bool(state)
        if self.autosave_enabled:
            try:
                self.text_edit.textChanged.connect(self.on_text_changed_autosave)
            except TypeError:
                pass
        else:
            try:
                self.text_edit.textChanged.disconnect(self.on_text_changed_autosave)
            except TypeError:
                pass
    def on_text_changed_autosave(self):
        pass

    def load_plugins(app, plugins_folder="Plugins"):
        if not os.path.exists(plugins_folder):
            os.makedirs(plugins_folder)
            return
        for filename in os.listdir(plugins_folder):
            if filename.endswith(".py"):
                plugin_path = os.path.join(plugins_folder, filename)
                spec = importlib.util.spec_from_file_location(filename[:-3], plugin_path)
                if spec is None:
                    continue
                plugin = importlib.util.module_from_spec(spec)
                try:
                    spec.loader.exec_module(plugin)
                    if hasattr(plugin, "register_plugin"):
                        plugin.register_plugin(app)
                except Exception as e:
                    print(f"Ошибка при загрузке плагина {filename}: {e}")


    def get_app_dir():
        return APPDIR


def create_default_config():
    config_path = os.path.join(PASSWORDS_DIR, "config.json")
    if not os.path.exists(config_path):
        default_config = {
            "salt_file": "salt.bin",
            "max_password_length": 32,
            "default_length": 18,
            "use_uppercase": True,
            "use_lowercase": True,
            "use_digits": True,
            "use_symbols": True,
            "excluded_chars": "O0DQl1I|i!S5Z2B8G6CGceaouvwxX",
            "passwords_file": "passwords.json",
        }
        try:
            with open(config_path, "w", encoding="utf-8") as f:
                json.dump(default_config, f, indent=4)
        except Exception as e:
            print(f"Ошибка создания конфигурации: {str(e)}")


class PasswordGenerator:
    def __init__(self):
        self.password_length = 15
        self.use_uppercase = True
        self.use_lowercase = True
        self.use_digits = True
        self.use_symbols = True
        self.excluded_chars = ""
        self.load_config()

    def load_config(self):
        config_path = os.path.join(PASSWORDS_DIR, "config.json")
        if os.path.exists(config_path):
            try:
                with open(config_path, "r", encoding="utf-8") as f:
                    config = json.load(f)
                    if "max_password_length" in config:
                        self.max_password_length = config["max_password_length"]
                    if "default_length" in config:
                        self.password_length = config["default_length"]
                    if "use_uppercase" in config:
                        self.use_uppercase = config["use_uppercase"]
                    if "use_lowercase" in config:
                        self.use_lowercase = config["use_lowercase"]
                    if "use_digits" in config:
                        self.use_digits = config["use_digits"]
                    if "use_symbols" in config:
                        self.use_symbols = config["use_symbols"]
                    if "excluded_chars" in config:
                        self.excluded_chars = config["excluded_chars"]
            except Exception as e:
                print(f"Ошибка загрузки конфигурации: {str(e)}")
        else:
            self.max_password_length = 32
            create_default_config()
            self.load_config()

    def generate_password(self):
        char_sets: list[list[str]] = []
        if self.use_uppercase:
            char_sets.append(list(string.ascii_uppercase))
        if self.use_lowercase:
            char_sets.append(list(string.ascii_lowercase))
        if self.use_digits:
            char_sets.append(list(string.digits))
        if self.use_symbols:
            char_sets.append(list(string.punctuation))
        if not char_sets:
            return ""
        charset = [c for subset in char_sets for c in subset]
        charset = [c for c in charset if c not in self.excluded_chars]
        if not charset:
            return ""
        password_chars: list[str] = []
        for char_set in char_sets:
            valid_chars = [c for c in char_set if c not in self.excluded_chars]
            if valid_chars:
                password_chars.append(random.choice(valid_chars))
        remaining_length = max(0, self.password_length - len(password_chars))
        if remaining_length > 0:
            password_chars += random.choices(charset, k=remaining_length)
        random.shuffle(password_chars)
        return "".join(password_chars)

    def evaluate_password_strength(self, password):
        score = 0
        if len(password) >= 8:
            score += 10
        if len(password) >= 12:
            score += 10
        if len(password) >= 16:
            score += 10
        if re.search(r"[A-Z]", password):
            score += 10
        if re.search(r"[a-z]", password):
            score += 10
        if re.search(r"[0-9]", password):
            score += 10
        if re.search(r"[^A-Za-z0-9]", password):
            score += 10
        if not re.search(r"(.)\1\1", password):
            score += 10
        unique_chars = len(set(password))
        score += min(20, unique_chars * 2)
        if re.search(r"[A-Z].*[0-9]|[0-9].*[A-Z]", password) and re.search(
            r"[a-z].*[0-9]|[0-9].*[a-z]", password
        ):
            score += 10
        return min(100, score)

class NumberGenerator:
    def __init__(self, start: int = 0, end: int = 100):
        self.start = start
        self.end = end

    def generate(self) -> int:
        if self.start > self.end:
            self.start, self.end = self.end, self.start
        return random.randint(self.start, self.end)


class PasswordManager:
    def __init__(self, master_password=None):
        self.master_password = master_password
        self.passwords = []
        self.key = None
        self.salt = None
        self._create_default_config()
        self.load_config()
        if self.master_password:
            self._initialize_encryption()
        self.load_passwords()

    def backup_passwords(self):
        backup_dir = os.path.join(PASSWORDS_DIR, "backups")
        backup_filename = os.path.join(backup_dir, "backup_passwords.bin")
        os.makedirs(backup_dir, exist_ok=True)
        try:
            data = json.dumps(self.passwords, indent=4, ensure_ascii=False)
            with open(backup_filename, "w", encoding="utf-8") as backup_file:
                backup_file.write(data)
            return True, f"Резервная копия создана: {backup_filename}"
        except Exception as e:
            return False, f"Ошибка резервного копирования: {str(e)}"

    def move_password(self, from_idx: int, to_idx: int) -> bool:
        n = len(self.passwords)
        if n == 0:
            return False
        to_idx = max(0, min(to_idx, n - 1))
        if not (0 <= from_idx < n) or from_idx == to_idx:
            return False
        item = self.passwords.pop(from_idx)
        self.passwords.insert(to_idx, item)
        self._save_passwords()
        return True

    def regenerate_salt(self, master_password):
        new_salt = os.urandom(16)
        backup_passwords = self.get_all_passwords()
        try:
            with open(SALT_PATH, "wb") as salt_file:
                salt_file.write(new_salt)
            self.salt = new_salt
            kdf = PBKDF2HMAC(
                algorithm=hashes.SHA256(),
                length=32,
                salt=self.salt,
                iterations=100000,
                backend=default_backend(),
            )
            key = kdf.derive(master_password.encode())
            self.key = base64.urlsafe_b64encode(key)
            self.passwords = []
            for item in backup_passwords:
                self.add_password(
                    item["password"],
                    item["description"],
                    item.get("tags", []),
                    item.get("url", ""),
                )
            return True, "Соль успешно обновлена. Все пароли перешифрованы."
        except Exception as e:
            self.passwords = backup_passwords
            return False, f"Ошибка обновления соли: {str(e)}"

    def _initialize_encryption(self):
        if not os.path.exists(SALT_PATH):
            self.salt = os.urandom(16)
            with open(SALT_PATH, "wb") as salt_file:
                salt_file.write(self.salt)
        else:
            with open(SALT_PATH, "rb") as salt_file:
                self.salt = salt_file.read()
        kdf = PBKDF2HMAC(
            algorithm=hashes.SHA256(),
            length=32,
            salt=self.salt,
            iterations=100000,
            backend=default_backend(),
        )
        key = kdf.derive(self.master_password.encode())
        self.key = base64.urlsafe_b64encode(key)

    def encrypt(self, data):
        fernet = Fernet(self.key)
        return fernet.encrypt(data.encode()).decode()

    def decrypt(self, encrypted_data):
        fernet = Fernet(self.key)
        return fernet.decrypt(encrypted_data.encode()).decode()

    def load_config(self):
        config_path = os.path.join(PASSWORDS_DIR, "config.json")
        if os.path.exists(config_path):
            try:
                with open(config_path, "r", encoding="utf-8") as f:
                    config = json.load(f)
                    filename = config.get("passwords_file", "passwords.json")
                    if not os.path.isabs(filename):
                        filename = os.path.join(PASSWORDS_DIR, filename)
                    self.filename = filename
            except Exception as e:
                self.filename = os.path.join(PASSWORDS_DIR, "passwords.json")
        else:
            self.filename = os.path.join(PASSWORDS_DIR, "passwords.json")

    def add_password(self, password, description, tags=None, url=None, login=""):
        encrypted_password = self.encrypt(password)
        self.passwords.append(
            {
                "password": encrypted_password,
                "description": description.strip(),
                "tags": tags or [],
                "url": url.strip() if url else "",
                "login": login.strip(),
                "encrypted": True,
                "updated_at": datetime.datetime.now().strftime("%d.%m.%Y %H:%M"),
            }
        )
        self._save_passwords()
        return True

    def _create_default_config(self):
        config_path = os.path.join(PASSWORDS_DIR, "config.json")
        if not os.path.exists(config_path):
            default_config = {
                "max_password_length": 32,
                "default_length": 15,
                "use_uppercase": True,
                "use_lowercase": True,
                "use_digits": True,
                "use_symbols": True,
                "excluded_chars": "O0oQl1I|i!S5Z2B8G6CQDceaouvwxX",
                "passwords_file": "passwords.json",
            }
            try:
                with open(config_path, "w", encoding="utf-8") as f:
                    json.dump(default_config, f, indent=4)
            except Exception as e:
                print(f"Ошибка создания конфигурации: {str(e)}")

    def update_password(
        self, index, password, description, tags=None, url=None, login=""
    ):
        if 0 <= index < len(self.passwords):
            encrypted_password = self.encrypt(password)
            self.passwords[index] = {
                "password": encrypted_password,
                "description": description,
                "tags": tags or [],
                "url": url.strip() if url else "",
                "login": login.strip(),
                "encrypted": True,
                "updated_at": datetime.datetime.now().strftime("%d.%m.%Y %H:%M"),
            }
            self._save_passwords()
            return True
        return False

    def export_to_txt(self, filename):
        try:
            decrypted_passwords = self.get_all_passwords()
            with open(filename, "w", encoding="utf-8") as f:
                for pwd in decrypted_passwords:
                    f.write(f"Описание: {pwd['description']}\n")
                    f.write(f"Пароль: {pwd['password']}\n")
                    f.write(f"Теги: {', '.join(pwd['tags'])}\n")
                    f.write(f"URL: {pwd.get('url', '')}\n")
                    f.write("-" * 40 + "\n")
            return True, "Пароли успешно экспортированы"
        except Exception as e:
            return False, f"Ошибка экспорта: {str(e)}"

    def import_from_txt(self, filename):
        try:
            with open(filename, "r", encoding="utf-8") as f:
                content = f.read().strip()
            password_blocks = []
            current_block = []
            for line in content.split("\n"):
                line = line.strip()
                if line.startswith("-----") or line == "":
                    if current_block:
                        password_blocks.append("\n".join(current_block))
                        current_block = []
                else:
                    current_block.append(line)
            if current_block:
                password_blocks.append("\n".join(current_block))
            imported_data = []
            for block in password_blocks:
                password_info = {"description": "", "password": "", "tags": []}
                lines = block.strip().split("\n")
                for line in lines:
                    line = line.strip()
                    if not line:
                        continue
                    if ":" in line:
                        key_part, value_part = line.split(":", 1)
                        key = key_part.strip().lower()
                        value = value_part.strip()
                        if key.startswith("опис"):
                            password_info["description"] = value
                        elif key.startswith("парол"):
                            password_info["password"] = value
                        elif key.startswith("url"):
                            password_info["url"] = value
                        elif key.startswith("тег"):
                            tags = [
                                t.strip() for t in re.split(r"[,;]", value) if t.strip()
                            ]
                            password_info["tags"] = tags
                if password_info["description"] and password_info["password"]:
                    imported_data.append(password_info)
                else:
                    print(f"Пропущен неполный блок: {block[:50]}...")
            if not imported_data:
                return False, "Не удалось найти валидные пароли в файле"
            imported_count = 0
            existing_descriptions = {p["description"] for p in self.passwords}
            for item in imported_data:
                if item["description"] in existing_descriptions:
                    print(f"Пропущен дубликат: {item['description']}")
                    continue
                encrypted_password = self.encrypt(item["password"])
                self.passwords.append(
                    {
                        "password": encrypted_password,
                        "description": item["description"],
                        "tags": item.get("tags", []),
                        "url": item.get("url", ""),
                        "encrypted": True,
                    }
                )
                imported_count += 1
            self._save_passwords()
            return True, f"Успешно импортировано {imported_count} паролей"
        except Exception as e:
            return False, f"Ошибка импорта: {str(e)}"

    def delete_password(self, index):
        if 0 <= index < len(self.passwords):
            del self.passwords[index]
            self._save_passwords()
            return True
        return False

    def get_password(self, index):
        if 0 <= index < len(self.passwords):
            item = self.passwords[index]
            try:
                return {
                    "password": self.decrypt(item["password"]),
                    "description": item["description"],
                    "tags": item.get("tags", []),
                    "url": item.get("url", ""),
                    "login": item.get("login", ""),
                }
            except:
                return {"error": "Неверный мастер-пароль для расшифровки"}
        return None

    def get_all_passwords(self):
        for p in self.passwords:
            if not p.get("updated_at"):
                p["updated_at"] = datetime.datetime.now().strftime("%d.%m.%Y %H:%M")
        return [
            {
                "password": self.decrypt(p["password"]),
                "description": p["description"],
                "tags": p.get("tags", []),
                "url": p.get("url", ""),
                "login": p.get("login", ""),
                "updated_at": p.get("updated_at", ""),
            }
            for p in self.passwords
        ]

    def get_all_metadata(self):
        for p in self.passwords:
            if not p.get("updated_at"):
                p["updated_at"] = datetime.datetime.now().strftime("%d.%m.%Y %H:%M")
        return [
            {
                "index": i,
                "description": p["description"],
                "tags": p.get("tags", []),
                "url": p.get("url", ""),
                "updated_at": p.get("updated_at", ""),
            }
            for i, p in enumerate(self.passwords)
        ]

    def load_passwords(self):
        if os.path.exists(self.filename):
            try:
                with open(self.filename, "r", encoding="utf-8") as f:
                    self.passwords = json.load(f)
                if self.passwords:
                    test_item = self.passwords[0]
                    self.decrypt(test_item["password"])
                if self.key and any(
                    not p.get("encrypted", False) for p in self.passwords
                ):
                    self._migrate_old_passwords()
            except (InvalidToken, json.JSONDecodeError):
                self.passwords = []
            except Exception:
                self.passwords = []

    def _save_passwords(self):
        data = [
            {
                "login": p.get("login", ""),
                "password": p["password"],
                "description": p["description"],
                "tags": p.get("tags", []),
                "url": p.get("url", ""),
                "encrypted": True,
                "updated_at": p.get("updated_at", ""),
            }
            for p in self.passwords
        ]
        try:
            with open(self.filename, "w", encoding="utf-8") as f:
                json.dump(data, f, indent=4, ensure_ascii=False)
                self.backup_passwords()
        except Exception as e:
            messagebox.showerror("Ошибка", f"Ошибка сохранения: {str(e)}")

    def _migrate_old_passwords(self):
        migrated = False
        for item in self.passwords:
            if not item.get("encrypted", False):
                try:
                    item["password"] = self.encrypt(item["password"])
                    item["encrypted"] = True
                    migrated = True
                except Exception as e:
                    messagebox.showerror(
                        "Ошибка шифрования",
                        f"Не удалось зашифровать пароль для '{item['description']}': {str(e)}",
                    )
        if migrated:
            self._save_passwords()
            messagebox.showinfo(
                "Миграция", "Существующие пароли были успешно зашифрованы"
            )

    def change_master_password(self, old_password, new_password):
        if old_password != self.master_password:
            return False, "Неверный текущий мастер-пароль"
        decrypted_passwords = []
        try:
            for item in self.passwords:
                decrypted_password = (
                    self.decrypt(item["password"])
                    if item.get("encrypted", False)
                    else item["password"]
                )
                decrypted_passwords.append(
                    {
                        "password": decrypted_password,
                        "description": item["description"],
                        "tags": item.get("tags", []),
                        "url": item.get("url", ""),
                    }
                )
        except Exception as e:
            return False, f"Ошибка расшифровки паролей: {str(e)}"
        self.master_password = new_password
        self._initialize_encryption()
        self.passwords = []
        for item in decrypted_passwords:
            self.add_password(
                item["password"], item["description"], item["tags"], item.get("url", "")
            )
        return True, "Мастер-пароль успешно изменен"


class AuthenticationDialog(tk.Toplevel):
    def __init__(self, parent, title="Аутентификация"):
        super().__init__(parent)
        self.title(title)
        self.transient(parent)
        self.grab_set()
        self.attributes("-topmost", True)
        self.parent = parent
        self.result = None
        self.master_password = None
        self.max_attempts = 3
        self.attempts = 0
        self._setup_ui()
        self._center_window()
        self.protocol("WM_DELETE_WINDOW", self._safe_destroy)
        self.after(100, lambda: self.parent.focus_force())

    def _center_window(self):
        self.update_idletasks()
        parent_x = self.parent.winfo_x()
        parent_y = self.parent.winfo_y()
        parent_width = self.parent.winfo_width()
        parent_height = self.parent.winfo_height()
        width = self.winfo_width()
        height = self.winfo_height()
        x = parent_x + (parent_width - width) // 2
        y = parent_y + (parent_height - height) // 2
        if parent_width == 1 or parent_height == 1:
            screen_width = self.winfo_screenwidth()
            screen_height = self.winfo_screenheight()
            x = (screen_width - width) // 2
            y = (screen_height - height) // 2
        self.geometry(f"+{x}+{y}")

    def _setup_ui(self):
        main_frame = ttk.Frame(self, padding=20)
        main_frame.pack(fill=tk.BOTH, expand=True)
        ttk.Label(
            main_frame, text="Введите мастер-пароль для шифрования/расшифрования:"
        ).pack(pady=(0, 10))
        self.password_var = tk.StringVar()
        self.password_entry = ttk.Entry(
            main_frame, textvariable=self.password_var, show="*", width=30
        )
        self.password_entry.pack(pady=(0, 10), fill=tk.X)
        self.password_entry.focus_force()
        self.attempts_label = ttk.Label(main_frame, text="", foreground="red")
        self.attempts_label.pack(pady=(0, 10), fill=tk.X)
        buttons_frame = ttk.Frame(main_frame)
        buttons_frame.pack(fill=tk.X, pady=(10, 0))
        ttk.Button(buttons_frame, text="ОК", command=self._on_ok).pack(
            side=tk.RIGHT, padx=5
        )
        ttk.Button(buttons_frame, text="Отмена", command=self._on_cancel).pack(
            side=tk.RIGHT, padx=5
        )
        self.bind("<Return>", lambda event: self._on_ok())
        self.minsize(300, 150)
        self.show_password_var = tk.BooleanVar(value=False)
        checkbox_frame = ttk.Frame(main_frame)
        checkbox_frame.pack(fill=tk.X, pady=(0, 10))
        self.show_password_var = tk.BooleanVar(value=False)
        self.show_checkbutton = ttk.Checkbutton(
            checkbox_frame,
            text="Показать пароль",
            variable=self.show_password_var,
            command=self._toggle_password_visibility,
        )
        self.show_checkbutton.pack(side=tk.LEFT, anchor=tk.W)
        self.context_menu = tk.Menu(self, tearoff=0)
        self.context_menu.add_command(
            label="Вставить",
            command=lambda: paste_from_clipboard(self.password_entry),
        )
        self.password_entry.bind("<Button-3>", self._show_context_menu)

    def _toggle_password_visibility(self):
        show = self.show_password_var.get()
        self.password_entry.config(show="" if show else "*")

    def _show_context_menu(self, event):
        try:
            self.context_menu.tk_popup(event.x_root, event.y_root)
        finally:
            self.context_menu.grab_release()

    def _validate_master_password(self, password):
        try:
            passwords_file = os.path.join(PASSWORDS_DIR, "passwords.json")
            if os.path.exists(passwords_file):
                with open(passwords_file, "r", encoding="utf-8") as f:
                    data = json.load(f)
                    if data:
                        test_item = data[0]
                        fernet = Fernet(self._generate_key(password))
                        fernet.decrypt(test_item["password"].encode())
            return True
        except Exception:
            return False

    def _generate_key(self, password):
        with open(SALT_PATH, "rb") as salt_file:
            salt = salt_file.read()
        kdf = PBKDF2HMAC(
            algorithm=hashes.SHA256(),
            length=32,
            salt=salt,
            iterations=100000,
            backend=default_backend(),
        )
        key = base64.urlsafe_b64encode(kdf.derive(password.encode()))
        return key

    def _on_ok(self):
        password = self.password_var.get()
        if not password:
            messagebox.showerror("Ошибка", "Пароль не может быть пустым")
            return
        self.attempts += 1
        try:
            if self._validate_master_password(password):
                self.master_password = password
                self.result = True
                if self.winfo_exists():
                    self.destroy()
            else:
                remaining = self.max_attempts - self.attempts
                if remaining > 0:
                    self.attempts_label.config(
                        text=f"Неверный пароль. Осталось попыток: {remaining}"
                    )
                    self.password_var.set("")
                    self.password_entry.focus_set()
                else:
                    if self.parent.winfo_exists():
                        self.parent.destroy()
                    self.destroy()
        except Exception as e:
            if self.parent.winfo_exists():
                self.parent.destroy()
            self.destroy()

    def _on_cancel(self):
        self.result = False
        self.destroy()
        self.parent.grab_release()

    def _safe_destroy(self):
        self.result = False
        self.destroy()
        self.parent.focus_set()


class BasePasswordDialog(tk.Toplevel):
    def __init__(self, parent, title):
        super().__init__(parent)
        self.title(title)
        self.transient(parent)
        self.grab_set()
        self.parent = parent
        self.result = None
        self._setup_ui()
        self._center_window()

    def _center_window(self):
        self.update_idletasks()
        width = self.winfo_width()
        height = self.winfo_height()
        x = (self.winfo_screenwidth() // 2) - (width // 2)
        y = (self.winfo_screenheight() // 2) - (height // 2)
        self.geometry(f"+{x}+{y}")

    def _setup_ui(self):
        raise NotImplementedError("Subclasses must implement _setup_ui")


class PasswordDialog(BasePasswordDialog):
    def __init__(self, parent, title, password_data=None):
        self.password_data = password_data or {
            "description": "",
            "password": "",
            "tags": [],
            "url": "",
            "login": "",
        }
        super().__init__(parent, title)
        self._setup_ui()
        self.attributes("-topmost", True)
        self._load_data()

    def _setup_ui(self):
        ttk.Label(self, text="Описание:").grid(
            row=0, column=0, padx=5, pady=5, sticky=tk.W
        )
        self.description_entry = ttk.Entry(self, width=30)
        self.description_entry.grid(row=0, column=1, padx=5, pady=5, sticky=tk.EW)
        ttk.Label(self, text="Логин:").grid(
            row=1, column=0, padx=5, pady=5, sticky=tk.W
        )
        self.login_entry = ttk.Entry(self, width=30)
        self.login_entry.grid(row=1, column=1, padx=5, pady=5, sticky=tk.EW)
        ttk.Label(self, text="Пароль:").grid(
            row=2, column=0, padx=5, pady=5, sticky=tk.W
        )
        password_frame = ttk.Frame(self)
        password_frame.grid(row=2, column=1, padx=5, pady=5, sticky=tk.EW)
        self.password_entry = ttk.Entry(password_frame, width=25, show="*")
        self.password_entry.pack(side=tk.LEFT, fill=tk.X, expand=True, ipady=2)
        self.paste_button = ttk.Button(
            password_frame,
            text="PASTE",
            command=lambda: paste_from_clipboard(self.password_entry),
            style="Small.TButton",
        )
        self.paste_button.pack(side=tk.RIGHT, padx=2, ipady=0)
        style = ttk.Style()
        style.configure("Small.TButton", width=6, font=("Arial", 8), padding=(1, 1))
        ttk.Label(self, text="Теги:").grid(row=3, column=0, padx=5, pady=5, sticky=tk.W)
        self.tags_entry = ttk.Entry(self, width=30)
        self.tags_entry.grid(row=3, column=1, padx=5, pady=5, sticky=tk.EW)
        ttk.Label(self, text="URL:").grid(row=4, column=0, padx=5, pady=5, sticky=tk.W)
        self.url_entry = ttk.Entry(self, width=30)
        self.url_entry.grid(row=4, column=1, padx=5, pady=5, sticky=tk.EW)
        self.show_password_var = tk.BooleanVar(value=False)
        ttk.Checkbutton(
            self,
            text="Показать пароль",
            variable=self.show_password_var,
            command=self._toggle_password_visibility,
        ).grid(row=5, column=1, padx=5, pady=5, sticky=tk.E)
        button_frame = ttk.Frame(self)
        button_frame.grid(row=6, column=0, columnspan=2, pady=10)
        ttk.Button(button_frame, text="Сохранить", command=self._save).pack(
            side=tk.RIGHT, padx=5
        )
        ttk.Button(button_frame, text="Отмена", command=self.destroy).pack(
            side=tk.RIGHT
        )

    def _load_data(self):
        self.description_entry.insert(0, self.password_data.get("description", ""))
        self.password_entry.insert(0, self.password_data.get("password", ""))
        self.tags_entry.insert(0, ", ".join(self.password_data.get("tags", [])))
        self.url_entry.insert(0, self.password_data.get("url", ""))
        self.login_entry.insert(0, self.password_data.get("login", ""))

    def _safe_copy_to_clipboard(self, text: str, clear_after_ms: int = 60000) -> bool:
        if not text:
            return False
        try:
            self.master.clipboard_clear()
            self.master.clipboard_append(text)
            self.master.update_idletasks()
            if clear_after_ms:
                self.master.after(
                    clear_after_ms,
                    lambda: self.master.winfo_exists()
                    and self.master.clipboard_clear(),
                )
            return True
        except Exception:
            import pyperclip

            pyperclip.copy(text)
            if clear_after_ms:
                self.master.after(clear_after_ms, lambda: pyperclip.copy(""))
            return True

    def _toggle_password_visibility(self):
        show = self.show_password_var.get()
        self.password_entry.config(show="" if show else "*")

    def _save(self):
        description = self.description_entry.get()
        password = self.password_entry.get()
        url = self.url_entry.get().strip()
        tags = [tag.strip() for tag in self.tags_entry.get().split(",") if tag.strip()]
        if not description or not password:
            messagebox.showerror("Ошибка", "Заполните описание и пароль!")
            return
        if len(password) < 4:
            messagebox.showerror("Ошибка", "Пароль должен быть не короче 4 символов!")
            return
        url = self.url_entry.get()
        login = self.login_entry.get()
        self.result = {
            "description": description,
            "login": login,
            "password": password,
            "tags": tags,
            "url": url.strip(),
        }
        self.destroy()


class ChangeMasterPasswordDialog(tk.Toplevel):
    def __init__(self, parent, password_manager):
        super().__init__(parent)
        self.attributes("-topmost", True)
        self.title("Смена мастер-пароля")
        self.transient(parent)
        self.grab_set()
        self.parent = parent
        self.password_manager = password_manager
        self.result = None
        self._setup_ui()
        self._center_window()
        self.context_menu = tk.Menu(self, tearoff=0)
        self.context_menu.add_command(
            label="Вставить",
            command=lambda: paste_from_clipboard(self.focus_get()),
        )
        self.old_password_entry.bind("<Button-3>", self._show_context_menu)
        self.new_password_entry.bind("<Button-3>", self._show_context_menu)
        self.confirm_password_entry.bind("<Button-3>", self._show_context_menu)

    def _center_window(self):
        self.update_idletasks()
        width = self.winfo_width()
        height = self.winfo_height()
        x = (self.winfo_screenwidth() // 2) - (width // 2)
        y = (self.winfo_screenheight() // 2) - (height // 2)
        self.geometry(f"+{x}+{y}")

    def _setup_ui(self):
        main_frame = ttk.Frame(self, padding=20)
        main_frame.pack(fill=tk.BOTH, expand=True)
        ttk.Label(main_frame, text="Текущий мастер-пароль:").pack(
            anchor=tk.W, pady=(0, 5)
        )
        self.old_password_var = tk.StringVar()
        self.old_password_entry = ttk.Entry(
            main_frame, textvariable=self.old_password_var, show="*", width=30
        )
        self.old_password_entry.pack(fill=tk.X, pady=(0, 10))
        ttk.Label(main_frame, text="Новый мастер-пароль:").pack(
            anchor=tk.W, pady=(0, 5)
        )
        self.new_password_var = tk.StringVar()
        self.new_password_entry = ttk.Entry(
            main_frame, textvariable=self.new_password_var, show="*", width=30
        )
        self.new_password_entry.pack(fill=tk.X, pady=(0, 10))
        ttk.Label(main_frame, text="Подтвердите новый пароль:").pack(
            anchor=tk.W, pady=(0, 5)
        )
        self.confirm_password_var = tk.StringVar()
        self.confirm_password_entry = ttk.Entry(
            main_frame, textvariable=self.confirm_password_var, show="*", width=30
        )
        self.confirm_password_entry.pack(fill=tk.X, pady=(0, 10))
        buttons_frame = ttk.Frame(main_frame)
        buttons_frame.pack(fill=tk.X, pady=(10, 0))
        ttk.Button(buttons_frame, text="Сменить пароль", command=self._on_change).pack(
            side=tk.RIGHT, padx=5
        )
        ttk.Button(buttons_frame, text="Отмена", command=self._on_cancel).pack(
            side=tk.RIGHT, padx=5
        )

    def _on_change(self):
        old_password = self.old_password_var.get()
        new_password = self.new_password_var.get()
        confirm_password = self.confirm_password_var.get()
        if not old_password or not new_password:
            messagebox.showerror("Ошибка", "Все поля должны быть заполнены")
            return
        if new_password != confirm_password:
            messagebox.showerror("Ошибка", "Новые пароли не совпадают")
            return
        success, message = self.password_manager.change_master_password(
            old_password, new_password
        )
        if success:
            messagebox.showinfo("Успех", message)
            self.result = True
            self.destroy()
        else:
            messagebox.showerror("Ошибка", message)

    def _on_cancel(self):
        self.result = False
        self.destroy()

    def _show_context_menu(self, event):
        try:
            self.context_menu.tk_popup(event.x_root, event.y_root)
        finally:
            self.context_menu.grab_release()


class ConfigEditorDialog(BasePasswordDialog):
    def __init__(self, parent, config_data):
        self.config_data = config_data
        super().__init__(parent, "Редактирование конфигурации")
        self.attributes("-topmost", True)

    def _setup_ui(self):
        ttk.Label(self, text="Макс. длина пароля:").grid(
            row=0, column=0, padx=5, pady=5, sticky=tk.W
        )
        self.max_length_entry = ttk.Spinbox(self, from_=8, to=128, width=10)
        self.max_length_entry.grid(row=0, column=1, padx=5, pady=5, sticky=tk.W)
        self.max_length_entry.set(self.config_data.get("max_password_length", 32))
        ttk.Label(self, text="Длина по умолчанию:").grid(
            row=1, column=0, padx=5, pady=5, sticky=tk.W
        )
        self.default_length_entry = ttk.Spinbox(
            self, from_=4, to=self.config_data.get("max_password_length", 32), width=10
        )
        self.default_length_entry.grid(row=1, column=1, padx=5, pady=5, sticky=tk.W)
        self.default_length_entry.set(self.config_data.get("default_length", 18))
        self.use_upper_var = tk.BooleanVar(
            value=self.config_data.get("use_uppercase", True)
        )
        ttk.Checkbutton(
            self, text="Исп. заглавные буквы", variable=self.use_upper_var
        ).grid(row=2, column=0, columnspan=2, padx=5, pady=2, sticky=tk.W)
        self.use_lower_var = tk.BooleanVar(
            value=self.config_data.get("use_lowercase", True)
        )
        ttk.Checkbutton(
            self, text="Исп. строчные буквы", variable=self.use_lower_var
        ).grid(row=3, column=0, columnspan=2, padx=5, pady=2, sticky=tk.W)
        self.use_digits_var = tk.BooleanVar(
            value=self.config_data.get("use_digits", True)
        )
        ttk.Checkbutton(self, text="Исп. цифры", variable=self.use_digits_var).grid(
            row=4, column=0, columnspan=2, padx=5, pady=2, sticky=tk.W
        )
        self.use_symbols_var = tk.BooleanVar(
            value=self.config_data.get("use_symbols", True)
        )
        ttk.Checkbutton(self, text="Исп. символы", variable=self.use_symbols_var).grid(
            row=5, column=0, columnspan=2, padx=5, pady=2, sticky=tk.W
        )
        ttk.Label(self, text="Исключенные символы:").grid(
            row=6, column=0, padx=5, pady=5, sticky=tk.W
        )
        self.excluded_chars_entry = ttk.Entry(self, width=30)
        self.excluded_chars_entry.grid(row=6, column=1, padx=5, pady=5)
        self.excluded_chars_entry.insert(0, self.config_data.get("excluded_chars", ""))
        ttk.Label(self, text="Файл для паролей:").grid(
            row=7, column=0, padx=5, pady=5, sticky=tk.W
        )
        self.pass_file_entry = ttk.Entry(self, width=30)
        self.pass_file_entry.grid(row=7, column=1, padx=5, pady=5)
        self.pass_file_entry.insert(
            0, self.config_data.get("passwords_file", "passwords.json")
        )
        button_frame = ttk.Frame(self)
        button_frame.grid(row=8, column=0, columnspan=2, pady=10)
        ttk.Button(button_frame, text="Сохранить", command=self._save).pack(
            side=tk.RIGHT, padx=5
        )
        ttk.Button(button_frame, text="Отмена", command=self.destroy).pack(
            side=tk.RIGHT
        )

    def _save(self):
        new_config = {
            "max_password_length": int(self.max_length_entry.get()),
            "default_length": int(self.default_length_entry.get()),
            "use_uppercase": self.use_upper_var.get(),
            "use_lowercase": self.use_lower_var.get(),
            "use_digits": self.use_digits_var.get(),
            "use_symbols": self.use_symbols_var.get(),
            "excluded_chars": self.excluded_chars_entry.get(),
            "passwords_file": self.pass_file_entry.get(),
        }
        if new_config["default_length"] > new_config["max_password_length"]:
            messagebox.showerror(
                "Ошибка", "Длина по умолчанию не может превышать максимальную"
            )
            return
        self.config_data.update(new_config)
        self.result = True
        self.destroy()


class RegenerateSaltDialog(tk.Toplevel):
    def __init__(self, parent, password_manager):
        super().__init__(parent)
        self.title("Обновление криптографической соли")
        self.password_manager = password_manager
        self._setup_ui()
        self.attributes("-topmost", True)
        self._center_window()
        self.context_menu = tk.Menu(self, tearoff=0)
        self.context_menu.add_command(
            label="Вставить",
            command=lambda: paste_from_clipboard(self.password_entry),
        )
        self.password_entry.bind("<Button-3>", self._show_context_menu)

    def _center_window(self):
        self.update_idletasks()
        width = self.winfo_width()
        height = self.winfo_height()
        x = (self.winfo_screenwidth() // 2) - (width // 2)
        y = (self.winfo_screenheight() // 2) - (height // 2)
        self.geometry(f"+{x}+{y}")

    def _show_context_menu(self, event):
        try:
            self.context_menu.tk_popup(event.x_root, event.y_root)
        finally:
            self.context_menu.grab_release()

    def _setup_ui(self):
        main_frame = ttk.Frame(self, padding=20)
        main_frame.pack(fill=tk.BOTH, expand=True)
        ttk.Label(main_frame, text="Введите текущий мастер-пароль:").pack(pady=(0, 10))
        self.password_var = tk.StringVar()
        self.password_entry = ttk.Entry(
            main_frame, textvariable=self.password_var, show="*", width=30
        )
        self.password_entry.pack(pady=(0, 10))
        buttons_frame = ttk.Frame(main_frame)
        buttons_frame.pack(fill=tk.X, pady=(10, 0))
        ttk.Button(buttons_frame, text="Обновить", command=self._regenerate).pack(
            side=tk.RIGHT, padx=5
        )
        ttk.Button(buttons_frame, text="Отмена", command=self.destroy).pack(
            side=tk.RIGHT, padx=5
        )

    def _regenerate(self):
        password = self.password_var.get()
        if not password:
            messagebox.showerror("Ошибка", "Введите мастер-пароль")
            return
        success, message = self.password_manager.regenerate_salt(password)
        if success:
            messagebox.showinfo("Успех", message)
            self.destroy()
        else:
            messagebox.showerror("Ошибка", message)


class PasswordGeneratorApp:
    def __init__(self, master):
        self.master = master
        master.title("Генератор паролей")
        master.geometry("800x600")
        self._backup_job = None
        self._closing = False
        self.master.protocol("WM_DELETE_WINDOW", self._on_close)
        self._backup_job = None
        self._closing = False
        self.master.protocol("WM_DELETE_WINDOW", self._on_close)
        master.minsize(800, 600)
        master.attributes("-topmost", True)
        master.lift()
        self.idle_timer = None
        self.idle_timeout = 120000
        self.setup_activity_tracking()
        if not self._initialize_password_manager():
            return
        self._full_ui_initialization()

    def _full_ui_initialization(self):
        self._cached_list = []
        self._cached_has_plain = False
        self.password_generator = PasswordGenerator()
        self._setup_styles()
        self._create_tabs()
        self._setup_context_menu()
        self._create_menu()
        self.length_slider.config(to=self.password_generator.max_password_length)
        self.master.deiconify()
        self.master.after_idle(self._refresh_password_list)
        self.master.after(2000, self.schedule_backup)

    def _safe_copy_to_clipboard(self, text: str, clear_after_ms: int = 60000) -> bool:
        if not text:
            return False
        try:
            self.master.clipboard_clear()
            self.master.clipboard_append(text)
            self.master.update_idletasks()
            if clear_after_ms:
                self.master.after(clear_after_ms, lambda: self.master.clipboard_clear())
            return True
        except Exception:
            pass
        pyperclip.copy(text)
        if clear_after_ms:
            self.master.after(clear_after_ms, lambda: pyperclip.copy(""))
        return True

    def _toast(self, text: str, ms: int = 1500):
        tw = tk.Toplevel(self.master)
        tw.overrideredirect(True)
        tw.attributes("-topmost", True)
        frm = ttk.Frame(tw, padding=(8, 6))
        frm.pack(fill="both", expand=True)
        ttk.Label(frm, text=text).pack()
        self.master.update_idletasks()
        tw.update_idletasks()
        mw, mh = self.master.winfo_width(), self.master.winfo_height()
        mx, my = self.master.winfo_rootx(), self.master.winfo_rooty()
        ww, wh = tw.winfo_width(), tw.winfo_height()
        x = mx + mw - ww - 12
        y = my + mh - wh - 12
        tw.geometry(f"+{x}+{y}")
        tw.after(ms, tw.destroy)

    def _switch_to_launcher(self, event=None):
        try:
            if QApplication.instance() is not None:
                self._on_close()
                return
        except Exception:
            pass

        subprocess.Popen([sys.executable, os.path.abspath(__file__)])
        self._on_close()

    def _add_apps_menu(self):
        apps_menu = tk.Menu(self.menu_bar, tearoff=0)
        self.menu_bar.add_cascade(label="Приложения", menu=apps_menu)
        apps_menu.add_command(
            label="Выбрать приложение…  (Ctrl+Alt+L)",
            command=self._switch_to_launcher,
        )
        self.master.bind_all("<Control-Alt-l>", lambda e: self._switch_to_launcher())
        self.master.bind_all("<Control-Alt-L>", lambda e: self._switch_to_launcher())

    def _edit_configuration(self):
        config_path = os.path.join(PASSWORDS_DIR, "config.json")
        try:
            with open(config_path, "r", encoding="utf-8") as f:
                config_data = json.load(f)
        except Exception as e:
            messagebox.showerror(
                "Ошибка", f"Не удалось загрузить конфигурацию: {str(e)}"
            )
            return
        dialog = ConfigEditorDialog(self.master, config_data)
        dialog.wait_window()
        if getattr(dialog, "result", False):
            try:
                with open(config_path, "w", encoding="utf-8") as f:
                    json.dump(config_data, f, indent=4, ensure_ascii=False)
                self.password_generator.load_config()
                self.length_slider.config(
                    to=self.password_generator.max_password_length
                )
                messagebox.showinfo("Успех", "Конфигурация успешно обновлена!")
            except Exception as e:
                messagebox.showerror(
                    "Ошибка", f"Ошибка сохранения конфигурации: {str(e)}"
                )

    def _export_passwords_txt(self):
        file_path = tk_filedialog.asksaveasfilename(
            title="Сохранить как...",
            defaultextension=".txt",
            filetypes=[("Text files", "*.txt"), ("All files", "*.*")],
        )
        if not file_path:
            return
        success, message = self.password_manager.export_to_txt(file_path)
        if success:
            messagebox.showinfo("Успех", message)
        else:
            messagebox.showerror("Ошибка", message)

    def _on_close(self):
        self._closing = True
        try:
            if getattr(self, "_backup_job", None):
                self.master.after_cancel(self._backup_job)
                self._backup_job = None
        except Exception:
            pass
        try:
            if getattr(self, "idle_timer", None):
                self.master.after_cancel(self.idle_timer)
                self.idle_timer = None
        except Exception:
            pass
        self.master.destroy()

    def schedule_backup(self):
        if getattr(self, "_closing", False) or not self.master.winfo_exists():
            return
        self.password_manager.backup_passwords()
        self._backup_job = self.master.after(3600000, self.schedule_backup)

    def _on_tree_drag_start(self, event):
        row_id = self.password_tree.identify_row(event.y)
        self._dragging_iid = row_id if row_id else None

    def _on_tree_drag_motion(self, event):
        if not getattr(self, "_dragging_iid", None):
            return
        y = event.y
        height = self.password_tree.winfo_height()
        if y < 20:
            self.password_tree.yview_scroll(-1, "units")
        elif y > height - 20:
            self.password_tree.yview_scroll(1, "units")
        over_iid = self.password_tree.identify_row(y)
        if over_iid:
            self.password_tree.selection_set(over_iid)

    def _on_tree_drag_release(self, event):
        try:
            if not getattr(self, "_dragging_iid", None):
                return
            from_idx = int(self._dragging_iid)

            over_iid = self.password_tree.identify_row(event.y)
            if over_iid:
                to_idx = int(over_iid)
            else:
                to_idx = len(self.password_manager.passwords) - 1

            if self.password_manager.move_password(from_idx, to_idx):
                self._refresh_password_list()
                self.password_tree.selection_set(str(to_idx))
                self.password_tree.see(str(to_idx))
        finally:
            self._dragging_iid = None

    def _initialize_password_manager(self):
        attempts = 3
        while attempts > 0:
            auth_dialog = AuthenticationDialog(self.master)
            auth_dialog.wait_window()
            if not auth_dialog.result:
                if self.master.winfo_exists():
                    self.master.destroy()
                return False
            try:
                self.password_manager = PasswordManager(auth_dialog.master_password)
                self.password_manager.load_passwords()
                self.master.withdraw()
                return True
            except Exception:
                attempts -= 1
                if attempts > 0:
                    messagebox.showwarning(
                        "Ошибка", f"Неверный пароль. Осталось попыток: {attempts}"
                    )
                else:
                    if self.master.winfo_exists():
                        self.master.destroy()
                    return

    def _show_change_master_password(self):
        dialog = ChangeMasterPasswordDialog(self.master, self.password_manager)
        dialog.wait_window()
        if dialog.result:
            self._refresh_password_list()

    def _create_default_config(self):
        config_path = os.path.join(PASSWORDS_DIR, "config.json")
        if not os.path.exists(config_path):
            default_config = {
                "max_password_length": 32,
                "default_length": 15,
                "use_uppercase": True,
                "use_lowercase": True,
                "use_digits": True,
                "use_symbols": True,
                "excluded_chars": "O0oQl1I|i!S5Z2B8G6CQDceaouvwxX",
                "passwords_file": "passwords.json",
            }
            try:
                with open(config_path, "w", encoding="utf-8") as f:
                    json.dump(default_config, f, indent=4)
            except Exception as e:
                print(f"Ошибка создания конфигурации: {str(e)}")

    def _setup_styles(self):
        style = ttk.Style()
        style.theme_use("clam")
        style.configure("TButton", font=("Arial", 10))
        style.configure("TLabel", font=("Arial", 10))
        style.configure("TEntry", font=("Arial", 10))
        style.configure("TCheckbutton", font=("Arial", 10))
        style.configure("red.Horizontal.TProgressbar", background="red")
        style.configure("yellow.Horizontal.TProgressbar", background="yellow")
        style.configure("green.Horizontal.TProgressbar", background="green")
        style.configure("Treeview.Cell", borderwidth=1, relief="solid", padding=(5, 2))
        style.layout(
            "Treeview.Item",
            [
                (
                    "Treeitem.padding",
                    {
                        "sticky": "nswe",
                        "children": [
                            ("Treeitem.indicator", {"side": "left", "sticky": ""}),
                            ("Treeitem.image", {"side": "left", "sticky": ""}),
                            ("Treeitem.text", {"side": "left", "sticky": ""}),
                            ("Treeitem.Cell", {"sticky": "nswe"}),
                        ],
                    },
                )
            ],
        )

    def _create_tabs(self):
        if not hasattr(self, "password_manager"):
            return
        self.tab_control = ttk.Notebook(self.master)
        self.generator_tab = ttk.Frame(self.tab_control)
        self.tab_control.add(self.generator_tab, text="Генератор")
        self._setup_generator_tab()
        self.manager_tab = ttk.Frame(self.tab_control)
        self.tab_control.add(self.manager_tab, text="Менеджер паролей")
        self._setup_manager_tab()
        self.tab_control.pack(expand=1, fill="both")

    def _setup_generator_tab(self):
        settings_frame = ttk.LabelFrame(
            self.generator_tab, text="Настройки генератора", padding=10
        )
        settings_frame.pack(padx=10, pady=10, fill=tk.BOTH)
        ttk.Label(settings_frame, text="Длина пароля:").grid(
            row=0, column=0, sticky=tk.W, pady=5
        )
        self.length_var = tk.IntVar(value=self.password_generator.password_length)
        length_entry = ttk.Entry(settings_frame, textvariable=self.length_var, width=5)
        length_entry.grid(row=0, column=1, sticky=tk.W, pady=5)
        self.length_slider = ttk.Scale(
            settings_frame,
            from_=4,
            to=32,
            orient=tk.HORIZONTAL,
            variable=self.length_var,
            length=200,
            command=lambda _: self.length_var.set(int(self.length_slider.get())),
        )
        self.length_slider.grid(row=0, column=2, sticky=tk.W, pady=5, padx=10)
        self.use_uppercase = tk.BooleanVar(value=self.password_generator.use_uppercase)
        ttk.Checkbutton(
            settings_frame, text="Заглавные буквы (A-Z)", variable=self.use_uppercase
        ).grid(row=1, column=0, columnspan=3, sticky=tk.W)
        self.use_lowercase = tk.BooleanVar(value=self.password_generator.use_lowercase)
        ttk.Checkbutton(
            settings_frame, text="Строчные буквы (a-z)", variable=self.use_lowercase
        ).grid(row=2, column=0, columnspan=3, sticky=tk.W)
        self.use_digits = tk.BooleanVar(value=self.password_generator.use_digits)
        ttk.Checkbutton(
            settings_frame, text="Цифры (0-9)", variable=self.use_digits
        ).grid(row=3, column=0, columnspan=3, sticky=tk.W)
        self.use_symbols = tk.BooleanVar(value=self.password_generator.use_symbols)
        ttk.Checkbutton(
            settings_frame,
            text="Специальные символы (!@#$%)",
            variable=self.use_symbols,
        ).grid(row=4, column=0, columnspan=3, sticky=tk.W)
        ttk.Label(settings_frame, text="Исключить:").grid(
            row=5, column=0, sticky=tk.W, pady=5
        )
        self.excluded_chars_var = tk.StringVar(
            value=self.password_generator.excluded_chars
        )
        ttk.Entry(settings_frame, textvariable=self.excluded_chars_var, width=30).grid(
            row=5, column=1, columnspan=2, sticky=tk.W, pady=5
        )
        ttk.Button(
            settings_frame, text="Сгенерировать пароль", command=self._generate_password
        ).grid(row=6, column=0, columnspan=3, pady=10)
        output_frame = ttk.LabelFrame(
            self.generator_tab, text="Сгенерированный пароль", padding=10
        )
        output_frame.pack(padx=10, pady=10, fill=tk.BOTH)
        self.password_var = tk.StringVar()
        password_entry_frame = ttk.Frame(output_frame)
        password_entry_frame.pack(fill=tk.X)
        self.password_entry = ttk.Entry(
            password_entry_frame, textvariable=self.password_var, width=30
        )
        self.password_entry.pack(side=tk.LEFT, expand=True, fill=tk.X)
        self.password_entry.config(show="*")
        self.show_password_var = tk.BooleanVar(value=False)
        ttk.Checkbutton(
            password_entry_frame,
            text="Показать",
            variable=self.show_password_var,
            command=lambda: self.password_entry.config(
                show="" if self.show_password_var.get() else "*"
            ),
        ).pack(side=tk.RIGHT)
        ttk.Label(output_frame, text="Сложность пароля:").pack(anchor=tk.W)
        self.strength_var = tk.IntVar()
        self.strength_bar = ttk.Progressbar(
            output_frame,
            orient=tk.HORIZONTAL,
            length=200,
            mode="determinate",
            variable=self.strength_var,
        )
        self.strength_bar.pack(fill=tk.X, pady=5)
        button_frame = ttk.Frame(output_frame)
        button_frame.pack(fill=tk.X, pady=5)
        ttk.Button(button_frame, text="Копировать", command=self._copy_password).pack(
            side=tk.LEFT, padx=5
        )
        ttk.Button(
            button_frame, text="Очистить", command=lambda: self.password_var.set("")
        ).pack(side=tk.LEFT, padx=5)
        ttk.Button(
            button_frame, text="Сохранить", command=self._save_password_dialog
        ).pack(side=tk.LEFT, padx=5)

    def _setup_manager_tab(self):
        control_frame = ttk.Frame(self.manager_tab, padding=10)
        control_frame.pack(fill=tk.X)
        self.hide_passwords_var = tk.BooleanVar(value=True)
        ttk.Checkbutton(
            control_frame,
            text="Скрыть пароли",
            variable=self.hide_passwords_var,
            command=self._refresh_password_list,
        ).pack(side=tk.LEFT, padx=5)
        ttk.Button(
            control_frame, text="Обновить список", command=self._refresh_password_list
        ).pack(side=tk.LEFT, padx=5)
        ttk.Button(
            control_frame, text="Добавить новый пароль", command=self._add_new_password
        ).pack(side=tk.LEFT, padx=5)
        ttk.Label(control_frame, text="Поиск:").pack(side=tk.LEFT, padx=(10, 0))
        self.search_var = tk.StringVar()
        self.search_var.trace("w", lambda *args: self._filter_passwords())
        ttk.Entry(control_frame, textvariable=self.search_var, width=20).pack(
            side=tk.LEFT, padx=5
        )
        ttk.Label(control_frame, text="Теги:").pack(side=tk.LEFT, padx=(10, 0))
        self.tag_filter_var = tk.StringVar(value="Все")
        self.tag_filter = ttk.Combobox(
            control_frame,
            textvariable=self.tag_filter_var,
            values=["Все"],
            state="readonly",
        )
        self.tag_filter.pack(side=tk.LEFT, padx=5)
        self.tag_filter.bind("<<ComboboxSelected>>", lambda e: self._filter_passwords())
        list_frame = ttk.Frame(self.manager_tab, padding=10)
        list_frame.pack(fill=tk.BOTH, expand=True)
        columns = ("description", "password", "tags", "updated_at")
        self.password_tree = ttk.Treeview(
            list_frame, columns=columns, show="headings", selectmode="browse"
        )
        headings = {
            "description": "Описание",
            "password": "Пароль",
            "tags": "Теги",
            "updated_at": "Обновлено",
        }
        for col in columns:
            self.password_tree.heading(col, text=headings[col], anchor="center")
            if col == "description":
                self.password_tree.column(col, anchor="e", stretch=True, width=150)
            elif col == "password":
                self.password_tree.column(col, anchor="center", stretch=True, width=200)
            elif col == "tags":
                self.password_tree.column(col, anchor="center", stretch=True, width=100)
            elif col == "updated_at":
                self.password_tree.column(col, anchor="center", stretch=True, width=120)
        vscroll = ttk.Scrollbar(
            list_frame, orient="vertical", command=self.password_tree.yview
        )
        hscroll = ttk.Scrollbar(
            list_frame, orient="horizontal", command=self.password_tree.xview
        )
        self.password_tree.configure(
            yscrollcommand=vscroll.set, xscrollcommand=hscroll.set
        )
        vscroll.pack(side=tk.RIGHT, fill=tk.Y)
        hscroll.pack(side=tk.BOTTOM, fill=tk.X)
        self.password_tree.pack(side=tk.TOP, fill=tk.BOTH, expand=True)
        self.password_tree.bind("<ButtonPress-1>", self._on_tree_drag_start)
        self.password_tree.bind("<B1-Motion>", self._on_tree_drag_motion)
        self.password_tree.bind("<ButtonRelease-1>", self._on_tree_drag_release)
        action_frame = ttk.Frame(self.manager_tab, padding=10)
        action_frame.pack(fill=tk.X)
        ttk.Button(
            action_frame, text="Копировать", command=self._copy_selected_password
        ).pack(side=tk.LEFT, padx=5)
        ttk.Button(
            action_frame, text="Просмотреть", command=self._view_selected_password
        ).pack(side=tk.LEFT, padx=5)
        ttk.Button(
            action_frame, text="Редактировать", command=self._edit_selected_password
        ).pack(side=tk.LEFT, padx=5)
        ttk.Button(
            action_frame, text="Удалить", command=self._delete_selected_password
        ).pack(side=tk.LEFT, padx=5)
        self._refresh_password_list()
        self.password_tree.bind("<Double-1>", lambda e: self._edit_selected_password())

    def _setup_context_menu(self):
        self.entry_context_menu = tk.Menu(self.master, tearoff=0)
        self.entry_context_menu.add_command(
            label="Копировать",
            command=lambda: self._copy_to_clipboard(self.master.focus_get()),
        )
        self.entry_context_menu.add_command(
            label="Вырезать",
            command=lambda: self._cut_to_clipboard(self.master.focus_get()),
        )
        self.entry_context_menu.add_command(
            label="Вставить",
            command=lambda: paste_from_clipboard(self.master.focus_get()),
        )
        self.entry_context_menu.add_command(
            label="Выбрать всё",
            command=lambda: self._select_all(self.master.focus_get()),
        )
        self.tree_context_menu = tk.Menu(self.master, tearoff=0)
        self.tree_context_menu.add_command(
            label="Копировать", command=self._copy_selected_password
        )
        self.tree_context_menu.add_command(
            label="Просмотреть", command=self._view_selected_password
        )
        self.tree_context_menu.add_command(
            label="Редактировать", command=self._edit_selected_password
        )
        self.tree_context_menu.add_command(
            label="Удалить", command=self._delete_selected_password
        )
        self.master.bind("<Button-3>", self._on_right_click)
        self.password_tree.bind("<Button-3>", self._on_tree_right_click)

    def setup_activity_tracking(self):
        def track_activity(event):
            self.reset_inactivity_timer()

        def bind_recursive(widget):
            widget.bind("<Motion>", track_activity)
            widget.bind("<KeyPress>", track_activity)
            for child in widget.winfo_children():
                bind_recursive(child)

        bind_recursive(self.master)
        self.reset_inactivity_timer()

    def reset_inactivity_timer(self, event=None):
        if self.idle_timer is not None:
            self.master.after_cancel(self.idle_timer)
        self.idle_timer = self.master.after(self.idle_timeout, self.lock_application)

    def lock_application(self):
        self.master.withdraw()
        if not hasattr(self, "password_manager") or not self.password_manager:
            return
        auth_window = tk.Toplevel(self.master)
        auth_window.title("Блокировка")
        auth_window.protocol("WM_DELETE_WINDOW", lambda: None)
        auth_window.grab_set()
        auth_window.update_idletasks()
        parent_width = self.master.winfo_width()
        parent_height = self.master.winfo_height()
        parent_x = self.master.winfo_x()
        parent_y = self.master.winfo_y()
        width = auth_window.winfo_width()
        height = auth_window.winfo_height()
        x = parent_x + (parent_width - width) // 2
        y = parent_y + (parent_height - height) // 2
        auth_window.geometry(f"+{x}+{y}")
        auth_dialog = AuthenticationDialog(auth_window)
        auth_dialog.wait_window()
        if auth_dialog.result:
            try:
                self.password_manager = PasswordManager(auth_dialog.master_password)
                self.password_manager.load_passwords()
                self._refresh_password_list()
                auth_window.destroy()
                self.master.deiconify()
            except Exception as e:
                messagebox.showerror("Ошибка", f"Ошибка разблокировки: {str(e)}")
                self.master.destroy()
        else:
            self.master.destroy()
        self.reset_inactivity_timer()

    def _create_menu(self):
        self.menu_bar = tk.Menu(self.master)
        self.master.config(menu=self.menu_bar)
        self._add_apps_menu()
        security_menu = tk.Menu(self.menu_bar, tearoff=0)
        self.menu_bar.add_cascade(label="Безопасность", menu=security_menu)
        security_menu.add_command(
            label="Обновить криптографическую соль", command=self._regenerate_salt
        )
        security_menu.add_command(
            label="Сменить мастер-пароль", command=self._show_change_master_password
        )
        security_menu.add_command(
            label="Заблокировать", command=lambda: self.lock_application()
        )
        file_menu = tk.Menu(self.menu_bar, tearoff=0)
        self.menu_bar.add_cascade(label="Файл", menu=file_menu)
        file_menu.add_command(
            label="Изменить конфигурацию", command=self._edit_configuration
        )
        file_menu.add_command(label="Экспорт в TXT", command=self._export_passwords_txt)
        file_menu.add_command(label="Импорт из TXT", command=self._import_passwords_txt)
        file_menu.add_command(
            label="Редактировать выбранный пароль", command=self._edit_selected_password
        )
        file_menu.add_command(
            label="Создать резервную копию", command=self._create_backup
        )
        file_menu.add_command(
            label="Открыть резервную копию", command=self._open_backup
        )
        file_menu.add_separator()
        file_menu.add_command(label="Выход", command=self.master.destroy)
        help_menu = tk.Menu(self.menu_bar, tearoff=0)
        self.menu_bar.add_cascade(label="Справка", menu=help_menu)
        help_menu.add_command(label="О программе", command=self._show_about)

    def _show_text_context_menu(self, event, menu):
        try:
            menu.tk_popup(event.x_root, event.y_root)
        finally:
            menu.grab_release()

    def _regenerate_salt(self):
        dialog = RegenerateSaltDialog(self.master, self.password_manager)
        dialog.wait_window()
        self._refresh_password_list()

    def _open_backup(self):
        backup_dir = os.path.join(PASSWORDS_DIR, "backups")
        backup_file = os.path.join(backup_dir, "backup_passwords.bin")
        if not os.path.exists(backup_file):
            messagebox.showerror("Ошибка", f"Файл не найден:\n{backup_file}")
            return
        try:
            with open(backup_file, "r", encoding="utf-8") as f:
                encrypted_data = f.read()
            backup_passwords = json.loads(encrypted_data)
            decrypted_passwords = []
            for item in backup_passwords:
                try:
                    decrypted_pwd = self.password_manager.decrypt(item["password"])
                    decrypted_item = {
                        "description": item["description"],
                        "password": decrypted_pwd,
                        "tags": item.get("tags", []),
                    }
                    decrypted_passwords.append(decrypted_item)
                except Exception as e:
                    messagebox.showerror(
                        "Ошибка расшифровки",
                        f"Не удалось расшифровать пароль для '{item['description']}': {str(e)}",
                    )
                    return
            text_window = tk.Toplevel(self.master)
            text_window.title("Резервная копия (расшифровано)")
            text = tk.Text(text_window, wrap=tk.WORD)
            text.pack(fill=tk.BOTH, expand=True)
            formatted_data = json.dumps(
                decrypted_passwords, indent=4, ensure_ascii=False
            )
            text.insert(tk.END, formatted_data)

            context_menu = tk.Menu(text_window, tearoff=0)
            context_menu.add_command(
                label="Копировать", command=lambda: self._copy_from_text_widget(text)
            )

            text.bind(
                "<Button-3>", lambda e: self._show_text_context_menu(e, context_menu)
            )
        except Exception as e:
            messagebox.showerror("Ошибка", f"Не удалось обработать файл: {str(e)}")

    def _copy_from_text_widget(self, text_widget):
        try:
            selected = text_widget.get("sel.first", "sel.last")
            self.master.clipboard_clear()
            self.master.clipboard_append(selected)
        except tk.TclError:
            pass

    def _create_backup(self):
        success, message = self.password_manager.backup_passwords()
        if success:
            messagebox.showinfo("Успех", message)
        else:
            messagebox.showerror("Ошибка", message)

    def _show_about(self):
        messagebox.showinfo("О программе", "Менеджер паролей с шифрованием\n\n")

    def _import_passwords_txt(self):
        file_path = tk_filedialog.askopenfilename(
            title="Выберите файл для импорта",
            filetypes=[("Text files", "*.txt"), ("All files", "*.*")],
        )
        if not file_path:
            return
        try:
            with open(file_path, "r", encoding="utf-8") as f:
                preview = f.read(1000)
            if not messagebox.askyesno(
                "Подтверждение",
                f"Импортировать пароли из файла?\n\nПревью:\n{preview[:200]}...",
            ):
                return
            success, message = self.password_manager.import_from_txt(file_path)
            if success:
                messagebox.showinfo("Успех", message)
                self._refresh_password_list()
            else:
                messagebox.showerror("Ошибка", message)
        except Exception as e:
            messagebox.showerror("Ошибка", f"Не удалось прочитать файл: {str(e)}")

    def _generate_password(self):
        self.password_generator.password_length = int(self.length_var.get())
        self.password_generator.use_uppercase = self.use_uppercase.get()
        self.password_generator.use_lowercase = self.use_lowercase.get()
        self.password_generator.use_digits = self.use_digits.get()
        self.password_generator.use_symbols = self.use_symbols.get()
        self.password_generator.excluded_chars = self.excluded_chars_var.get().replace(
            " ", ""
        )
        new_password = self.password_generator.generate_password()
        self.password_var.set(new_password)
        self._update_strength_bar(new_password)
        try:
            if self._safe_copy_to_clipboard(new_password, clear_after_ms=60000):
                self._toast("Пароль скопирован")
            else:
                import pyperclip

                pyperclip.copy(new_password)
                self._toast("Пароль скопирован (pyperclip)")
        except Exception:
            pass

    def _update_strength_bar(self, password):
        score = self.password_generator.evaluate_password_strength(password)
        self.strength_var.set(score)
        if score < 40:
            self.strength_bar.config(style="red.Horizontal.TProgressbar")
        elif score < 70:
            self.strength_bar.config(style="yellow.Horizontal.TProgressbar")
        else:
            self.strength_bar.config(style="green.Horizontal.TProgressbar")

    def _copy_password(self):
        password = self.password_var.get()
        if not password:
            messagebox.showerror(
                "Ошибка", "Пароль отсутствует. Сначала сгенерируйте пароль."
            )
            return
        try:
            if self._safe_copy_to_clipboard(password, clear_after_ms=60000):
                self._toast("Пароль скопирован (очистится через 60 с)")
                return
        except Exception:
            pass
        try:
            import pyperclip

            pyperclip.copy(password)
            self._toast("Пароль скопирован (pyperclip)")
        except Exception as e:
            messagebox.showerror("Ошибка", f"Не удалось скопировать пароль: {e}")

    def _save_password_dialog(self):
        password = self.password_var.get()
        if not password:
            messagebox.showerror("Ошибка", "Сначала сгенерируйте пароль!")
            return
        initial_data = {"password": password, "description": "", "tags": [], "url": ""}
        dialog = PasswordDialog(self.master, "Сохранение пароля", initial_data)
        dialog.wait_window()
        if hasattr(dialog, "result") and dialog.result:
            data = dialog.result
            if not data["description"]:
                messagebox.showerror("Ошибка", "Заполните описание!")
                return
            self.password_manager.add_password(
                data["password"],
                data["description"],
                data.get("tags", []),
                data.get("url", ""),
            )
            messagebox.showinfo("Успешно", "Пароль сохранен!")
            self._refresh_password_list()

    def _add_new_password(self):
        dialog = PasswordDialog(self.master, "Добавить новый пароль")
        dialog.wait_window()
        if hasattr(dialog, "result") and dialog.result:
            data = dialog.result
            self.password_manager.add_password(
                data["password"],
                data["description"],
                data.get("tags", []),
                data.get("url", ""),
            )
            self._refresh_password_list()

    def _refresh_password_list(self):
        self.password_tree.delete(*self.password_tree.get_children())
        show_plain = not self.hide_passwords_var.get()
        if show_plain:
            self._cached_list = [
                {**p, "index": i}
                for i, p in enumerate(self.password_manager.get_all_passwords())
            ]
            self._cached_has_plain = True
        else:
            self._cached_list = self.password_manager.get_all_metadata()
            self._cached_has_plain = False
        for item in self._cached_list:
            display_password = item.get("password", "") if show_plain else "•" * 36
            self.password_tree.insert(
                "",
                tk.END,
                iid=str(item["index"]),
                values=(
                    item["description"],
                    display_password,
                    ", ".join(item.get("tags", [])),
                    item.get("updated_at", "—"),
                ),
            )
        self._update_tag_filter_options(self._cached_list)
        self._filter_passwords()

    def _update_tag_filter_options(self, items=None):
        items = items or self._cached_list
        tags = set()
        for it in items:
            tags.update(it.get("tags", []))
        all_tags = ["Все"] + sorted(tags)
        self.tag_filter["values"] = all_tags
        self.tag_filter.config(
            width=max(
                10, max(len(str(tag)) for tag in all_tags) + 2 if all_tags else 10
            )
        )

    def _filter_passwords(self):
        search_term = self.search_var.get().lower()
        selected_tag = self.tag_filter_var.get()
        search_in_passwords = (not self.hide_passwords_var.get()) and bool(search_term)
        if search_in_passwords and not self._cached_has_plain:
            self._cached_list = [
                {**p, "index": i}
                for i, p in enumerate(self.password_manager.get_all_passwords())
            ]
            self._cached_has_plain = True

        self.password_tree.delete(*self.password_tree.get_children())

        for it in self._cached_list:
            tags = [t.lower() for t in it.get("tags", [])]
            match_search = (
                not search_term
                or search_term in it["description"].lower()
                or (
                    search_in_passwords
                    and search_term in it.get("password", "").lower()
                )
                or any(search_term in t for t in tags)
            )
            match_tag = selected_tag == "Все" or selected_tag.lower() in tags
            if match_search and match_tag:
                display_password = (
                    ("•" * 64)
                    if self.hide_passwords_var.get()
                    else it.get("password", "")
                )
                self.password_tree.insert(
                    "",
                    tk.END,
                    iid=str(it["index"]),
                    values=(
                        it["description"],
                        display_password,
                        ", ".join(it.get("tags", [])),
                        it.get("updated_at", "—"),
                    ),
                )

    def _get_selected_password_index(self):
        selection = self.password_tree.selection()
        return int(selection[0]) if selection else None

    def _copy_selected_password(self):
        index = self._get_selected_password_index()
        if index is None:
            messagebox.showerror("Ошибка", "Выберите пароль для копирования.")
            return
        pwd = self.password_manager.get_password(index)
        if not pwd:
            messagebox.showerror("Ошибка", "Пароль не найден.")
            return
        if self._safe_copy_to_clipboard(pwd["password"], clear_after_ms=60000):
            self._toast("Пароль скопирован")
        else:
            messagebox.showerror("Ошибка", "Не удалось скопировать пароль.")

    def _view_selected_password(self):
        index = self._get_selected_password_index()
        if index is None:
            messagebox.showerror("Ошибка", "Выберите пароль для просмотра.")
            return
        pwd = self.password_manager.get_password(index)
        if pwd:
            dialog = PasswordDialog(self.master, "Просмотр пароля", pwd)
            dialog.paste_button.pack_forget()
            dialog.paste_button.destroy()
            dialog.password_entry.config(state="readonly")
            dialog.description_entry.config(state="readonly")
            dialog.tags_entry.config(state="readonly")
            dialog.url_entry.config(state="readonly")
            for child in dialog.winfo_children():
                if isinstance(child, ttk.Button):
                    child.destroy()
            url = pwd.get("url", "")
            if url.startswith(("http://", "https://")):
                link_frame = ttk.Frame(dialog)
                link_frame.grid(row=5, column=0, columnspan=2, pady=10, sticky="ew")
                ttk.Button(
                    link_frame,
                    text="Открыть ссылку",
                    command=lambda: QDesktopServices.openUrl(QUrl(url)),
                    cursor="hand2",
                ).pack(side=tk.TOP, fill=tk.X)
            ttk.Button(dialog, text="Закрыть", command=dialog.destroy).grid(
                row=6, column=0, columnspan=2, pady=10, sticky="ew"
            )

    def _edit_selected_password(self):
        index = self._get_selected_password_index()
        if index is None:
            messagebox.showerror("Ошибка", "Выберите пароль для редактирования.")
            return
        pwd = self.password_manager.get_password(index)
        if pwd:
            dialog = PasswordDialog(self.master, "Редактировать пароль", pwd)
            dialog.wait_window()
            if hasattr(dialog, "result") and dialog.result:
                data = dialog.result
                self.password_manager.update_password(
                    index,
                    data["password"],
                    data["description"],
                    data.get("tags", []),
                    data.get("url", ""),
                    data.get("login", ""),
                )
                self._refresh_password_list()

    def _delete_selected_password(self):
        index = self._get_selected_password_index()
        if index is None:
            messagebox.showerror("Ошибка", "Выберите пароль для удаления.")
            return
        if messagebox.askyesno(
            "Подтверждение", "Вы уверены, что хотите удалить этот пароль?"
        ):
            if self.password_manager.delete_password(index):
                messagebox.showinfo("Успех", "Пароль удален.")
                self._refresh_password_list()
            else:
                messagebox.showerror("Ошибка", "Не удалось удалить пароль.")

    def _copy_to_clipboard(self, widget):
        if widget.selection_present():
            self.master.clipboard_clear()
            self.master.clipboard_append(widget.selection_get())

    def _cut_to_clipboard(self, widget):
        if widget.selection_present():
            self._copy_to_clipboard(widget)
            widget.delete("sel.first", "sel.last")

    def _select_all(self, widget):
        if isinstance(widget, (tk.Entry, ttk.Entry)):
            widget.select_range(0, tk.END)
            widget.icursor(tk.END)
        elif isinstance(widget, tk.Text):
            widget.tag_add("sel", "1.0", tk.END)
            widget.mark_set("insert", tk.END)

    def _on_right_click(self, event):
        widget = event.widget
        if isinstance(widget, (tk.Entry, ttk.Entry)):
            widget.focus_set()
            try:
                widget.icursor(widget.index(f"@{event.x}"))
            except tk.TclError:
                pass
            self.entry_context_menu.tk_popup(event.x_root, event.y_root)

    def _on_tree_right_click(self, event):
        item = self.password_tree.identify_row(event.y)
        if item:
            self.password_tree.selection_set(item)
            self.tree_context_menu.tk_popup(event.x_root, event.y_root)


def heading_to_path(head):
    return head.replace(" ", "_").lower()


def run_password_manager():
    global messagebox
    _old_messagebox = messagebox
    messagebox = tk_messagebox
    try:
        root = tk.Tk()
        root.attributes("-topmost", True)
        root.lift()
        root.after(100, root.focus_force)
        app = PasswordGeneratorApp(root)
        root.mainloop()
    finally:
        messagebox = _old_messagebox


def run_notes_app():
    app = QApplication(sys.argv)
    window = NotesApp()
    window.show()
    app.exec()


class LauncherWindow(QMainWindow):
    def __init__(self):
        super().__init__()
        self.setWindowFlag(Qt.WindowStaysOnTopHint, True)
        self.settings = QSettings(SETTINGS_PATH, QSettings.IniFormat)
        sticky = self.settings.value("ui/launcher_always_on_top", True, type=bool)
        self.setWindowFlag(Qt.WindowStaysOnTopHint, sticky)
        self.notes_window = None
        self.notes_widget = None
        self.setWindowIcon(QIcon(ICON_PATH))
        central_widget = QWidget(self)
        layout = QVBoxLayout(central_widget)
        layout.setSpacing(20)
        layout.setContentsMargins(40, 40, 40, 40)
        btn_notes = QPushButton("Заметки")
        btn_notes.setMinimumHeight(40)
        btn_notes.clicked.connect(self.launch_notes)
        btn_notes.setFocusPolicy(Qt.NoFocus)
        layout.addWidget(btn_notes)
        btn_manager = QPushButton("Менеджер паролей")
        btn_manager.setMinimumHeight(40)
        btn_manager.clicked.connect(self.launch_password_manager)
        btn_manager.setFocusPolicy(Qt.NoFocus)
        layout.addWidget(btn_manager)
        btn_screenshot = QPushButton("Скриншотер")
        btn_screenshot.setMinimumHeight(40)
        btn_screenshot.setFocusPolicy(Qt.NoFocus)
        btn_screenshot.clicked.connect(self.launch_screenshoter)
        layout.addWidget(btn_screenshot)
        btn_tg_outgoing = QPushButton("Отчет: исходящие в Telegram")
        btn_tg_outgoing.setMinimumHeight(40)
        btn_tg_outgoing.setFocusPolicy(Qt.NoFocus)
        btn_tg_outgoing.clicked.connect(self.launch_tg_outgoing)
        layout.addWidget(btn_tg_outgoing)
        btn_testdata = QPushButton("Тестовые данные")
        btn_testdata.setMinimumHeight(40)
        btn_testdata.setFocusPolicy(Qt.NoFocus)
        btn_testdata.clicked.connect(self.launch_testdata)
        layout.addWidget(btn_testdata)
        btn_notes_widget = QPushButton("Виджет заметок", self)
        btn_notes_widget.setMinimumHeight(40)
        btn_notes_widget.setFocusPolicy(Qt.NoFocus)
        btn_notes_widget.clicked.connect(self.launch_desktop_notes)
        layout.addWidget(btn_notes_widget)
        btn_exit = QPushButton("Выход")
        btn_exit.setMinimumHeight(40)
        btn_exit.clicked.connect(sys.exit)
        btn_exit.setFocusPolicy(Qt.NoFocus)
        layout.addWidget(btn_exit)
        self.setCentralWidget(central_widget)
        self.setWindowTitle("Выберите приложение")
        self.setMinimumSize(400, 250)
        self.setStyleSheet(
            """
            QPushButton {
                font-size: 20px;
                border-radius: 10px;
                padding: 8px 0;
                outline: none;
                border: none;
                background-color: #f6f6f6;
                color: #222;
            }
            QPushButton:focus {
                outline: none;
                border: none;
                background-color: #e0e0e0;
                color: #222;
            }
            QPushButton:hover {
                background-color: #e0e0e0;
                color: #222;
            }
            """
        )

    def _start_external(self, title: str, candidates: list[str]) -> None:
        path = self._find_first_existing(candidates)
        if not path:
            QMessageBox.warning(
                self,
                title,
                "Не найден исполняемый файл.\nИскал:\n- " + "\n- ".join(candidates),
            )
            return

        proc = QProcess(self)
        proc.setProgram(path)
        proc.setArguments([])
        proc.setWorkingDirectory(os.path.dirname(path))

        def _on_finish(_exit_code: int, _exit_status):
            try:
                self._children.remove(proc)
            except ValueError:
                pass
            self.show()
            self.raise_()
            self.activateWindow()
            proc.deleteLater()

        proc.finished.connect(_on_finish)
        proc.errorOccurred.connect(lambda *_: _on_finish(0, None))
        proc.start()
        if not proc.waitForStarted(3000):
            QMessageBox.warning(self, title, "Не удалось запустить процесс.")
            try:
                proc.kill()
            except Exception:
                pass
            proc.deleteLater()
            return
        if not hasattr(self, "_children"):
            self._children = []
        self._children.append(proc)
        self.hide()

    def launch_notes(self):
        if getattr(self, "notes_window", None) is not None:
            try:
                if self.notes_window.isVisible():
                    if self.notes_window.isMinimized():
                        self.notes_window.showNormal()
                    self.notes_window.raise_()
                    self.notes_window.activateWindow()
                    return
            except RuntimeError:
                self.notes_window = None
        for w in QApplication.topLevelWidgets():
            try:
                if isinstance(w, NotesApp):
                    self.notes_window = w
                    if self.isVisible():
                        self.hide()
                    if w.isMinimized():
                        w.showNormal()
                    w.raise_()
                    w.activateWindow()
                    w.destroyed.connect(lambda *_: setattr(self, "notes_window", None))
                    return
            except Exception:
                pass
        if self.isVisible():
            self.hide()
        win = NotesApp()
        self.notes_window = win
        win.show()
        win.raise_()
        win.activateWindow()
        win.destroyed.connect(lambda *_: setattr(self, "notes_window", None))

    def on_notes_hidden(self):
        self.show()
        self.raise_()
        self.activateWindow()

    def launch_password_manager(self):
        self.hide()
        run_password_manager()
        self.show()
        self.raise_()
        self.activateWindow()

    def launch_desktop_notes(self):
        try:
            existing_widget = None
            if self.notes_widget is not None and self.notes_widget.isVisible():
                existing_widget = self.notes_widget
            else:
                self.notes_widget = None
                for w in QApplication.topLevelWidgets():
                    try:
                        if isinstance(w, DesktopNotesWidget):
                            existing_widget = w
                            self.notes_widget = w
                            w.destroyed.connect(lambda: setattr(self, "notes_widget", None))
                            break
                    except Exception:
                        pass
            if existing_widget is not None:
                if self.isVisible():
                    self.hide()
                if existing_widget.isMinimized():
                    existing_widget.showNormal()
                existing_widget.raise_()
                existing_widget.activateWindow()
                return
            self.notes_widget = DesktopNotesWidget()
            self.notes_widget.destroyed.connect(lambda: setattr(self, "notes_widget", None))
            if self.isVisible():
                self.hide()
            self.notes_widget.show()
            self.notes_widget.raise_()
            self.notes_widget.activateWindow()
        except Exception as e:
            try:
                with open(os.path.join(DATA_DIR, "last_run.log"), "a", encoding="utf-8") as f:
                    f.write(f"DesktopNotesWidget error: {e}\n")
                    import traceback; traceback.print_exc(file=f)
            except Exception:
                pass
            QMessageBox.critical(self, "Ошибка", f"Не удалось открыть Заметки:\n{e}")
        
    def launch_screenshoter(self):
        self._start_external("Скриншотер", ["ScreenshotPN.exe"])

    def launch_tg_outgoing(self):
        self._start_external(
            "Отчет: исходящие в Telegram", ["TelegramOutgoingReporter.exe"]
        )

    def launch_testdata(self):
        self._start_external("Тестовые данные", ["TestData.exe"])

    def _find_first_existing(self, names: list[str]) -> str | None:
        dirs = []
        if APPDIR:
            dirs.append(APPDIR)
        try:
            if getattr(sys, "frozen", False):
                dirs.append(os.path.dirname(sys.executable))
            else:
                dirs.append(os.path.dirname(os.path.abspath(sys.argv[0])))
        except Exception:
            pass
        try:
            dirs.append(os.path.dirname(os.path.abspath(__file__)))
        except Exception:
            pass
        try:
            dirs.append(os.getcwd())
        except Exception:
            pass
        norm_dirs = []
        seen = set()
        for d in dirs:
            if not d:
                continue
            nd = os.path.abspath(os.path.realpath(d))
            if nd not in seen:
                seen.add(nd)
                norm_dirs.append(nd)
        for d in norm_dirs:
            for name in names:
                path = os.path.join(d, name)
                if os.path.isfile(path):
                    return path
        return None


if __name__ == "__main__":
    app = QApplication(sys.argv)
    settings = QSettings("NotesPM", "NotesPM")
    default_start = "notes" if getattr(sys, "frozen", False) else "launcher"
    start_app = settings.value("startup_app", default_start)
    if any(arg.lower() == "--notes" for arg in sys.argv):
        start_app = "notes"
    elif any(arg.lower() == "--launcher" for arg in sys.argv):
        start_app = "launcher"
    if start_app == "notes":
        win = NotesApp()
        win.show()
    else:
        win = LauncherWindow()
        app.launcher_window = win
        win.show()
    sys.exit(app.exec())

# UPD 20.11.2025
